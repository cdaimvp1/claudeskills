"""
evaluation_report_generator.py
Lilly Procurement Skills, evaluation-engine-1c344a deterministic DOCX generator.

Purpose (mirrors should-cost-builder-1c344a/should_cost_generator.py and
pro-forma-builder-1c344a/pro_forma_generator.py's "generator scripts, not
freehand authoring" pattern): this module takes a validated evaluation
register (suppliers, weighted scoring criteria, per-criterion per-supplier
scores, risk, commercial figures, and recommendation inputs, shaped per
evaluation-engine-1c344a/SKILL.md's "Evaluation Schemas" reference) and
MECHANICALLY produces a professional evaluation_report.docx: title page,
Executive Summary, Evaluation Methodology, Supplier Summaries, Score
Comparison, Risk Assessment Summary, Commercial Comparison, Recommendation,
and Appendices, exactly the eight sections SKILL.md's evaluation-schemas
reference requires, in the same fixed order every run.

Two-stage discipline (same as should-cost/pro-forma):
  1. Every supplier's Grand Total is computed by calling the vendored
     numeric_kernel.py's weighted_score() (never re-derived by hand or model
     arithmetic). Per SKILL.md's "HARD RULE, kernel-computed weighted score"
     (Scoring Rules section): "The grand total... and the weight-sum-to-1.0
     validation... are computed by calling weighted_score() in the vendored
     numeric_kernel.py... never by model arithmetic. weighted_score() refuses
     ... if the weights passed to it do not sum to 1.0 within tolerance."
  2. The per-criterion weighted contribution shown in the Scoring Detail /
     Appendix A table (Final_Score x Effective_Weight_Frac for that single
     row) is, per the same HARD RULE paragraph, "the deterministic
     single-term readout of the same already-validated score and weight that
     fed the kernel call, not an independently estimated figure." This
     module computes that single multiplication directly, then asserts
     (before saving) that the sum of every supplier's per-criterion
     contributions reconciles exactly to that supplier's kernel-computed
     Grand Total, the same "numbers-reconcile assertion" the Canonical
     Dashboard section requires before rendering.

Hard-coded invariants (code-level checks, not comments; mirrors should-cost's
run_hardcoded_invariant_checks() pattern):
  - Effective-weight-sums-to-one invariant: Score Validation Checks table,
    "Effective weight sum: All Effective_Weight_Frac values must sum to 1.0
    +/- 0.001 across every criterion."
  - Weighted-contributions-reconcile invariant: the per-criterion weighted
    contributions for a supplier must sum exactly to that supplier's
    kernel-computed Grand Total (Canonical Dashboard "Numbers-reconcile
    assertion").
  - Grand-total-range invariant: "Grand total range: All Grand Total scores
    must be 0.0-5.0" (Score Validation Checks table).
  - Category-subtotals-reconcile invariant: a supplier's category subtotals
    (CATEGORY_TOTAL_{category} footer rows in final_score_rollup.csv) must
    sum exactly to that supplier's Grand Total.
  - Ranking-sorted invariant: the compensatory ranking is sorted strictly by
    Grand Total, descending.
  - Must-Have-gate-not-buried invariant: Recommendation Logic's
    gate-vs-compensatory reconciliation, "If the highest weighted-total
    supplier carries a Must Have zero, do NOT silently recommend it on total
    alone... present both the compensatory ranking AND the gate-pass
    ranking."
  - Requirements-coverage-counts-foot invariant: every supplier's Fully
    Meets / Partially Meets / Does Not Meet / Not Answered tally must sum to
    the total criterion count (Section 3's "Requirements coverage summary").
  - Sensitivity-rank1-base-matches-ground-truth invariant: every row of the
    (upstream-supplied) sensitivity table's "Rank 1 (Base)" column must
    equal this module's own computed compensatory rank-1 supplier, so a
    stale or drifted sensitivity table cannot silently disagree with the
    scoring this report actually shows.

If any invariant fails, this script RAISES rather than writing a document
that fails its own reconciliation.

Scope note: this script builds the evaluation_report.docx generator ONLY,
per SKILL.md's Evaluation Schemas "1. evaluation_report.docx" section. The
other mandated artifacts (final_score_rollup.csv, ai_scoring.csv,
sensitivity_matrix.csv, supplier_evaluation_ui.json, the six-tab canonical
dashboard, and the communication DOCX outputs such as award_notification.docx)
are separate deliverables and are NOT produced here.

============================================================================
Evaluation register input contract (JUDGMENT CALL, flagged): SKILL.md's
Evaluation Schemas reference gives the OUTPUT column shapes for
final_score_rollup.csv, ai_scoring.csv, requirements_response_mapping.csv,
and supplier_evaluation_ui.json, but not a single raw INPUT contract for a
report generator. This module resolves that ambiguity as follows.

Score provenance (JUDGMENT CALL, flagged): the SKILL.md Aggregation
Formulas define Final_Score as the already-blended combination of AI and
stakeholder scores (default weights AI 0.4 / Stakeholder 0.6, adjustable),
computed in evaluation-engine's own Phase 6. That blend is this skill's own
upstream analytical judgment call, not something a deterministic document
renderer should re-derive from raw AI/stakeholder columns (the same
"consume, don't re-derive" discipline should-cost-builder and pro-forma-
builder already apply to each other's sourced figures). This generator
therefore accepts one already-aggregated `score` (0.0-5.0) per criterion per
supplier, i.e. the final_score_rollup.csv `{supplier_id}_Final_Score` column,
not separate AI/stakeholder columns.

Sensitivity provenance (JUDGMENT CALL, flagged): per Phase 7, the
weight-perturbation matrix is computed upstream (test +/-5pp per category,
redistributed proportionally, re-scoring every supplier under the perturbed
weights). Re-implementing that full recomputation inside a document
renderer would duplicate evaluation-engine's own Phase 7 logic outside this
generator's stated scope (build the DOCX from already-decided figures).
This module accepts the already-computed sensitivity rows as input and
enforces exactly one reconciliation check on them: every row's "Rank 1
(Base)" supplier must match this module's own independently kernel-computed
compensatory rank-1 supplier (protects against upstream drift without
re-deriving Phase 7's weight-redistribution math).

Register shape:
{
  "rfx_name": "string", "category": "string", "case_id": "string (optional)",
  "evaluation_date": "YYYY-MM-DD",
  "scoring_matrix_source": "Provided" | "Generated",
  "ai_scoring_enabled": true | false,
  "stakeholder_scoring_included": true | false,
  "stakeholder_scorer_count": int (required if stakeholder_scoring_included),
  "report_mode": "Brief" | "Full" (optional, default "Full"),
  "suppliers": [ {"supplier_id": "SUP_A", "supplier_name": "string", "snapshot": "free text (optional)"}, ... ],
  "criteria": [
    {
      "criterion_id": "REQ-001", "category": "Functional Fit",
      "criterion_text": "full requirement text",
      "priority": "Must Have" | "Should Have" | "Nice to Have",
      "weight_in_category_frac": float,   # fractions within a category sum to 1.0
      "category_weight_frac": float,      # identical across every criterion sharing this category; distinct categories sum to 1.0
      "entries": {
        "SUP_A": {"score": float 0.0-5.0, "coverage": "Fully Meets"|"Partially Meets"|"Does Not Meet"|"Not Answered",
                   "confidence": "High"|"Medium"|"Low", "evidence": "doc + location", "rationale": "<=2 sentences"},
        ...  # exactly one entry per declared supplier_id, no omissions
      }
    }, ...
  ],
  "commercial": {"SUP_A": {"year1_cost_usd": float, "tcv_3yr_usd": float, "pricing_model": "string", "notes": "string"}, ...},
  "risk_matrix": {"Legal/Contractual": {"SUP_A": "Low"|"Medium"|"High", ...}, "Cyber/Security": {...}, ...},
  "risk_notes": {"SUP_A": "free text", ...},
  "strengths": {"SUP_A": ["...", "...", "..."], ...},
  "weaknesses": {"SUP_A": ["...", "...", "..."], ...},
  "sensitivity_rows": [
    {"category_adjusted": "Functional Fit", "base_weight_frac": 0.40, "test_weight_frac": 0.45,
     "rank1_base": "SUP_A", "rank1_adjusted": "SUP_A", "ranking_changed": false}, ...
  ],
  "recommendation": {"secondary_supplier_id": "SUP_C" | null, "secondary_tradeoffs": "string",
                      "key_tradeoffs": "string", "next_steps": ["string", ...]},
  "clarifications_note": "string (optional)"
}

The primary recommended supplier, its rationale text, Grand Totals, ranking,
and every Must Have gate flag are COMPUTED by this generator, never accepted
as input, per SKILL.md's Scoring Authority statement ("this skill is the
sole owner of the official score, the official weighted total... and the
official award recommendation").
============================================================================
"""

from __future__ import annotations

import copy
import os
import sys
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Sequence, Tuple

# ---------------------------------------------------------------------------
# Vendored numeric kernel (same directory). Per SKILL.md's "HARD RULE,
# kernel-computed weighted score": Grand Total MUST be computed by calling
# weighted_score() in the vendored numeric_kernel.py, never hand-computed.
# ---------------------------------------------------------------------------
_KERNEL_IMPORT_ERROR: Optional[Exception] = None
try:
    _THIS_DIR = os.path.dirname(os.path.abspath(__file__))
    if _THIS_DIR not in sys.path:
        sys.path.insert(0, _THIS_DIR)
    from numeric_kernel import weighted_score as kernel_weighted_score
    from numeric_kernel import WeightSumError as KernelWeightSumError
    from numeric_kernel import InvalidInputError as KernelInvalidInputError
    KERNEL_AVAILABLE = True
except Exception as _exc:  # pragma: no cover - defensive, disclosed at runtime
    KERNEL_AVAILABLE = False
    _KERNEL_IMPORT_ERROR = _exc

    def kernel_weighted_score(*args, **kwargs):  # type: ignore
        raise RuntimeError(
            "numeric_kernel.py unavailable; cannot compute the kernel-required "
            f"weighted score. Import error: {_KERNEL_IMPORT_ERROR}"
        )

    class KernelWeightSumError(Exception):  # type: ignore
        pass

    class KernelInvalidInputError(Exception):  # type: ignore
        pass

# ---------------------------------------------------------------------------
# python-docx detection. Mirrors should_cost_generator.py's openpyxl
# handling: try the import; if unavailable, raise a clear ImportError at
# document-build time (not at module-import time), so validation and
# ground-truth logic (which need no DOCX library) remain usable regardless.
# ---------------------------------------------------------------------------
try:
    import docx  # noqa: F401
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
    DOCX_VERSION = getattr(docx, "__version__", "unknown")
except Exception as _exc:  # pragma: no cover
    DOCX_AVAILABLE = False
    DOCX_VERSION = None
    _DOCX_IMPORT_ERROR = _exc


class EvaluationValidationError(Exception):
    """Raised when the evaluation register is missing a required field or
    carries a value the skill's Accuracy and Anti-Drift Rules forbid. Per
    Rule 1 ("Never fabricate scores or evaluation data") this module refuses
    rather than guessing a missing or malformed value."""


class ReconciliationError(Exception):
    """Raised when a hard-coded invariant fails (effective-weight-sum,
    weighted-contributions-reconcile, grand-total-range, category-subtotals-
    reconcile, ranking-sorted, must-have-gate-not-buried, requirements-
    coverage-counts-foot, or sensitivity-rank1-base-matches). The document
    is not saved when this fires."""


# ===========================================================================
# House style constants (quoted verbatim from lilly-brand-assets-1c344a's
# brand-colors.md and docx-design-system.md, hardcoded here exactly as
# should_cost_generator.py hardcodes "Lilly Black #212121... Bold Blue
# #0F3A85" from its own SKILL.md's house-style block; this keeps the
# generator self-contained and deterministic rather than reading an external
# file at build time).
# ===========================================================================
LILLY_RED = "E1251B"      # table header rows, title-page accents, High severity
BOLD_BLUE = "0F3A85"      # section header text (H1), Low severity ("safe" token, no-green rule)
LILLY_BLACK = "212121"    # all body text
AMBER = "B45309"          # Medium severity
MUTED_GREY = "8A969E"     # footer/header chrome text
WHITE = "FFFFFF"          # table header text on colored fill

SEVERITY_COLOR = {"Low": BOLD_BLUE, "Medium": AMBER, "High": LILLY_RED}

VALID_PRIORITIES = ("Must Have", "Should Have", "Nice to Have")
VALID_COVERAGE = ("Fully Meets", "Partially Meets", "Does Not Meet", "Not Answered")
VALID_CONFIDENCE = ("High", "Medium", "Low")
VALID_SEVERITY = ("Low", "Medium", "High")
VALID_MATRIX_SOURCE = ("Provided", "Generated")
VALID_REPORT_MODE = ("Brief", "Full")


# ===========================================================================
# 1. Evaluation register: typed input + validation
# ===========================================================================

@dataclass
class SupplierProfile:
    supplier_id: str
    supplier_name: str
    snapshot: str = ""


@dataclass
class CriterionEntry:
    score: float
    coverage: str
    confidence: str
    evidence: str = ""
    rationale: str = ""


@dataclass
class Criterion:
    criterion_id: str
    category: str
    criterion_text: str
    priority: str
    weight_in_category_frac: float
    category_weight_frac: float
    entries: Dict[str, CriterionEntry]


@dataclass
class CommercialFigures:
    year1_cost_usd: float
    tcv_3yr_usd: float
    pricing_model: str
    notes: str = ""


@dataclass
class SensitivityRow:
    category_adjusted: str
    base_weight_frac: float
    test_weight_frac: float
    rank1_base: str
    rank1_adjusted: str
    ranking_changed: bool


@dataclass
class RecommendationInputs:
    secondary_supplier_id: Optional[str]
    secondary_tradeoffs: str
    key_tradeoffs: str
    next_steps: List[str]


@dataclass
class EvaluationInput:
    rfx_name: str
    category: str
    case_id: str
    evaluation_date: str
    scoring_matrix_source: str
    ai_scoring_enabled: bool
    stakeholder_scoring_included: bool
    stakeholder_scorer_count: int
    report_mode: str
    suppliers: List[SupplierProfile]
    criteria: List[Criterion]
    commercial: Dict[str, CommercialFigures]
    risk_matrix: Dict[str, Dict[str, str]]
    risk_notes: Dict[str, str]
    strengths: Dict[str, List[str]]
    weaknesses: Dict[str, List[str]]
    sensitivity_rows: List[SensitivityRow]
    recommendation: RecommendationInputs
    clarifications_note: Optional[str]


REQUIRED_TOP_LEVEL_FIELDS = [
    "rfx_name", "category", "evaluation_date", "scoring_matrix_source",
    "ai_scoring_enabled", "stakeholder_scoring_included", "suppliers", "criteria",
]
REQUIRED_ENTRY_FIELDS = ["score", "coverage", "confidence"]


def validate_evaluation_input(register: Dict[str, Any]) -> EvaluationInput:
    """Validate a raw evaluation register dict and return a typed
    EvaluationInput. Refuses (raises EvaluationValidationError) rather than
    guessing when a required field is missing or a value violates a
    documented, fixed convention, per Rule 1 ("Never fabricate scores or
    evaluation data") and Rule 2 ("Never invent evaluation criteria")."""
    errors: List[str] = []

    for f in REQUIRED_TOP_LEVEL_FIELDS:
        if f not in register:
            errors.append(f"Missing required top-level field: '{f}'")
    if errors:
        raise EvaluationValidationError(
            "Evaluation register is missing required fields; refusing to "
            "guess. " + "; ".join(errors)
        )

    rfx_name = str(register["rfx_name"]).strip()
    category = str(register["category"]).strip()
    case_id = str(register.get("case_id", "")).strip()
    evaluation_date = register["evaluation_date"]
    if not isinstance(evaluation_date, str) or not evaluation_date:
        errors.append("'evaluation_date' must be a non-empty string (e.g. '2026-07-23').")

    matrix_source_raw = str(register["scoring_matrix_source"]).strip()
    matrix_source = next((v for v in VALID_MATRIX_SOURCE if v.lower() == matrix_source_raw.lower()), None)
    if matrix_source is None:
        errors.append(f"'scoring_matrix_source' must be one of {VALID_MATRIX_SOURCE}; got {matrix_source_raw!r}.")

    ai_scoring_enabled = register["ai_scoring_enabled"]
    if not isinstance(ai_scoring_enabled, bool):
        errors.append("'ai_scoring_enabled' must be a boolean.")

    stakeholder_scoring_included = register["stakeholder_scoring_included"]
    if not isinstance(stakeholder_scoring_included, bool):
        errors.append("'stakeholder_scoring_included' must be a boolean.")

    stakeholder_scorer_count = register.get("stakeholder_scorer_count", 0)
    if stakeholder_scoring_included:
        if not isinstance(stakeholder_scorer_count, int) or stakeholder_scorer_count < 1:
            errors.append(
                "'stakeholder_scorer_count' must be a positive integer when "
                "'stakeholder_scoring_included' is true."
            )
    else:
        stakeholder_scorer_count = 0

    report_mode_raw = str(register.get("report_mode", "Full")).strip()
    report_mode = next((v for v in VALID_REPORT_MODE if v.lower() == report_mode_raw.lower()), None)
    if report_mode is None:
        errors.append(f"'report_mode' must be one of {VALID_REPORT_MODE}; got {report_mode_raw!r}.")

    # --- Suppliers ---------------------------------------------------------
    raw_suppliers = register.get("suppliers", [])
    if not isinstance(raw_suppliers, list) or len(raw_suppliers) == 0:
        errors.append("'suppliers' must be a non-empty list.")
        raw_suppliers = []
    suppliers: List[SupplierProfile] = []
    supplier_ids_seen: List[str] = []
    for i, s in enumerate(raw_suppliers):
        if not isinstance(s, dict) or "supplier_id" not in s or "supplier_name" not in s:
            errors.append(f"'suppliers[{i}]' must be an object with 'supplier_id' and 'supplier_name'.")
            continue
        sid = str(s["supplier_id"]).strip()
        if sid in supplier_ids_seen:
            errors.append(f"Duplicate supplier_id '{sid}' in 'suppliers'.")
        supplier_ids_seen.append(sid)
        suppliers.append(SupplierProfile(
            supplier_id=sid, supplier_name=str(s["supplier_name"]).strip(),
            snapshot=str(s.get("snapshot", "")).strip(),
        ))
    supplier_id_set = set(supplier_ids_seen)

    # --- Criteria (and per-criterion, per-supplier entries) ----------------
    raw_criteria = register.get("criteria", [])
    if not isinstance(raw_criteria, list) or len(raw_criteria) == 0:
        errors.append(
            "'criteria' must be a non-empty list of scoring-matrix rows "
            "(criterion_id, category, criterion_text, priority, "
            "weight_in_category_frac, category_weight_frac, entries)."
        )
        raw_criteria = []

    criteria: List[Criterion] = []
    criterion_ids_seen: List[str] = []
    category_weight_by_category: Dict[str, float] = {}
    weight_in_category_sum: Dict[str, float] = {}

    for i, c in enumerate(raw_criteria):
        if not isinstance(c, dict):
            errors.append(f"'criteria[{i}]' must be an object.")
            continue
        required = ["criterion_id", "category", "criterion_text", "priority",
                    "weight_in_category_frac", "category_weight_frac", "entries"]
        missing = [f for f in required if f not in c]
        if missing:
            errors.append(f"'criteria[{i}]' is missing required field(s): {missing}.")
            continue

        cid = str(c["criterion_id"]).strip()
        if cid in criterion_ids_seen:
            errors.append(f"Duplicate criterion_id '{cid}' in 'criteria'.")
        criterion_ids_seen.append(cid)

        cat = str(c["category"]).strip()
        priority = str(c["priority"]).strip()
        if priority not in VALID_PRIORITIES:
            errors.append(f"'criteria[{i}]' ({cid}) priority must be one of {VALID_PRIORITIES}; got {priority!r}.")

        wic = c["weight_in_category_frac"]
        cw = c["category_weight_frac"]
        if not isinstance(wic, (int, float)) or not isinstance(cw, (int, float)):
            errors.append(f"'criteria[{i}]' ({cid}) weight_in_category_frac / category_weight_frac must be numeric.")
            wic, cw = 0.0, 0.0
        else:
            if wic > 1.0 or cw > 1.0:
                errors.append(
                    f"'criteria[{i}]' ({cid}) carries a weight > 1.0 (weight_in_category_frac={wic}, "
                    f"category_weight_frac={cw}). Weights must be stored as FRACTIONS in [0.0, 1.0], never "
                    "percentages (Score Validation Checks: 'Weight format: No weight stored as a percentage')."
                )
            if wic < 0.0 or cw < 0.0:
                errors.append(f"'criteria[{i}]' ({cid}) carries a negative weight; weights must be >= 0.0.")

        if cat in category_weight_by_category and abs(category_weight_by_category[cat] - cw) > 1e-9:
            errors.append(
                f"'criteria[{i}]' ({cid}) category_weight_frac ({cw}) disagrees with an earlier criterion in "
                f"category {cat!r} (previously {category_weight_by_category[cat]}). category_weight_frac must "
                "be identical for every criterion sharing a category."
            )
        category_weight_by_category.setdefault(cat, cw)
        weight_in_category_sum[cat] = weight_in_category_sum.get(cat, 0.0) + wic

        raw_entries = c["entries"]
        entries: Dict[str, CriterionEntry] = {}
        if not isinstance(raw_entries, dict):
            errors.append(f"'criteria[{i}]' ({cid}) 'entries' must be an object keyed by supplier_id.")
        else:
            entry_ids = set(raw_entries.keys())
            if entry_ids != supplier_id_set:
                missing_ids = supplier_id_set - entry_ids
                extra_ids = entry_ids - supplier_id_set
                errors.append(
                    f"'criteria[{i}]' ({cid}) 'entries' must have exactly one entry per declared supplier_id. "
                    f"Missing: {sorted(missing_ids) or 'none'}; unexpected: {sorted(extra_ids) or 'none'}."
                )
            for sid, e in raw_entries.items():
                if not isinstance(e, dict):
                    errors.append(f"'criteria[{i}]' ({cid}) entries[{sid!r}] must be an object.")
                    continue
                emissing = [f for f in REQUIRED_ENTRY_FIELDS if f not in e]
                if emissing:
                    errors.append(f"'criteria[{i}]' ({cid}) entries[{sid!r}] is missing field(s): {emissing}.")
                    continue
                score = e["score"]
                coverage = str(e["coverage"]).strip()
                confidence = str(e["confidence"]).strip()
                if not isinstance(score, (int, float)) or not (0.0 <= float(score) <= 5.0):
                    errors.append(
                        f"'criteria[{i}]' ({cid}) entries[{sid!r}].score must be numeric in [0.0, 5.0]; got {score!r}."
                    )
                    score = 0.0
                if coverage not in VALID_COVERAGE:
                    errors.append(
                        f"'criteria[{i}]' ({cid}) entries[{sid!r}].coverage must be one of {VALID_COVERAGE}; "
                        f"got {coverage!r}."
                    )
                if confidence not in VALID_CONFIDENCE:
                    errors.append(
                        f"'criteria[{i}]' ({cid}) entries[{sid!r}].confidence must be one of {VALID_CONFIDENCE}; "
                        f"got {confidence!r}."
                    )
                entries[sid] = CriterionEntry(
                    score=float(score), coverage=coverage, confidence=confidence,
                    evidence=str(e.get("evidence", "")), rationale=str(e.get("rationale", "")),
                )

        criteria.append(Criterion(
            criterion_id=cid, category=cat, criterion_text=str(c["criterion_text"]).strip(),
            priority=priority, weight_in_category_frac=float(wic), category_weight_frac=float(cw),
            entries=entries,
        ))

    for cat, total in weight_in_category_sum.items():
        if abs(total - 1.0) > 0.001:
            errors.append(
                f"Criterion weight sum FAILED for category {cat!r}: weight_in_category_frac values sum to "
                f"{total:.4f}, not 1.0 (+/- 0.001 per Score Validation Checks)."
            )
    category_weight_total = sum(category_weight_by_category.values())
    if category_weight_by_category and abs(category_weight_total - 1.0) > 0.001:
        errors.append(
            f"Category weight sum FAILED: category_weight_frac values sum to {category_weight_total:.4f}, "
            "not 1.0 (+/- 0.001 per Score Validation Checks)."
        )

    # --- Commercial ----------------------------------------------------
    commercial: Dict[str, CommercialFigures] = {}
    raw_commercial = register.get("commercial", {})
    if not isinstance(raw_commercial, dict):
        errors.append("'commercial', when present, must be an object keyed by supplier_id.")
    else:
        for sid, cf in raw_commercial.items():
            if sid not in supplier_id_set:
                errors.append(f"'commercial' references unknown supplier_id {sid!r}.")
                continue
            if not isinstance(cf, dict) or not all(k in cf for k in ("year1_cost_usd", "tcv_3yr_usd", "pricing_model")):
                errors.append(f"'commercial[{sid!r}]' must be an object with year1_cost_usd, tcv_3yr_usd, pricing_model.")
                continue
            commercial[sid] = CommercialFigures(
                year1_cost_usd=float(cf["year1_cost_usd"]), tcv_3yr_usd=float(cf["tcv_3yr_usd"]),
                pricing_model=str(cf["pricing_model"]), notes=str(cf.get("notes", "")),
            )

    # --- Risk ------------------------------------------------------------
    risk_matrix: Dict[str, Dict[str, str]] = {}
    raw_risk = register.get("risk_matrix", {})
    if not isinstance(raw_risk, dict):
        errors.append("'risk_matrix', when present, must be an object keyed by risk category.")
    else:
        for rcat, per_supplier in raw_risk.items():
            if not isinstance(per_supplier, dict):
                errors.append(f"'risk_matrix[{rcat!r}]' must be an object keyed by supplier_id.")
                continue
            row: Dict[str, str] = {}
            for sid, sev in per_supplier.items():
                if sid not in supplier_id_set:
                    errors.append(f"'risk_matrix[{rcat!r}]' references unknown supplier_id {sid!r}.")
                    continue
                sev_s = str(sev).strip()
                if sev_s not in VALID_SEVERITY:
                    errors.append(f"'risk_matrix[{rcat!r}][{sid!r}]' must be one of {VALID_SEVERITY}; got {sev_s!r}.")
                    continue
                row[sid] = sev_s
            risk_matrix[str(rcat)] = row

    risk_notes = {str(k): str(v) for k, v in register.get("risk_notes", {}).items()}
    strengths = {str(k): [str(x) for x in v][:3] for k, v in register.get("strengths", {}).items()}
    weaknesses = {str(k): [str(x) for x in v][:3] for k, v in register.get("weaknesses", {}).items()}

    # --- Sensitivity -----------------------------------------------------
    sensitivity_rows: List[SensitivityRow] = []
    for i, sr in enumerate(register.get("sensitivity_rows", [])):
        if not isinstance(sr, dict):
            errors.append(f"'sensitivity_rows[{i}]' must be an object.")
            continue
        required = ["category_adjusted", "base_weight_frac", "test_weight_frac", "rank1_base", "rank1_adjusted", "ranking_changed"]
        missing = [f for f in required if f not in sr]
        if missing:
            errors.append(f"'sensitivity_rows[{i}]' is missing field(s): {missing}.")
            continue
        if sr["rank1_base"] not in supplier_id_set or sr["rank1_adjusted"] not in supplier_id_set:
            errors.append(f"'sensitivity_rows[{i}]' rank1_base/rank1_adjusted must reference declared supplier_ids.")
            continue
        sensitivity_rows.append(SensitivityRow(
            category_adjusted=str(sr["category_adjusted"]), base_weight_frac=float(sr["base_weight_frac"]),
            test_weight_frac=float(sr["test_weight_frac"]), rank1_base=str(sr["rank1_base"]),
            rank1_adjusted=str(sr["rank1_adjusted"]), ranking_changed=bool(sr["ranking_changed"]),
        ))

    # --- Recommendation inputs --------------------------------------------
    raw_rec = register.get("recommendation", {})
    if not isinstance(raw_rec, dict):
        errors.append("'recommendation' must be an object.")
        raw_rec = {}
    secondary_id = raw_rec.get("secondary_supplier_id")
    if secondary_id is not None and secondary_id not in supplier_id_set:
        errors.append(f"'recommendation.secondary_supplier_id' ({secondary_id!r}) is not a declared supplier_id.")
    secondary_tradeoffs = str(raw_rec.get("secondary_tradeoffs", ""))
    if secondary_id and not secondary_tradeoffs.strip():
        errors.append("'recommendation.secondary_tradeoffs' is required when secondary_supplier_id is set.")
    key_tradeoffs = str(raw_rec.get("key_tradeoffs", ""))
    if not key_tradeoffs.strip():
        errors.append("'recommendation.key_tradeoffs' must be a non-empty string.")
    next_steps = raw_rec.get("next_steps", [])
    if not isinstance(next_steps, list) or len(next_steps) == 0:
        errors.append("'recommendation.next_steps' must be a non-empty list of strings.")
        next_steps = []

    clarifications_note = register.get("clarifications_note")
    if clarifications_note is not None:
        clarifications_note = str(clarifications_note)

    if errors:
        raise EvaluationValidationError(
            "Evaluation register failed validation; refusing to guess missing "
            "or invalid fields. Issues found:\n  - " + "\n  - ".join(errors)
        )

    return EvaluationInput(
        rfx_name=rfx_name, category=category, case_id=case_id, evaluation_date=evaluation_date,
        scoring_matrix_source=matrix_source, ai_scoring_enabled=ai_scoring_enabled,
        stakeholder_scoring_included=stakeholder_scoring_included,
        stakeholder_scorer_count=stakeholder_scorer_count, report_mode=report_mode,
        suppliers=suppliers, criteria=criteria, commercial=commercial, risk_matrix=risk_matrix,
        risk_notes=risk_notes, strengths=strengths, weaknesses=weaknesses,
        sensitivity_rows=sensitivity_rows,
        recommendation=RecommendationInputs(
            secondary_supplier_id=secondary_id, secondary_tradeoffs=secondary_tradeoffs,
            key_tradeoffs=key_tradeoffs, next_steps=[str(s) for s in next_steps],
        ),
        clarifications_note=clarifications_note,
    )


# ===========================================================================
# Ground-truth computation (via the vendored kernel only, per the HARD RULE)
# ===========================================================================

@dataclass
class GroundTruth:
    effective_weight: Dict[str, float]                       # criterion_id -> effective weight (fraction)
    category_weight: Dict[str, float]                        # category -> category weight (fraction)
    weighted_contribution: Dict[str, Dict[str, float]]        # supplier_id -> {criterion_id: contribution}
    category_subtotal: Dict[str, Dict[str, float]]            # supplier_id -> {category: subtotal}
    grand_total: Dict[str, float]                             # supplier_id -> kernel-computed grand total
    must_have_zero_flag: Dict[str, bool]                      # supplier_id -> flag
    gate_detail: Dict[str, List[str]]                         # supplier_id -> [criterion_id, ...]
    compensatory_rank: List[str]                              # supplier_ids, sorted by grand_total desc
    gate_clear_rank: List[str]                                # subset, sorted by grand_total desc
    recommended_supplier_id: str
    gate_conflict: bool
    all_suppliers_gate_failed: bool
    coverage_counts: Dict[str, Dict[str, int]]                # supplier_id -> {coverage_label: count, "Total": n}
    delta_to_top: Dict[str, float]                            # supplier_id -> grand_total[top] - grand_total[sid]
    tie_flag: Dict[str, bool]                                 # supplier_id -> within 0.1 of the rank immediately above
    sensitivity_verdict: str


def compute_ground_truth(reg: EvaluationInput) -> GroundTruth:
    """Compute every derived figure this report shows, calling ONLY
    numeric_kernel.weighted_score() for each supplier's Grand Total, per
    SKILL.md's Scoring Rules HARD RULE. This is the reference against which
    the hard-coded invariants are checked before the document is saved."""
    if not KERNEL_AVAILABLE:
        raise RuntimeError(
            "numeric_kernel.py could not be imported; this generator refuses "
            "to hand-compute the weighted score in its place (SKILL.md "
            f"'HARD RULE, kernel-computed weighted score'). Import error: {_KERNEL_IMPORT_ERROR}"
        )

    supplier_ids = [s.supplier_id for s in reg.suppliers]

    category_weight: Dict[str, float] = {}
    for c in reg.criteria:
        category_weight.setdefault(c.category, c.category_weight_frac)

    effective_weight: Dict[str, float] = {
        c.criterion_id: c.weight_in_category_frac * c.category_weight_frac for c in reg.criteria
    }

    weighted_contribution: Dict[str, Dict[str, float]] = {sid: {} for sid in supplier_ids}
    grand_total: Dict[str, float] = {}
    for sid in supplier_ids:
        scores = {c.criterion_id: c.entries[sid].score for c in reg.criteria}
        weights = {c.criterion_id: effective_weight[c.criterion_id] for c in reg.criteria}
        try:
            gt_total = kernel_weighted_score(scores, weights)
        except KernelWeightSumError as e:
            raise ReconciliationError(
                f"Effective-weight-sums-to-one invariant FAILED for supplier {sid!r}: numeric_kernel.weighted_score() "
                f"refused because the effective weights do not sum to 1.0 (Score Validation Checks: 'Effective weight "
                f"sum ... 1.0 +/- 0.001'). Kernel message: {e}"
            )
        except KernelInvalidInputError as e:
            raise ReconciliationError(f"weighted_score() call for supplier {sid!r} failed: {e}")
        grand_total[sid] = gt_total
        for c in reg.criteria:
            weighted_contribution[sid][c.criterion_id] = c.entries[sid].score * effective_weight[c.criterion_id]

    category_subtotal: Dict[str, Dict[str, float]] = {sid: {} for sid in supplier_ids}
    for sid in supplier_ids:
        for cat in category_weight:
            category_subtotal[sid][cat] = sum(
                weighted_contribution[sid][c.criterion_id] for c in reg.criteria if c.category == cat
            )

    must_have_zero_flag: Dict[str, bool] = {}
    gate_detail: Dict[str, List[str]] = {}
    for sid in supplier_ids:
        failed = [c.criterion_id for c in reg.criteria if c.priority == "Must Have" and c.entries[sid].score == 0.0]
        gate_detail[sid] = failed
        must_have_zero_flag[sid] = len(failed) > 0

    compensatory_rank = sorted(supplier_ids, key=lambda sid: grand_total[sid], reverse=True)
    gate_clear_rank = [sid for sid in compensatory_rank if not must_have_zero_flag[sid]]
    all_suppliers_gate_failed = len(gate_clear_rank) == 0

    top_total = grand_total[compensatory_rank[0]]
    delta_to_top = {sid: top_total - grand_total[sid] for sid in supplier_ids}
    tie_flag: Dict[str, bool] = {sid: False for sid in supplier_ids}
    for i in range(1, len(compensatory_rank)):
        prev_sid, cur_sid = compensatory_rank[i - 1], compensatory_rank[i]
        if abs(grand_total[prev_sid] - grand_total[cur_sid]) < 0.1:
            tie_flag[cur_sid] = True
            tie_flag[prev_sid] = True

    if all_suppliers_gate_failed:
        recommended_supplier_id = compensatory_rank[0]
        gate_conflict = False
    else:
        recommended_supplier_id = gate_clear_rank[0]
        gate_conflict = recommended_supplier_id != compensatory_rank[0]

    coverage_counts: Dict[str, Dict[str, int]] = {}
    for sid in supplier_ids:
        counts = {label: 0 for label in VALID_COVERAGE}
        for c in reg.criteria:
            counts[c.entries[sid].coverage] += 1
        counts["Total"] = len(reg.criteria)
        coverage_counts[sid] = counts

    if reg.sensitivity_rows:
        changed = [r for r in reg.sensitivity_rows if r.ranking_changed]
        if changed:
            cats = ", ".join(sorted({r.category_adjusted for r in changed}))
            sensitivity_verdict = (
                f"Recommendation is sensitive to {cats} weighting. At least one tested weight perturbation "
                "changes the top-ranked supplier by compensatory total; evaluation team should confirm alignment "
                "on weight priorities before finalizing."
            )
        else:
            sensitivity_verdict = "Recommendation is robust, ranking holds under every tested weight variation."
    else:
        sensitivity_verdict = "Sensitivity analysis not yet run (no sensitivity_rows provided)."

    return GroundTruth(
        effective_weight=effective_weight, category_weight=category_weight,
        weighted_contribution=weighted_contribution, category_subtotal=category_subtotal,
        grand_total=grand_total, must_have_zero_flag=must_have_zero_flag, gate_detail=gate_detail,
        compensatory_rank=compensatory_rank, gate_clear_rank=gate_clear_rank,
        recommended_supplier_id=recommended_supplier_id, gate_conflict=gate_conflict,
        all_suppliers_gate_failed=all_suppliers_gate_failed, coverage_counts=coverage_counts,
        delta_to_top=delta_to_top, tie_flag=tie_flag, sensitivity_verdict=sensitivity_verdict,
    )


# ===========================================================================
# Hard-coded invariant checks (run BEFORE saving)
# ===========================================================================

def _assert_effective_weights_sum_to_one(gt: GroundTruth, tolerance: float = 0.001) -> None:
    total = sum(gt.effective_weight.values())
    if abs(total - 1.0) > tolerance:
        raise ReconciliationError(
            f"Effective-weight-sums-to-one invariant FAILED: Effective_Weight_Frac values sum to {total:.4f}, "
            f"not 1.0 (+/- {tolerance}), per Score Validation Checks."
        )


def _assert_weighted_contributions_reconcile(gt: GroundTruth, tolerance: float = 1e-6) -> None:
    for sid, contributions in gt.weighted_contribution.items():
        summed = sum(contributions.values())
        if abs(summed - gt.grand_total[sid]) > tolerance:
            raise ReconciliationError(
                f"Weighted-contributions-reconcile invariant FAILED for supplier {sid!r}: the sum of every "
                f"per-criterion weighted contribution ({summed}) does not equal the kernel-computed Grand Total "
                f"({gt.grand_total[sid]}). Canonical Dashboard 'Numbers-reconcile assertion'."
            )


def _assert_grand_total_range(gt: GroundTruth) -> None:
    for sid, total in gt.grand_total.items():
        if not (0.0 <= total <= 5.0):
            raise ReconciliationError(
                f"Grand-total-range invariant FAILED for supplier {sid!r}: Grand Total {total} is outside "
                "[0.0, 5.0] (Score Validation Checks: 'All Grand Total scores must be 0.0-5.0')."
            )


def _assert_category_subtotals_reconcile(gt: GroundTruth, tolerance: float = 1e-6) -> None:
    for sid, subtotals in gt.category_subtotal.items():
        summed = sum(subtotals.values())
        if abs(summed - gt.grand_total[sid]) > tolerance:
            raise ReconciliationError(
                f"Category-subtotals-reconcile invariant FAILED for supplier {sid!r}: the sum of every "
                f"CATEGORY_TOTAL subtotal ({summed}) does not equal the Grand Total ({gt.grand_total[sid]})."
            )


def _assert_ranking_sorted(gt: GroundTruth) -> None:
    totals = [gt.grand_total[sid] for sid in gt.compensatory_rank]
    if totals != sorted(totals, reverse=True):
        raise ReconciliationError(
            f"Ranking-sorted invariant FAILED: compensatory_rank {gt.compensatory_rank} is not sorted by "
            f"Grand Total descending (totals: {totals})."
        )


def _assert_must_have_gate_not_buried(gt: GroundTruth) -> None:
    """Recommendation Logic's gate-vs-compensatory reconciliation: 'If the
    highest weighted-total supplier carries a Must Have zero, do NOT
    silently recommend it on total alone.' Asserts the recommended supplier
    is never a gate-failed supplier unless every supplier failed the gate
    (the disclosed edge case, see compute_ground_truth's
    all_suppliers_gate_failed handling)."""
    top_sid = gt.compensatory_rank[0]
    if gt.must_have_zero_flag[top_sid] and not gt.all_suppliers_gate_failed:
        if gt.recommended_supplier_id == top_sid:
            raise ReconciliationError(
                f"Must-Have-gate-not-buried invariant FAILED: the top compensatory-total supplier ({top_sid}) "
                f"carries a Must Have zero flag, and a gate-clear alternative exists "
                f"({gt.gate_clear_rank}), but recommended_supplier_id still points at the gate-failed supplier."
            )
        if not gt.gate_conflict:
            raise ReconciliationError(
                "Must-Have-gate-not-buried invariant FAILED: a gate conflict exists (top compensatory supplier "
                "failed a Must Have) but gate_conflict was not flagged True."
            )


def _assert_requirements_coverage_counts_foot(reg: EvaluationInput, gt: GroundTruth) -> None:
    n_criteria = len(reg.criteria)
    for sid, counts in gt.coverage_counts.items():
        summed = sum(counts[label] for label in VALID_COVERAGE)
        if summed != n_criteria or counts["Total"] != n_criteria:
            raise ReconciliationError(
                f"Requirements-coverage-counts-foot invariant FAILED for supplier {sid!r}: coverage tally sums "
                f"to {summed} (Total field {counts['Total']}), not the criteria count ({n_criteria})."
            )


def _assert_sensitivity_rank1_base_matches(reg: EvaluationInput, gt: GroundTruth) -> None:
    top_sid = gt.compensatory_rank[0]
    for row in reg.sensitivity_rows:
        if row.rank1_base != top_sid:
            raise ReconciliationError(
                f"Sensitivity-rank1-base-matches-ground-truth invariant FAILED: sensitivity row for "
                f"{row.category_adjusted!r} states Rank 1 (Base) = {row.rank1_base!r}, but this module's own "
                f"kernel-computed compensatory rank-1 supplier is {top_sid!r}."
            )


def run_hardcoded_invariant_checks(reg: EvaluationInput, gt: GroundTruth) -> None:
    """Run every hard-coded invariant. Raises ReconciliationError on any
    failure; callers must not save the document if this raises."""
    _assert_effective_weights_sum_to_one(gt)
    _assert_weighted_contributions_reconcile(gt)
    _assert_grand_total_range(gt)
    _assert_category_subtotals_reconcile(gt)
    _assert_ranking_sorted(gt)
    _assert_must_have_gate_not_buried(gt)
    _assert_requirements_coverage_counts_foot(reg, gt)
    _assert_sensitivity_rank1_base_matches(reg, gt)


# ===========================================================================
# 2. Document builder (python-docx)
# ===========================================================================

def _require_docx() -> None:
    if not DOCX_AVAILABLE:
        raise ImportError(
            "python-docx is not installed in this Python environment, so no "
            ".docx file can be written. Install it (`pip install python-docx`) "
            "or point this script at an interpreter that already has it. "
            f"Original import error: {_DOCX_IMPORT_ERROR}"
        )


def _set_cell_background(cell, hex_color: str) -> None:
    """Set a table cell's background fill via raw OOXML; python-docx has no
    high-level cell-shading API."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_text(cell, text: str, bold: bool = False, color_hex: str = LILLY_BLACK,
                    size_pt: int = 9, align=None) -> None:
    cell.text = str(text)
    para = cell.paragraphs[0]
    if align is not None:
        para.alignment = align
    run = para.runs[0] if para.runs else para.add_run("")
    run.font.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor.from_string(color_hex)


def _add_heading(doc, text: str, level: int) -> None:
    """Manually-styled heading (no dependency on a template's built-in
    Heading styles), per docx-design-system.md: H1 Calibri 14pt Bold
    #0F3A85, H2 Calibri 12pt Bold #212121, H3 Calibri 11pt Bold #212121."""
    spec = {1: (14, BOLD_BLUE, 18, 8), 2: (12, LILLY_BLACK, 14, 6), 3: (11, LILLY_BLACK, 10, 4)}
    size_pt, color_hex, before_pt, after_pt = spec.get(level, spec[3])
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(before_pt)
    para.paragraph_format.space_after = Pt(after_pt)
    run = para.add_run(text)
    run.font.bold = True
    run.font.size = Pt(size_pt)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor.from_string(color_hex)


def _add_body_paragraph(doc, text: str, bold: bool = False, italic: bool = False,
                         size_pt: int = 10, color_hex: str = LILLY_BLACK) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.font.bold = bold
    run.font.italic = italic
    run.font.size = Pt(size_pt)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor.from_string(color_hex)


def _add_bullet_list(doc, items: Sequence[str]) -> None:
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(item)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor.from_string(LILLY_BLACK)


def _add_table(doc, headers: Sequence[str], rows: Sequence[Sequence[Any]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell_background(hdr_cells[i], LILLY_RED)
        _set_cell_text(hdr_cells[i], h, bold=True, color_hex=WHITE, size_pt=9)
    for row_values in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_values):
            _set_cell_text(cells[i], str(val), size_pt=9)
    return table


def _pct(frac: float) -> str:
    return f"{frac * 100:.1f}%"


def _usd(value: float) -> str:
    return f"${value:,.0f}"


def build_document(reg: EvaluationInput, gt: GroundTruth):
    """Build the evaluation_report.docx per SKILL.md's Evaluation Schemas
    reference: Sections 1-8, exactly the fixed order and required content
    described there, respecting the Brief vs Full depth table. This
    function does NOT save the file and does NOT run the reconciliation
    checks; call generate_evaluation_report() for the full validated
    pipeline."""
    _require_docx()

    supplier_by_id = {s.supplier_id: s for s in reg.suppliers}
    supplier_ids_ordered = [s.supplier_id for s in reg.suppliers]
    full_mode = reg.report_mode == "Full"

    doc = Document()

    # Footer (all pages), per docx-design-system.md's Footer text spec
    # (Calibri 8pt, Bold Grey #8A969E).
    footer_para = doc.sections[0].footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        "CONFIDENTIAL, internal use only. Lilly Procurement Skills Suite, Evaluation Engine."
    )
    footer_run.font.size = Pt(8)
    footer_run.font.name = "Calibri"
    footer_run.font.color.rgb = RGBColor.from_string(MUTED_GREY)

    # --- Title page ----------------------------------------------------
    title_para = doc.add_paragraph()
    title_para.paragraph_format.space_after = Pt(4)
    title_run = title_para.add_run("Supplier Evaluation Report")
    title_run.font.bold = True
    title_run.font.size = Pt(26)
    title_run.font.name = "Calibri"
    title_run.font.color.rgb = RGBColor.from_string(LILLY_BLACK)

    scope_para = doc.add_paragraph()
    scope_para.paragraph_format.space_after = Pt(10)
    scope_run = scope_para.add_run(f"{reg.rfx_name}, {reg.category}")
    scope_run.font.size = Pt(11)
    scope_run.font.name = "Calibri"
    scope_run.font.color.rgb = RGBColor.from_string(BOLD_BLUE)

    conf_para = doc.add_paragraph()
    conf_para.paragraph_format.space_after = Pt(10)
    conf_run = conf_para.add_run("CONFIDENTIAL, INTERNAL USE ONLY")
    conf_run.font.bold = True
    conf_run.font.size = Pt(10)
    conf_run.font.name = "Calibri"
    conf_run.font.color.rgb = RGBColor.from_string(LILLY_RED)

    meta_lines = [
        f"Case ID: {reg.case_id or 'N/A'}",
        f"Evaluation date: {reg.evaluation_date}",
        f"Suppliers evaluated: {len(reg.suppliers)}",
        f"Report mode: {reg.report_mode}",
    ]
    for line in meta_lines:
        _add_body_paragraph(doc, line)
    doc.add_page_break()

    # --- Section 1: Executive Summary -----------------------------------
    _add_heading(doc, "1. Executive Summary", level=1)
    recommended = supplier_by_id[gt.recommended_supplier_id]
    _add_body_paragraph(
        doc,
        f"{reg.rfx_name} ({reg.category}) evaluated {len(reg.suppliers)} suppliers against a "
        f"{len(reg.criteria)}-criterion weighted scoring matrix. Recommended supplier: "
        f"{recommended.supplier_name}, Grand Total {gt.grand_total[gt.recommended_supplier_id]:.2f} of 5.0.",
        bold=True,
    )
    if gt.gate_conflict:
        top_sid = gt.compensatory_rank[0]
        top_supplier = supplier_by_id[top_sid]
        _add_body_paragraph(
            doc,
            f"Note: {top_supplier.supplier_name} carries the highest compensatory Grand Total "
            f"({gt.grand_total[top_sid]:.2f}) but failed a Must Have requirement "
            f"({', '.join(gt.gate_detail[top_sid])}). Compensatory scoring cannot offset a hard requirement "
            f"gap; {recommended.supplier_name} is recommended instead. See Section 7 for the full "
            "gate-versus-compensatory reconciliation.",
        )
    if full_mode:
        _add_body_paragraph(
            doc,
            f"Scoring matrix source: {reg.scoring_matrix_source}. AI scoring: "
            f"{'enabled' if reg.ai_scoring_enabled else 'disabled'}. Stakeholder scoring: "
            f"{'included, ' + str(reg.stakeholder_scorer_count) + ' scorer(s)' if reg.stakeholder_scoring_included else 'not included'}."
        )
    _add_body_paragraph(doc, f"Key next steps: {reg.recommendation.key_tradeoffs}")
    _add_body_paragraph(doc, f"Date of evaluation completion: {reg.evaluation_date}.")

    # --- Section 2: Evaluation Methodology --------------------------------
    _add_heading(doc, "2. Evaluation Methodology", level=1)
    _add_body_paragraph(doc, f"Scoring matrix source: {reg.scoring_matrix_source}.")
    _add_body_paragraph(doc, f"AI scoring: {'Enabled' if reg.ai_scoring_enabled else 'Disabled'}.")
    _add_body_paragraph(
        doc,
        f"Stakeholder scoring: {'Included' if reg.stakeholder_scoring_included else 'Not included'}"
        + (f", {reg.stakeholder_scorer_count} scorer(s)." if reg.stakeholder_scoring_included else "."),
    )
    _add_body_paragraph(
        doc,
        "Aggregation method: weighted average. Each criterion's Final_Score (0.0 to 5.0) is multiplied by "
        "its Effective_Weight_Frac (a fraction; effective weights sum to 1.0 across all criteria), and the "
        "products are summed for each supplier's Grand Total, computed via numeric_kernel.weighted_score().",
    )
    _add_body_paragraph(doc, "Scoring scale: 0.0 to 5.0, decimals allowed.")
    _add_table(
        doc, ["Score", "Label", "Definition"],
        [
            ["5.0", "Exceptional", "Fully meets with evidence of excellence beyond baseline"],
            ["4.0-4.9", "Exceeds", "Clearly meets; minor enhancements beyond baseline"],
            ["3.0-3.9", "Meets", "Adequately meets; no gaps, no standout capability"],
            ["2.0-2.9", "Partially Meets", "Meets some aspects; notable gaps"],
            ["1.0-1.9", "Minimally Meets", "Barely addresses the requirement"],
            ["0.0-0.9", "Does Not Meet", "Does not address or explicitly cannot meet"],
        ],
    )

    # --- Section 3: Supplier Summaries ------------------------------------
    _add_heading(doc, "3. Supplier Summaries", level=1)
    for idx, sid in enumerate(supplier_ids_ordered, start=1):
        supplier = supplier_by_id[sid]
        _add_heading(doc, f"3.{idx} {supplier.supplier_name}", level=2)
        if supplier.snapshot:
            _add_body_paragraph(doc, supplier.snapshot)
        counts = gt.coverage_counts[sid]
        _add_body_paragraph(
            doc,
            f"Requirements coverage: {counts['Fully Meets']} Fully Meets, {counts['Partially Meets']} Partially "
            f"Meets, {counts['Does Not Meet']} Does Not Meet, {counts['Not Answered']} Not Answered "
            f"(of {counts['Total']} criteria).",
        )
        if sid in reg.commercial:
            cf = reg.commercial[sid]
            _add_body_paragraph(
                doc,
                f"Commercial summary: Year 1 {_usd(cf.year1_cost_usd)}, 3-year TCV {_usd(cf.tcv_3yr_usd)}, "
                f"pricing model: {cf.pricing_model}.",
            )
        else:
            _add_body_paragraph(doc, "Commercial summary: NEEDS_INPUT, Year 1 and 3-year TCV figures not gathered for this supplier.")
        if sid in reg.strengths and reg.strengths[sid]:
            _add_body_paragraph(doc, "Top strengths:", bold=True)
            _add_bullet_list(doc, reg.strengths[sid])
        if sid in reg.weaknesses and reg.weaknesses[sid]:
            _add_body_paragraph(doc, "Top weaknesses / risks:", bold=True)
            _add_bullet_list(doc, reg.weaknesses[sid])
        gate_text = "GATE FAIL, " + ", ".join(gt.gate_detail[sid]) if gt.must_have_zero_flag[sid] else "Gate clear"
        _add_body_paragraph(
            doc, f"Final weighted score (Grand Total): {gt.grand_total[sid]:.2f} of 5.0. {gate_text}.", bold=True,
        )

    # --- Section 4: Score Comparison --------------------------------------
    _add_heading(doc, "4. Score Comparison", level=1)
    ranked_rows = []
    for rank, sid in enumerate(gt.compensatory_rank, start=1):
        supplier = supplier_by_id[sid]
        gate_label = "GATE FAIL" if gt.must_have_zero_flag[sid] else "Clear"
        tie_label = " (TIE)" if gt.tie_flag[sid] else ""
        ranked_rows.append([
            rank, supplier.supplier_name, f"{gt.grand_total[sid]:.2f}",
            f"{gt.delta_to_top[sid]:.2f}{tie_label}", gate_label,
        ])
    _add_table(doc, ["Rank", "Supplier", "Grand Total (0.0-5.0)", "Delta to #1", "Must Have Gate"], ranked_rows)

    if full_mode:
        _add_body_paragraph(doc, "Category-by-category score breakdown:", bold=True)
        cat_headers = ["Category", "Category Weight"] + [supplier_by_id[sid].supplier_name for sid in supplier_ids_ordered]
        cat_rows = []
        for cat, weight in gt.category_weight.items():
            row = [cat, _pct(weight)] + [f"{gt.category_subtotal[sid][cat]:.2f}" for sid in supplier_ids_ordered]
            cat_rows.append(row)
        _add_table(doc, cat_headers, cat_rows)

    top_sid = gt.compensatory_rank[0]
    second_sid = gt.compensatory_rank[1] if len(gt.compensatory_rank) > 1 else None
    if second_sid:
        driver_cat = max(gt.category_weight, key=lambda cat: gt.category_subtotal[top_sid][cat] - gt.category_subtotal[second_sid][cat])
        _add_body_paragraph(
            doc,
            f"Score drivers: {supplier_by_id[top_sid].supplier_name}'s largest margin over "
            f"{supplier_by_id[second_sid].supplier_name} is in {driver_cat} "
            f"({gt.category_subtotal[top_sid][driver_cat]:.2f} vs {gt.category_subtotal[second_sid][driver_cat]:.2f}).",
        )

    # --- Section 5: Risk Assessment Summary -------------------------------
    _add_heading(doc, "5. Risk Assessment Summary", level=1)
    if len(reg.suppliers) < 3:
        _add_body_paragraph(doc, "NOT APPLICABLE, fewer than 3 suppliers evaluated; risk heatmap requires 3 or more.")
    elif not reg.risk_matrix:
        _add_body_paragraph(doc, "NEEDS_INPUT, no risk data was provided for this evaluation.")
    else:
        table = doc.add_table(rows=1, cols=1 + len(supplier_ids_ordered))
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        _set_cell_background(hdr[0], LILLY_RED)
        _set_cell_text(hdr[0], "Category", bold=True, color_hex=WHITE, size_pt=9)
        for i, sid in enumerate(supplier_ids_ordered, start=1):
            _set_cell_background(hdr[i], LILLY_RED)
            _set_cell_text(hdr[i], supplier_by_id[sid].supplier_name, bold=True, color_hex=WHITE, size_pt=9)
        for rcat, per_supplier in reg.risk_matrix.items():
            cells = table.add_row().cells
            _set_cell_text(cells[0], rcat, size_pt=9)
            for i, sid in enumerate(supplier_ids_ordered, start=1):
                sev = per_supplier.get(sid, "NEEDS_INPUT")
                if sev in SEVERITY_COLOR:
                    _set_cell_background(cells[i], SEVERITY_COLOR[sev])
                    _set_cell_text(cells[i], sev, bold=True, color_hex=WHITE, size_pt=9,
                                   align=WD_ALIGN_PARAGRAPH.CENTER)
                else:
                    _set_cell_text(cells[i], sev, size_pt=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        _add_body_paragraph(doc, "Key risk notes:", bold=True)
        for sid in supplier_ids_ordered:
            if sid in reg.risk_notes:
                _add_body_paragraph(doc, f"{supplier_by_id[sid].supplier_name}: {reg.risk_notes[sid]}")

    # --- Section 6: Commercial Comparison ---------------------------------
    _add_heading(doc, "6. Commercial Comparison", level=1)
    if reg.commercial:
        comm_rows = []
        for sid in supplier_ids_ordered:
            supplier = supplier_by_id[sid]
            if sid in reg.commercial:
                cf = reg.commercial[sid]
                comm_rows.append([supplier.supplier_name, _usd(cf.year1_cost_usd), _usd(cf.tcv_3yr_usd), cf.pricing_model])
            else:
                comm_rows.append([supplier.supplier_name, "NEEDS_INPUT", "NEEDS_INPUT", "NEEDS_INPUT"])
        _add_table(doc, ["Supplier", "Year 1 Cost", "3-Year TCV", "Pricing Model"], comm_rows)
        for sid in supplier_ids_ordered:
            if sid in reg.commercial and reg.commercial[sid].notes:
                _add_body_paragraph(doc, f"{supplier_by_id[sid].supplier_name}: {reg.commercial[sid].notes}")
    else:
        _add_body_paragraph(doc, "NEEDS_INPUT, no commercial figures were provided for this evaluation.")

    # --- Section 7: Recommendation ----------------------------------------
    _add_heading(doc, "7. Recommendation", level=1)
    _add_body_paragraph(
        doc,
        f"Primary recommended supplier: {recommended.supplier_name} (Grand Total "
        f"{gt.grand_total[gt.recommended_supplier_id]:.2f} of 5.0, "
        f"{'gate clear' if not gt.must_have_zero_flag[gt.recommended_supplier_id] else 'GATE FAIL'}).",
        bold=True,
    )
    if gt.gate_conflict:
        top_sid = gt.compensatory_rank[0]
        top_supplier = supplier_by_id[top_sid]
        _add_body_paragraph(
            doc,
            f"Supplier {top_supplier.supplier_name} has the top weighted total "
            f"({gt.grand_total[top_sid]:.2f} on 0.0-5.0) but failed Must Have "
            f"{', '.join(gt.gate_detail[top_sid])}. Compensatory scoring cannot offset a hard requirement gap. "
            "Both rankings are shown below; the choice is routed to the evaluation team as an ASK moment.",
        )
        gate_clear_rows = [
            [rank, supplier_by_id[sid].supplier_name, f"{gt.grand_total[sid]:.2f}"]
            for rank, sid in enumerate(gt.gate_clear_rank, start=1)
        ]
        _add_body_paragraph(doc, "Gate-pass ranking (Must Have gate clear suppliers only):", bold=True)
        _add_table(doc, ["Rank", "Supplier", "Grand Total (0.0-5.0)"], gate_clear_rows)
    if gt.all_suppliers_gate_failed:
        _add_body_paragraph(
            doc,
            "All suppliers carry at least one Must Have zero flag. No gate-clear alternative exists; the "
            "compensatory top-ranked supplier is shown as recommended, subject to leadership review of every "
            "Must Have gap before proceeding.",
        )
    for sid in supplier_ids_ordered:
        if gt.must_have_zero_flag[sid]:
            _add_body_paragraph(
                doc,
                f"Supplier {sid} scored 0.0 on Must Have requirement {', '.join(gt.gate_detail[sid])}. "
                "Recommend leadership review before proceeding to award.",
            )

    if full_mode and reg.recommendation.secondary_supplier_id:
        secondary = supplier_by_id[reg.recommendation.secondary_supplier_id]
        _add_heading(doc, "Secondary Option", level=2)
        _add_body_paragraph(
            doc,
            f"{secondary.supplier_name} (Grand Total {gt.grand_total[reg.recommendation.secondary_supplier_id]:.2f}). "
            f"Tradeoffs: {reg.recommendation.secondary_tradeoffs}",
        )

    _add_heading(doc, "Key Tradeoffs", level=2)
    _add_body_paragraph(doc, reg.recommendation.key_tradeoffs)

    _add_heading(doc, "Sensitivity Analysis", level=2)
    _add_body_paragraph(doc, gt.sensitivity_verdict)
    if full_mode and reg.sensitivity_rows:
        sens_rows = [
            [
                r.category_adjusted, _pct(r.base_weight_frac), _pct(r.test_weight_frac),
                supplier_by_id[r.rank1_base].supplier_name, supplier_by_id[r.rank1_adjusted].supplier_name,
                "YES" if r.ranking_changed else "No",
            ]
            for r in reg.sensitivity_rows
        ]
        _add_table(
            doc, ["Category Adjusted", "Base Weight", "Test Weight", "Rank 1 (Base)", "Rank 1 (Adjusted)", "Ranking Changed?"],
            sens_rows,
        )

    _add_heading(doc, "Required Next Steps", level=2)
    _add_bullet_list(doc, reg.recommendation.next_steps)

    # --- Section 8: Appendices ---------------------------------------------
    _add_heading(doc, "8. Appendices", level=1)
    _add_heading(doc, "Appendix A: Full Score Rollup", level=2)
    if full_mode:
        _add_body_paragraph(
            doc,
            "DOCX-embedded equivalent of final_score_rollup.csv; see that CSV for the machine-readable version.",
            italic=True,
        )
        headers = ["Criterion ID", "Category", "Priority", "Effective Weight"]
        for sid in supplier_ids_ordered:
            headers += [f"{supplier_by_id[sid].supplier_name} Score", f"{supplier_by_id[sid].supplier_name} Weighted"]
        rollup_rows = []
        for c in reg.criteria:
            row = [c.criterion_id, c.category, c.priority, _pct(gt.effective_weight[c.criterion_id])]
            for sid in supplier_ids_ordered:
                row += [f"{c.entries[sid].score:.1f}", f"{gt.weighted_contribution[sid][c.criterion_id]:.3f}"]
            rollup_rows.append(row)
        grand_row = ["GRAND TOTAL", "", "", ""]
        for sid in supplier_ids_ordered:
            grand_row += ["", f"{gt.grand_total[sid]:.3f}"]
        rollup_rows.append(grand_row)
        _add_table(doc, headers, rollup_rows)
    else:
        _add_body_paragraph(doc, "Full criterion-level detail is available in the Full report and in final_score_rollup.csv.")

    _add_heading(doc, "Appendix B: Requirements Coverage Matrix Summary", level=2)
    cov_headers = ["Supplier"] + list(VALID_COVERAGE) + ["Total"]
    cov_rows = [
        [supplier_by_id[sid].supplier_name] + [gt.coverage_counts[sid][label] for label in VALID_COVERAGE]
        + [gt.coverage_counts[sid]["Total"]]
        for sid in supplier_ids_ordered
    ]
    _add_table(doc, cov_headers, cov_rows)

    _add_heading(doc, "Appendix C: Clarifications Requested and Received", level=2)
    if reg.clarifications_note:
        _add_body_paragraph(doc, reg.clarifications_note)
    else:
        _add_body_paragraph(doc, "No clarifications were requested or received as part of this evaluation.")

    return doc


# ===========================================================================
# Full pipeline: validate -> compute ground truth -> hard-invariant checks
# -> build document -> save
# ===========================================================================

def generate_evaluation_report(raw_register: Dict[str, Any], output_path: str,
                                mode_override: Optional[str] = None) -> EvaluationInput:
    """End-to-end: validate the raw evaluation register, compute ground truth
    via numeric_kernel, run the hard-coded invariant checks, build the
    document, and only then save it. Raises rather than saving a document
    that fails validation or reconciliation."""
    register = copy.deepcopy(raw_register)
    if mode_override is not None:
        register["report_mode"] = mode_override
    reg = validate_evaluation_input(register)
    gt = compute_ground_truth(reg)
    run_hardcoded_invariant_checks(reg, gt)
    doc = build_document(reg, gt)
    doc.save(output_path)
    return reg


# ===========================================================================
# Demo data (clearly illustrative; SKILL.md carries no numeric worked
# example for evaluation-engine, unlike should-cost-builder's quoted
# "60+30+10=100" illustration, so this dataset is authored for this
# generator, not quoted verbatim from any skill file)
# ===========================================================================

def _demo_evaluation_register() -> Dict[str, Any]:
    """ILLUSTRATIVE DEMO DATA, not a real sourcing event. Three fictional
    suppliers competing for an illustrative ITSM platform replacement.
    Deliberately exercises every code path this generator implements: a
    Must Have zero flag on the compensatory top-ranked supplier (forcing
    the gate-versus-compensatory reconciliation in Section 7), a full
    risk heatmap (3 suppliers), a NEEDS_INPUT-free commercial table, and a
    sensitivity row where the ranking flips (exercising the 'sensitive to'
    narrative branch)."""
    return {
        "rfx_name": "ITSM Platform Replacement RFP (Illustrative Demo)",
        "category": "IT Service Management (ITSM) Platform",
        "case_id": "DEMO-EVAL-0001",
        "evaluation_date": "2026-07-23",
        "scoring_matrix_source": "Provided",
        "ai_scoring_enabled": True,
        "stakeholder_scoring_included": True,
        "stakeholder_scorer_count": 4,
        "report_mode": "Full",
        "suppliers": [
            {"supplier_id": "SUP_A", "supplier_name": "Northbridge Service Solutions (illustrative)",
             "snapshot": "Illustrative mid-market ITSM vendor, privately held, cited by this demo dataset only."},
            {"supplier_id": "SUP_B", "supplier_name": "Solstice Cloud Systems (illustrative)",
             "snapshot": "Illustrative enterprise ITSM vendor, publicly listed in this demo scenario only."},
            {"supplier_id": "SUP_C", "supplier_name": "Meridian Platform Group (illustrative)",
             "snapshot": "Illustrative challenger ITSM vendor, smaller reference base in this demo scenario."},
        ],
        "criteria": [
            {"criterion_id": "REQ-001", "category": "Functional Fit",
             "criterion_text": "Native ITSM ticketing and incident management, no third-party plugin required.",
             "priority": "Must Have", "weight_in_category_frac": 0.6, "category_weight_frac": 0.40,
             "entries": {
                 "SUP_A": {"score": 5.0, "coverage": "Fully Meets", "confidence": "High",
                           "evidence": "Proposal Section 2.1", "rationale": "Native module confirmed in demo."},
                 "SUP_B": {"score": 4.0, "coverage": "Fully Meets", "confidence": "High",
                           "evidence": "Proposal Section 2.1", "rationale": "Native module, minor config needed."},
                 "SUP_C": {"score": 3.5, "coverage": "Partially Meets", "confidence": "Medium",
                           "evidence": "Proposal Section 2.2", "rationale": "Native module, limited workflow depth."},
             }},
            {"criterion_id": "REQ-002", "category": "Functional Fit",
             "criterion_text": "Self-service portal with SSO for end users.",
             "priority": "Should Have", "weight_in_category_frac": 0.4, "category_weight_frac": 0.40,
             "entries": {
                 "SUP_A": {"score": 4.8, "coverage": "Fully Meets", "confidence": "High",
                           "evidence": "Proposal Section 2.4", "rationale": "SAML 2.0 SSO confirmed."},
                 "SUP_B": {"score": 3.5, "coverage": "Partially Meets", "confidence": "Medium",
                           "evidence": "Proposal Section 2.4", "rationale": "SSO available, portal UI dated."},
                 "SUP_C": {"score": 3.0, "coverage": "Partially Meets", "confidence": "Medium",
                           "evidence": "Proposal Section 2.5", "rationale": "SSO via third-party IdP only."},
             }},
            {"criterion_id": "REQ-010", "category": "Security & Compliance",
             "criterion_text": "Current SOC 2 Type II certification.",
             "priority": "Must Have", "weight_in_category_frac": 0.5, "category_weight_frac": 0.35,
             "entries": {
                 "SUP_A": {"score": 0.0, "coverage": "Does Not Meet", "confidence": "High",
                           "evidence": "Security questionnaire Q14", "rationale": "No current SOC 2 Type II certification."},
                 "SUP_B": {"score": 4.5, "coverage": "Fully Meets", "confidence": "High",
                           "evidence": "Security questionnaire Q14", "rationale": "Current SOC 2 Type II, report on file."},
                 "SUP_C": {"score": 3.5, "coverage": "Partially Meets", "confidence": "Medium",
                           "evidence": "Security questionnaire Q14", "rationale": "SOC 2 Type I only, Type II in progress."},
             }},
            {"criterion_id": "REQ-011", "category": "Security & Compliance",
             "criterion_text": "Data residency in the US region.",
             "priority": "Should Have", "weight_in_category_frac": 0.5, "category_weight_frac": 0.35,
             "entries": {
                 "SUP_A": {"score": 4.8, "coverage": "Fully Meets", "confidence": "High",
                           "evidence": "Security questionnaire Q9", "rationale": "US-only hosting confirmed."},
                 "SUP_B": {"score": 4.0, "coverage": "Fully Meets", "confidence": "High",
                           "evidence": "Security questionnaire Q9", "rationale": "US region available and selected."},
                 "SUP_C": {"score": 3.0, "coverage": "Partially Meets", "confidence": "Medium",
                           "evidence": "Security questionnaire Q9", "rationale": "US region available, not default."},
             }},
            {"criterion_id": "REQ-020", "category": "Commercial",
             "criterion_text": "Transparent per-agent pricing model.",
             "priority": "Should Have", "weight_in_category_frac": 0.7, "category_weight_frac": 0.25,
             "entries": {
                 "SUP_A": {"score": 4.8, "coverage": "Fully Meets", "confidence": "High",
                           "evidence": "Pricing sheet p.1", "rationale": "Flat per-agent rate, no tiers."},
                 "SUP_B": {"score": 3.5, "coverage": "Partially Meets", "confidence": "Medium",
                           "evidence": "Pricing sheet p.1", "rationale": "Tiered per-agent rate, moderate complexity."},
                 "SUP_C": {"score": 4.0, "coverage": "Fully Meets", "confidence": "Medium",
                           "evidence": "Pricing sheet p.1", "rationale": "Platform fee plus per-agent, clearly stated."},
             }},
            {"criterion_id": "REQ-021", "category": "Commercial",
             "criterion_text": "Multi-year price lock option.",
             "priority": "Nice to Have", "weight_in_category_frac": 0.3, "category_weight_frac": 0.25,
             "entries": {
                 "SUP_A": {"score": 4.5, "coverage": "Fully Meets", "confidence": "Medium",
                           "evidence": "Pricing sheet p.2", "rationale": "3-year lock offered."},
                 "SUP_B": {"score": 3.0, "coverage": "Partially Meets", "confidence": "Medium",
                           "evidence": "Pricing sheet p.2", "rationale": "1-year lock only, annual escalation after."},
                 "SUP_C": {"score": 2.0, "coverage": "Does Not Meet", "confidence": "Medium",
                           "evidence": "Pricing sheet p.2", "rationale": "No multi-year lock offered."},
             }},
        ],
        "commercial": {
            "SUP_A": {"year1_cost_usd": 185000, "tcv_3yr_usd": 510000,
                      "pricing_model": "Per-agent subscription, annual, flat rate", "notes": "Illustrative demo figures."},
            "SUP_B": {"year1_cost_usd": 210000, "tcv_3yr_usd": 560000,
                      "pricing_model": "Per-agent subscription, annual, tiered", "notes": "Illustrative demo figures."},
            "SUP_C": {"year1_cost_usd": 165000, "tcv_3yr_usd": 470000,
                      "pricing_model": "Platform fee plus per-agent, annual", "notes": "Illustrative demo figures."},
        },
        "risk_matrix": {
            "Legal/Contractual": {"SUP_A": "Low", "SUP_B": "Low", "SUP_C": "Medium"},
            "Cyber/Security": {"SUP_A": "High", "SUP_B": "Low", "SUP_C": "Medium"},
            "Operational": {"SUP_A": "Low", "SUP_B": "Medium", "SUP_C": "Medium"},
            "Geopolitical": {"SUP_A": "Low", "SUP_B": "Low", "SUP_C": "Low"},
            "ESG": {"SUP_A": "Medium", "SUP_B": "Low", "SUP_C": "Medium"},
        },
        "risk_notes": {
            "SUP_A": "No current SOC 2 Type II certification, the primary driver of the Must Have gate failure on REQ-010.",
            "SUP_B": "No material risk flags identified in the illustrative demo inputs.",
            "SUP_C": "Moderate implementation risk given a smaller reference base in this illustrative scenario.",
        },
        "strengths": {
            "SUP_A": ["Highest functional fit score among the three suppliers",
                      "Strong self-service portal capability with native SSO",
                      "Most flexible multi-year price lock terms"],
            "SUP_B": ["Only supplier with current SOC 2 Type II certification",
                      "Balanced performance across every scored category",
                      "No High-severity risk flags in this illustrative scenario"],
            "SUP_C": ["Lowest Year 1 cost of the three suppliers",
                      "Transparent platform-plus-per-agent pricing model",
                      "Competitive commercial pricing criterion score"],
        },
        "weaknesses": {
            "SUP_A": ["No current SOC 2 Type II certification, triggering a Must Have gate failure",
                      "Cyber/Security risk rated High pending certification",
                      "Recommendation cannot rely on total score alone given the gate failure"],
            "SUP_B": ["Highest Year 1 cost of the three suppliers",
                      "Self-service portal scored lower than Supplier A",
                      "Multi-year price lock terms less flexible than Supplier A"],
            "SUP_C": ["Lowest functional fit score among the three suppliers",
                      "No criterion scored above 4.0",
                      "Smaller illustrative reference base increases implementation risk"],
        },
        "sensitivity_rows": [
            {"category_adjusted": "Functional Fit", "base_weight_frac": 0.40, "test_weight_frac": 0.45,
             "rank1_base": "SUP_A", "rank1_adjusted": "SUP_A", "ranking_changed": False},
            {"category_adjusted": "Functional Fit", "base_weight_frac": 0.40, "test_weight_frac": 0.35,
             "rank1_base": "SUP_A", "rank1_adjusted": "SUP_B", "ranking_changed": True},
            {"category_adjusted": "Security & Compliance", "base_weight_frac": 0.35, "test_weight_frac": 0.40,
             "rank1_base": "SUP_A", "rank1_adjusted": "SUP_A", "ranking_changed": False},
            {"category_adjusted": "Security & Compliance", "base_weight_frac": 0.35, "test_weight_frac": 0.30,
             "rank1_base": "SUP_A", "rank1_adjusted": "SUP_A", "ranking_changed": False},
        ],
        "recommendation": {
            "secondary_supplier_id": "SUP_C",
            "secondary_tradeoffs": "Lower Year 1 cost and a simpler commercial structure, but a weaker "
                                   "functional fit score and a lower Security & Compliance subtotal than "
                                   "the recommended supplier.",
            "key_tradeoffs": "Solstice Cloud Systems is the only gate-clear supplier with current SOC 2 Type II "
                              "certification and the highest gate-clear Grand Total; Northbridge scores higher "
                              "on total but is disqualified from the compensatory recommendation by its Must "
                              "Have gate failure pending certification.",
            "next_steps": [
                "Confirm Solstice Cloud Systems' SOC 2 Type II scope covers the proposed hosting environment.",
                "Negotiate Year 1 pricing given the gap to Meridian Platform Group's lower cost position.",
                "Request Northbridge provide a dated SOC 2 Type II roadmap; revisit eligibility if "
                "certification completes before contract execution.",
                "Route the Must Have gate finding on Northbridge to the evaluation lead for confirmation, "
                "per the gate reconciliation rule.",
            ],
        },
        "clarifications_note": "Clarification requests were issued to Northbridge and Meridian regarding "
                                "certification scope and implementation timeline; responses are on file and "
                                "are reflected in the scores above.",
    }


# ===========================================================================
# Self-test / CLI
# ===========================================================================

def _run_self_test() -> int:
    import tempfile
    import zipfile

    print("=" * 78)
    print("evaluation_report_generator.py self-test")
    print("=" * 78)
    print(f"numeric_kernel.py available: {KERNEL_AVAILABLE}")
    print(f"python-docx available: {DOCX_AVAILABLE}"
          + (f" (version {DOCX_VERSION})" if DOCX_AVAILABLE else ""))
    print()

    results: List[tuple] = []

    def check(label, condition, detail=""):
        results.append((label, bool(condition), detail))
        status = "PASS" if condition else "FAIL"
        line = f"[{status}] {label}"
        if detail:
            line += f"  ({detail})"
        print(line)

    demo = _demo_evaluation_register()

    # --- Step 0: demo register validates and computes as hand-derived -----
    try:
        reg = validate_evaluation_input(copy.deepcopy(demo))
        check("validate_evaluation_input accepts the demo register", True)
    except Exception as e:
        check("validate_evaluation_input accepts the demo register", False, str(e))
        raise

    try:
        gt = compute_ground_truth(reg)
        check("compute_ground_truth runs via numeric_kernel.weighted_score()", True,
              f"grand_total={ {sid: round(v, 4) for sid, v in gt.grand_total.items()} }")
    except Exception as e:
        check("compute_ground_truth runs via numeric_kernel.weighted_score()", False, str(e))
        raise

    # Hand-derived expected values (independently computed from the demo
    # register's own numbers, not quoted from any skill file; evaluation-
    # engine's SKILL.md carries no numeric worked example, unlike should-
    # cost-builder's quoted '60+30+10=100' illustration).
    expected_totals = {"SUP_A": 3.9855, "SUP_B": 3.8450, "SUP_C": 3.3075}
    for sid, expected in expected_totals.items():
        check(
            f"Grand Total for {sid} matches hand-derived expectation ({expected})",
            abs(gt.grand_total[sid] - expected) < 1e-9,
            f"got {gt.grand_total[sid]}",
        )
    check("Effective weights sum to 1.0", abs(sum(gt.effective_weight.values()) - 1.0) < 1e-9)
    check("SUP_A carries the Must Have zero flag (REQ-010)", gt.must_have_zero_flag["SUP_A"] is True
          and gt.gate_detail["SUP_A"] == ["REQ-010"])
    check("SUP_B and SUP_C are gate clear", gt.must_have_zero_flag["SUP_B"] is False and gt.must_have_zero_flag["SUP_C"] is False)
    check("Compensatory rank-1 is SUP_A (highest total, gate-failed)", gt.compensatory_rank[0] == "SUP_A")
    check("Gate-clear rank-1 is SUP_B", gt.gate_clear_rank[0] == "SUP_B")
    check("Gate conflict correctly flagged", gt.gate_conflict is True)
    check("Recommended supplier is the gate-clear leader, SUP_B, not the gate-failed SUP_A",
          gt.recommended_supplier_id == "SUP_B")
    check("Sensitivity verdict flags Functional Fit as the sensitive category",
          "Functional Fit" in gt.sensitivity_verdict and "sensitive" in gt.sensitivity_verdict)

    try:
        run_hardcoded_invariant_checks(reg, gt)
        check("All hard-coded invariant checks PASS on the demo register", True)
    except ReconciliationError as e:
        check("All hard-coded invariant checks PASS on the demo register", False, str(e))
        raise

    # --- Step 1: validation refusal tests ----------------------------------
    def _expect_validation_error(label, broken_register):
        try:
            validate_evaluation_input(broken_register)
            check(label, False, "did not raise")
        except EvaluationValidationError as e:
            check(label, True, str(e)[:160])

    b1 = copy.deepcopy(demo)
    del b1["criteria"]
    _expect_validation_error("validate_evaluation_input refuses a register missing 'criteria'", b1)

    b2 = copy.deepcopy(demo)
    b2["criteria"][0]["entries"]["SUP_A"]["score"] = 7.5
    _expect_validation_error("validate_evaluation_input refuses a score outside [0.0, 5.0]", b2)

    b3 = copy.deepcopy(demo)
    b3["criteria"][0]["weight_in_category_frac"] = 0.9  # REQ-001+REQ-002 in-category no longer sums to 1.0
    _expect_validation_error("validate_evaluation_input refuses a category whose weight_in_category_frac does not sum to 1.0", b3)

    b4 = copy.deepcopy(demo)
    b4["criteria"][0]["category_weight_frac"] = 0.99  # disagrees with REQ-002's category_weight_frac for the same category
    _expect_validation_error("validate_evaluation_input refuses inconsistent category_weight_frac within one category", b4)

    b5 = copy.deepcopy(demo)
    b5["criteria"][2]["category_weight_frac"] = 30  # a percentage-mistake value, > 1.0
    b5["criteria"][3]["category_weight_frac"] = 30
    _expect_validation_error("validate_evaluation_input refuses a weight stored as a percentage (> 1.0)", b5)

    b6 = copy.deepcopy(demo)
    del b6["criteria"][0]["entries"]["SUP_C"]  # missing entry for a declared supplier
    _expect_validation_error("validate_evaluation_input refuses a criterion missing an entry for a declared supplier", b6)

    b7 = copy.deepcopy(demo)
    b7["criteria"][0]["priority"] = "Critical"  # unknown priority enum
    _expect_validation_error("validate_evaluation_input refuses an unknown priority enum", b7)

    b8 = copy.deepcopy(demo)
    b8["criteria"][0]["entries"]["SUP_A"]["coverage"] = "Sort Of Meets"  # unknown coverage enum
    _expect_validation_error("validate_evaluation_input refuses an unknown coverage enum", b8)

    b9 = copy.deepcopy(demo)
    b9["risk_matrix"]["Legal/Contractual"]["SUP_A"] = "Severe"  # unknown severity enum
    _expect_validation_error("validate_evaluation_input refuses an unknown risk severity enum", b9)

    # --- Step 2: invariant-catches-a-bug tests -----------------------------
    def _expect_reconciliation_error(label, fn):
        try:
            fn()
            check(label, False, "did not raise, but should have")
        except ReconciliationError as e:
            check(label, True, str(e)[:160])

    bad_gt1 = copy.deepcopy(gt)
    bad_gt1.grand_total["SUP_A"] = 999.0  # break the reconcile-to-grand-total identity
    _expect_reconciliation_error(
        "Weighted-contributions-reconcile invariant CORRECTLY REJECTS a corrupted grand_total",
        lambda: _assert_weighted_contributions_reconcile(bad_gt1),
    )

    bad_gt2 = copy.deepcopy(gt)
    bad_gt2.grand_total["SUP_B"] = 5.5  # out of [0.0, 5.0]
    _expect_reconciliation_error(
        "Grand-total-range invariant CORRECTLY REJECTS an out-of-range total",
        lambda: _assert_grand_total_range(bad_gt2),
    )

    bad_gt3 = copy.deepcopy(gt)
    bad_gt3.category_subtotal["SUP_C"]["Commercial"] += 1.0  # break category-subtotal footing
    _expect_reconciliation_error(
        "Category-subtotals-reconcile invariant CORRECTLY REJECTS a broken category subtotal",
        lambda: _assert_category_subtotals_reconcile(bad_gt3),
    )

    bad_gt4 = copy.deepcopy(gt)
    bad_gt4.compensatory_rank = list(reversed(bad_gt4.compensatory_rank))  # break sort order
    _expect_reconciliation_error(
        "Ranking-sorted invariant CORRECTLY REJECTS a reversed ranking",
        lambda: _assert_ranking_sorted(bad_gt4),
    )

    bad_gt5 = copy.deepcopy(gt)
    bad_gt5.recommended_supplier_id = "SUP_A"  # silently recommend the gate-failed leader
    _expect_reconciliation_error(
        "Must-Have-gate-not-buried invariant CORRECTLY REJECTS burying a gate failure behind a high total",
        lambda: _assert_must_have_gate_not_buried(bad_gt5),
    )

    bad_gt6 = copy.deepcopy(gt)
    bad_gt6.coverage_counts["SUP_A"]["Fully Meets"] -= 1  # counts no longer foot to the criteria count
    _expect_reconciliation_error(
        "Requirements-coverage-counts-foot invariant CORRECTLY REJECTS a broken tally",
        lambda: _assert_requirements_coverage_counts_foot(reg, bad_gt6),
    )

    b_sens = copy.deepcopy(demo)
    b_sens["sensitivity_rows"][0]["rank1_base"] = "SUP_C"  # disagrees with the ground truth's own rank-1
    reg_sens = validate_evaluation_input(b_sens)
    _expect_reconciliation_error(
        "Sensitivity-rank1-base-matches-ground-truth invariant CORRECTLY REJECTS a drifted rank1_base",
        lambda: _assert_sensitivity_rank1_base_matches(reg_sens, gt),
    )

    # --- Step 3: build + save + reopen the real document -------------------
    if not DOCX_AVAILABLE:
        check("python-docx available to write a real .docx file", False,
              "python-docx is NOT installed in this interpreter; DOCX writing could not be exercised.")
        print()
        print("Cannot proceed past this point without python-docx.")
    else:
        check("python-docx available to write a real .docx file", True, f"version {DOCX_VERSION}")

        tmp_dir = tempfile.gettempdir()
        full_path = os.path.join(tmp_dir, "evaluation_report_selftest_full.docx")
        brief_path = os.path.join(tmp_dir, "evaluation_report_selftest_brief.docx")

        try:
            generate_evaluation_report(demo, full_path, mode_override="Full")
            check("generate_evaluation_report() ran end-to-end (Full mode) without raising", True)
        except Exception as e:
            check("generate_evaluation_report() ran end-to-end (Full mode) without raising", False, str(e))
            raise

        demo_no_clarifications = copy.deepcopy(demo)
        del demo_no_clarifications["clarifications_note"]
        try:
            generate_evaluation_report(demo_no_clarifications, brief_path, mode_override="Brief")
            check("generate_evaluation_report() ran end-to-end (Brief mode, no clarifications_note) without raising", True)
        except Exception as e:
            check("generate_evaluation_report() ran end-to-end (Brief mode) without raising", False, str(e))
            raise

        for label_path, path in (("Full", full_path), ("Brief", brief_path)):
            exists = os.path.exists(path)
            size = os.path.getsize(path) if exists else 0
            check(f"{label_path} DOCX file written to {path}", exists and size > 0, f"size={size} bytes")
            check(f"{label_path} DOCX unzips cleanly (valid OOXML zip container)", zipfile.is_zipfile(path))

        try:
            reopened_full = Document(full_path)
            all_text_full = "\n".join(p.text for p in reopened_full.paragraphs)
            for tbl in reopened_full.tables:
                for row in tbl.rows:
                    all_text_full += "\n" + "\n".join(c.text for c in row.cells)

            expected_headings = [
                "1. Executive Summary", "2. Evaluation Methodology", "3. Supplier Summaries",
                "4. Score Comparison", "5. Risk Assessment Summary", "6. Commercial Comparison",
                "7. Recommendation", "8. Appendices",
            ]
            check("Re-opened Full DOCX contains all 8 Evaluation Schemas section headings",
                  all(h in all_text_full for h in expected_headings))
            check("Re-opened Full DOCX names the recommended supplier (Solstice Cloud Systems)",
                  "Solstice Cloud Systems (illustrative)" in all_text_full)
            check("Re-opened Full DOCX carries the Must Have gate conflict language",
                  "Compensatory scoring cannot offset a hard requirement gap" in all_text_full)
            check("Re-opened Full DOCX shows SUP_A's Grand Total (3.99, rounded)", "3.99" in all_text_full)
            check("Re-opened Full DOCX shows SUP_B's Grand Total (3.85, rounded)", "3.85" in all_text_full or "3.84" in all_text_full)
            check("Re-opened Full DOCX includes the Secondary Option subsection", "Secondary Option" in all_text_full)
            check("Re-opened Full DOCX includes the Weight Perturbation sensitivity table header", "Ranking Changed?" in all_text_full)
            check("Re-opened Full DOCX includes Appendix A's per-criterion rollup header", "GRAND TOTAL" in all_text_full)
            check("Re-opened Full DOCX includes the populated clarifications note", "SOC 2 Type II" in all_text_full)

            n_tables_full = len(reopened_full.tables)
            check("Re-opened Full DOCX contains multiple tables", n_tables_full >= 6, f"n_tables={n_tables_full}")
        except Exception as e:
            check("Re-opened Full DOCX structural spot-checks", False, str(e))

        try:
            reopened_brief = Document(brief_path)
            all_text_brief = "\n".join(p.text for p in reopened_brief.paragraphs)
            for tbl in reopened_brief.tables:
                for row in tbl.rows:
                    all_text_brief += "\n" + "\n".join(c.text for c in row.cells)

            check("Re-opened Brief DOCX still contains all 8 section headings",
                  all(h in all_text_brief for h in expected_headings))
            check("Re-opened Brief DOCX OMITS the Secondary Option subsection", "Secondary Option" not in all_text_brief)
            check("Re-opened Brief DOCX OMITS the full sensitivity perturbation table", "Ranking Changed?" not in all_text_brief)
            check("Re-opened Brief DOCX still shows the one-line sensitivity verdict", "Functional Fit" in all_text_brief)
            check("Re-opened Brief DOCX falls back to the no-clarifications default text",
                  "No clarifications were requested or received" in all_text_brief)
            check("Brief DOCX has fewer tables than Full DOCX (category breakdown and Appendix A rollup omitted)",
                  len(reopened_brief.tables) < n_tables_full,
                  f"brief={len(reopened_brief.tables)}, full={n_tables_full}")
        except Exception as e:
            check("Re-opened Brief DOCX structural spot-checks", False, str(e))

    print()
    print("=" * 78)
    passed = sum(1 for _, ok, _ in results if ok)
    failed = sum(1 for _, ok, _ in results if not ok)
    total = len(results)
    print(f"SUMMARY: {passed}/{total} passed, {failed}/{total} failed")
    if failed:
        print("FAILED CASES:")
        for label, ok, detail in results:
            if not ok:
                print(f"  - {label}: {detail}")
    print("=" * 78)
    return 1 if failed else 0


def main(argv: Optional[List[str]] = None) -> int:
    """CLI entry point. Usage:
        python evaluation_report_generator.py --input register.json --output evaluation_report.docx [--mode brief|full]
        python evaluation_report_generator.py --demo
        python evaluation_report_generator.py --self-test
        python evaluation_report_generator.py                 (no args -> runs the self-test)
    """
    import argparse

    parser = argparse.ArgumentParser(
        description="Generate evaluation_report.docx from a validated evaluation register "
                    "(evaluation-engine-1c344a). See the module docstring for the input JSON schema."
    )
    parser.add_argument("--input", "-i", help="Path to a JSON file containing the evaluation register.")
    parser.add_argument("--output", "-o", default="evaluation_report.docx",
                         help="Output .docx path (default: evaluation_report.docx in the current directory).")
    parser.add_argument("--mode", choices=["brief", "full"], default=None,
                         help="Override report_mode (Brief or Full). Defaults to the register's own report_mode, or Full.")
    parser.add_argument("--demo", action="store_true",
                         help="Run the built-in self-test suite (generates the illustrative demo register's "
                              "DOCX, reopens it, and asserts every expected section, table, and value).")
    parser.add_argument("--self-test", action="store_true", dest="self_test",
                         help="Alias for --demo.")
    args = parser.parse_args(argv)

    if args.demo or args.self_test or not args.input:
        return _run_self_test()

    import json
    try:
        with open(args.input, "r", encoding="utf-8") as f:
            raw = json.load(f)
    except (OSError, ValueError) as e:
        print(f"ERROR: could not read/parse --input {args.input!r}: {e}", file=sys.stderr)
        return 1

    mode_override = args.mode.capitalize() if args.mode else None
    try:
        generate_evaluation_report(raw, args.output, mode_override=mode_override)
    except ImportError as e:
        print(f"ERROR: {e}", file=sys.stderr)
        print("The evaluation register was valid and ground truth computed cleanly; "
              "only the DOCX-writing step could not run because python-docx is missing.", file=sys.stderr)
        return 1
    except (EvaluationValidationError, ReconciliationError) as e:
        print(f"ERROR: {e}", file=sys.stderr)
        return 1

    print(f"Wrote {args.output}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
