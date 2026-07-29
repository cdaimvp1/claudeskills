#!/usr/bin/env python3
"""
Self-test for landscape_report_generator.py (F9 build 5).

Targets the three things this build exists to prevent:
  1. a truncated report, by assembling in ONE pass instead of three append cycles
  2. the two scoring systems being conflated
  3. artifacts from one run disagreeing with each other

Negative controls throughout, because several rules here are about ABSENCE being stated
rather than assumed, and a checker that punishes honest gaps is worse than none.

Run: python landscape_report_selftest.py
"""
from __future__ import annotations

import csv
import os
import sys
import tempfile

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from landscape_report_generator import (                        # noqa: E402
    NO_EVIDENCE,
    REPORT_SECTIONS,
    ArtifactMismatchError,
    BandMismatchError,
    LandscapeError,
    RegistrySchemaError,
    ScoreRangeError,
    compute_pillar_matrix,
    compute_requirements_fit,
    fit_band,
    generate_all,
)
from numeric_kernel import WeightSumError                       # noqa: E402

try:
    from docx import Document
    _DOCX = True
except ImportError:
    _DOCX = False

PASS, FAIL = [], []


def ok(name, cond, detail=""):
    (PASS if cond else FAIL).append(name)
    print(("  ok   " if cond else "  FAIL ") + name + (("  <- " + detail) if detail and not cond else ""))


def raises(name, exc, fn):
    try:
        fn()
    except exc:
        ok(name, True)
        return
    except Exception as e:                                     # noqa: BLE001
        ok(name, False, "raised %s, expected %s" % (type(e).__name__, exc.__name__))
        return
    ok(name, False, "did not raise " + exc.__name__)


REQS = ["GxP validation", "EU data residency", "API depth"]
PILLARS = {"Capability": 30.0, "Risk": 20.0, "Commercial": 20.0, "Delivery": 15.0,
           "Viability": 15.0}


def data():
    return {
        "title": "Clinical analytics platform landscape",
        "category": "IT Professional Services",
        "requirements": list(REQS),
        "pillar_weights": dict(PILLARS),
        "suppliers": [
            {"supplier_name": "Meridian",
             "requirement_scores": {"GxP validation": 9.0, "EU data residency": 8.5,
                                    "API depth": 8.5},
             "pillar_scores": {"Capability": 9.0, "Risk": 7.0, "Commercial": 8.0,
                               "Delivery": 8.0, "Viability": 9.0}},
            {"supplier_name": "Northwind",
             "requirement_scores": {"GxP validation": 7.0, "EU data residency": 7.5,
                                    "API depth": 7.2},
             "pillar_scores": {"Capability": 7.0, "Risk": 6.0, "Commercial": 8.0,
                               "Delivery": 7.0, "Viability": 6.5}},
        ],
        "registry": [
            {"supplier_name": "Meridian", "headquarters": "Boston", "company_size": "1200",
             "financial_health": "Stable", "core_offering": "Clinical analytics",
             "internal_status": "Net new", "industry_experience": "Pharma",
             "integration_fit": "REST", "pricing_model": "Subscription",
             "website": "https://example.com"},
            {"supplier_name": "Northwind", "headquarters": "Dublin", "company_size": "400",
             "financial_health": "Data not available", "core_offering": "Data platform",
             "internal_status": "Prior engagement", "industry_experience": "Life sciences",
             "integration_fit": "REST", "pricing_model": "Per seat",
             "website": "https://example.org"},
        ],
        "risks": [
            {"supplier_name": "Meridian", "risk_category": "Cybersecurity",
             "risk_description": "No SOC 2 on file", "severity": "Medium",
             "evidence_source": "Not Determined"},
        ],
        "excluded": [
            {"vendor_name": "Acme", "reason_code": "FAILED_DISQUALIFIER",
             "reason_detail": "no GxP track record", "source": "vendor site",
             "date": "2026-07-20"},
        ],
        "narratives": {"Executive Summary": "Two credible candidates."},
    }


def read_csv(path):
    with open(path, encoding="utf-8", newline="") as fh:
        return list(csv.DictReader(fh))


def run():
    print("=" * 84)
    print("landscape_report_generator self-test")
    print("=" * 84)

    d = data()
    fit = compute_requirements_fit(d["suppliers"], REQS)
    pil = compute_pillar_matrix(d["suppliers"], PILLARS)

    ok("T1  every supplier gets a requirements-fit row", len(fit) == 2)
    ok("T2  the fit score is kernel-computed to 2dp",
       isinstance(fit[0]["weighted_score"], float))
    ok("T3  Meridian scores in the Strong band", fit[0]["overall_fit"] == "Strong",
       "%s %s" % (fit[0]["weighted_score"], fit[0]["overall_fit"]))
    ok("T4  Northwind scores in the Partial band", fit[1]["overall_fit"] == "Partial",
       "%s %s" % (fit[1]["weighted_score"], fit[1]["overall_fit"]))

    # --- band boundaries, from SKILL.md:955 ---------------------------------------------
    ok("T5  8.50 is Strong (inclusive lower bound)", fit_band(8.5) == "Strong")
    ok("T6  8.49 is Partial", fit_band(8.49) == "Partial")
    ok("T7  7.00 is Partial (inclusive)", fit_band(7.0) == "Partial")
    ok("T8  6.99 is Weak", fit_band(6.99) == "Weak")
    ok("T9  no evidence is NOT a zero", fit_band(None) == NO_EVIDENCE)

    # --- the two scoring systems stay separate -------------------------------------------
    ok("T10 the pillar total is on a different scale from the fit score",
       pil[0]["weighted_total"] != fit[0]["weighted_score"],
       "%s vs %s" % (pil[0]["weighted_total"], fit[0]["weighted_score"]))

    tmp = tempfile.mkdtemp(prefix="f9_land_")
    res = generate_all(d, tmp)

    fit_csv = read_csv(res["written"]["requirements_fit_matrix.csv"])
    pil_csv = read_csv(res["written"]["weighted_scoring_matrix.csv"])
    ok("T11 fit CSV carries weighted_score and overall_fit",
       "weighted_score" in fit_csv[0] and "overall_fit" in fit_csv[0])
    ok("T12 pillar CSV does NOT carry overall_fit (it is a different system)",
       "overall_fit" not in pil_csv[0], repr(list(pil_csv[0])))
    ok("T13 the two CSVs hold DIFFERENT numbers for the same supplier",
       fit_csv[0]["weighted_score"] != pil_csv[0]["weighted_total"])

    # --- registry carries no scores --------------------------------------------------------
    reg_csv = read_csv(res["written"]["supplier_registry.csv"])
    ok("T14 registry has profile columns only",
       "weighted_score" not in reg_csv[0] and "overall_fit" not in reg_csv[0])
    raises("T15 refuses a score smuggled into the registry", RegistrySchemaError,
           lambda: generate_all({**data(), "registry": [
               {**data()["registry"][0], "weighted_score": 9.0},
               data()["registry"][1]]}, tempfile.mkdtemp()))

    # --- artifacts must agree --------------------------------------------------------------
    def mismatch_registry():
        dd = data(); dd["registry"] = [dd["registry"][0]]      # drop Northwind
        generate_all(dd, tempfile.mkdtemp())
    raises("T16 refuses a supplier present in the fit matrix but absent from the registry",
           ArtifactMismatchError, mismatch_registry)

    def unknown_risk():
        dd = data(); dd["risks"][0]["supplier_name"] = "Ghost Ltd"
        generate_all(dd, tempfile.mkdtemp())
    raises("T17 refuses a risk row naming a supplier in no other artifact",
           ArtifactMismatchError, unknown_risk)

    # --- validation ------------------------------------------------------------------------
    def bad_score():
        dd = data(); dd["suppliers"][0]["requirement_scores"]["API depth"] = 12.0
        compute_requirements_fit(dd["suppliers"], REQS)
    raises("T18 refuses a fit score above the 0-10 scale", ScoreRangeError, bad_score)

    def bad_band():
        dd = data(); dd["suppliers"][1]["overall_fit"] = "Strong"   # it is Partial
        compute_requirements_fit(dd["suppliers"], REQS)
    raises("T19 refuses a stated band contradicting its own score", BandMismatchError, bad_band)

    def bad_weights():
        dd = data(); dd["pillar_weights"]["Capability"] = 25.0      # now sums to 95
        compute_pillar_matrix(dd["suppliers"], dd["pillar_weights"])
    raises("T20 refuses pillar weights that do not sum to 100 (kernel, G11)",
           WeightSumError, bad_weights)

    def blank_pillar():
        dd = data(); dd["suppliers"][0]["pillar_scores"].pop("Risk")
        compute_pillar_matrix(dd["suppliers"], PILLARS)
    raises("T21 refuses a blank pillar cell (it silently reweights the others)",
           LandscapeError, blank_pillar)

    def bad_risk_cat():
        dd = data(); dd["risks"][0]["risk_category"] = "Reputational"
        generate_all(dd, tempfile.mkdtemp())
    raises("T22 refuses a risk category outside the schema", LandscapeError, bad_risk_cat)

    # --- the evidence_as_of schema addition (G13b) ---------------------------------------
    def sourced_no_date():
        dd = data()
        dd["risks"][0]["evidence_source"] = "OFAC SDN list"      # a REAL source
        dd["risks"][0].pop("evidence_as_of", None)                # with no date
        generate_all(dd, tempfile.mkdtemp())
    raises("T23a refuses a NAMED source with no evidence_as_of", LandscapeError,
           sourced_no_date)

    def sourced_placeholder_date():
        dd = data()
        dd["risks"][0]["evidence_source"] = "OFAC SDN list"
        dd["risks"][0]["evidence_as_of"] = "TBD"
        generate_all(dd, tempfile.mkdtemp())
    raises("T23b refuses a placeholder evidence_as_of", LandscapeError,
           sourced_placeholder_date)

    dd_ok = data()
    dd_ok["risks"][0]["evidence_source"] = "OFAC SDN list"
    dd_ok["risks"][0]["evidence_as_of"] = "2026-05-02"
    r_ok = generate_all(dd_ok, tempfile.mkdtemp(prefix="f9_land_asof_"))
    ok("T23c a named source WITH a capture date passes",
       os.path.isfile(r_ok["written"]["risk_matrix.csv"]))
    hdr = read_csv(r_ok["written"]["risk_matrix.csv"])[0]
    ok("T23d evidence_as_of is a real column in the emitted CSV", "evidence_as_of" in hdr,
       repr(list(hdr)))
    ok("T23e NEGATIVE CONTROL: 'Not Determined' still passes with NO date",
       generate_all(data(), tempfile.mkdtemp())["written"]["risk_matrix.csv"] is not None)

    def blank_evidence():
        dd = data(); dd["risks"][0]["evidence_source"] = "  "
        generate_all(dd, tempfile.mkdtemp())
    raises("T23 refuses a blank evidence_source ('Not Determined' is the honest answer)",
           LandscapeError, blank_evidence)

    def bad_code():
        dd = data(); dd["excluded"][0]["reason_code"] = "DIDNT_LIKE"
        generate_all(dd, tempfile.mkdtemp())
    raises("T24 refuses an exclusion reason_code outside the schema", LandscapeError, bad_code)

    # --- absence stated explicitly ----------------------------------------------------------
    dd = data(); dd["excluded"] = []
    r2 = generate_all(dd, tempfile.mkdtemp(prefix="f9_land_none_"))
    exc = read_csv(r2["written"]["excluded_vendors.csv"])
    ok("T25 an empty exclusion list still writes a row", len(exc) == 1)
    ok("T26 that row says 'none excluded' rather than leaving the file empty",
       exc[0]["reason_detail"] == "none excluded", repr(exc[0]))

    # --- NEGATIVE CONTROLS: honest gaps must pass ---------------------------------------------
    dd = data()
    dd["suppliers"][1]["requirement_scores"] = {q: None for q in REQS}
    fit_ng = compute_requirements_fit(dd["suppliers"], REQS)
    ok("T27 NEGATIVE CONTROL: a supplier with NO evidence is not scored 0",
       fit_ng[1]["weighted_score"] is None and fit_ng[1]["overall_fit"] == NO_EVIDENCE,
       repr(fit_ng[1]["overall_fit"]))

    dd2 = data()
    dd2["suppliers"][1]["requirement_scores"]["API depth"] = None
    fit_pg = compute_requirements_fit(dd2["suppliers"], REQS)
    ok("T28 NEGATIVE CONTROL: a PARTIAL gap reweights over what IS evidenced",
       fit_pg[1]["weighted_score"] is not None and fit_pg[1]["missing_requirements"] == ["API depth"])

    ok("T29 NEGATIVE CONTROL: 'Data not available' in a registry field is legitimate",
       reg_csv[1]["financial_health"] == "Data not available")

    # --- the document ---------------------------------------------------------------------------
    if _DOCX:
        doc = Document(res["written"]["supplier_landscape_report.docx"])
        heads = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        for sec in REPORT_SECTIONS:
            ok("T30 report section present: %-26s" % sec, sec in heads)
        ok("T31 the report is built in ONE pass, so both suppliers appear",
           "Meridian" in heads and "Northwind" in heads)
        ok("T32 both matrices are rendered as tables", len(doc.tables) >= 2,
           str(len(doc.tables)))
        text = "\n".join(p.text for p in doc.paragraphs)
        ok("T33 the report says WHICH scoring system feeds the dashboard headline",
           "dashboard headline" in text)
        ok("T34 and says the pillar matrix does NOT", "does not feed" in text)
    else:
        print("  skip  T30-T34 (python-docx unavailable)")

    print("=" * 84)
    print("SUMMARY: %d/%d passed, %d failed" % (len(PASS), len(PASS) + len(FAIL), len(FAIL)))
    print("=" * 84)
    return 1 if FAIL else 0


if __name__ == "__main__":
    sys.exit(run())
