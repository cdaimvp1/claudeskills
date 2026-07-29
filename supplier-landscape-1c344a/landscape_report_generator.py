#!/usr/bin/env python3
"""
landscape_report_generator.py — the landscape report and its CSVs (F9 build 5).

WHY THIS EXISTS
---------------
`SKILL.md` told the model to build the report in THREE PASSES: write a partial document,
reopen it, append, save, reopen, append, save. F2 removed exactly that pattern from
rfp-response-analysis, and deferred this one by name because its generator did not yet
cover the document.

**This is deliberately NOT fixed the way F2 was.** Collapsing three appends into a single
model-authored write is the truncation failure G10 warns about: the document silently comes
out short and looks finished. The fix is a real generator, so the document is ASSEMBLED
rather than written, and length stops being a generation-time risk at all.

THE CSVs COME FROM THE SAME CALL (the F6 lesson)
------------------------------------------------
`weighted_scoring_matrix.csv`, `requirements_fit_matrix.csv`, `supplier_registry.csv`,
`risk_matrix.csv` and `excluded_vendors.csv` are emitted by this generator, from the same
object that builds the report. Emitting them separately is how a report and its own
appendix end up disagreeing, and a reader who spots the disagreement cannot tell which one
is wrong.

TWO SCORING SYSTEMS, AND THEY ARE NOT INTERCHANGEABLE
-----------------------------------------------------
SKILL.md:381 is explicit and this generator enforces it:

  requirements-fit, requirement-count-weighted (0-10)   -> requirements_fit_matrix.csv
                                                        -> the dashboard headline score
  8-pillar, percentage-weighted                         -> weighted_scoring_matrix.csv
                                                        -> a report table ONLY

Writing one into the other's artifact produces a number that is individually correct and
completely wrong in context. Both are computed by `weighted_score()` in the vendored kernel
(G11), never by model arithmetic.

WHAT IT REFUSES TO DO
---------------------
  * pillar weights that do not sum to 100          -> the kernel's WeightSumError
  * a fit score outside 0-10                       -> ScoreRangeError
  * an overall_fit band contradicting its score    -> BandMismatchError
  * a supplier in one artifact and missing from another -> ArtifactMismatchError
  * scores written into supplier_registry.csv      -> RegistrySchemaError
  * an empty excluded_vendors list emitted silently -> writes the explicit "none excluded" row
"""
from __future__ import annotations

import csv
import json
import os
import sys
from datetime import datetime

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from numeric_kernel import assert_weight_sum, weighted_score    # noqa: E402  (G11)

try:
    from docx import Document
    from docx.shared import Pt
    _DOCX = True
except ImportError:                                             # pragma: no cover
    _DOCX = False

# From SKILL.md:955. Bands are inclusive at the lower bound.
FIT_BANDS = [(8.5, "Strong"), (7.0, "Partial"), (0.0, "Weak")]
NO_EVIDENCE = "Information Not Provided"

# Sources that honestly say "no evidence" rather than naming one. These need no capture
# date, because there is nothing to date.
EVIDENCE_ABSTENTIONS = ("Not Determined", "Not determined", "not determined")

# Accepted capture-date formats. Same rule and same reasoning as the research-table
# generators (#32): named months are unambiguous and accepted, slash formats are not.
_DATE_PLACEHOLDERS = frozenset({"", "tbd", "tba", "n/a", "na", "none", "null",
                                "unknown", "recent", "current", "various", "-", "?"})
_DATE_FORMATS = ("%Y-%m-%d", "%Y-%m", "%Y", "%b %d, %Y", "%B %d, %Y",
                 "%b %Y", "%B %Y", "%d %b %Y", "%d %B %Y")


def _valid_as_of(value):
    if not isinstance(value, str):
        return False
    raw = value.strip()
    if raw.lower() in _DATE_PLACEHOLDERS:
        return False
    for fmt in _DATE_FORMATS:
        try:
            parsed = datetime.strptime(raw, fmt)
        except ValueError:
            continue
        return 1990 <= parsed.year <= 2100
    return False


RISK_CATEGORIES = ("Legal", "Cybersecurity", "Operational", "Geopolitical", "Financial")
SEVERITIES = ("Low", "Medium", "High")

EXCLUSION_CODES = ("FAILED_DISQUALIFIER", "OUT_OF_SCOPE", "INSUFFICIENT_EVIDENCE",
                   "DUPLICATE", "BUYER_EXCLUDED")

REGISTRY_COLUMNS = ["supplier_name", "headquarters", "company_size", "financial_health",
                    "core_offering", "internal_status", "industry_experience",
                    "integration_fit", "pricing_model", "website"]

# The report's sections, in order. The generator owns the STRUCTURE; the model supplies the
# narrative for each section in the data object. Assembling in one pass is the whole point.
REPORT_SECTIONS = [
    "Executive Summary",
    "Market Context",
    "Supplier Profiles",
    "Cross-Vendor Comparison",
    "Requirements Fit Matrix",
    "Risk Assessment",
    "Recommendation",
]

NOT_AVAILABLE = "Data not available"


class LandscapeError(Exception):
    """Base for every refusal."""


class ScoreRangeError(LandscapeError):
    pass


class BandMismatchError(LandscapeError):
    pass


class ArtifactMismatchError(LandscapeError):
    pass


class RegistrySchemaError(LandscapeError):
    pass


class DocxUnavailableError(LandscapeError):
    pass


def fit_band(score):
    """0-10 score -> categorical band. None means no evidence, NOT a zero."""
    if score is None:
        return NO_EVIDENCE
    for floor, label in FIT_BANDS:
        if score >= floor:
            return label
    return "Weak"


def _check_score(value, label, lo=0.0, hi=10.0):
    if value is None:
        return None
    try:
        v = float(value)
    except (TypeError, ValueError):
        raise ScoreRangeError("%s is %r, which is not a number" % (label, value))
    if not (lo <= v <= hi):
        raise ScoreRangeError(
            "%s is %.2f, outside the documented %g-%g scale. A score off its own scale is "
            "not a strong result, it is a unit error, and it will rank this supplier against "
            "others measured differently." % (label, v, lo, hi)
        )
    return v


def compute_requirements_fit(suppliers, requirements):
    """Requirement-count-weighted score per supplier (the DASHBOARD headline system).

    Equal weight per requirement, hence "requirement-count-weighted". Computed by the
    kernel, never by hand.
    """
    if not requirements:
        raise LandscapeError("no requirements supplied; there is no fit matrix to compute")
    n = len(requirements)

    rows = []
    for s in suppliers:
        name = (s.get("supplier_name") or "").strip()
        if not name:
            raise LandscapeError("a supplier has no supplier_name")
        raw = s.get("requirement_scores") or {}
        missing = [r for r in requirements if raw.get(r) is None]
        scores = {}
        for r in requirements:
            v = _check_score(raw.get(r), "%s / %s" % (name, r))
            scores[r] = v

        if len(missing) == n:
            # No evidence at all. Reporting 0.0 would rank this supplier last on merit
            # when in fact nothing was measured.
            weighted, band = None, NO_EVIDENCE
        else:
            present = {r: v for r, v in scores.items() if v is not None}
            pw = {r: 1.0 / len(present) for r in present}
            weighted = round(weighted_score(present, pw), 2)
            band = fit_band(weighted)

        stated = (s.get("overall_fit") or "").strip()
        if stated and stated != band:
            raise BandMismatchError(
                "supplier %r states overall_fit %r but its weighted score %s puts it in %r. "
                "The band is derived from the score; a stated band that disagrees means one "
                "of the two was edited by hand, and this value feeds the rfp-engine handoff."
                % (name, stated, "None" if weighted is None else "%.2f" % weighted, band)
            )

        rows.append({"supplier_name": name, "scores": scores,
                     "weighted_score": weighted, "overall_fit": band,
                     "missing_requirements": missing})
    return rows


def compute_pillar_matrix(suppliers, pillar_weights):
    """The 8-pillar PERCENTAGE-weighted matrix. A REPORT TABLE ONLY (SKILL.md:381).

    Deliberately separate from the requirements-fit system above. This number must never be
    written into requirements_fit_matrix.csv, nor used as the dashboard headline.
    """
    if not pillar_weights:
        raise LandscapeError("no pillar weights supplied")
    rows = []
    for s in suppliers:
        name = s["supplier_name"]
        raw = s.get("pillar_scores") or {}
        missing = [p for p in pillar_weights if raw.get(p) is None]
        if missing:
            raise LandscapeError(
                "supplier %r has no pillar score for %s. The 8-pillar matrix is a complete "
                "grid; a blank cell silently reweights every other pillar for that supplier."
                % (name, ", ".join(missing))
            )
        scores = {p: _check_score(raw[p], "%s / %s" % (name, p)) for p in pillar_weights}

        # TWO different weight conventions, and mixing them is a silent scaling bug.
        # SKILL.md states the 8 pillars as PERCENTAGES summing to 100, but
        # weighted_score() refuses any weight set that does not sum to 1.0. So the
        # percentage set is validated on its own 100-scale by assert_weight_sum (which
        # names the over- or under-allocation), then converted to fractions for scoring.
        #
        # Calling weighted_score() with the raw percentages would raise on every valid
        # input; multiplying without validating would silently score a 95-point weight set
        # as though it footed. Both checks are needed, in this order.
        assert_weight_sum(pillar_weights, expected=100.0)
        fractions = {p: w / 100.0 for p, w in pillar_weights.items()}
        total = round(weighted_score(scores, fractions), 2)
        rows.append({"supplier_name": name, "scores": scores, "weighted_total": total})
    return rows


def check_artifact_consistency(fit_rows, pillar_rows, registry, risks):
    """Every artifact must describe the same supplier set. The F6 lesson, enforced."""
    fit = {r["supplier_name"] for r in fit_rows}
    pillar = {r["supplier_name"] for r in pillar_rows}
    reg = {(r.get("supplier_name") or "").strip() for r in registry}
    risked = {(r.get("supplier_name") or "").strip() for r in risks}

    for label, s in (("weighted_scoring_matrix", pillar), ("supplier_registry", reg)):
        if s != fit:
            raise ArtifactMismatchError(
                "%s covers %s but requirements_fit_matrix covers %s. Artifacts built from "
                "one run must describe one supplier set: a reader who spots the difference "
                "cannot tell which artifact is wrong."
                % (label, sorted(s), sorted(fit))
            )
    unknown = risked - fit
    if unknown:
        raise ArtifactMismatchError(
            "risk_matrix names supplier(s) %s that appear in no other artifact"
            % sorted(unknown)
        )
    return True


def check_registry_schema(registry):
    """`supplier_registry.csv` carries PROFILE data only, never scores (SKILL.md:379)."""
    banned = ("weighted_score", "overall_fit", "score", "fit_score", "rank")
    for i, row in enumerate(registry or []):
        for k in row:
            if k in banned:
                raise RegistrySchemaError(
                    "supplier_registry[%d] carries %r. The registry is profile data only; "
                    "scores live in the two scoring artifacts. A score duplicated here is a "
                    "second copy that will eventually disagree with the first." % (i, k)
                )
    return True


# --------------------------------------------------------------------------- CSV emitters

def _write_csv(path, columns, rows):
    with open(path, "w", encoding="utf-8", newline="") as fh:
        w = csv.DictWriter(fh, fieldnames=columns, extrasaction="ignore")
        w.writeheader()
        for r in rows:
            w.writerow(r)
    return path


def write_requirements_fit_csv(fit_rows, requirements, path):
    cols = ["supplier_name"] + list(requirements) + ["weighted_score", "overall_fit"]
    out = []
    for r in fit_rows:
        row = {"supplier_name": r["supplier_name"]}
        for q in requirements:
            v = r["scores"].get(q)
            row[q] = "" if v is None else "%.2f" % v
        row["weighted_score"] = "" if r["weighted_score"] is None else "%.2f" % r["weighted_score"]
        row["overall_fit"] = r["overall_fit"]
        out.append(row)
    return _write_csv(path, cols, out)


def write_pillar_csv(pillar_rows, pillars, path):
    cols = ["supplier_name"] + list(pillars) + ["weighted_total"]
    out = []
    for r in pillar_rows:
        row = {"supplier_name": r["supplier_name"]}
        for p in pillars:
            row[p] = "%.2f" % r["scores"][p]
        row["weighted_total"] = "%.2f" % r["weighted_total"]
        out.append(row)
    return _write_csv(path, cols, out)


def write_registry_csv(registry, path):
    check_registry_schema(registry)
    return _write_csv(path, REGISTRY_COLUMNS, registry)


def write_risk_csv(risks, path):
    """risk_matrix.csv. Carries a per-row capture date as of the G13b schema change.

    `evidence_as_of` was added because a source without a date is not provenance: a risk
    citation from 2019 and one from last week were previously indistinguishable, and a
    stale risk read is the kind that gets a supplier approved.

    Required WHENEVER a real source is named. NOT required for `Not Determined`, because
    there is no evidence to date, and demanding one there would push a caller toward
    inventing a date to satisfy the column.
    """
    cols = ["supplier_name", "risk_category", "risk_description", "severity",
            "evidence_source", "evidence_as_of"]
    for i, r in enumerate(risks or []):
        if r.get("risk_category") not in RISK_CATEGORIES:
            raise LandscapeError(
                "risk_matrix[%d] category %r is outside %s"
                % (i, r.get("risk_category"), list(RISK_CATEGORIES)))
        if r.get("severity") not in SEVERITIES:
            raise LandscapeError(
                "risk_matrix[%d] severity %r is outside %s"
                % (i, r.get("severity"), list(SEVERITIES)))
        source = (r.get("evidence_source") or "").strip()
        if not source:
            raise LandscapeError(
                "risk_matrix[%d] has no evidence_source. The schema requires a URL, a "
                "document name, or the literal 'Not Determined'. Blank is not one of them: "
                "it reads as evidence that was never looked for." % i)
        if source not in EVIDENCE_ABSTENTIONS and not _valid_as_of(r.get("evidence_as_of")):
            raise LandscapeError(
                "risk_matrix[%d] cites %r but has evidence_as_of=%r, which is not a usable "
                "capture date. A source with no date cannot be aged or re-verified, so a "
                "2019 risk read and a current one look identical. Use 'Not Determined' as "
                "the source if there is genuinely no evidence."
                % (i, source, r.get("evidence_as_of")))
    return _write_csv(path, cols, risks or [])


def write_excluded_csv(excluded, path):
    """The disqualifier audit trail. An EMPTY list still writes a row (SKILL.md:977)."""
    cols = ["vendor_name", "reason_code", "reason_detail", "source", "date"]
    rows = list(excluded or [])
    for i, r in enumerate(rows):
        if r.get("reason_code") not in EXCLUSION_CODES:
            raise LandscapeError(
                "excluded_vendors[%d] reason_code %r is outside %s"
                % (i, r.get("reason_code"), list(EXCLUSION_CODES)))
    if not rows:
        # Absence stated explicitly. An empty file is indistinguishable from a step that
        # never ran, and this file exists to make the shortlist defensible.
        rows = [{"vendor_name": "", "reason_code": "", "reason_detail": "none excluded",
                 "source": "", "date": ""}]
    return _write_csv(path, cols, rows)


# --------------------------------------------------------------------------- the document

def build_report(data, fit_rows, pillar_rows, path):
    """Assemble the whole report in ONE pass. No open/append/save cycle, ever."""
    if not _DOCX:
        raise DocxUnavailableError(
            "python-docx is unavailable, so the report cannot be built. Refusing to emit "
            "markdown under a .docx name.")

    doc = Document()
    doc.add_heading(data.get("title") or "Supplier Landscape", level=0)
    if data.get("category"):
        doc.add_paragraph(str(data["category"]))

    narratives = data.get("narratives") or {}
    for section in REPORT_SECTIONS:
        doc.add_heading(section, level=1)

        if section == "Supplier Profiles":
            for s in data.get("suppliers") or []:
                doc.add_heading(s["supplier_name"], level=2)
                doc.add_paragraph(narratives.get(s["supplier_name"])
                                  or "[PROSE: profile for %s]" % s["supplier_name"])
            continue

        if section == "Requirements Fit Matrix":
            reqs = data["requirements"]
            t = doc.add_table(rows=1, cols=len(reqs) + 3)
            t.style = "Table Grid"
            for c, h in enumerate(["Supplier"] + list(reqs) + ["Weighted", "Fit"]):
                cell = t.rows[0].cells[c]
                cell.text = str(h)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for r in fit_rows:
                cells = t.add_row().cells
                cells[0].text = r["supplier_name"]
                for i, q in enumerate(reqs, start=1):
                    v = r["scores"].get(q)
                    cells[i].text = NOT_AVAILABLE if v is None else "%.2f" % v
                cells[len(reqs) + 1].text = (
                    NOT_AVAILABLE if r["weighted_score"] is None else "%.2f" % r["weighted_score"])
                cells[len(reqs) + 2].text = r["overall_fit"]
            doc.add_paragraph(
                "Requirement-count-weighted (0-10). This is the scoring system the dashboard "
                "headline uses. It is NOT the 8-pillar matrix below.")
            continue

        if section == "Cross-Vendor Comparison":
            pillars = list(data["pillar_weights"].keys())
            t = doc.add_table(rows=1, cols=len(pillars) + 2)
            t.style = "Table Grid"
            for c, h in enumerate(["Supplier"] + pillars + ["Total"]):
                cell = t.rows[0].cells[c]
                cell.text = str(h)
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for r in pillar_rows:
                cells = t.add_row().cells
                cells[0].text = r["supplier_name"]
                for i, p in enumerate(pillars, start=1):
                    cells[i].text = "%.2f" % r["scores"][p]
                cells[len(pillars) + 1].text = "%.2f" % r["weighted_total"]
            doc.add_paragraph(
                "8-pillar percentage-weighted. A report table only: this figure does not feed "
                "the dashboard headline and must not be written into the fit matrix.")
            continue

        p = doc.add_paragraph(narratives.get(section) or "[PROSE: %s]" % section)
        if not narratives.get(section):
            for run in p.runs:
                run.italic = True
                run.font.size = Pt(10)

    doc.save(path)
    return path


def generate_all(data, outdir):
    """Report + all five CSVs, from ONE object, validated against each other first."""
    requirements = data["requirements"]
    suppliers = data["suppliers"]

    fit_rows = compute_requirements_fit(suppliers, requirements)
    pillar_rows = compute_pillar_matrix(suppliers, data["pillar_weights"])
    registry = data.get("registry") or []
    risks = data.get("risks") or []

    check_registry_schema(registry)
    check_artifact_consistency(fit_rows, pillar_rows, registry, risks)

    os.makedirs(outdir, exist_ok=True)
    written = {
        "requirements_fit_matrix.csv": write_requirements_fit_csv(
            fit_rows, requirements, os.path.join(outdir, "requirements_fit_matrix.csv")),
        "weighted_scoring_matrix.csv": write_pillar_csv(
            pillar_rows, list(data["pillar_weights"].keys()),
            os.path.join(outdir, "weighted_scoring_matrix.csv")),
        "supplier_registry.csv": write_registry_csv(
            registry, os.path.join(outdir, "supplier_registry.csv")),
        "risk_matrix.csv": write_risk_csv(risks, os.path.join(outdir, "risk_matrix.csv")),
        "excluded_vendors.csv": write_excluded_csv(
            data.get("excluded"), os.path.join(outdir, "excluded_vendors.csv")),
    }
    if _DOCX:
        written["supplier_landscape_report.docx"] = build_report(
            data, fit_rows, pillar_rows,
            os.path.join(outdir, "supplier_landscape_report.docx"))
    return {"fit_rows": fit_rows, "pillar_rows": pillar_rows, "written": written}


def main(argv):
    if not argv:
        print(__doc__.strip())
        print("\nusage: landscape_report_generator.py <landscape.json> <outdir>")
        return 0
    with open(argv[0], encoding="utf-8") as fh:
        data = json.load(fh)
    try:
        res = generate_all(data, argv[1] if len(argv) > 1 else ".")
    except LandscapeError as e:
        print("REFUSED: %s: %s" % (type(e).__name__, e), file=sys.stderr)
        return 2
    print(json.dumps(sorted(res["written"]), indent=2))
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv[1:]))
