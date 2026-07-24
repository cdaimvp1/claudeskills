"""
rfp_analysis_report_generator.py
Lilly Procurement Skills, rfp-response-analysis-1c344a deterministic DOCX generator.

Purpose (mirrors should-cost-builder-1c344a/should_cost_generator.py and
evaluation-engine-1c344a/evaluation_report_generator.py's "generator scripts,
not freehand authoring" pattern): this module takes a validated RFP response
analysis register (per-supplier profiles, response analysis narrative, N.3
adequacy scores, strengths/risks, the gated Bid Leveling worksheet, and the
cross-cutting comparison/coverage/scoring/pricing/legal/inconsistency/
clarification/recommendation data SKILL.md's "Complete Section Structure"
and "comparison-patterns.md" describe) and MECHANICALLY produces
analysis_summary.docx: title page, table of contents, executive summary, one
consolidated section per supplier (N.1 Profile through N.5 Risks, including
the fixed nine N.2 Response Summary & Analysis subsections and the N.3
13-row adequacy table), the gated Bid Leveling & Normalization section, and
the eight cross-cutting sections (Cross-Supplier Comparison Matrix,
Requirements Coverage Heatmap, Weighted Scoring Matrix, Commercial & Pricing
Analysis, MSA / Legal Risk Assessment, Inconsistency Register, Clarification
Questions, Final Recommendation (Proposed)), in that fixed order, every run.

Two-stage discipline (same as should-cost/pro-forma/evaluation-engine):
  1. Every supplier's Weighted Scoring Matrix total (Section 8) is computed
     by calling the vendored numeric_kernel.py's weighted_score() over that
     supplier's SCORED (non-PENDING) dimensions, with weights renormalized
     to sum to 1.0 across just that covered subset (see "Weighted-scoring
     renormalization" judgment call below). Never hand-computed. Per
     SKILL.md Section 8's "Kernel-computed arithmetic (HARD RULE)": "The
     weighted-sum arithmetic behind this matrix ... is performed by calling
     weighted_score() in the vendored numeric_kernel.py ... never by model
     arithmetic."
  2. This is the PROPOSED weighted total per this skill's own Rule 5 and
     Boundaries section ("This skill's scores ... are proposed, evidence-
     cited starting points, not the RFx's authoritative outcome"), never
     the official evaluation-engine figure.

Bid Leveling gate (HARD RULE, this skill's own named gate). Per Workflow
Phase 4 / "GATE CHECK: Bid Leveling Complete": "no ranking, weighted score,
or recommendation may be produced from unleveled figures." This generator
enforces that gate as a hard-coded invariant BEFORE building any content:
if any supplier that submitted pricing lacks a normalized price, a one-time/
recurring split, a reported-vs-normalized TCO, or an assumption & exclusion
register entry, or if a logged Bid Leveling clarification gap has no
matching clarification_questions.csv row, this module raises
ReconciliationError and refuses to build the document at all (see "All-or-
nothing gate" judgment call below for why this generator gates the WHOLE
document rather than only Sections 6-13).

Hard-coded invariants (code-level checks, not comments; mirrors should-
cost's run_hardcoded_invariant_checks() / evaluation-engine's pattern):
  - Weighted-scoring-weights-sum-to-one invariant (full framework, static):
    every weighted_scoring.dimensions[].weight_frac must sum to 1.0 +/- 0.001
    across ALL rows (scored and PENDING alike), mirroring SKILL.md's own
    illustrative matrix ("| TOTAL | | 100% | | | |").
  - Per-supplier-covered-weight-sums-to-one invariant: enforced by
    numeric_kernel.weighted_score()'s own refusal on the renormalized subset.
  - Weighted-total-range invariant: every non-null weighted total is 0.0-5.0.
  - Weighted-ranking-sorted invariant: the computed ranking is sorted by
    weighted total, descending, non-null entries first.
  - Adequacy-overall-range invariant: every computed OVERALL adequacy score
    is 0.0-5.0.
  - Requirements-coverage-counts-foot invariant: each supplier's top-level
    Fully/Partially/Does Not/Not Answered tally sums to that supplier's
    declared coverage total, which must equal the register's
    requirement_count.
  - Coverage-heatmap-category-counts-foot invariant: each category's
    per-supplier Fully/Partially/Does Not/Not Answered tally sums to that
    category's requirement_count.
  - Bid-Leveling-gate invariant (the named gate above): comparison basis
    stated; every pricing-submitting supplier has a normalized pricing row,
    a one-time/recurring split, a reported-and-normalized TCO, at least one
    assumption & exclusion register entry, and a leveling_status of
    "Complete"; every logged leveling gap has a matching Bid_Leveling_Gap
    clarification question.
  - Pricing-dimension-not-fabricated invariant: no Weighted Scoring Matrix
    "Pricing" category subcategory may carry a non-null score for a supplier
    whose pricing_submitted is False (Rule 1, "never fabricate scores").
  - Adequacy-score-set-complete invariant: every supplier's N.3 table covers
    exactly the twelve canonical RFP sections, no more, no fewer, no
    duplicates (OVERALL is computed by this generator, never accepted as
    input, exactly as evaluation-engine computes Grand Total rather than
    accepting it).

If any invariant fails, this script RAISES rather than writing a document
that fails its own reconciliation.

Scope note: this script builds the analysis_summary.docx generator ONLY, per
SKILL.md's "Primary User-Facing Deliverables" table. The interactive
dashboard (response_analysis_dashboard.jsx), the pipeline artifacts
({supplier}_profile.json, requirements_coverage_matrix.csv,
bid_leveling_worksheet.csv, bid_leveling_register.csv,
evaluation_engine_handoff.json, etc.), and the debrief email drafts are
separate deliverables and are NOT produced here.

============================================================================
RFP analysis register input contract (JUDGMENT CALLS, flagged): SKILL.md's
"Complete Section Structure" and comparison-patterns.md give the DOCUMENT
shape (sections, tables, narrative subsections) but, like evaluation-
engine-1c344a before this skill added a generator, no single raw INPUT
contract for a report generator. This module resolves that ambiguity as
follows, each flagged inline where it matters most in the code below too.

1. Narrative content is supplied, not authored here. SKILL.md's "Content
   Writing Rules" mandate magazine-quality, multi-paragraph connected prose
   (Rule: "The document must read like a magazine feature, not a database
   export"). A deterministic renderer cannot originate that prose without
   fabricating analytical content the skill's own Rule 1-4 (Accuracy and
   Anti-Drift Rules) forbid. This generator therefore accepts every
   narrative field (profile introductions, the nine N.2 subsections, N.4/
   N.5 evidence bullets, Bid Leveling analysis paragraphs, per-dimension and
   per-supplier cross-cutting analysis, the Final Recommendation prose) as
   ALREADY-AUTHORED text supplied by the calling skill run, and is
   responsible only for assembling it into the fixed skeleton, validating
   its shape, computing the one number set the skill mandates a kernel
   compute (Section 8), and enforcing the named invariants above. This
   mirrors exactly how evaluation_report_generator.py consumes
   strengths/weaknesses/risk_notes as pre-authored text rather than writing
   them itself.

2. Section numbering is computed, not fixed. comparison-patterns.md's
   worked illustration numbers the cross-cutting sections 6-13 assuming
   (implicitly) a fixed small supplier count; the Multi-Pass Generation
   Guidance's own "sections 4-5" phrasing for supplier sections conflicts
   with "5. Bid Leveling & Normalization (Gate)" once more than two
   suppliers are evaluated. This generator resolves it the only way that
   stays internally consistent for ANY supplier count: Executive Summary is
   always Section 3, per-supplier sections occupy Sections 4..(3+N) for N
   suppliers, Bid Leveling is Section (4+N), and the eight cross-cutting
   sections occupy (5+N)..(12+N), preserving comparison-patterns.md's
   section TITLES and their relative order exactly.

3. TOC field / Heading styles. SKILL.md's Branding and Document Design
   section quotes docx-title-page-spec.md directly: "The TOC indexes
   Heading2 paragraphs only; the two recurring per-vendor sub-headings that
   must use HeadingLevel.HEADING_2 are 'Response Summary & Analysis' and
   'RFP Section Adequacy Scores'." This generator therefore applies python-
   docx's built-in "Heading 2" PARAGRAPH STYLE (not just custom run
   formatting, unlike every other heading level in this document, which
   follows evaluation_report_generator.py's "no dependency on a template's
   built-in Heading styles" convention) to every top-level numbered section
   heading and to those two named per-vendor sub-headings specifically, and
   inserts a real Word TOC field (\\o "2-2" \\h \\z \\u) referencing that
   style level. python-docx cannot force Word to compute page numbers at
   write time; the field carries a placeholder run Word replaces once the
   user updates the field (F9 or "Update Table"), which is the standard,
   disclosed limitation of building a TOC field outside Word itself.

4. Weighted-scoring renormalization for PENDING dimensions. Section 8's
   "Weighted-total discipline (HARD RULE)" states two allowed options:
   "Exclude pending dimensions from the weighted-total denominator and
   compute the total over the submitted-and-scored weight only, OR render
   the total as labeled INCOMPLETE with the missing weight stated." This
   generator implements the first (renormalize the scored subset's weights
   to sum to 1.0 and call weighted_score() on it) AND always surfaces the
   second (the missing weight fraction is shown alongside every total), so
   neither allowed reading is silently dropped.

5. OVERALL adequacy score. profile-schema.md defines `adequacy_overall` as
   "weighted average of section scores" but names no weights and no worked
   example (unlike should-cost-builder's quoted 60+30+10 illustration).
   This generator computes it as the unweighted arithmetic mean of the
   twelve canonical N.3 section scores and, exactly as evaluation-engine
   computes Grand Total rather than accepting it as caller input, NEVER
   accepts an OVERALL score as input; it is always this generator's own
   computed figure, so the report can never show a stale or drifted OVERALL
   next to its own section table.

6. All-or-nothing Bid Leveling gate. SKILL.md's own GATE CHECK language
   only blocks "Phase 5 (Cross-Vendor Comparison, Sections 6-13) or ...
   any ranking, weighted score, or recommendation," leaving Sections 1-5
   producible even when the gate fails for some supplier. A one-shot
   document renderer that has already been handed a complete register for
   Sections 6-13 has no use for a half-built document in that state; this
   generator therefore raises and refuses the ENTIRE document (mirroring
   should-cost/evaluation-engine's "raise, don't write a document that
   fails its own reconciliation" discipline) rather than emitting a partial
   file. A caller that wants a Sections-1-5-only interim document should
   build its register only through Section 5 and re-run once leveling is
   complete for every pricing-submitting supplier.

7. Comparison / commercial / legal table shape. The 11-dimension Cross-
   Supplier Comparison Matrix (Section 6) is a named, fixed list in both
   SKILL.md and comparison-patterns.md, so it is validated as a closed set
   (the "Adequacy Score" row is always this generator's own computed
   figure, per judgment call 5, never accepted as caller input). The
   Commercial & Pricing (Section 9, "12+ rows") and MSA/Legal (Section 10,
   "15+ rows") tables are explicitly category-varying ("This skill works
   across ALL procurement categories"), so they are validated only by
   minimum row count, with dimension names supplied by the caller.

Register shape: see validate_rfp_analysis_input() and its dataclasses below
for the complete field-by-field contract; the demo register at the bottom
of this file (_demo_rfp_analysis_register()) is a fully worked, schema-
valid example.
============================================================================
"""

from __future__ import annotations

import copy
import os
import sys
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Sequence, Tuple

# ---------------------------------------------------------------------------
# Vendored numeric kernel (same directory). Per SKILL.md Section 8's "Kernel-
# computed arithmetic (HARD RULE)": the Weighted Scoring Matrix total MUST be
# computed by calling weighted_score() in the vendored numeric_kernel.py,
# never by model arithmetic.
# ---------------------------------------------------------------------------
_KERNEL_IMPORT_ERROR: Optional[Exception] = None
try:
    _THIS_DIR = os.path.dirname(os.path.abspath(__file__))
    if _THIS_DIR not in sys.path:
        sys.path.insert(0, _THIS_DIR)
    from numeric_kernel import weighted_score as kernel_weighted_score
    from numeric_kernel import WeightSumError as KernelWeightSumError
    from numeric_kernel import InvalidInputError as KernelInvalidInputError
    KERNEL_AVAILABLE = True
except Exception as _exc:  # pragma: no cover - defensive, disclosed at runtime
    KERNEL_AVAILABLE = False
    _KERNEL_IMPORT_ERROR = _exc

    def kernel_weighted_score(*args, **kwargs):  # type: ignore
        raise RuntimeError(
            "numeric_kernel.py unavailable; cannot compute the kernel-required "
            f"weighted score. Import error: {_KERNEL_IMPORT_ERROR}"
        )

    class KernelWeightSumError(Exception):  # type: ignore
        pass

    class KernelInvalidInputError(Exception):  # type: ignore
        pass

# ---------------------------------------------------------------------------
# python-docx detection. Mirrors evaluation_report_generator.py's handling:
# try the import; if unavailable, raise a clear ImportError at document-build
# time (not at module-import time), so validation and ground-truth logic
# (which need no DOCX library) remain usable regardless.
# ---------------------------------------------------------------------------
try:
    import docx  # noqa: F401
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
    DOCX_VERSION = getattr(docx, "__version__", "unknown")
except Exception as _exc:  # pragma: no cover
    DOCX_AVAILABLE = False
    DOCX_VERSION = None
    _DOCX_IMPORT_ERROR = _exc


class RfpAnalysisValidationError(Exception):
    """Raised when the RFP analysis register is missing a required field or
    carries a value the skill's Accuracy and Anti-Drift Rules forbid. Per
    Rule 1 ("Never attribute capabilities a supplier did not claim") and
    Rule 2 ("Never fabricate competitive comparisons") this module refuses
    rather than guessing a missing or malformed value."""


class ReconciliationError(Exception):
    """Raised when a hard-coded invariant fails, including the named Bid
    Leveling gate ("GATE CHECK: Bid Leveling Complete"). The document is not
    saved when this fires."""


# ===========================================================================
# House style constants (quoted verbatim from lilly-brand-assets-1c344a's
# brand-colors.md and docx-design-system.md, the same shared foundation
# should_cost_generator.py and evaluation_report_generator.py hardcode their
# own copies of, since SKILL.md's Branding and Document Design section
# points at the identical two files: "pull exact values from them so it
# matches every other suite report"; keeping the same hex values here means
# this DOCX matches every other suite report byte-for-byte on color).
# ===========================================================================
LILLY_RED = "E1251B"      # table header rows, title-page accents, Critical severity
BOLD_BLUE = "0F3A85"      # section header text (H1/Heading 2), Low priority ("safe" token, no-green rule)
LILLY_BLACK = "212121"    # all body text
AMBER = "B45309"          # Moderate severity / MEDIUM priority
MUTED_GREY = "8A969E"     # footer/header chrome text
WHITE = "FFFFFF"          # table header text on colored fill
POSITIVE_BLUE = "1668B3"  # positive-status azure (matches this skill's own dashboard POS token); distinct from BOLD_BLUE, never green

SEVERITY_COLOR = {"Critical": LILLY_RED, "Moderate": AMBER, "Minor": MUTED_GREY}
CLARIFICATION_PRIORITY_COLOR = {"GATING": LILLY_RED, "HIGH": AMBER, "MEDIUM": BOLD_BLUE}

VALID_COVERAGE = ("Fully Meets", "Partially Meets", "Does Not Meet", "Not Answered")
VALID_CONFIDENCE = ("High", "Medium", "Low")
VALID_CONTRACT_POSTURE = ("Collaborative", "Standard", "Aggressive")
VALID_LEVELING_STATUS = ("Complete", "Pending Pricing", "Pending Normalization", "Pending Clarification")
VALID_COST_IMPACT = ("Included", "Additional Cost", "Excluded", "Unknown")
VALID_SCOPE_STATUS = ("Included", "Additional Cost", "Excluded", "Silent")
VALID_ASSUMPTION_ROW_TYPE = ("Assumption", "Exclusion")
VALID_INCONSISTENCY_SEVERITY = ("Critical", "Moderate", "Minor")
VALID_CLARIFICATION_PRIORITY = ("GATING", "HIGH", "MEDIUM")
VALID_SOURCE_TYPE = (
    "Missing_Response", "Low_Confidence", "Internal_Inconsistency",
    "Pricing_Anomaly", "Unverified_Claim", "Bid_Leveling_Gap",
)
VALID_ANALYSIS_MODE = ("A", "B", "C")
VALID_REPORT_MODE = ("Brief", "Full")
VALID_LEVELING_PRIORITY = ("GATING", "HIGH", "MEDIUM")

# The twelve canonical RFP sections N.3 scores every supplier against, per
# SKILL.md Phase 2 N.3: "Cover Letter, Vendor Profile, Financial Statements,
# References, Ability to Meet Objectives, Functional Requirements Matrix,
# Architecture, Solution & Pricing, Demo, MSA/Legal, Implementation Plan,
# Training, OVERALL." OVERALL is computed by this generator (judgment call
# 5 above), so it is excluded from the required INPUT set below.
CANONICAL_RFP_SECTIONS = (
    "Cover Letter", "Vendor Profile", "Financial Statements", "References",
    "Ability to Meet Objectives", "Functional Requirements Matrix", "Architecture",
    "Solution & Pricing", "Demo", "MSA/Legal", "Implementation Plan", "Training",
)

# The fixed nine N.2 Response Summary & Analysis subsections, in SKILL.md's
# own listed order (Phase 2 / comparison-patterns.md).
RESPONSE_ANALYSIS_SUBSECTIONS = (
    ("submission_volume", "Submission Volume"),
    ("understanding_of_requirements", "Understanding of Requirements"),
    ("proposed_solution_architecture", "Proposed Solution & Architecture"),
    ("implementation_approach", "Implementation Approach"),
    ("integration_strategy", "Integration Strategy"),
    ("references_domain_evidence", "References & Domain Evidence"),
    ("legal_contract_posture", "Legal / Contract Posture"),
    ("key_concerns", "Key Concerns"),
    ("overall_assessment", "Overall Assessment"),
)

# The eleven Cross-Supplier Comparison Matrix dimensions (Section 6), a
# fixed named list per SKILL.md Section 6 and comparison-patterns.md.
# "Adequacy Score" is always this generator's own computed figure (judgment
# call 5), so it is excluded from the required INPUT set below.
COMPARISON_DIMENSIONS_INPUT = (
    "Requirements Fit", "Financial Health", "Risk Level", "Pricing Clarity",
    "Pricing Competitiveness", "Contract Complexity", "Vendor Status",
    "Pharma Experience", "Implementation Readiness", "Technology Differentiation",
)
COMPARISON_DIMENSIONS_ALL = COMPARISON_DIMENSIONS_INPUT + ("Adequacy Score",)

# Section 9 and Section 10 minimum row counts, per comparison-patterns.md
# ("full pricing model comparison table (12+ rows...)" / "full risk heatmap
# table (15+ rows...)").
COMMERCIAL_MIN_ROWS = 12
LEGAL_MIN_ROWS = 15

# The eight cross-cutting section titles, in fixed order, per comparison-
# patterns.md Sections 6-13 (see judgment call 2 above for numbering).
CROSSCUT_SECTION_TITLES = (
    "Cross-Supplier Comparison Matrix",
    "Requirements Coverage Heatmap",
    "Weighted Scoring Matrix",
    "Commercial & Pricing Analysis",
    "MSA / Legal Risk Assessment",
    "Inconsistency Register",
    "Clarification Questions",
    "Final Recommendation (Proposed)",
)


# ===========================================================================
# 1. RFP analysis register: typed input dataclasses
# ===========================================================================

@dataclass
class SupplierProfile:
    supplier_id: str
    supplier_name: str
    narrative: str                      # N.1's 2-3 paragraph narrative introduction
    data_card: Dict[str, str]           # 3-6 critical fields (Headquarters, Employees, Revenue, ...)
    pharma_clients: List[str]           # up to 10; empty means "Not Stated"
    deployment_model: str
    contract_posture: str               # Collaborative | Standard | Aggressive
    lilly_vendor_status: str
    pricing_summary: str
    pricing_submitted: bool
    leveling_status: str                # Complete | Pending Pricing | Pending Normalization | Pending Clarification


@dataclass
class ResponseAnalysis:
    submission_volume: str
    submission_volume_pages: int
    understanding_of_requirements: str
    proposed_solution_architecture: str
    implementation_approach: str
    integration_strategy: str
    references_domain_evidence: str
    legal_contract_posture: str
    key_concerns: str
    overall_assessment: str


@dataclass
class AdequacyScore:
    rfp_section: str
    score: float
    notes: str


@dataclass
class SupplierSection:
    profile: SupplierProfile
    response_analysis: ResponseAnalysis
    coverage: Dict[str, int]            # fully_meets/partially_meets/does_not_meet/not_answered/total
    adequacy_scores: List[AdequacyScore]  # exactly the 12 canonical sections
    strengths: List[str]                # N.4
    risks: List[str]                    # N.5


@dataclass
class RecommendationPrimary:
    supplier_id: str
    evidence: List[str]
    conditions: List[str]
    caveats: List[str]


@dataclass
class RecommendationSecondary:
    supplier_id: str
    triggering_scenario: str
    sensitivity_analysis: str


@dataclass
class RecommendationConditional:
    supplier_id: str
    prerequisites: List[str]


@dataclass
class RecommendationNotRecommended:
    supplier_id: str
    rationale: str


@dataclass
class FinalRecommendation:
    primary: RecommendationPrimary
    secondary: Optional[RecommendationSecondary]
    conditional: List[RecommendationConditional]
    not_recommended: List[RecommendationNotRecommended]
    standard_caveats: List[str]


@dataclass
class ScopeComplianceRow:
    scope_line: str
    status: Dict[str, str]              # supplier_id -> Included | Additional Cost | Excluded | Silent


@dataclass
class AssumptionExclusionRow:
    supplier_id: str
    row_type: str                       # Assumption | Exclusion
    item: str
    cost_impact: str                    # Included | Additional Cost | Excluded | Unknown
    confidence: str
    source: str


@dataclass
class NormalizedPricingRow:
    supplier_id: str
    scenario: str
    comparison_basis_unit: str
    reported_price: float
    one_time_total: float
    recurring_total: float
    reported_tco: float
    normalized_price_per_unit: float
    normalized_tco: float
    term_years: int
    confidence: str
    source: str


@dataclass
class MissingCostPlaceholder:
    scope_line: str
    note: str                           # a should-cost estimate, or "Not priced by any supplier"


@dataclass
class LevelingQuestion:
    supplier_id: str
    question: str
    priority: str                       # GATING | HIGH | MEDIUM


@dataclass
class BidLeveling:
    comparison_basis_text: str
    scope_compliance: List[ScopeComplianceRow]
    assumption_exclusion_register: List[AssumptionExclusionRow]
    normalized_pricing: List[NormalizedPricingRow]
    missing_cost_placeholders: List[MissingCostPlaceholder]
    questions_before_final_evaluation: List[LevelingQuestion]
    analysis_paragraphs: List[str]


@dataclass
class ComparisonMatrix:
    values: Dict[str, Dict[str, str]]   # dimension (excl. Adequacy Score) -> {supplier_id: str}
    analysis: Dict[str, str]            # dimension -> paragraph (all 11, incl. Adequacy Score)


@dataclass
class CoverageCategory:
    category: str
    requirement_count: int
    per_supplier: Dict[str, Dict[str, int]]  # supplier_id -> {fully_meets, partially_meets, does_not_meet, not_answered}


@dataclass
class CoverageHeatmap:
    categories: List[CoverageCategory]
    intro_paragraph: str
    analysis_paragraphs: List[str]


@dataclass
class ScoringDimension:
    category: str
    subcategory: str
    weight_frac: float
    scores: Dict[str, Optional[float]]  # supplier_id -> 0.0-5.0, or None for PENDING
    rationale: Dict[str, str]           # supplier_id -> 1-3 sentence rationale (incl. PENDING reason)


@dataclass
class WeightedScoring:
    dimensions: List[ScoringDimension]
    sensitivity_notes: List[str]
    improvement_suggestions: List[str]  # <= 5


@dataclass
class CommercialRow:
    dimension: str
    values: Dict[str, str]


@dataclass
class CommercialPricing:
    rows: List[CommercialRow]
    per_supplier_analysis: Dict[str, List[str]]  # supplier_id -> list of paragraphs
    normalization_recommendation: str


@dataclass
class LegalRow:
    clause: str
    values: Dict[str, str]


@dataclass
class LegalRisk:
    rows: List[LegalRow]
    per_supplier_analysis: Dict[str, str]


@dataclass
class InconsistencyItem:
    inconsistency_id: str
    supplier_id: str
    severity: str                       # Critical | Moderate | Minor
    description: str
    action_required: str


@dataclass
class ClarificationQuestion:
    question_id: str
    supplier_id: str
    source_type: str
    req_id: Optional[str]
    description: str
    priority: str                       # GATING | HIGH | MEDIUM
    recommended_format: str


@dataclass
class RfpAnalysisInput:
    rfx_name: str
    category: str
    case_id: str
    analysis_date: str
    analysis_mode: str                  # A | B | C
    report_mode: str                    # Brief | Full
    requirement_count: int
    suppliers: List[SupplierSection]
    executive_summary_scope: str
    executive_summary_overviews: Dict[str, str]
    executive_summary_key_findings: List[str]
    bid_leveling: BidLeveling
    comparison_matrix: ComparisonMatrix
    coverage_heatmap: CoverageHeatmap
    weighted_scoring: WeightedScoring
    commercial_pricing: CommercialPricing
    legal_risk: LegalRisk
    inconsistency_register: List[InconsistencyItem]
    clarification_questions: List[ClarificationQuestion]
    final_recommendation: FinalRecommendation


REQUIRED_TOP_LEVEL_FIELDS = [
    "rfx_name", "category", "case_id", "analysis_date", "analysis_mode", "report_mode",
    "requirement_count", "suppliers", "executive_summary_scope", "executive_summary_overviews",
    "executive_summary_key_findings", "bid_leveling", "comparison_matrix", "coverage_heatmap",
    "weighted_scoring", "commercial_pricing", "legal_risk", "inconsistency_register",
    "clarification_questions", "final_recommendation",
]


def _enum(value: Any, valid: Sequence[str], label: str, errors: List[str]) -> str:
    v = str(value).strip()
    if v not in valid:
        errors.append(f"{label} must be one of {valid}; got {v!r}.")
    return v


def _nonempty_str(value: Any, label: str, errors: List[str]) -> str:
    v = str(value).strip() if value is not None else ""
    if not v:
        errors.append(f"{label} must be a non-empty string.")
    return v


def _keys_match(d: Dict[str, Any], expected: set, label: str, errors: List[str]) -> None:
    got = set(d.keys())
    if got != expected:
        missing = expected - got
        extra = got - expected
        errors.append(
            f"{label} keys must exactly match {sorted(expected)}. "
            f"Missing: {sorted(missing) or 'none'}; unexpected: {sorted(extra) or 'none'}."
        )


# ===========================================================================
# 2. Validation
# ===========================================================================

def validate_rfp_analysis_input(register: Dict[str, Any]) -> RfpAnalysisInput:
    """Validate a raw RFP analysis register dict and return a typed
    RfpAnalysisInput. Refuses (raises RfpAnalysisValidationError) rather
    than guessing when a required field is missing or a value violates a
    documented, fixed convention, per Rule 1 ("Never attribute capabilities
    a supplier did not claim") and Rule 2 ("Never fabricate competitive
    comparisons")."""
    errors: List[str] = []

    for f in REQUIRED_TOP_LEVEL_FIELDS:
        if f not in register:
            errors.append(f"Missing required top-level field: '{f}'")
    if errors:
        raise RfpAnalysisValidationError(
            "RFP analysis register is missing required fields; refusing to "
            "guess. " + "; ".join(errors)
        )

    rfx_name = _nonempty_str(register["rfx_name"], "rfx_name", errors)
    category = _nonempty_str(register["category"], "category", errors)
    case_id = _nonempty_str(register["case_id"], "case_id", errors)
    analysis_date = register["analysis_date"]
    if not isinstance(analysis_date, str) or not analysis_date:
        errors.append("'analysis_date' must be a non-empty string (e.g. '2026-07-23').")
    analysis_mode = _enum(register["analysis_mode"], VALID_ANALYSIS_MODE, "analysis_mode", errors)
    report_mode = _enum(register["report_mode"], VALID_REPORT_MODE, "report_mode", errors)

    requirement_count = register["requirement_count"]
    if not isinstance(requirement_count, int) or requirement_count <= 0:
        errors.append("'requirement_count' must be a positive integer.")
        requirement_count = 0

    # --- Suppliers -----------------------------------------------------
    raw_suppliers = register.get("suppliers", [])
    if not isinstance(raw_suppliers, list) or len(raw_suppliers) == 0:
        errors.append("'suppliers' must be a non-empty list.")
        raw_suppliers = []

    supplier_ids_seen: List[str] = []
    for i, s in enumerate(raw_suppliers):
        if not isinstance(s, dict) or "profile" not in s:
            errors.append(f"'suppliers[{i}]' must be an object with a 'profile' key.")
            continue
        prof = s.get("profile", {})
        sid = str(prof.get("supplier_id", "")).strip()
        if not sid:
            errors.append(f"'suppliers[{i}].profile.supplier_id' is required.")
            continue
        if sid in supplier_ids_seen:
            errors.append(f"Duplicate supplier_id '{sid}' in 'suppliers'.")
        supplier_ids_seen.append(sid)
    supplier_id_set = set(supplier_ids_seen)

    suppliers: List[SupplierSection] = []
    for i, s in enumerate(raw_suppliers):
        if not isinstance(s, dict):
            continue
        prof_raw = s.get("profile", {})
        sid = str(prof_raw.get("supplier_id", f"UNKNOWN_{i}")).strip()

        pricing_submitted = prof_raw.get("pricing_submitted")
        if not isinstance(pricing_submitted, bool):
            errors.append(f"suppliers[{sid}].profile.pricing_submitted must be a boolean.")
            pricing_submitted = False
        leveling_status = _enum(
            prof_raw.get("leveling_status", ""), VALID_LEVELING_STATUS,
            f"suppliers[{sid}].profile.leveling_status", errors,
        )
        if not pricing_submitted and leveling_status != "Pending Pricing":
            errors.append(
                f"suppliers[{sid}].profile: pricing_submitted is False but leveling_status is "
                f"{leveling_status!r}, not 'Pending Pricing'. A supplier that has not submitted "
                "pricing must be labeled Pending Pricing (SKILL.md's non-fabrication rule)."
            )
        if pricing_submitted and leveling_status == "Pending Pricing":
            errors.append(
                f"suppliers[{sid}].profile: pricing_submitted is True but leveling_status is still "
                "'Pending Pricing', which is inconsistent."
            )

        data_card = prof_raw.get("data_card", {})
        if not isinstance(data_card, dict) or not (3 <= len(data_card) <= 6):
            errors.append(
                f"suppliers[{sid}].profile.data_card must be an object with 3-6 fields "
                "(SKILL.md: 'a compact data card for the 4-5 most critical numeric fields only')."
            )
            data_card = {}

        pharma_clients = prof_raw.get("pharma_clients", [])
        if not isinstance(pharma_clients, list) or len(pharma_clients) > 10:
            errors.append(f"suppliers[{sid}].profile.pharma_clients must be a list of at most 10 items.")
            pharma_clients = []

        profile = SupplierProfile(
            supplier_id=sid,
            supplier_name=_nonempty_str(prof_raw.get("supplier_name"), f"suppliers[{sid}].profile.supplier_name", errors),
            narrative=_nonempty_str(prof_raw.get("narrative"), f"suppliers[{sid}].profile.narrative", errors),
            data_card={str(k): str(v) for k, v in data_card.items()},
            pharma_clients=[str(x) for x in pharma_clients],
            deployment_model=_nonempty_str(prof_raw.get("deployment_model"), f"suppliers[{sid}].profile.deployment_model", errors),
            contract_posture=_enum(prof_raw.get("contract_posture", ""), VALID_CONTRACT_POSTURE,
                                    f"suppliers[{sid}].profile.contract_posture", errors),
            lilly_vendor_status=_nonempty_str(prof_raw.get("lilly_vendor_status"), f"suppliers[{sid}].profile.lilly_vendor_status", errors),
            pricing_summary=_nonempty_str(prof_raw.get("pricing_summary"), f"suppliers[{sid}].profile.pricing_summary", errors),
            pricing_submitted=pricing_submitted,
            leveling_status=leveling_status,
        )

        ra_raw = s.get("response_analysis", {})
        if not isinstance(ra_raw, dict):
            ra_raw = {}
        ra_kwargs: Dict[str, Any] = {}
        for key, title in RESPONSE_ANALYSIS_SUBSECTIONS:
            ra_kwargs[key] = _nonempty_str(ra_raw.get(key), f"suppliers[{sid}].response_analysis.{key} ({title})", errors)
        pages = ra_raw.get("submission_volume_pages", 0)
        if not isinstance(pages, int) or pages <= 0:
            errors.append(f"suppliers[{sid}].response_analysis.submission_volume_pages must be a positive integer.")
            pages = 0
        response_analysis = ResponseAnalysis(submission_volume_pages=pages, **ra_kwargs)

        cov_raw = s.get("coverage", {})
        required_cov = {"fully_meets", "partially_meets", "does_not_meet", "not_answered", "total"}
        if not isinstance(cov_raw, dict) or not required_cov.issubset(cov_raw.keys()):
            errors.append(f"suppliers[{sid}].coverage must have keys {sorted(required_cov)}.")
            coverage = {k: 0 for k in required_cov}
        else:
            coverage = {k: int(cov_raw[k]) for k in required_cov}
            footed = coverage["fully_meets"] + coverage["partially_meets"] + coverage["does_not_meet"] + coverage["not_answered"]
            if footed != coverage["total"]:
                errors.append(
                    f"suppliers[{sid}].coverage tally FAILED: Fully Meets + Partially Meets + Does Not Meet + "
                    f"Not Answered = {footed}, not the declared total ({coverage['total']})."
                )
            if requirement_count and coverage["total"] != requirement_count:
                errors.append(
                    f"suppliers[{sid}].coverage.total ({coverage['total']}) does not match the register's "
                    f"requirement_count ({requirement_count})."
                )

        adeq_raw = s.get("adequacy_scores", [])
        adequacy_scores: List[AdequacyScore] = []
        seen_sections: List[str] = []
        if not isinstance(adeq_raw, list):
            adeq_raw = []
        for a in adeq_raw:
            if not isinstance(a, dict) or "rfp_section" not in a or "score" not in a:
                errors.append(f"suppliers[{sid}].adequacy_scores has a malformed row (needs rfp_section, score).")
                continue
            rsec = str(a["rfp_section"]).strip()
            seen_sections.append(rsec)
            sc = a["score"]
            if not isinstance(sc, (int, float)) or not (0.0 <= float(sc) <= 5.0):
                errors.append(f"suppliers[{sid}].adequacy_scores[{rsec!r}].score must be numeric in [0.0, 5.0]; got {sc!r}.")
                sc = 0.0
            adequacy_scores.append(AdequacyScore(rfp_section=rsec, score=float(sc), notes=str(a.get("notes", ""))))
        if set(seen_sections) != set(CANONICAL_RFP_SECTIONS) or len(seen_sections) != len(CANONICAL_RFP_SECTIONS):
            errors.append(
                f"suppliers[{sid}].adequacy_scores must cover exactly the twelve canonical RFP sections "
                f"{CANONICAL_RFP_SECTIONS}, no more, no fewer, no duplicates (OVERALL is computed, not "
                f"accepted as input). Got: {seen_sections}."
            )

        strengths = s.get("strengths", [])
        risks = s.get("risks", [])
        if not isinstance(strengths, list) or len(strengths) == 0:
            errors.append(f"suppliers[{sid}].strengths (N.4) must be a non-empty list.")
            strengths = []
        if not isinstance(risks, list) or len(risks) == 0:
            errors.append(f"suppliers[{sid}].risks (N.5) must be a non-empty list.")
            risks = []

        suppliers.append(SupplierSection(
            profile=profile, response_analysis=response_analysis, coverage=coverage,
            adequacy_scores=adequacy_scores, strengths=[str(x) for x in strengths],
            risks=[str(x) for x in risks],
        ))

    # --- Executive Summary -------------------------------------------------
    executive_summary_scope = _nonempty_str(register["executive_summary_scope"], "executive_summary_scope", errors)
    exec_overviews = register.get("executive_summary_overviews", {})
    if not isinstance(exec_overviews, dict):
        errors.append("'executive_summary_overviews' must be an object keyed by supplier_id.")
        exec_overviews = {}
    else:
        _keys_match(exec_overviews, supplier_id_set, "'executive_summary_overviews'", errors)
    key_findings = register.get("executive_summary_key_findings", [])
    if not isinstance(key_findings, list) or not (5 <= len(key_findings) <= 6):
        errors.append(
            "'executive_summary_key_findings' must be a list of 5-6 items (SKILL.md: 'Key findings "
            "(5-6 numbered items)')."
        )
        key_findings = list(key_findings) if isinstance(key_findings, list) else []

    # --- Bid Leveling --------------------------------------------------
    bl_raw = register.get("bid_leveling", {})
    if not isinstance(bl_raw, dict):
        bl_raw = {}
    comparison_basis_text = _nonempty_str(bl_raw.get("comparison_basis_text"), "bid_leveling.comparison_basis_text", errors)

    scope_compliance: List[ScopeComplianceRow] = []
    for i, row in enumerate(bl_raw.get("scope_compliance", [])):
        if not isinstance(row, dict) or "scope_line" not in row or "status" not in row:
            errors.append(f"bid_leveling.scope_compliance[{i}] must have scope_line and status.")
            continue
        status = row["status"]
        if not isinstance(status, dict):
            errors.append(f"bid_leveling.scope_compliance[{i}].status must be an object keyed by supplier_id.")
            continue
        _keys_match(status, supplier_id_set, f"bid_leveling.scope_compliance[{row['scope_line']!r}].status", errors)
        clean_status = {}
        for k, v in status.items():
            clean_status[k] = _enum(v, VALID_SCOPE_STATUS, f"bid_leveling.scope_compliance[{row['scope_line']!r}].status[{k!r}]", errors)
        scope_compliance.append(ScopeComplianceRow(scope_line=str(row["scope_line"]), status=clean_status))
    if not scope_compliance:
        errors.append("bid_leveling.scope_compliance must be a non-empty list covering every major RFP scope line.")

    assumption_exclusion_register: List[AssumptionExclusionRow] = []
    for i, row in enumerate(bl_raw.get("assumption_exclusion_register", [])):
        if not isinstance(row, dict):
            errors.append(f"bid_leveling.assumption_exclusion_register[{i}] must be an object.")
            continue
        rsid = str(row.get("supplier_id", "")).strip()
        if rsid not in supplier_id_set:
            errors.append(f"bid_leveling.assumption_exclusion_register[{i}] references unknown supplier_id {rsid!r}.")
            continue
        assumption_exclusion_register.append(AssumptionExclusionRow(
            supplier_id=rsid,
            row_type=_enum(row.get("row_type", ""), VALID_ASSUMPTION_ROW_TYPE, f"bid_leveling.assumption_exclusion_register[{i}].row_type", errors),
            item=_nonempty_str(row.get("item"), f"bid_leveling.assumption_exclusion_register[{i}].item", errors),
            cost_impact=_enum(row.get("cost_impact", ""), VALID_COST_IMPACT, f"bid_leveling.assumption_exclusion_register[{i}].cost_impact", errors),
            confidence=_enum(row.get("confidence", ""), VALID_CONFIDENCE, f"bid_leveling.assumption_exclusion_register[{i}].confidence", errors),
            source=str(row.get("source", "")),
        ))

    normalized_pricing: List[NormalizedPricingRow] = []
    for i, row in enumerate(bl_raw.get("normalized_pricing", [])):
        if not isinstance(row, dict):
            errors.append(f"bid_leveling.normalized_pricing[{i}] must be an object.")
            continue
        rsid = str(row.get("supplier_id", "")).strip()
        if rsid not in supplier_id_set:
            errors.append(f"bid_leveling.normalized_pricing[{i}] references unknown supplier_id {rsid!r}.")
            continue
        numeric_fields = ["reported_price", "one_time_total", "recurring_total", "reported_tco",
                           "normalized_price_per_unit", "normalized_tco"]
        vals = {}
        for nf in numeric_fields:
            v = row.get(nf)
            if not isinstance(v, (int, float)) or v < 0:
                errors.append(f"bid_leveling.normalized_pricing[{i}] ({rsid}).{nf} must be a non-negative number.")
                v = 0.0
            vals[nf] = float(v)
        term_years = row.get("term_years", 0)
        if not isinstance(term_years, int) or term_years <= 0:
            errors.append(f"bid_leveling.normalized_pricing[{i}] ({rsid}).term_years must be a positive integer.")
            term_years = 1
        normalized_pricing.append(NormalizedPricingRow(
            supplier_id=rsid, scenario=str(row.get("scenario", "Base")),
            comparison_basis_unit=_nonempty_str(row.get("comparison_basis_unit"), f"bid_leveling.normalized_pricing[{i}] ({rsid}).comparison_basis_unit", errors),
            term_years=term_years,
            confidence=_enum(row.get("confidence", ""), VALID_CONFIDENCE, f"bid_leveling.normalized_pricing[{i}] ({rsid}).confidence", errors),
            source=str(row.get("source", "")), **vals,
        ))

    missing_cost_placeholders = []
    for i, row in enumerate(bl_raw.get("missing_cost_placeholders", [])):
        if not isinstance(row, dict):
            errors.append(f"bid_leveling.missing_cost_placeholders[{i}] must be an object.")
            continue
        missing_cost_placeholders.append(MissingCostPlaceholder(
            scope_line=_nonempty_str(row.get("scope_line"), f"bid_leveling.missing_cost_placeholders[{i}].scope_line", errors),
            note=_nonempty_str(row.get("note"), f"bid_leveling.missing_cost_placeholders[{i}].note", errors),
        ))

    leveling_questions: List[LevelingQuestion] = []
    for i, row in enumerate(bl_raw.get("questions_before_final_evaluation", [])):
        if not isinstance(row, dict):
            errors.append(f"bid_leveling.questions_before_final_evaluation[{i}] must be an object.")
            continue
        rsid = str(row.get("supplier_id", "")).strip()
        if rsid not in supplier_id_set:
            errors.append(f"bid_leveling.questions_before_final_evaluation[{i}] references unknown supplier_id {rsid!r}.")
            continue
        leveling_questions.append(LevelingQuestion(
            supplier_id=rsid, question=_nonempty_str(row.get("question"), f"bid_leveling.questions_before_final_evaluation[{i}].question", errors),
            priority=_enum(row.get("priority", ""), VALID_LEVELING_PRIORITY, f"bid_leveling.questions_before_final_evaluation[{i}].priority", errors),
        ))

    bl_analysis = bl_raw.get("analysis_paragraphs", [])
    if not isinstance(bl_analysis, list) or len(bl_analysis) == 0:
        errors.append("bid_leveling.analysis_paragraphs must be a non-empty list of paragraphs.")
        bl_analysis = []

    bid_leveling = BidLeveling(
        comparison_basis_text=comparison_basis_text, scope_compliance=scope_compliance,
        assumption_exclusion_register=assumption_exclusion_register, normalized_pricing=normalized_pricing,
        missing_cost_placeholders=missing_cost_placeholders, questions_before_final_evaluation=leveling_questions,
        analysis_paragraphs=[str(p) for p in bl_analysis],
    )

    # --- Comparison Matrix (Section 6) --------------------------------
    cm_raw = register.get("comparison_matrix", {})
    if not isinstance(cm_raw, dict):
        cm_raw = {}
    cm_values_raw = cm_raw.get("values", {})
    if not isinstance(cm_values_raw, dict):
        cm_values_raw = {}
    _keys_match(cm_values_raw, set(COMPARISON_DIMENSIONS_INPUT), "'comparison_matrix.values'", errors)
    cm_values: Dict[str, Dict[str, str]] = {}
    for dim, per_supplier in cm_values_raw.items():
        if not isinstance(per_supplier, dict):
            errors.append(f"comparison_matrix.values[{dim!r}] must be an object keyed by supplier_id.")
            continue
        _keys_match(per_supplier, supplier_id_set, f"comparison_matrix.values[{dim!r}]", errors)
        cm_values[dim] = {str(k): str(v) for k, v in per_supplier.items()}
    cm_analysis_raw = cm_raw.get("analysis", {})
    if not isinstance(cm_analysis_raw, dict):
        cm_analysis_raw = {}
    _keys_match(cm_analysis_raw, set(COMPARISON_DIMENSIONS_ALL), "'comparison_matrix.analysis'", errors)
    cm_analysis = {str(k): str(v) for k, v in cm_analysis_raw.items()}
    comparison_matrix = ComparisonMatrix(values=cm_values, analysis=cm_analysis)

    # --- Coverage Heatmap (Section 7) ---------------------------------
    ch_raw = register.get("coverage_heatmap", {})
    if not isinstance(ch_raw, dict):
        ch_raw = {}
    ch_categories_raw = ch_raw.get("categories", [])
    if not isinstance(ch_categories_raw, list) or len(ch_categories_raw) == 0:
        errors.append("coverage_heatmap.categories must be a non-empty list.")
        ch_categories_raw = []
    ch_categories: List[CoverageCategory] = []
    coverage_keys = {"fully_meets", "partially_meets", "does_not_meet", "not_answered"}
    for i, cat in enumerate(ch_categories_raw):
        if not isinstance(cat, dict) or "category" not in cat or "requirement_count" not in cat or "per_supplier" not in cat:
            errors.append(f"coverage_heatmap.categories[{i}] must have category, requirement_count, per_supplier.")
            continue
        cname = str(cat["category"])
        rc = cat["requirement_count"]
        if not isinstance(rc, int) or rc <= 0:
            errors.append(f"coverage_heatmap.categories[{cname!r}].requirement_count must be a positive integer.")
            rc = 0
        per_supplier_raw = cat["per_supplier"]
        if not isinstance(per_supplier_raw, dict):
            errors.append(f"coverage_heatmap.categories[{cname!r}].per_supplier must be an object keyed by supplier_id.")
            continue
        _keys_match(per_supplier_raw, supplier_id_set, f"coverage_heatmap.categories[{cname!r}].per_supplier", errors)
        per_supplier: Dict[str, Dict[str, int]] = {}
        for sid2, counts in per_supplier_raw.items():
            if not isinstance(counts, dict) or not coverage_keys.issubset(counts.keys()):
                errors.append(f"coverage_heatmap.categories[{cname!r}].per_supplier[{sid2!r}] must have keys {sorted(coverage_keys)}.")
                continue
            vals = {k: int(counts[k]) for k in coverage_keys}
            footed = sum(vals.values())
            if footed != rc:
                errors.append(
                    f"coverage_heatmap.categories[{cname!r}].per_supplier[{sid2!r}] tally FAILED: "
                    f"sums to {footed}, not the category's requirement_count ({rc})."
                )
            per_supplier[sid2] = vals
        ch_categories.append(CoverageCategory(category=cname, requirement_count=rc, per_supplier=per_supplier))
    ch_intro = _nonempty_str(ch_raw.get("intro_paragraph"), "coverage_heatmap.intro_paragraph", errors)
    ch_analysis = ch_raw.get("analysis_paragraphs", [])
    if not isinstance(ch_analysis, list) or len(ch_analysis) == 0:
        errors.append("coverage_heatmap.analysis_paragraphs must be a non-empty list.")
        ch_analysis = []
    coverage_heatmap = CoverageHeatmap(categories=ch_categories, intro_paragraph=ch_intro, analysis_paragraphs=[str(p) for p in ch_analysis])

    # --- Weighted Scoring Matrix (Section 8) --------------------------
    ws_raw = register.get("weighted_scoring", {})
    if not isinstance(ws_raw, dict):
        ws_raw = {}
    dims_raw = ws_raw.get("dimensions", [])
    if not isinstance(dims_raw, list) or len(dims_raw) == 0:
        errors.append("weighted_scoring.dimensions must be a non-empty list.")
        dims_raw = []
    pricing_submitted_by_sid = {sp.profile.supplier_id: sp.profile.pricing_submitted for sp in suppliers}
    dimensions: List[ScoringDimension] = []
    total_weight = 0.0
    for i, d in enumerate(dims_raw):
        if not isinstance(d, dict):
            errors.append(f"weighted_scoring.dimensions[{i}] must be an object.")
            continue
        cat = str(d.get("category", "")).strip()
        subcat = str(d.get("subcategory", "")).strip()
        label = f"weighted_scoring.dimensions[{i}] ({cat}/{subcat})"
        if not cat or not subcat:
            errors.append(f"{label}: category and subcategory are required.")
        wf = d.get("weight_frac")
        if not isinstance(wf, (int, float)) or wf < 0 or wf > 1.0:
            errors.append(f"{label}.weight_frac must be a fraction in [0.0, 1.0]; got {wf!r}.")
            wf = 0.0
        total_weight += float(wf)
        scores_raw = d.get("scores", {})
        rationale_raw = d.get("rationale", {})
        if not isinstance(scores_raw, dict):
            errors.append(f"{label}.scores must be an object keyed by supplier_id.")
            scores_raw = {}
        else:
            _keys_match(scores_raw, supplier_id_set, f"{label}.scores", errors)
        if not isinstance(rationale_raw, dict):
            errors.append(f"{label}.rationale must be an object keyed by supplier_id.")
            rationale_raw = {}
        else:
            _keys_match(rationale_raw, supplier_id_set, f"{label}.rationale", errors)
        scores: Dict[str, Optional[float]] = {}
        for sid2, sc in scores_raw.items():
            if sc is None:
                scores[sid2] = None
            elif isinstance(sc, (int, float)) and 0.0 <= float(sc) <= 5.0:
                scores[sid2] = float(sc)
            else:
                errors.append(f"{label}.scores[{sid2!r}] must be null (PENDING) or numeric in [0.0, 5.0]; got {sc!r}.")
                scores[sid2] = None
            if cat.strip().lower() == "pricing" and scores[sid2] is not None and not pricing_submitted_by_sid.get(sid2, False):
                errors.append(
                    f"{label}.scores[{sid2!r}] is a scored Pricing-category dimension, but supplier {sid2!r} has "
                    "pricing_submitted=False. Pricing-dimension scores are read from the Bid Leveling normalized "
                    "TCO and must be PENDING (null) until pricing is submitted (Rule 1, never fabricate)."
                )
        rationale = {sid2: str(txt) for sid2, txt in rationale_raw.items()}
        for sid2 in supplier_id_set:
            if not rationale.get(sid2, "").strip():
                errors.append(f"{label}.rationale[{sid2!r}] must be a non-empty string, even to state why a dimension is PENDING.")
        dimensions.append(ScoringDimension(category=cat, subcategory=subcat, weight_frac=float(wf), scores=scores, rationale=rationale))
    if dims_raw and abs(total_weight - 1.0) > 0.001:
        errors.append(
            f"weighted_scoring.dimensions weight_frac values sum to {total_weight:.4f}, not 1.0 (+/- 0.001), "
            "across the full framework (scored and PENDING dimensions alike; SKILL.md's illustrative matrix "
            "totals 100%)."
        )
    sensitivity_notes = ws_raw.get("sensitivity_notes", [])
    if not isinstance(sensitivity_notes, list) or len(sensitivity_notes) == 0:
        errors.append("weighted_scoring.sensitivity_notes must be a non-empty list.")
        sensitivity_notes = []
    improvement_suggestions = ws_raw.get("improvement_suggestions", [])
    if not isinstance(improvement_suggestions, list) or len(improvement_suggestions) > 5:
        errors.append("weighted_scoring.improvement_suggestions must be a list of at most 5 items.")
        improvement_suggestions = list(improvement_suggestions)[:5] if isinstance(improvement_suggestions, list) else []
    weighted_scoring = WeightedScoring(
        dimensions=dimensions, sensitivity_notes=[str(x) for x in sensitivity_notes],
        improvement_suggestions=[str(x) for x in improvement_suggestions],
    )

    # --- Commercial & Pricing Analysis (Section 9) --------------------
    cp_raw = register.get("commercial_pricing", {})
    if not isinstance(cp_raw, dict):
        cp_raw = {}
    cp_rows_raw = cp_raw.get("rows", [])
    if not isinstance(cp_rows_raw, list) or len(cp_rows_raw) < COMMERCIAL_MIN_ROWS:
        errors.append(f"commercial_pricing.rows must be a list of at least {COMMERCIAL_MIN_ROWS} rows (SKILL.md Section 9: '12+ rows').")
        cp_rows_raw = list(cp_rows_raw) if isinstance(cp_rows_raw, list) else []
    cp_rows: List[CommercialRow] = []
    for i, row in enumerate(cp_rows_raw):
        if not isinstance(row, dict) or "dimension" not in row or "values" not in row:
            errors.append(f"commercial_pricing.rows[{i}] must have dimension and values.")
            continue
        vals = row["values"]
        if not isinstance(vals, dict):
            errors.append(f"commercial_pricing.rows[{i}].values must be an object keyed by supplier_id.")
            continue
        _keys_match(vals, supplier_id_set, f"commercial_pricing.rows[{row['dimension']!r}].values", errors)
        cp_rows.append(CommercialRow(dimension=str(row["dimension"]), values={str(k): str(v) for k, v in vals.items()}))
    cp_analysis_raw = cp_raw.get("per_supplier_analysis", {})
    if not isinstance(cp_analysis_raw, dict):
        cp_analysis_raw = {}
    _keys_match(cp_analysis_raw, supplier_id_set, "'commercial_pricing.per_supplier_analysis'", errors)
    cp_analysis: Dict[str, List[str]] = {}
    for sid2, paras in cp_analysis_raw.items():
        if not isinstance(paras, list) or len(paras) == 0:
            errors.append(f"commercial_pricing.per_supplier_analysis[{sid2!r}] must be a non-empty list of paragraphs.")
            paras = []
        cp_analysis[sid2] = [str(p) for p in paras]
    cp_norm_rec = _nonempty_str(cp_raw.get("normalization_recommendation"), "commercial_pricing.normalization_recommendation", errors)
    commercial_pricing = CommercialPricing(rows=cp_rows, per_supplier_analysis=cp_analysis, normalization_recommendation=cp_norm_rec)

    # --- MSA / Legal Risk Assessment (Section 10) ---------------------
    lr_raw = register.get("legal_risk", {})
    if not isinstance(lr_raw, dict):
        lr_raw = {}
    lr_rows_raw = lr_raw.get("rows", [])
    if not isinstance(lr_rows_raw, list) or len(lr_rows_raw) < LEGAL_MIN_ROWS:
        errors.append(f"legal_risk.rows must be a list of at least {LEGAL_MIN_ROWS} rows (SKILL.md Section 10: '15+ rows').")
        lr_rows_raw = list(lr_rows_raw) if isinstance(lr_rows_raw, list) else []
    lr_rows: List[LegalRow] = []
    for i, row in enumerate(lr_rows_raw):
        if not isinstance(row, dict) or "clause" not in row or "values" not in row:
            errors.append(f"legal_risk.rows[{i}] must have clause and values.")
            continue
        vals = row["values"]
        if not isinstance(vals, dict):
            errors.append(f"legal_risk.rows[{i}].values must be an object keyed by supplier_id.")
            continue
        _keys_match(vals, supplier_id_set, f"legal_risk.rows[{row['clause']!r}].values", errors)
        lr_rows.append(LegalRow(clause=str(row["clause"]), values={str(k): str(v) for k, v in vals.items()}))
    lr_analysis_raw = lr_raw.get("per_supplier_analysis", {})
    if not isinstance(lr_analysis_raw, dict):
        lr_analysis_raw = {}
    _keys_match(lr_analysis_raw, supplier_id_set, "'legal_risk.per_supplier_analysis'", errors)
    lr_analysis = {str(k): str(v) for k, v in lr_analysis_raw.items()}
    legal_risk = LegalRisk(rows=lr_rows, per_supplier_analysis=lr_analysis)

    # --- Inconsistency Register (Section 11) --------------------------
    inconsistency_register: List[InconsistencyItem] = []
    seen_incon_ids: List[str] = []
    for i, row in enumerate(register.get("inconsistency_register", [])):
        if not isinstance(row, dict):
            errors.append(f"inconsistency_register[{i}] must be an object.")
            continue
        iid = str(row.get("inconsistency_id", "")).strip()
        if iid in seen_incon_ids:
            errors.append(f"Duplicate inconsistency_id {iid!r} in inconsistency_register.")
        seen_incon_ids.append(iid)
        rsid = str(row.get("supplier_id", "")).strip()
        if rsid not in supplier_id_set:
            errors.append(f"inconsistency_register[{i}] ({iid}) references unknown supplier_id {rsid!r}.")
            continue
        inconsistency_register.append(InconsistencyItem(
            inconsistency_id=iid, supplier_id=rsid,
            severity=_enum(row.get("severity", ""), VALID_INCONSISTENCY_SEVERITY, f"inconsistency_register[{iid}].severity", errors),
            description=_nonempty_str(row.get("description"), f"inconsistency_register[{iid}].description", errors),
            action_required=_nonempty_str(row.get("action_required"), f"inconsistency_register[{iid}].action_required", errors),
        ))

    # --- Clarification Questions (Section 12) -------------------------
    clarification_questions: List[ClarificationQuestion] = []
    seen_q_ids: List[str] = []
    for i, row in enumerate(register.get("clarification_questions", [])):
        if not isinstance(row, dict):
            errors.append(f"clarification_questions[{i}] must be an object.")
            continue
        qid = str(row.get("question_id", "")).strip()
        if qid in seen_q_ids:
            errors.append(f"Duplicate question_id {qid!r} in clarification_questions.")
        seen_q_ids.append(qid)
        rsid = str(row.get("supplier_id", "")).strip()
        if rsid not in supplier_id_set:
            errors.append(f"clarification_questions[{i}] ({qid}) references unknown supplier_id {rsid!r}.")
            continue
        source_type = _enum(row.get("source_type", ""), VALID_SOURCE_TYPE, f"clarification_questions[{qid}].source_type", errors)
        req_id = row.get("req_id")
        if req_id is not None:
            req_id = str(req_id)
        elif source_type != "Bid_Leveling_Gap":
            errors.append(f"clarification_questions[{qid}].req_id is required unless source_type is 'Bid_Leveling_Gap'.")
        clarification_questions.append(ClarificationQuestion(
            question_id=qid, supplier_id=rsid, source_type=source_type, req_id=req_id,
            description=_nonempty_str(row.get("description"), f"clarification_questions[{qid}].description", errors),
            priority=_enum(row.get("priority", ""), VALID_CLARIFICATION_PRIORITY, f"clarification_questions[{qid}].priority", errors),
            recommended_format=_nonempty_str(row.get("recommended_format"), f"clarification_questions[{qid}].recommended_format", errors),
        ))

    # --- Final Recommendation (Section 13) ----------------------------
    fr_raw = register.get("final_recommendation", {})
    if not isinstance(fr_raw, dict):
        fr_raw = {}
    primary_raw = fr_raw.get("primary", {})
    primary_sid = str(primary_raw.get("supplier_id", "")).strip() if isinstance(primary_raw, dict) else ""
    if primary_sid not in supplier_id_set:
        errors.append(f"final_recommendation.primary.supplier_id ({primary_sid!r}) is not a declared supplier_id.")
    primary_evidence = primary_raw.get("evidence", []) if isinstance(primary_raw, dict) else []
    if not isinstance(primary_evidence, list) or len(primary_evidence) == 0:
        errors.append("final_recommendation.primary.evidence must be a non-empty list.")
        primary_evidence = []
    primary_caveats = primary_raw.get("caveats", []) if isinstance(primary_raw, dict) else []
    if not isinstance(primary_caveats, list) or len(primary_caveats) == 0:
        errors.append("final_recommendation.primary.caveats must be a non-empty list.")
        primary_caveats = []
    primary_conditions = primary_raw.get("conditions", []) if isinstance(primary_raw, dict) else []
    primary = RecommendationPrimary(
        supplier_id=primary_sid, evidence=[str(x) for x in primary_evidence],
        conditions=[str(x) for x in primary_conditions] if isinstance(primary_conditions, list) else [],
        caveats=[str(x) for x in primary_caveats],
    )

    secondary_raw = fr_raw.get("secondary")
    secondary: Optional[RecommendationSecondary] = None
    if secondary_raw is not None:
        if not isinstance(secondary_raw, dict):
            errors.append("final_recommendation.secondary, when present, must be an object.")
        else:
            ssid = str(secondary_raw.get("supplier_id", "")).strip()
            if ssid not in supplier_id_set:
                errors.append(f"final_recommendation.secondary.supplier_id ({ssid!r}) is not a declared supplier_id.")
            if ssid == primary_sid:
                errors.append("final_recommendation.secondary.supplier_id must differ from the primary recommendation.")
            secondary = RecommendationSecondary(
                supplier_id=ssid,
                triggering_scenario=_nonempty_str(secondary_raw.get("triggering_scenario"), "final_recommendation.secondary.triggering_scenario", errors),
                sensitivity_analysis=_nonempty_str(secondary_raw.get("sensitivity_analysis"), "final_recommendation.secondary.sensitivity_analysis", errors),
            )

    conditional: List[RecommendationConditional] = []
    for i, row in enumerate(fr_raw.get("conditional", [])):
        if not isinstance(row, dict):
            errors.append(f"final_recommendation.conditional[{i}] must be an object.")
            continue
        csid = str(row.get("supplier_id", "")).strip()
        if csid not in supplier_id_set:
            errors.append(f"final_recommendation.conditional[{i}].supplier_id ({csid!r}) is not a declared supplier_id.")
            continue
        prereqs = row.get("prerequisites", [])
        if not isinstance(prereqs, list) or len(prereqs) == 0:
            errors.append(f"final_recommendation.conditional[{i}] ({csid}).prerequisites must be a non-empty list.")
            prereqs = []
        conditional.append(RecommendationConditional(supplier_id=csid, prerequisites=[str(x) for x in prereqs]))

    not_recommended: List[RecommendationNotRecommended] = []
    for i, row in enumerate(fr_raw.get("not_recommended", [])):
        if not isinstance(row, dict):
            errors.append(f"final_recommendation.not_recommended[{i}] must be an object.")
            continue
        nsid = str(row.get("supplier_id", "")).strip()
        if nsid not in supplier_id_set:
            errors.append(f"final_recommendation.not_recommended[{i}].supplier_id ({nsid!r}) is not a declared supplier_id.")
            continue
        not_recommended.append(RecommendationNotRecommended(
            supplier_id=nsid, rationale=_nonempty_str(row.get("rationale"), f"final_recommendation.not_recommended[{i}] ({nsid}).rationale", errors),
        ))

    standard_caveats = fr_raw.get("standard_caveats", [])
    if not isinstance(standard_caveats, list) or len(standard_caveats) == 0:
        errors.append("final_recommendation.standard_caveats must be a non-empty list.")
        standard_caveats = []

    final_recommendation = FinalRecommendation(
        primary=primary, secondary=secondary, conditional=conditional,
        not_recommended=not_recommended, standard_caveats=[str(x) for x in standard_caveats],
    )

    if errors:
        raise RfpAnalysisValidationError(
            "RFP analysis register failed validation; refusing to guess missing or invalid fields. "
            "Issues found:\n  - " + "\n  - ".join(errors)
        )

    return RfpAnalysisInput(
        rfx_name=rfx_name, category=category, case_id=case_id, analysis_date=analysis_date,
        analysis_mode=analysis_mode, report_mode=report_mode, requirement_count=requirement_count,
        suppliers=suppliers, executive_summary_scope=executive_summary_scope,
        executive_summary_overviews={str(k): str(v) for k, v in exec_overviews.items()},
        executive_summary_key_findings=[str(x) for x in key_findings],
        bid_leveling=bid_leveling, comparison_matrix=comparison_matrix, coverage_heatmap=coverage_heatmap,
        weighted_scoring=weighted_scoring, commercial_pricing=commercial_pricing, legal_risk=legal_risk,
        inconsistency_register=inconsistency_register, clarification_questions=clarification_questions,
        final_recommendation=final_recommendation,
    )


# ===========================================================================
# 3. Ground-truth computation (via the vendored kernel only, per the HARD RULE)
# ===========================================================================

@dataclass
class GroundTruth:
    adequacy_overall: Dict[str, float]                  # supplier_id -> mean of the 12 canonical scores
    weighted_total: Dict[str, Optional[float]]          # supplier_id -> kernel total over covered dims, or None
    weighted_missing_frac: Dict[str, float]              # supplier_id -> sum of PENDING dims' weight_frac
    weighted_rank: List[str]                            # supplier_ids with a non-null total, sorted desc
    category_fm_pct: Dict[str, Dict[str, float]]        # category -> {supplier_id: Fully-Meets %}
    category_leader: Dict[str, str]                      # category -> supplier_id with the highest FM %
    bid_leveling_gate_passed: bool
    bid_leveling_gate_failures: List[str]
    n_suppliers: int
    section_numbers: Dict[str, Any]                      # see _compute_section_numbers()


def _compute_section_numbers(n_suppliers: int) -> Dict[str, Any]:
    """Compute the fixed-order section numbering for N suppliers (judgment
    call 2 in the module docstring): Executive Summary is always Section 3,
    per-supplier sections occupy 4..(3+N), Bid Leveling is (4+N), and the
    eight cross-cutting sections occupy (5+N)..(12+N), preserving
    comparison-patterns.md's section titles and relative order exactly."""
    exec_no = 3
    supplier_first = 4
    supplier_last = 3 + n_suppliers
    bid_leveling_no = supplier_last + 1
    crosscut_first = bid_leveling_no + 1
    crosscut_numbers = list(range(crosscut_first, crosscut_first + len(CROSSCUT_SECTION_TITLES)))
    return {
        "executive_summary": exec_no,
        "supplier_first": supplier_first,
        "supplier_last": supplier_last,
        "bid_leveling": bid_leveling_no,
        "crosscut": dict(zip(CROSSCUT_SECTION_TITLES, crosscut_numbers)),
    }


def compute_ground_truth(reg: RfpAnalysisInput) -> GroundTruth:
    """Compute every derived figure this report shows. The Weighted Scoring
    Matrix total is computed by calling ONLY numeric_kernel.weighted_score()
    per supplier, over that supplier's scored (non-PENDING) dimensions with
    weights renormalized to sum to 1.0 (judgment call 4). This is the
    reference against which the hard-coded invariants are checked before
    the document is saved."""
    if not KERNEL_AVAILABLE:
        raise RuntimeError(
            "numeric_kernel.py could not be imported; this generator refuses to hand-compute the "
            f"weighted score in its place (SKILL.md Section 8 'Kernel-computed arithmetic (HARD RULE)'). "
            f"Import error: {_KERNEL_IMPORT_ERROR}"
        )

    supplier_ids = [sp.profile.supplier_id for sp in reg.suppliers]

    # --- N.3 OVERALL adequacy score (computed, never accepted as input) ---
    adequacy_overall: Dict[str, float] = {}
    for sp in reg.suppliers:
        scores_by_section = {a.rfp_section: a.score for a in sp.adequacy_scores}
        vals = [scores_by_section[s] for s in CANONICAL_RFP_SECTIONS]
        adequacy_overall[sp.profile.supplier_id] = sum(vals) / len(vals)

    # --- Section 8 Weighted Scoring Matrix, per-supplier kernel total -----
    weighted_total: Dict[str, Optional[float]] = {}
    weighted_missing_frac: Dict[str, float] = {}
    for sid in supplier_ids:
        covered = [d for d in reg.weighted_scoring.dimensions if d.scores.get(sid) is not None]
        covered_weight = sum(d.weight_frac for d in covered)
        weighted_missing_frac[sid] = 1.0 - covered_weight
        if not covered:
            weighted_total[sid] = None
            continue
        scores_map = {f"DIM_{i}": d.scores[sid] for i, d in enumerate(covered)}
        renorm_weights = {f"DIM_{i}": d.weight_frac / covered_weight for i, d in enumerate(covered)}
        try:
            weighted_total[sid] = kernel_weighted_score(scores_map, renorm_weights)
        except KernelWeightSumError as e:
            raise ReconciliationError(
                f"Per-supplier-covered-weight-sums-to-one invariant FAILED for supplier {sid!r}: "
                f"numeric_kernel.weighted_score() refused the renormalized covered-dimension weights. "
                f"Kernel message: {e}"
            )
        except KernelInvalidInputError as e:
            raise ReconciliationError(f"weighted_score() call for supplier {sid!r} failed: {e}")

    weighted_rank = sorted(
        [sid for sid in supplier_ids if weighted_total[sid] is not None],
        key=lambda sid: weighted_total[sid], reverse=True,
    )

    # --- Section 7 Coverage Heatmap: Fully-Meets % and category leader ----
    category_fm_pct: Dict[str, Dict[str, float]] = {}
    category_leader: Dict[str, str] = {}
    for cat in reg.coverage_heatmap.categories:
        pct_row: Dict[str, float] = {}
        for sid in supplier_ids:
            counts = cat.per_supplier[sid]
            pct_row[sid] = (counts["fully_meets"] / cat.requirement_count) * 100.0 if cat.requirement_count else 0.0
        category_fm_pct[cat.category] = pct_row
        category_leader[cat.category] = max(pct_row, key=lambda sid: pct_row[sid])

    # --- Bid Leveling gate (the named HARD RULE gate) ----------------------
    gate_failures: List[str] = []
    pricing_suppliers = [sp.profile.supplier_id for sp in reg.suppliers if sp.profile.pricing_submitted]
    normalized_by_sid: Dict[str, List[NormalizedPricingRow]] = {}
    for row in reg.bid_leveling.normalized_pricing:
        normalized_by_sid.setdefault(row.supplier_id, []).append(row)
    register_by_sid: Dict[str, List[AssumptionExclusionRow]] = {}
    for row in reg.bid_leveling.assumption_exclusion_register:
        register_by_sid.setdefault(row.supplier_id, []).append(row)

    for sid in pricing_suppliers:
        sp = next(s for s in reg.suppliers if s.profile.supplier_id == sid)
        if sp.profile.leveling_status != "Complete":
            gate_failures.append(
                f"Supplier {sid!r} submitted pricing but leveling_status is {sp.profile.leveling_status!r}, not 'Complete'."
            )
        rows = normalized_by_sid.get(sid, [])
        if not rows:
            gate_failures.append(f"Supplier {sid!r} submitted pricing but has no bid_leveling.normalized_pricing row.")
        else:
            for r in rows:
                if r.one_time_total is None or r.recurring_total is None:
                    gate_failures.append(f"Supplier {sid!r} scenario {r.scenario!r} is missing its one-time/recurring split.")
                if r.reported_tco is None or r.normalized_tco is None:
                    gate_failures.append(f"Supplier {sid!r} scenario {r.scenario!r} is missing reported_tco or normalized_tco.")
        if not register_by_sid.get(sid):
            gate_failures.append(f"Supplier {sid!r} submitted pricing but has no assumption_exclusion_register entry.")

    matched_gap_suppliers = {
        (q.supplier_id) for q in reg.clarification_questions if q.source_type == "Bid_Leveling_Gap"
    }
    for gap in reg.bid_leveling.questions_before_final_evaluation:
        if gap.supplier_id not in matched_gap_suppliers:
            gate_failures.append(
                f"Bid Leveling question for supplier {gap.supplier_id!r} ({gap.question!r}) has no matching "
                "clarification_questions entry with source_type='Bid_Leveling_Gap'."
            )

    if not reg.bid_leveling.comparison_basis_text.strip():
        gate_failures.append("bid_leveling.comparison_basis_text is empty; a common comparison basis is mandatory.")

    bid_leveling_gate_passed = len(gate_failures) == 0

    return GroundTruth(
        adequacy_overall=adequacy_overall, weighted_total=weighted_total,
        weighted_missing_frac=weighted_missing_frac, weighted_rank=weighted_rank,
        category_fm_pct=category_fm_pct, category_leader=category_leader,
        bid_leveling_gate_passed=bid_leveling_gate_passed, bid_leveling_gate_failures=gate_failures,
        n_suppliers=len(supplier_ids), section_numbers=_compute_section_numbers(len(supplier_ids)),
    )


# ===========================================================================
# 4. Hard-coded invariant checks (run BEFORE saving)
# ===========================================================================

def _assert_adequacy_overall_range(gt: GroundTruth) -> None:
    for sid, val in gt.adequacy_overall.items():
        if not (0.0 <= val <= 5.0):
            raise ReconciliationError(
                f"Adequacy-overall-range invariant FAILED for supplier {sid!r}: computed OVERALL {val} "
                "is outside [0.0, 5.0]."
            )


def _assert_weighted_total_range(gt: GroundTruth) -> None:
    for sid, val in gt.weighted_total.items():
        if val is not None and not (0.0 <= val <= 5.0):
            raise ReconciliationError(
                f"Weighted-total-range invariant FAILED for supplier {sid!r}: weighted total {val} "
                "is outside [0.0, 5.0]."
            )


def _assert_weighted_ranking_sorted(gt: GroundTruth) -> None:
    totals = [gt.weighted_total[sid] for sid in gt.weighted_rank]
    if totals != sorted(totals, reverse=True):
        raise ReconciliationError(
            f"Weighted-ranking-sorted invariant FAILED: weighted_rank {gt.weighted_rank} is not sorted "
            f"by weighted total descending (totals: {totals})."
        )


def _assert_bid_leveling_gate(gt: GroundTruth) -> None:
    """The named HARD RULE gate: 'GATE CHECK: Bid Leveling Complete.' If any
    check fails, this generator refuses the entire document (judgment call
    6 in the module docstring)."""
    if not gt.bid_leveling_gate_passed:
        raise ReconciliationError(
            "Bid-Leveling-gate invariant FAILED (SKILL.md 'GATE CHECK: Bid Leveling Complete'): no ranking, "
            "weighted score, or recommendation may be produced from unleveled figures. Failures:\n  - "
            + "\n  - ".join(gt.bid_leveling_gate_failures)
        )


def run_hardcoded_invariant_checks(reg: RfpAnalysisInput, gt: GroundTruth) -> None:
    """Run every hard-coded invariant. Raises ReconciliationError on any
    failure; callers must not save the document if this raises."""
    _assert_adequacy_overall_range(gt)
    _assert_weighted_total_range(gt)
    _assert_weighted_ranking_sorted(gt)
    _assert_bid_leveling_gate(gt)


# ===========================================================================
# 5. Document builder (python-docx)
# ===========================================================================

def _require_docx() -> None:
    if not DOCX_AVAILABLE:
        raise ImportError(
            "python-docx is not installed in this Python environment, so no "
            ".docx file can be written. Install it (`pip install python-docx`) "
            "or point this script at an interpreter that already has it. "
            f"Original import error: {_DOCX_IMPORT_ERROR}"
        )


def _set_cell_background(cell, hex_color: str) -> None:
    """Set a table cell's background fill via raw OOXML; python-docx has no
    high-level cell-shading API."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_text(cell, text: str, bold: bool = False, color_hex: str = LILLY_BLACK,
                    size_pt: int = 9, align=None) -> None:
    cell.text = str(text)
    para = cell.paragraphs[0]
    if align is not None:
        para.alignment = align
    run = para.runs[0] if para.runs else para.add_run("")
    run.font.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor.from_string(color_hex)


def _add_heading(doc, text: str, level: int, toc: bool = False):
    """Manually-styled heading (no dependency on a template's built-in
    Heading styles, mirroring evaluation_report_generator.py), UNLESS
    toc=True: per judgment call 3 in the module docstring (docx-title-
    page-spec.md, quoted in SKILL.md: 'The TOC indexes Heading2 paragraphs
    only'), every heading that must appear in the Table of Contents field
    carries python-docx's built-in 'Heading 2' PARAGRAPH STYLE (for Word's
    TOC field to find it) in addition to this module's own run-level visual
    formatting (so top-level sections and the two named per-vendor sub-
    headings can still look visually distinct from each other and from
    body text)."""
    spec = {1: (14, BOLD_BLUE, 18, 8), 2: (12, LILLY_BLACK, 14, 6), 3: (11, LILLY_BLACK, 10, 4)}
    size_pt, color_hex, before_pt, after_pt = spec.get(level, spec[3])
    para = doc.add_paragraph()
    if toc:
        para.style = doc.styles["Heading 2"]
    para.paragraph_format.space_before = Pt(before_pt)
    para.paragraph_format.space_after = Pt(after_pt)
    run = para.add_run(text)
    run.font.bold = True
    run.font.size = Pt(size_pt)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor.from_string(color_hex)
    return para


def _add_toc_field(doc) -> None:
    """Insert a real Word TOC field indexing Heading2 paragraphs only
    ('\\o "2-2" \\h \\z \\u'), per judgment call 3 in the module docstring.
    python-docx has no high-level TOC API; this is the standard raw-OOXML
    technique (fldChar begin / instrText / fldChar separate / placeholder
    run / fldChar end). python-docx cannot force Word to compute page
    numbers at write time, so the field carries a placeholder run Word
    replaces once the user updates the field (F9 or right-click ->
    Update Field), the standard, disclosed limitation of building a TOC
    field outside Word itself."""
    para = doc.add_paragraph()
    run = para.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "2-2" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "Right-click and choose Update Field to populate the Table of Contents."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r_element = run._r
    r_element.append(fld_begin)
    r_element.append(instr)
    r_element.append(fld_sep)
    r_element.append(fld_text)
    r_element.append(fld_end)


def _add_body_paragraph(doc, text: str, bold: bool = False, italic: bool = False,
                         size_pt: int = 10, color_hex: str = LILLY_BLACK) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.font.bold = bold
    run.font.italic = italic
    run.font.size = Pt(size_pt)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor.from_string(color_hex)


def _add_body_paragraphs(doc, text: str, **kwargs) -> None:
    """Split already-authored multi-paragraph narrative on blank lines and
    render each as its own paragraph (Content Writing Rule: 'at least one
    full paragraph of connected prose')."""
    parts = [p.strip() for p in text.split("\n\n") if p.strip()]
    if not parts:
        parts = [text]
    for p in parts:
        _add_body_paragraph(doc, p, **kwargs)


def _add_bullet_list(doc, items: Sequence[str]) -> None:
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        run = para.add_run(item)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor.from_string(LILLY_BLACK)


def _add_numbered_list(doc, items: Sequence[str]) -> None:
    """Manually-numbered list (avoids depending on the default template
    carrying a 'List Number' style; mirrors this module's own manual-
    heading convention of not depending on template styles beyond the
    'List Bullet' style evaluation_report_generator.py already relies on)."""
    for i, item in enumerate(items, start=1):
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run(f"{i}. {item}")
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor.from_string(LILLY_BLACK)


def _add_table(doc, headers: Sequence[str], rows: Sequence[Sequence[Any]],
                cell_colors: Optional[Dict[Tuple[int, int], str]] = None):
    """cell_colors, if given, maps (row_index_0_based, col_index_0_based) ->
    hex fill color for data rows (header row is always LILLY_RED)."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell_background(hdr_cells[i], LILLY_RED)
        _set_cell_text(hdr_cells[i], h, bold=True, color_hex=WHITE, size_pt=9)
    for r, row_values in enumerate(rows):
        cells = table.add_row().cells
        for c, val in enumerate(row_values):
            fill = (cell_colors or {}).get((r, c))
            if fill:
                _set_cell_background(cells[c], fill)
                _set_cell_text(cells[c], val, bold=True, color_hex=WHITE, size_pt=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                _set_cell_text(cells[c], str(val), size_pt=9)
    return table


def _add_two_column(doc, left_title: str, left_items: Sequence[str],
                     right_title: str, right_items: Sequence[str]) -> None:
    """Strengths/risks-style 2-column layout (Content Writing Rule 4: 'Use
    columns where appropriate. Strengths and risks can be presented in a
    2-column layout'). Implemented as a borderless 1-row x 2-col table."""
    table = doc.add_table(rows=1, cols=2)
    left_cell, right_cell = table.rows[0].cells
    for cell, title, items, marker in ((left_cell, left_title, left_items, "+ "), (right_cell, right_title, right_items, "! ")):
        cell.text = ""
        p0 = cell.paragraphs[0]
        run = p0.add_run(title)
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor.from_string(BOLD_BLUE)
        for item in items:
            p = cell.add_paragraph()
            r = p.add_run(marker + item)
            r.font.size = Pt(10)
            r.font.name = "Calibri"
            r.font.color.rgb = RGBColor.from_string(LILLY_BLACK)


def _pct(frac: float) -> str:
    return f"{frac * 100:.1f}%"


def _usd(value: float) -> str:
    return f"${value:,.0f}"


def _score1(v: Optional[float]) -> str:
    return f"{v:.1f}" if v is not None else "PENDING"


def _total_or_label(v: Optional[float]) -> str:
    return f"{v:.2f}" if v is not None else "INSUFFICIENT DATA (no scored dimensions)"


RISK_LEVEL_COLOR = {"Low": POSITIVE_BLUE, "Medium": AMBER, "High": LILLY_RED}


def build_document(reg: RfpAnalysisInput, gt: GroundTruth):
    """Build analysis_summary.docx per SKILL.md's 'Complete Section
    Structure' / comparison-patterns.md: title page, TOC, Executive
    Summary, one consolidated section per supplier, the gated Bid Leveling
    section, and the eight cross-cutting sections, in that fixed order.
    This function does NOT save the file and does NOT run the
    reconciliation checks; call generate_rfp_analysis_report() for the
    full validated pipeline."""
    _require_docx()

    supplier_by_id = {sp.profile.supplier_id: sp for sp in reg.suppliers}
    supplier_ids_ordered = [sp.profile.supplier_id for sp in reg.suppliers]
    supplier_name = {sid: supplier_by_id[sid].profile.supplier_name for sid in supplier_ids_ordered}
    full_mode = reg.report_mode == "Full"
    sn = gt.section_numbers

    doc = Document()

    # Footer (all pages), per docx-design-system.md's Footer text spec
    # (Calibri 8pt, Bold Grey #8A969E), mirroring evaluation_report_generator.py.
    footer_para = doc.sections[0].footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        "CONFIDENTIAL, internal use only. Lilly Procurement Skills Suite, RFP Response Analysis."
    )
    footer_run.font.size = Pt(8)
    footer_run.font.name = "Calibri"
    footer_run.font.color.rgb = RGBColor.from_string(MUTED_GREY)

    # --- 1. Title page (this skill's report title is LOCKED per SKILL.md's
    # Branding and Document Design section: "This skill's report title is
    # 'SUPPLIER RESPONSE ANALYSIS'") ---------------------------------------
    title_para = doc.add_paragraph()
    title_para.paragraph_format.space_after = Pt(4)
    title_run = title_para.add_run("SUPPLIER RESPONSE ANALYSIS")
    title_run.font.bold = True
    title_run.font.size = Pt(26)
    title_run.font.name = "Calibri"
    title_run.font.color.rgb = RGBColor.from_string(LILLY_BLACK)

    scope_para = doc.add_paragraph()
    scope_para.paragraph_format.space_after = Pt(10)
    scope_run = scope_para.add_run(f"{reg.rfx_name}, {reg.category}")
    scope_run.font.size = Pt(11)
    scope_run.font.name = "Calibri"
    scope_run.font.color.rgb = RGBColor.from_string(BOLD_BLUE)

    conf_para = doc.add_paragraph()
    conf_para.paragraph_format.space_after = Pt(10)
    conf_run = conf_para.add_run("CONFIDENTIAL, INTERNAL USE ONLY")
    conf_run.font.bold = True
    conf_run.font.size = Pt(10)
    conf_run.font.name = "Calibri"
    conf_run.font.color.rgb = RGBColor.from_string(LILLY_RED)

    meta_lines = [
        f"Case ID: {reg.case_id}",
        f"Analysis date: {reg.analysis_date}",
        f"Prepared by Eli Lilly and Company | Procurement",
        f"Suppliers evaluated: {len(reg.suppliers)}",
        f"Requirements evaluated: {reg.requirement_count}",
        f"Analysis mode: {reg.analysis_mode}    Report mode: {reg.report_mode}",
    ]
    for line in meta_lines:
        _add_body_paragraph(doc, line, color_hex=BOLD_BLUE if "Prepared by" in line else LILLY_BLACK)
    doc.add_page_break()

    # --- 2. Table of Contents ----------------------------------------------
    _add_heading(doc, "Table of Contents", level=1, toc=False)
    _add_toc_field(doc)
    doc.add_page_break()

    # --- Section 3: Executive Summary --------------------------------------
    _add_heading(doc, f"{sn['executive_summary']}. Executive Summary", level=1, toc=True)
    _add_body_paragraphs(doc, reg.executive_summary_scope)

    _add_heading(doc, "Overview of Evaluated Suppliers", level=3, toc=False)
    for sid in supplier_ids_ordered:
        _add_body_paragraph(doc, f"{supplier_name[sid]}: {reg.executive_summary_overviews[sid]}")

    _add_heading(doc, "Key Findings", level=3, toc=False)
    _add_numbered_list(doc, reg.executive_summary_key_findings)

    _add_heading(doc, "Recommendation Summary (Preliminary)", level=3, toc=False)
    primary = reg.final_recommendation.primary
    _add_body_paragraph(
        doc,
        f"Primary recommendation: {supplier_name[primary.supplier_id]} "
        f"(proposed weighted total {_total_or_label(gt.weighted_total[primary.supplier_id])} of 5.0). "
        + (primary.caveats[0] if primary.caveats else ""),
        bold=True,
    )
    if reg.final_recommendation.secondary and full_mode:
        sec = reg.final_recommendation.secondary
        _add_body_paragraph(
            doc, f"Secondary recommendation: {supplier_name[sec.supplier_id]}. {sec.triggering_scenario}"
        )
    if reg.final_recommendation.conditional and full_mode:
        for c in reg.final_recommendation.conditional:
            _add_body_paragraph(
                doc, f"Conditional: {supplier_name[c.supplier_id]}, subject to: {'; '.join(c.prerequisites)}."
            )
    _add_body_paragraph(
        doc,
        "This is a preliminary, AI-proposed recommendation per this skill's own Boundaries and Rule 5; the "
        "official score, weighted total, and award recommendation belong to evaluation-engine.",
        italic=True, size_pt=9,
    )

    # --- Sections 4..(3+N): one consolidated section per supplier ----------
    for offset, sid in enumerate(supplier_ids_ordered):
        num = sn["supplier_first"] + offset
        sp = supplier_by_id[sid]
        prof = sp.profile

        _add_heading(doc, f"{num}. {prof.supplier_name}", level=1, toc=True)

        # N.1 Profile
        _add_heading(doc, f"{num}.1 Profile", level=3, toc=False)
        _add_body_paragraphs(doc, prof.narrative)
        _add_table(doc, ["Field", "Value"], [[k, v] for k, v in prof.data_card.items()])
        clients_line = ", ".join(prof.pharma_clients) if prof.pharma_clients else "Not Stated"
        _add_body_paragraph(doc, f"Named pharma / life sciences clients: {clients_line}")
        _add_body_paragraph(
            doc,
            f"Deployment model: {prof.deployment_model}. Contract posture: {prof.contract_posture}. "
            f"Lilly vendor status: {prof.lilly_vendor_status}.",
        )
        _add_body_paragraph(doc, f"Pricing summary: {prof.pricing_summary}")
        cov = sp.coverage
        _add_body_paragraph(
            doc,
            f"Requirements coverage summary: {cov['fully_meets']} Fully Meets, {cov['partially_meets']} "
            f"Partially Meets, {cov['does_not_meet']} Does Not Meet, {cov['not_answered']} Not Answered "
            f"(of {cov['total']} requirements).",
        )
        _add_body_paragraph(
            doc,
            f"Bid Leveling status: {prof.leveling_status}"
            + (" (pricing not yet submitted)." if not prof.pricing_submitted else "."),
            italic=not prof.pricing_submitted,
        )

        # N.2 Response Summary & Analysis (TOC-indexed per docx-title-page-spec.md)
        _add_heading(doc, "Response Summary & Analysis", level=2, toc=True)
        ra = sp.response_analysis
        for key, subtitle in RESPONSE_ANALYSIS_SUBSECTIONS:
            _add_heading(doc, subtitle, level=3, toc=False)
            text = getattr(ra, key)
            if key == "submission_volume":
                text = f"{text} (Approximately {ra.submission_volume_pages} pages submitted.)"
            _add_body_paragraphs(doc, text)

        # N.3 RFP Section Adequacy Scores (TOC-indexed per docx-title-page-spec.md)
        _add_heading(doc, "RFP Section Adequacy Scores", level=2, toc=True)
        scores_by_section = {a.rfp_section: a for a in sp.adequacy_scores}
        adeq_rows = [
            [section, f"{scores_by_section[section].score:.1f}", scores_by_section[section].notes]
            for section in CANONICAL_RFP_SECTIONS
        ]
        adeq_rows.append(["OVERALL", f"{gt.adequacy_overall[sid]:.2f}", "Computed: mean of the twelve section scores above."])
        _add_table(doc, ["RFP Section", "Score (0.0-5.0)", "Notes"], adeq_rows)

        # N.4 Strengths / N.5 Risks
        _add_two_column(doc, f"{num}.4 Key Strengths", sp.strengths, f"{num}.5 Key Risks", sp.risks)

    # --- Section (4+N): Bid Leveling & Normalization (Gate) -----------------
    bl_num = sn["bid_leveling"]
    _add_heading(doc, f"{bl_num}. Bid Leveling & Normalization (Gate)", level=1, toc=True)
    _add_body_paragraph(
        doc,
        "GATE CHECK: Bid Leveling Complete, PASSED for every supplier that submitted pricing. Every "
        "downstream comparison, the Weighted Scoring Matrix, the pricing analysis, and the Final "
        "Recommendation below read from the normalized figures in this section, never the raw reported price.",
        bold=True,
    )
    _add_heading(doc, "Comparison Basis", level=3, toc=False)
    _add_body_paragraph(doc, reg.bid_leveling.comparison_basis_text)

    _add_heading(doc, "Scope-Compliance Map", level=3, toc=False)
    scope_headers = ["Scope Line"] + [supplier_name[sid] for sid in supplier_ids_ordered]
    scope_status_color = {"Included": POSITIVE_BLUE, "Additional Cost": AMBER, "Excluded": LILLY_RED, "Silent": MUTED_GREY}
    scope_rows, scope_colors = [], {}
    for r, row in enumerate(reg.bid_leveling.scope_compliance):
        vals = [row.status[sid] for sid in supplier_ids_ordered]
        scope_rows.append([row.scope_line] + vals)
        for c, sid in enumerate(supplier_ids_ordered, start=1):
            scope_colors[(r, c)] = scope_status_color.get(row.status[sid])
    _add_table(doc, scope_headers, scope_rows, cell_colors=scope_colors)

    _add_heading(doc, "Assumption & Exclusion Register", level=3, toc=False)
    if full_mode:
        ae_rows = [
            [supplier_name[a.supplier_id], a.row_type, a.item, a.cost_impact, a.confidence, a.source]
            for a in reg.bid_leveling.assumption_exclusion_register
        ]
        _add_table(doc, ["Supplier", "Type", "Item", "Cost Impact", "Confidence", "Source"], ae_rows)
    else:
        by_sid_count: Dict[str, int] = {}
        for a in reg.bid_leveling.assumption_exclusion_register:
            by_sid_count[a.supplier_id] = by_sid_count.get(a.supplier_id, 0) + 1
        summary = "; ".join(f"{supplier_name[sid]}: {n} item(s)" for sid, n in by_sid_count.items())
        _add_body_paragraph(doc, f"Assumption & exclusion counts by supplier (Full report carries the itemized register): {summary}.")

    _add_heading(doc, "Normalized Pricing", level=3, toc=False)
    np_rows = [
        [supplier_name[r.supplier_id], r.scenario, r.comparison_basis_unit, _usd(r.reported_price),
         _usd(r.one_time_total), _usd(r.recurring_total), _usd(r.reported_tco), _usd(r.normalized_tco)]
        for r in reg.bid_leveling.normalized_pricing
    ]
    _add_table(doc, ["Supplier", "Scenario", "Unit", "Reported Price", "One-Time", "Recurring", "Reported TCO", "Normalized TCO"], np_rows)

    _add_heading(doc, "Missing-Cost Placeholders", level=3, toc=False)
    if reg.bid_leveling.missing_cost_placeholders:
        _add_bullet_list(doc, [f"{m.scope_line}: {m.note}" for m in reg.bid_leveling.missing_cost_placeholders])
    else:
        _add_body_paragraph(doc, "No missing-cost placeholders were identified; every priced scope line was addressed by at least one supplier.")

    _add_heading(doc, "Questions Before Final Evaluation", level=3, toc=False)
    if reg.bid_leveling.questions_before_final_evaluation:
        _add_numbered_list(doc, [
            f"[{q.priority}] {supplier_name[q.supplier_id]}: {q.question}"
            for q in reg.bid_leveling.questions_before_final_evaluation
        ])
    else:
        _add_body_paragraph(doc, "No leveling gaps were identified; no clarification questions were generated from this phase.")

    for para_text in reg.bid_leveling.analysis_paragraphs:
        _add_body_paragraph(doc, para_text)

    # --- Section (5+N): Cross-Supplier Comparison Matrix --------------------
    cc = sn["crosscut"]
    _add_heading(doc, f"{cc['Cross-Supplier Comparison Matrix']}. Cross-Supplier Comparison Matrix", level=1, toc=True)
    cmp_headers = ["Dimension"] + [supplier_name[sid] for sid in supplier_ids_ordered]
    cmp_rows = []
    for dim in COMPARISON_DIMENSIONS_ALL:
        if dim == "Adequacy Score":
            vals = [f"{gt.adequacy_overall[sid]:.2f}" for sid in supplier_ids_ordered]
        else:
            vals = [reg.comparison_matrix.values[dim][sid] for sid in supplier_ids_ordered]
        cmp_rows.append([dim] + vals)
    _add_table(doc, cmp_headers, cmp_rows)
    for dim in COMPARISON_DIMENSIONS_ALL:
        _add_body_paragraph(doc, f"{dim}: {reg.comparison_matrix.analysis[dim]}")

    # --- Section (6+N): Requirements Coverage Heatmap -----------------------
    _add_heading(doc, f"{cc['Requirements Coverage Heatmap']}. Requirements Coverage Heatmap", level=1, toc=True)
    _add_body_paragraph(doc, reg.coverage_heatmap.intro_paragraph)
    heat_headers = ["Category", "Reqs"]
    for sid in supplier_ids_ordered:
        heat_headers += [f"{supplier_name[sid]} FM", f"{supplier_name[sid]} %"]
    heat_headers.append("Leader")
    heat_rows = []
    for cat in reg.coverage_heatmap.categories:
        row = [cat.category, cat.requirement_count]
        for sid in supplier_ids_ordered:
            row += [cat.per_supplier[sid]["fully_meets"], f"{gt.category_fm_pct[cat.category][sid]:.0f}%"]
        row.append(supplier_name[gt.category_leader[cat.category]])
        heat_rows.append(row)
    _add_table(doc, heat_headers, heat_rows)
    for p in reg.coverage_heatmap.analysis_paragraphs:
        _add_body_paragraph(doc, p)

    # --- Section (7+N): Weighted Scoring Matrix ------------------------------
    _add_heading(doc, f"{cc['Weighted Scoring Matrix']}. Weighted Scoring Matrix", level=1, toc=True)
    _add_body_paragraph(
        doc,
        "The weighted-sum arithmetic behind this matrix is computed by calling weighted_score() in the "
        "vendored numeric_kernel.py over each supplier's scored (non-PENDING) dimensions, with those "
        "dimensions' weights renormalized to sum to 1.0; a PENDING dimension is excluded from a supplier's "
        "denominator rather than counted as a real zero, and the resulting missing weight fraction is shown "
        "alongside every total. This is a PROPOSED total per this skill's own Rule 5 and Boundaries, not "
        "evaluation-engine's official figure.",
    )
    ws_headers = ["Category", "Subcategory", "Weight"] + [supplier_name[sid] for sid in supplier_ids_ordered]
    ws_rows = [
        [d.category, d.subcategory, _pct(d.weight_frac)] + [_score1(d.scores[sid]) for sid in supplier_ids_ordered]
        for d in reg.weighted_scoring.dimensions
    ]
    total_weight_all = sum(d.weight_frac for d in reg.weighted_scoring.dimensions)
    ws_rows.append(["TOTAL", "", _pct(total_weight_all)] + [_total_or_label(gt.weighted_total[sid]) for sid in supplier_ids_ordered])
    ws_rows.append(["Missing weight (PENDING dimensions)", "", ""] + [_pct(gt.weighted_missing_frac[sid]) for sid in supplier_ids_ordered])
    _add_table(doc, ws_headers, ws_rows)

    if full_mode:
        _add_heading(doc, "Per-Dimension Scoring Rationale", level=3, toc=False)
        for d in reg.weighted_scoring.dimensions:
            _add_body_paragraph(doc, f"{d.category} / {d.subcategory}:", bold=True)
            _add_bullet_list(doc, [f"{supplier_name[sid]}: {d.rationale[sid]}" for sid in supplier_ids_ordered])

    _add_heading(doc, "Sensitivity Analysis", level=3, toc=False)
    if full_mode:
        _add_bullet_list(doc, reg.weighted_scoring.sensitivity_notes)
    else:
        _add_body_paragraph(doc, reg.weighted_scoring.sensitivity_notes[0])
        if len(reg.weighted_scoring.sensitivity_notes) > 1:
            _add_body_paragraph(doc, "See the Full report for the complete sensitivity analysis.", italic=True, size_pt=9)

    if full_mode and reg.weighted_scoring.improvement_suggestions:
        _add_heading(doc, "Scoring Matrix Improvement Suggestions", level=3, toc=False)
        _add_bullet_list(doc, reg.weighted_scoring.improvement_suggestions)

    # --- Section (8+N): Commercial & Pricing Analysis ------------------------
    _add_heading(doc, f"{cc['Commercial & Pricing Analysis']}. Commercial & Pricing Analysis", level=1, toc=True)
    cp_headers = ["Dimension"] + [supplier_name[sid] for sid in supplier_ids_ordered]
    cp_rows = [[row.dimension] + [row.values[sid] for sid in supplier_ids_ordered] for row in reg.commercial_pricing.rows]
    _add_table(doc, cp_headers, cp_rows)
    for sid in supplier_ids_ordered:
        _add_heading(doc, supplier_name[sid], level=3, toc=False)
        for p in reg.commercial_pricing.per_supplier_analysis[sid]:
            _add_body_paragraph(doc, p)
    _add_heading(doc, "Normalization Recommendation", level=3, toc=False)
    _add_body_paragraph(doc, reg.commercial_pricing.normalization_recommendation)

    # --- Section (9+N): MSA / Legal Risk Assessment --------------------------
    _add_heading(doc, f"{cc['MSA / Legal Risk Assessment']}. MSA / Legal Risk Assessment", level=1, toc=True)
    lr_headers = ["Clause"] + [supplier_name[sid] for sid in supplier_ids_ordered]
    lr_rows, lr_colors = [], {}
    for r, row in enumerate(reg.legal_risk.rows):
        vals = [row.values[sid] for sid in supplier_ids_ordered]
        lr_rows.append([row.clause] + vals)
        if row.clause.strip().lower() == "risk level":
            for c, sid in enumerate(supplier_ids_ordered, start=1):
                color = RISK_LEVEL_COLOR.get(row.values[sid])
                if color:
                    lr_colors[(r, c)] = color
    _add_table(doc, lr_headers, lr_rows, cell_colors=lr_colors)
    for sid in supplier_ids_ordered:
        _add_body_paragraph(doc, f"{supplier_name[sid]}: {reg.legal_risk.per_supplier_analysis[sid]}")

    # --- Section (10+N): Inconsistency Register -------------------------------
    _add_heading(doc, f"{cc['Inconsistency Register']}. Inconsistency Register", level=1, toc=True)
    if reg.inconsistency_register:
        inc_rows, inc_colors = [], {}
        for r, item in enumerate(reg.inconsistency_register):
            inc_rows.append([item.inconsistency_id, supplier_name[item.supplier_id], item.severity, item.description, item.action_required])
            inc_colors[(r, 2)] = SEVERITY_COLOR.get(item.severity)
        _add_table(doc, ["ID", "Supplier", "Severity", "Description", "Action Required"], inc_rows, cell_colors=inc_colors)
    else:
        _add_body_paragraph(doc, "No inconsistencies were identified across the evaluated supplier responses.")

    # --- Section (11+N): Clarification Questions -------------------------------
    _add_heading(doc, f"{cc['Clarification Questions']}. Clarification Questions", level=1, toc=True)
    if reg.clarification_questions:
        clar_rows, clar_colors = [], {}
        for r, q in enumerate(reg.clarification_questions):
            clar_rows.append([q.question_id, supplier_name[q.supplier_id], q.source_type, q.priority, q.description, q.recommended_format])
            clar_colors[(r, 3)] = CLARIFICATION_PRIORITY_COLOR.get(q.priority)
        _add_table(doc, ["ID", "Supplier", "Source Type", "Priority", "Description", "Recommended Format"], clar_rows, cell_colors=clar_colors)
        gating = [q for q in reg.clarification_questions if q.priority == "GATING"]
        if gating:
            _add_body_paragraph(doc, f"{len(gating)} question(s) are GATING and must be resolved before final evaluation.", bold=True)
    else:
        _add_body_paragraph(doc, "No clarification questions were logged for this evaluation.")

    # --- Section (12+N): Final Recommendation (Proposed) ------------------------
    _add_heading(doc, f"{cc['Final Recommendation (Proposed)']}. Final Recommendation (Proposed)", level=1, toc=True)
    fr = reg.final_recommendation
    _add_heading(doc, "Primary Recommendation", level=3, toc=False)
    _add_body_paragraph(
        doc,
        f"{supplier_name[fr.primary.supplier_id]} (proposed weighted total "
        f"{_total_or_label(gt.weighted_total[fr.primary.supplier_id])} of 5.0).",
        bold=True,
    )
    _add_numbered_list(doc, fr.primary.evidence)
    if fr.primary.conditions:
        _add_body_paragraph(doc, "Conditions:", bold=True)
        _add_bullet_list(doc, fr.primary.conditions)
    _add_body_paragraph(doc, "Caveats:", bold=True)
    _add_bullet_list(doc, fr.primary.caveats)

    if fr.secondary:
        _add_heading(doc, "Secondary Recommendation", level=3, toc=False)
        _add_body_paragraph(
            doc,
            f"{supplier_name[fr.secondary.supplier_id]} (proposed weighted total "
            f"{_total_or_label(gt.weighted_total[fr.secondary.supplier_id])} of 5.0). "
            f"Triggering scenario: {fr.secondary.triggering_scenario}",
        )
        _add_body_paragraph(doc, f"Sensitivity: {fr.secondary.sensitivity_analysis}")

    if fr.conditional:
        _add_heading(doc, "Conditional Recommendations", level=3, toc=False)
        for c in fr.conditional:
            _add_body_paragraph(doc, f"{supplier_name[c.supplier_id]}:", bold=True)
            _add_bullet_list(doc, c.prerequisites)

    if fr.not_recommended:
        _add_heading(doc, "Not Recommended", level=3, toc=False)
        for nr in fr.not_recommended:
            _add_body_paragraph(doc, f"{supplier_name[nr.supplier_id]}: {nr.rationale}")

    _add_heading(doc, "Standard Caveats", level=3, toc=False)
    _add_bullet_list(doc, fr.standard_caveats)
    _add_body_paragraph(
        doc,
        "Subject to evaluation-engine's official scoring and award decision. This section is a preliminary, "
        "AI-proposed recommendation per this skill's Boundaries and Recommendation Rules; it does not replace "
        "evaluation-engine's own scoring authority.",
        italic=True, size_pt=9,
    )

    return doc


# ===========================================================================
# Full pipeline: validate -> compute ground truth -> hard-invariant checks
# -> build document -> save
# ===========================================================================

def generate_rfp_analysis_report(raw_register: Dict[str, Any], output_path: str,
                                  mode_override: Optional[str] = None) -> RfpAnalysisInput:
    """End-to-end: validate the raw RFP analysis register, compute ground
    truth via numeric_kernel, run the hard-coded invariant checks (including
    the named Bid Leveling gate), build the document, and only then save it.
    Raises rather than saving a document that fails validation or
    reconciliation."""
    register = copy.deepcopy(raw_register)
    if mode_override is not None:
        register["report_mode"] = mode_override
    reg = validate_rfp_analysis_input(register)
    gt = compute_ground_truth(reg)
    run_hardcoded_invariant_checks(reg, gt)
    doc = build_document(reg, gt)
    doc.save(output_path)
    return reg


# ===========================================================================
# Demo data (clearly illustrative; SKILL.md carries no numeric worked example
# for rfp-response-analysis, unlike should-cost-builder's quoted "60+30+10=100"
# illustration, so this dataset is authored for this generator, not quoted
# verbatim from any skill file. Deliberately exercises: a supplier that has
# NOT submitted pricing (SUP_B, PENDING throughout, the normal allowed state
# per SKILL.md's non-fabrication rule), the Bid Leveling gate passing (only
# SUP_A must be fully leveled), PENDING weighted-scoring dimensions with a
# non-zero missing-weight fraction, a Moderate and a Critical inconsistency,
# a GATING Bid_Leveling_Gap clarification question reconciling to the Bid
# Leveling worksheet, and a Conditional (not Secondary) recommendation.)
# ===========================================================================

def _demo_rfp_analysis_register() -> Dict[str, Any]:
    """ILLUSTRATIVE DEMO DATA, not a real RFx. Two fictional suppliers
    responding to an illustrative Contract Lifecycle Management (CLM)
    platform RFP."""
    weighted_scoring_dims = [
        {"category": "Requirements Fit", "subcategory": "Functional Alignment", "weight_frac": 0.20,
         "scores": {"SUP_A": 4.6, "SUP_B": 3.4},
         "rationale": {"SUP_A": "Covers every core CLM workflow natively; strongest functional coverage in the field.",
                       "SUP_B": "Covers core authoring and approvals; several advanced clause-library requirements only partially addressed."}},
        {"category": "Requirements Fit", "subcategory": "Technical Alignment", "weight_frac": 0.10,
         "scores": {"SUP_A": 4.2, "SUP_B": 3.6},
         "rationale": {"SUP_A": "REST API and SSO confirmed against Lilly's stated integration list.",
                       "SUP_B": "API available but integration detail is thinner than Ashcroft's response."}},
        {"category": "Financial Stability", "subcategory": "Revenue & Growth", "weight_frac": 0.05,
         "scores": {"SUP_A": 4.0, "SUP_B": 4.5},
         "rationale": {"SUP_A": "Stable, profitable, privately held; smaller revenue base than Meridian.",
                       "SUP_B": "Larger, publicly listed, stronger growth trajectory."}},
        {"category": "Financial Stability", "subcategory": "Credit & Solvency", "weight_frac": 0.05,
         "scores": {"SUP_A": 4.0, "SUP_B": 4.8},
         "rationale": {"SUP_A": "No solvency concerns identified in the financial statements provided.",
                       "SUP_B": "Publicly reported balance sheet shows strong solvency ratios."}},
        {"category": "Risk Posture", "subcategory": "Legal Risk", "weight_frac": 0.05,
         "scores": {"SUP_A": 4.3, "SUP_B": 3.0},
         "rationale": {"SUP_A": "Collaborative redline posture, minimal Hard Stop conflicts.",
                       "SUP_B": "Standard posture; MSA redlines not yet reviewed given pending pricing."}},
        {"category": "Risk Posture", "subcategory": "Operational Risk", "weight_frac": 0.05,
         "scores": {"SUP_A": 3.8, "SUP_B": 3.5},
         "rationale": {"SUP_A": "New to Lilly, but references indicate comparable pharma-scale deployments.",
                       "SUP_B": "Existing Lilly relationship in a different category reduces onboarding risk."}},
        {"category": "Risk Posture", "subcategory": "Cyber / Compliance", "weight_frac": 0.05,
         "scores": {"SUP_A": 4.0, "SUP_B": 3.8},
         "rationale": {"SUP_A": "SOC 2 Type II claimed in narrative; questionnaire shows Type I only (flagged, see Inconsistency Register).",
                       "SUP_B": "SOC 2 Type II current and evidenced."}},
        {"category": "Pricing", "subcategory": "Clarity & Structure", "weight_frac": 0.05,
         "scores": {"SUP_A": 4.5, "SUP_B": None},
         "rationale": {"SUP_A": "Fully itemized pricing sheet with per-user rate, implementation, and term stated.",
                       "SUP_B": "PENDING, no pricing template submitted; scope excluded from the weighted-total denominator."}},
        {"category": "Pricing", "subcategory": "Competitiveness", "weight_frac": 0.05,
         "scores": {"SUP_A": 4.0, "SUP_B": None},
         "rationale": {"SUP_A": "Normalized $1,680 per named user per year, mid-market for this category.",
                       "SUP_B": "PENDING, cannot compare until a normalized figure exists."}},
        {"category": "Pricing", "subcategory": "Transparency", "weight_frac": 0.05,
         "scores": {"SUP_A": 4.2, "SUP_B": None},
         "rationale": {"SUP_A": "No hidden fees identified; assumptions and exclusions clearly itemized.",
                       "SUP_B": "PENDING, no pricing narrative to assess for transparency."}},
        {"category": "Implementation Readiness", "subcategory": "Plan Quality", "weight_frac": 0.05,
         "scores": {"SUP_A": 4.4, "SUP_B": 3.2},
         "rationale": {"SUP_A": "Detailed 3-phase plan with named milestones and a staffing plan.",
                       "SUP_B": "High-level plan only; phases named but durations not itemized."}},
        {"category": "Implementation Readiness", "subcategory": "Timeline Realism", "weight_frac": 0.05,
         "scores": {"SUP_A": 4.0, "SUP_B": 3.0},
         "rationale": {"SUP_A": "20-week timeline is consistent with comparable pharma deployments referenced.",
                       "SUP_B": "16-week timeline stated with no supporting reference deployment of similar scale."}},
        {"category": "Technology Differentiation", "subcategory": "Architecture Fit", "weight_frac": 0.05,
         "scores": {"SUP_A": 4.5, "SUP_B": 3.5},
         "rationale": {"SUP_A": "Single-tenant AWS architecture matches Lilly's stated data-residency preference exactly.",
                       "SUP_B": "Multi-tenant Azure architecture meets residency requirements via a configuration option."}},
        {"category": "Technology Differentiation", "subcategory": "Roadmap Strength", "weight_frac": 0.05,
         "scores": {"SUP_A": 4.0, "SUP_B": 4.0},
         "rationale": {"SUP_A": "Roadmap includes generative clause drafting, currently in beta.",
                       "SUP_B": "Roadmap includes expanded analytics, currently in beta; comparable maturity to Ashcroft."}},
        {"category": "Legal / Contract", "subcategory": "Lock-in Risk", "weight_frac": 0.05,
         "scores": {"SUP_A": 4.0, "SUP_B": 3.5},
         "rationale": {"SUP_A": "Standard data-export rights on termination; no proprietary format lock-in identified.",
                       "SUP_B": "Data export rights stated but export format not yet confirmed."}},
        {"category": "Legal / Contract", "subcategory": "MSA Deviation", "weight_frac": 0.05,
         "scores": {"SUP_A": 4.5, "SUP_B": 3.0},
         "rationale": {"SUP_A": "Accepted Lilly's standard MSA with only two minor redlines.",
                       "SUP_B": "MSA redlines not yet reviewed; posture assessed as Standard pending full legal review."}},
    ]

    return {
        "rfx_name": "Contract Lifecycle Management (CLM) Platform RFP (Illustrative Demo)",
        "category": "Contract Lifecycle Management (CLM) Platform",
        "case_id": "DEMO-RFP-ANALYSIS-0001",
        "analysis_date": "2026-07-23",
        "analysis_mode": "C",
        "report_mode": "Full",
        "requirement_count": 240,
        "executive_summary_scope": (
            "This RFP sought a replacement Contract Lifecycle Management platform covering authoring, "
            "approval workflow, clause library management, and reporting across Lilly's global procurement "
            "organization. Two suppliers responded to the 240-requirement matrix: Ashcroft Vantage CLM "
            "(illustrative), a privately held CLM specialist, and Meridian Contract Systems (illustrative), "
            "a larger, publicly listed platform vendor with an existing Lilly relationship in a different "
            "category. Meridian has not yet submitted a completed pricing template; its figures are labeled "
            "PENDING throughout this report per the skill's non-fabrication rule, not estimated or assumed."
        ),
        "executive_summary_overviews": {
            "SUP_A": "Ashcroft Vantage CLM (illustrative) submitted the most complete response of the two suppliers, with the highest requirements coverage and a fully itemized pricing sheet.",
            "SUP_B": "Meridian Contract Systems (illustrative) is a larger, financially stronger incumbent-adjacent vendor whose pricing template has not yet arrived, leaving its commercial position PENDING.",
        },
        "executive_summary_key_findings": [
            "Ashcroft leads on requirements coverage (95.1% Fully Meets) and is the only supplier with a complete, normalized pricing position.",
            "Meridian has not submitted a pricing template; per the Bid Leveling gate, it remains excluded from the normalized pricing comparison and cannot be primary-recommended on cost.",
            "A Moderate inconsistency was flagged in Ashcroft's SOC 2 certification claim (Type II claimed in narrative, Type I evidenced in the security questionnaire).",
            "A Critical inconsistency was flagged in Meridian's pricing narrative, which references a dollar figure that cannot be reconciled to any submitted template.",
            "Both suppliers propose a named-user pricing model on a comparable basis once Meridian's template arrives, so the ranking is expected to remain data-comparable, not structurally different, once leveling completes.",
        ],
        "suppliers": [
            {
                "profile": {
                    "supplier_id": "SUP_A", "supplier_name": "Ashcroft Vantage CLM (illustrative)",
                    "narrative": (
                        "Ashcroft Vantage CLM (illustrative) is a 14-year-old, privately held contract "
                        "lifecycle management specialist headquartered in Austin, Texas. The company has "
                        "grown to 640 employees and reports $210 million in FY2025 revenue, entirely from "
                        "its CLM product line. Ashcroft has not previously held a Lilly contract but lists "
                        "three named pharma or life sciences clients in its response.\n\n"
                        "Ashcroft's submission proposes a single-tenant AWS deployment with US-only data "
                        "residency and a bundled AI-assisted clause-redlining module. The 500-page "
                        "submission covers every RFP section with a completed requirements matrix, a fully "
                        "itemized pricing sheet, and a security questionnaire."
                    ),
                    "data_card": {
                        "Headquarters": "Austin, Texas, USA", "Employees": "640",
                        "Revenue": "$210M (FY2025)", "Years in Category": "14",
                        "Financial Health": "Stable, privately held, profitable",
                    },
                    "pharma_clients": [
                        "MedGen Therapeutics (illustrative)", "NovaCarta Biosciences (illustrative)",
                        "Aurelia Pharma Group (illustrative)",
                    ],
                    "deployment_model": "Single-tenant cloud (AWS), US data residency",
                    "contract_posture": "Collaborative", "lilly_vendor_status": "New to Lilly",
                    "pricing_summary": "Per-named-user annual subscription, $840,000/year at 500 users, 3-year term, no stated escalator.",
                    "pricing_submitted": True, "leveling_status": "Complete",
                },
                "response_analysis": {
                    "submission_volume": "Ashcroft submitted a proposal, a completed requirements matrix, an itemized pricing sheet, a security questionnaire, and MSA redlines.",
                    "submission_volume_pages": 500,
                    "understanding_of_requirements": "The response addresses Lilly's stated pharma-specific requirements directly, including GxP-relevant audit trail language, rather than generic CLM capabilities copy.",
                    "proposed_solution_architecture": "A single-tenant AWS deployment with modular authoring, approvals, clause library, and reporting components, plus a bundled AI-assisted clause-redlining module currently used by two of Ashcroft's named pharma clients.",
                    "implementation_approach": "A 3-phase, 20-week plan (Discovery, Configuration & Data Migration, Go-Live & Hypercare) with a named delivery lead and a staffing plan consistent with Ashcroft's own services team, not a third-party systems integrator.",
                    "integration_strategy": "REST API with documented endpoints for Lilly's existing e-signature and identity provider (SAML 2.0 SSO), validated against the integration list in the RFP instructions.",
                    "references_domain_evidence": "Three named pharma or life sciences clients are cited with tenure ranging from 2 to 6 years; one reference call was completed during evaluation and confirmed the stated go-live timeline was met within two weeks of plan.",
                    "legal_contract_posture": "Collaborative; only two redlines to Lilly's standard MSA (data-return timeline, audit notice period), both Minor per the negotiation-difficulty assessment.",
                    "key_concerns": "The narrative claims current SOC 2 Type II certification, but the completed security questionnaire lists SOC 2 Type I only with Type II 'in progress'; this is logged as a Moderate inconsistency (see Inconsistency Register) and a clarification question has been issued.",
                    "overall_assessment": "Ashcroft is the strongest response on functional fit, implementation readiness, and pricing clarity, with one certification inconsistency to resolve before award.",
                },
                "coverage": {"fully_meets": 190, "partially_meets": 30, "does_not_meet": 15, "not_answered": 5, "total": 240},
                "adequacy_scores": [
                    {"rfp_section": "Cover Letter", "score": 4.5, "notes": "Addressed to the correct stakeholder, referenced the RFP number correctly."},
                    {"rfp_section": "Vendor Profile", "score": 4.5, "notes": "Complete financials and company history provided."},
                    {"rfp_section": "Financial Statements", "score": 4.0, "notes": "Two years of audited statements provided; privately held so disclosure is more limited than Meridian's."},
                    {"rfp_section": "References", "score": 4.0, "notes": "Three references provided; one reference call completed."},
                    {"rfp_section": "Ability to Meet Objectives", "score": 4.5, "notes": "Objectives mapped explicitly to proposed modules."},
                    {"rfp_section": "Functional Requirements Matrix", "score": 4.6, "notes": "95.1% Fully Meets coverage, highest of the two suppliers."},
                    {"rfp_section": "Architecture", "score": 4.5, "notes": "Single-tenant AWS architecture diagram provided with data-flow detail."},
                    {"rfp_section": "Solution & Pricing", "score": 4.5, "notes": "Fully itemized pricing sheet, normalized cleanly to the common comparison basis."},
                    {"rfp_section": "Demo", "score": 4.0, "notes": "Live demo completed, covered 8 of 10 requested workflows."},
                    {"rfp_section": "MSA/Legal", "score": 4.3, "notes": "Collaborative posture, two minor redlines only."},
                    {"rfp_section": "Implementation Plan", "score": 4.4, "notes": "Detailed 3-phase, 20-week plan with named milestones."},
                    {"rfp_section": "Training", "score": 4.0, "notes": "Standard training program included in the base fee."},
                ],
                "strengths": [
                    "Highest requirements coverage of the two suppliers (95.1% Fully Meets).",
                    "Only supplier with a complete, normalized pricing position.",
                    "Bundled AI-assisted clause-redlining module already in production at two pharma clients.",
                ],
                "risks": [
                    "SOC 2 Type II certification claimed in narrative but not yet evidenced in the security questionnaire (Moderate inconsistency).",
                    "New to Lilly; no existing relationship or master data on file.",
                    "Privately held with more limited financial disclosure than Meridian.",
                ],
            },
            {
                "profile": {
                    "supplier_id": "SUP_B", "supplier_name": "Meridian Contract Systems (illustrative)",
                    "narrative": (
                        "Meridian Contract Systems (illustrative) is a 22-year-old, publicly listed "
                        "enterprise software vendor headquartered in Toronto, Ontario, reporting $780 "
                        "million in FY2025 revenue across several product lines including CLM. Meridian "
                        "already holds a Lilly vendor relationship in a different category (clinical data "
                        "management), which shortens onboarding but does not itself establish CLM-specific "
                        "trust.\n\n"
                        "Meridian's submission proposes a multi-tenant Azure deployment with a named-user "
                        "pricing model in principle, but the pricing template itself was not included with "
                        "the 210-page submission; per the RFP instructions this leaves the commercial "
                        "position PENDING rather than assumed."
                    ),
                    "data_card": {
                        "Headquarters": "Toronto, Ontario, Canada", "Employees": "2,300",
                        "Revenue": "$780M (FY2025)", "Years in Category": "22",
                        "Financial Health": "Stable, publicly listed",
                    },
                    "pharma_clients": ["Bridgeway Biopharma (illustrative)"],
                    "deployment_model": "Multi-tenant cloud (Azure), US and EU data residency options",
                    "contract_posture": "Standard",
                    "lilly_vendor_status": "Existing Lilly vendor (different category: clinical data management)",
                    "pricing_summary": "Pricing template not yet submitted; narrative references a per-named-user model but no figures are confirmed.",
                    "pricing_submitted": False, "leveling_status": "Pending Pricing",
                },
                "response_analysis": {
                    "submission_volume": "Meridian submitted a proposal, a partially completed requirements matrix, and a security questionnaire; no pricing template or MSA redlines were included.",
                    "submission_volume_pages": 210,
                    "understanding_of_requirements": "The response addresses most stated requirements but leans on general CLM capabilities language for several pharma-specific items rather than confirming them directly against Lilly's stated context.",
                    "proposed_solution_architecture": "A multi-tenant Azure deployment with configurable data-residency options; core authoring and approvals modules are proposed, with clause-library scope left silent in the narrative.",
                    "implementation_approach": "A 16-week timeline is stated across three named phases, but durations per phase and a staffing plan were not itemized, so realism cannot be fully assessed.",
                    "integration_strategy": "An API is referenced but no endpoint-level integration detail was provided against Lilly's stated integration list.",
                    "references_domain_evidence": "One named pharma client is cited (Bridgeway Biopharma); no reference call was completed during this evaluation window.",
                    "legal_contract_posture": "MSA redlines were not submitted with this response; contract posture is assessed as Standard pending a full legal review once redlines arrive.",
                    "key_concerns": "No pricing template was submitted despite the RFP instructions requesting one; the pricing narrative references a dollar figure that cannot be reconciled to any template, a Critical inconsistency (see Inconsistency Register).",
                    "overall_assessment": "Meridian is financially larger and has an existing Lilly relationship, but its response is materially less complete than Ashcroft's on pricing, integration detail, and implementation planning.",
                },
                "coverage": {"fully_meets": 165, "partially_meets": 45, "does_not_meet": 20, "not_answered": 10, "total": 240},
                "adequacy_scores": [
                    {"rfp_section": "Cover Letter", "score": 4.0, "notes": "Standard cover letter, correctly addressed."},
                    {"rfp_section": "Vendor Profile", "score": 4.5, "notes": "Complete, publicly disclosed financials provided."},
                    {"rfp_section": "Financial Statements", "score": 4.8, "notes": "Full audited public financials, strongest of the two suppliers."},
                    {"rfp_section": "References", "score": 2.5, "notes": "Only one reference provided; no reference call completed."},
                    {"rfp_section": "Ability to Meet Objectives", "score": 3.5, "notes": "Objectives addressed at a general level, less specific than Ashcroft's mapping."},
                    {"rfp_section": "Functional Requirements Matrix", "score": 3.4, "notes": "68.8% Fully Meets coverage; several pharma-specific items partially addressed."},
                    {"rfp_section": "Architecture", "score": 3.5, "notes": "Multi-tenant architecture described narratively; no data-flow diagram provided."},
                    {"rfp_section": "Solution & Pricing", "score": 0.0, "notes": "No pricing template submitted; Information Not Provided per N.3's scoring rubric."},
                    {"rfp_section": "Demo", "score": 3.5, "notes": "Live demo completed, covered 6 of 10 requested workflows."},
                    {"rfp_section": "MSA/Legal", "score": 3.0, "notes": "No redlines submitted; assessed on narrative posture only."},
                    {"rfp_section": "Implementation Plan", "score": 3.2, "notes": "High-level plan only, phase durations not itemized."},
                    {"rfp_section": "Training", "score": 4.0, "notes": "Standard training program described, comparable to Ashcroft's."},
                ],
                "strengths": [
                    "Strongest financial profile of the two suppliers (publicly listed, $780M revenue).",
                    "Existing Lilly vendor relationship shortens onboarding and master-data setup.",
                    "Data residency options include both US and EU, broader than Ashcroft's US-only offering.",
                ],
                "risks": [
                    "No pricing template submitted; commercial position is PENDING and excluded from the normalized comparison.",
                    "Pricing narrative references a figure that cannot be reconciled to any submitted template (Critical inconsistency).",
                    "Lowest reference-check completion of the two suppliers (one named reference, no call completed).",
                ],
            },
        ],
        "bid_leveling": {
            "comparison_basis_text": (
                "$ per named user per year, based on the RFP's stated 500-named-user base scenario; a fully "
                "loaded 3-year normalized total cost of ownership is used as the secondary comparison figure. "
                "No supplier's proposal states a multi-year escalator, so the flat reported-TCO shorthand "
                "(recurring x term + one-time) is used rather than numeric_kernel.escalate()."
            ),
            "scope_compliance": [
                {"scope_line": "Core contract authoring module", "status": {"SUP_A": "Included", "SUP_B": "Included"}},
                {"scope_line": "E-signature integration", "status": {"SUP_A": "Included", "SUP_B": "Additional Cost"}},
                {"scope_line": "Clause library seeding (Lilly standard clauses)", "status": {"SUP_A": "Additional Cost", "SUP_B": "Silent"}},
                {"scope_line": "Implementation & data migration", "status": {"SUP_A": "Included", "SUP_B": "Silent"}},
                {"scope_line": "Ongoing support (Tier 1)", "status": {"SUP_A": "Included", "SUP_B": "Included"}},
            ],
            "assumption_exclusion_register": [
                {"supplier_id": "SUP_A", "row_type": "Assumption", "item": "500 named users assumed for Year 1 pricing.",
                 "cost_impact": "Included", "confidence": "High", "source": "Pricing Sheet p.2"},
                {"supplier_id": "SUP_A", "row_type": "Exclusion", "item": "Data migration from the legacy CLM tool is excluded from the base fee.",
                 "cost_impact": "Additional Cost", "confidence": "High", "source": "SOW Section 4"},
                {"supplier_id": "SUP_B", "row_type": "Exclusion", "item": "Implementation services excluded pending a follow-on statement of work.",
                 "cost_impact": "Unknown", "confidence": "Medium", "source": "Proposal Section 6 (pricing template not yet submitted)"},
            ],
            "normalized_pricing": [
                {"supplier_id": "SUP_A", "scenario": "Base (500 named users)", "comparison_basis_unit": "$ per named user per year",
                 "reported_price": 840000, "one_time_total": 180000, "recurring_total": 840000,
                 "reported_tco": 2700000, "normalized_price_per_unit": 1680, "normalized_tco": 2700000,
                 "term_years": 3, "confidence": "High", "source": "Pricing Sheet p.1-2"},
            ],
            "missing_cost_placeholders": [
                {"scope_line": "Ongoing support beyond Tier 1 (Tier 2/3 escalation)",
                 "note": "Not priced by any supplier; a Lilly should-cost estimate has not yet been developed."},
            ],
            "questions_before_final_evaluation": [
                {"supplier_id": "SUP_B", "question": "Please submit a completed pricing template consistent with the RFP's per-named-user-per-year comparison basis.", "priority": "GATING"},
            ],
            "analysis_paragraphs": [
                "Ashcroft's reported and normalized TCO are identical ($2,700,000 over the 3-year term) because no escalator was stated; the normalization work still mattered at the per-unit level, translating the reported annual fee into a $1,680 per named user per year figure that can be compared once Meridian's template arrives.",
                "Meridian is the largest source of missing-cost risk in this evaluation: its e-signature integration is flagged Additional Cost while clause library seeding and implementation are Silent, meaning none of those costs are reflected anywhere in the comparison yet. The GATING clarification question above must be resolved before Meridian can be ranked, scored on the Pricing dimension, or named in the Award Recommendation.",
            ],
        },
        "comparison_matrix": {
            "values": {
                "Requirements Fit": {"SUP_A": "Strong (95.1% Fully Meets)", "SUP_B": "Moderate (68.8% Fully Meets)"},
                "Financial Health": {"SUP_A": "Stable, profitable, privately held", "SUP_B": "Stable, publicly listed, strong balance sheet"},
                "Risk Level": {"SUP_A": "Low, one Moderate inconsistency flagged", "SUP_B": "Medium, pricing not yet submitted"},
                "Pricing Clarity": {"SUP_A": "Clear, itemized pricing sheet provided", "SUP_B": "PENDING, no pricing submitted"},
                "Pricing Competitiveness": {"SUP_A": "Normalized $1,680 per named user per year", "SUP_B": "PENDING, cannot compare"},
                "Contract Complexity": {"SUP_A": "Standard MSA posture, two minor redlines", "SUP_B": "Not yet assessed, MSA redlines outstanding"},
                "Vendor Status": {"SUP_A": "New to Lilly", "SUP_B": "Existing Lilly vendor (different category)"},
                "Pharma Experience": {"SUP_A": "3 named pharma clients", "SUP_B": "1 named pharma client"},
                "Implementation Readiness": {"SUP_A": "Detailed 3-phase, 20-week plan", "SUP_B": "High-level plan only, durations not itemized"},
                "Technology Differentiation": {"SUP_A": "Bundled AI-assisted clause redlining", "SUP_B": "Standard clause library, no AI features stated"},
            },
            "analysis": {
                "Requirements Fit": "Ashcroft's 95.1% Fully-Meets coverage is materially ahead of Meridian's 68.8%, driven mainly by the Functional Requirements Matrix and pharma-specific audit-trail language Meridian addressed only generically.",
                "Financial Health": "Both suppliers are financially sound; Meridian's public disclosure and larger revenue base give it a marginally stronger financial-stability score, though this does not offset its pricing and coverage gaps.",
                "Risk Level": "Ashcroft carries one Moderate inconsistency (SOC 2 certification claim); Meridian carries a Critical inconsistency (unreconciled pricing figure) and an incomplete legal review, netting a higher overall risk rating.",
                "Pricing Clarity": "Ashcroft's itemized pricing sheet allowed full normalization; Meridian's absence of a pricing template is the single largest gap in this evaluation and the reason it cannot be ranked on cost.",
                "Pricing Competitiveness": "Only Ashcroft's figure is available for comparison; no competitiveness judgment can be made about Meridian until its template arrives.",
                "Contract Complexity": "Ashcroft's two minor redlines suggest a fast negotiation; Meridian's MSA position is simply unknown pending redlines.",
                "Vendor Status": "Meridian's existing Lilly relationship (clinical data management) may ease onboarding logistics but does not carry over CLM-specific trust or master data.",
                "Pharma Experience": "Ashcroft's three named pharma clients, versus Meridian's one, favor Ashcroft on domain evidence, though both reference sets are limited enough to warrant caution either way.",
                "Implementation Readiness": "Ashcroft's named milestones and staffing plan are materially more implementation-ready than Meridian's high-level, non-itemized timeline.",
                "Technology Differentiation": "Ashcroft's bundled AI-assisted clause redlining module, already in production at two pharma clients, is a genuine differentiator Meridian's roadmap does not yet match.",
                "Adequacy Score": "Ashcroft's computed OVERALL adequacy score leads Meridian's across every N.3 section except Financial Statements and References is the one section where the gap reverses.",
            },
        },
        "coverage_heatmap": {
            "intro_paragraph": (
                "Cells show each supplier's Fully-Meets count and percentage within a requirement category "
                "(out of that category's requirement count), all from actual submitted responses, none "
                "inferred. The Leader column names whichever supplier holds the higher Fully-Meets percentage "
                "in that category."
            ),
            "categories": [
                {"category": "Core Contract Authoring", "requirement_count": 80,
                 "per_supplier": {"SUP_A": {"fully_meets": 68, "partially_meets": 8, "does_not_meet": 3, "not_answered": 1},
                                   "SUP_B": {"fully_meets": 58, "partially_meets": 14, "does_not_meet": 6, "not_answered": 2}}},
                {"category": "Workflow & Approvals", "requirement_count": 70,
                 "per_supplier": {"SUP_A": {"fully_meets": 55, "partially_meets": 10, "does_not_meet": 4, "not_answered": 1},
                                   "SUP_B": {"fully_meets": 48, "partially_meets": 15, "does_not_meet": 5, "not_answered": 2}}},
                {"category": "Clause Library & Playbooks", "requirement_count": 50,
                 "per_supplier": {"SUP_A": {"fully_meets": 40, "partially_meets": 7, "does_not_meet": 2, "not_answered": 1},
                                   "SUP_B": {"fully_meets": 35, "partially_meets": 10, "does_not_meet": 4, "not_answered": 1}}},
                {"category": "Reporting & Analytics", "requirement_count": 40,
                 "per_supplier": {"SUP_A": {"fully_meets": 27, "partially_meets": 5, "does_not_meet": 6, "not_answered": 2},
                                   "SUP_B": {"fully_meets": 24, "partially_meets": 6, "does_not_meet": 5, "not_answered": 5}}},
            ],
            "analysis_paragraphs": [
                "Ashcroft leads in all four categories, with the widest margin in Core Contract Authoring (85.0% vs. 72.5% Fully Meets) and the narrowest in Reporting & Analytics (67.5% vs. 60.0%), the weakest category for both suppliers.",
                "Reporting & Analytics is the category at closest parity and the one most likely to be resolved through configuration rather than a fundamental capability gap; Clause Library & Playbooks shows the most pharma-specific differentiation, driven by Ashcroft's Lilly-clause-seeding proposal.",
            ],
        },
        "weighted_scoring": {
            "dimensions": weighted_scoring_dims,
            "sensitivity_notes": [
                "If Meridian's pricing template arrives and scores at or above Ashcroft's Pricing-dimension average (4.23), Meridian's weighted total would rise from 3.56 toward the mid-3.7s, still short of Ashcroft's 4.25 total.",
                "If Ashcroft's SOC 2 Type II inconsistency resolves unfavorably (certification not achieved), its Cyber / Compliance score would drop toward 2.5, reducing its weighted total by roughly 0.075 and narrowing, but not closing, the gap to Meridian.",
            ],
            "improvement_suggestions": [
                "Add an explicit Data Migration Complexity subcategory under Implementation Readiness; both suppliers' migration effort is currently folded into Plan Quality, obscuring a real cost and risk driver.",
                "Separate Pharma-Specific Compliance (GxP audit trail, Part 11 signature requirements) from the general Cyber / Compliance subcategory; the current single score blends two materially different capabilities.",
            ],
        },
        "commercial_pricing": {
            "rows": [
                {"dimension": "Model", "values": {"SUP_A": "Per-named-user annual subscription", "SUP_B": "Proposed per-named-user annual subscription (pricing pending)"}},
                {"dimension": "Annual Fee", "values": {"SUP_A": "$840,000", "SUP_B": "NEEDS_INPUT"}},
                {"dimension": "List Price", "values": {"SUP_A": "$1,680 per named user per year", "SUP_B": "NEEDS_INPUT"}},
                {"dimension": "Discount", "values": {"SUP_A": "0% (list price quoted)", "SUP_B": "NEEDS_INPUT"}},
                {"dimension": "Implementation", "values": {"SUP_A": "$180,000 one-time", "SUP_B": "NEEDS_INPUT"}},
                {"dimension": "Setup", "values": {"SUP_A": "Included in implementation fee", "SUP_B": "NEEDS_INPUT"}},
                {"dimension": "Escalator", "values": {"SUP_A": "None stated for the 3-year term", "SUP_B": "NEEDS_INPUT"}},
                {"dimension": "Scope", "values": {"SUP_A": "Core CLM, e-signature, clause library seeding", "SUP_B": "Core CLM and e-signature proposed; clause library scope silent"}},
                {"dimension": "Term", "values": {"SUP_A": "3 years", "SUP_B": "3 years proposed"}},
                {"dimension": "User Model", "values": {"SUP_A": "Named user, 500 users assumed", "SUP_B": "Named user, count not yet confirmed"}},
                {"dimension": "Add-ons", "values": {"SUP_A": "AI clause-redlining module, bundled at $0", "SUP_B": "NEEDS_INPUT"}},
                {"dimension": "Binding?", "values": {"SUP_A": "Yes, valid 90 days", "SUP_B": "NEEDS_INPUT"}},
            ],
            "per_supplier_analysis": {
                "SUP_A": [
                    "Ashcroft's pricing is fully itemized and internally consistent: the $840,000 annual fee divides cleanly to $1,680 per named user per year at the stated 500-user base, and the $180,000 implementation fee is broken out separately rather than folded into the subscription.",
                    "No escalator is stated for the 3-year term, which the evaluation team should confirm is intentional (a genuinely flat 3-year rate) rather than an omission, since most SaaS agreements of this size carry some annual increase.",
                ],
                "SUP_B": [
                    "Meridian's commercial position cannot be assessed on figures because no pricing template was submitted; the narrative's reference to a per-named-user model is directionally comparable to Ashcroft's structure, which is a mild positive signal for normalization once figures arrive, but it is not a substitute for the template itself.",
                    "The GATING clarification question in the Bid Leveling section must be resolved before any Commercial & Pricing conclusion can be drawn about Meridian.",
                ],
            },
            "normalization_recommendation": "Once Meridian's pricing template arrives, apply the same $-per-named-user-per-year and 3-year normalized-TCO basis used for Ashcroft above (see Bid Leveling, Section 5) before updating this table or the Weighted Scoring Matrix's Pricing dimension.",
        },
        "legal_risk": {
            "rows": [
                {"clause": "MSA Approach", "values": {"SUP_A": "Accepted Lilly's standard MSA with minor redlines", "SUP_B": "Not yet submitted"}},
                {"clause": "Redline Tone", "values": {"SUP_A": "Collaborative", "SUP_B": "Not yet assessed"}},
                {"clause": "Indemnification", "values": {"SUP_A": "Mutual, standard caps", "SUP_B": "Not yet assessed"}},
                {"clause": "Liability Cap", "values": {"SUP_A": "12 months' fees", "SUP_B": "Not yet assessed"}},
                {"clause": "Confidentiality", "values": {"SUP_A": "Standard, 5-year survival", "SUP_B": "Not yet assessed"}},
                {"clause": "Acceptance Testing", "values": {"SUP_A": "30-day acceptance window proposed", "SUP_B": "Not yet assessed"}},
                {"clause": "Source Code Escrow", "values": {"SUP_A": "Available on request, additional cost", "SUP_B": "Not yet assessed"}},
                {"clause": "Accuracy Warranty", "values": {"SUP_A": "Standard warranty accepted", "SUP_B": "Not yet assessed"}},
                {"clause": "Termination", "values": {"SUP_A": "Termination for convenience accepted, 90-day notice", "SUP_B": "Not yet assessed"}},
                {"clause": "Data Privacy", "values": {"SUP_A": "Standard DPA accepted, US processing only", "SUP_B": "Standard DPA referenced, US and EU processing"}},
                {"clause": "IP Ownership", "values": {"SUP_A": "Lilly data ownership confirmed", "SUP_B": "Lilly data ownership referenced in narrative only"}},
                {"clause": "Audit Rights", "values": {"SUP_A": "Annual audit rights accepted", "SUP_B": "Not yet assessed"}},
                {"clause": "Subcontracting", "values": {"SUP_A": "Disclosed, one hosting subcontractor (AWS)", "SUP_B": "Disclosed, one hosting subcontractor (Azure)"}},
                {"clause": "Governing Law", "values": {"SUP_A": "Delaware, accepted Lilly's preferred venue", "SUP_B": "Ontario proposed, not yet negotiated"}},
                {"clause": "Negotiation Difficulty", "values": {"SUP_A": "Low, two minor redlines only", "SUP_B": "Unknown, redlines not yet submitted"}},
                {"clause": "Protection Score (0-100, higher = better)", "values": {"SUP_A": "82", "SUP_B": "Not yet assessed"}},
                {"clause": "Risk Level", "values": {"SUP_A": "Low", "SUP_B": "Medium"}},
                {"clause": "Hard Stop Conflicts", "values": {"SUP_A": "None identified", "SUP_B": "Not yet assessed, redlines outstanding"}},
            ],
            "per_supplier_analysis": {
                "SUP_A": "Ashcroft's legal posture is Collaborative with only two Minor redlines (data-return timeline, audit notice period), both resolvable without escalation; the 82 Protection Score reflects a near-complete acceptance of Lilly's standard MSA. Estimated negotiation timeline: 2-3 weeks.",
                "SUP_B": "Meridian did not submit MSA redlines with this response, so its legal risk profile is Not yet assessed for most rows; the proposed Ontario governing law (versus Lilly's preferred Delaware venue) is worth flagging early, once redlines do arrive, as a likely negotiation point. Estimated negotiation timeline: cannot be assessed until redlines are received.",
            },
        },
        "inconsistency_register": [
            {"inconsistency_id": "INC-001", "supplier_id": "SUP_A", "severity": "Moderate",
             "description": "Proposal narrative states current SOC 2 Type II certification, but the completed security questionnaire lists SOC 2 Type I only, with Type II described as 'in progress'.",
             "action_required": "Request Ashcroft's current SOC 2 Type II report or a dated certification roadmap before award."},
            {"inconsistency_id": "INC-002", "supplier_id": "SUP_B", "severity": "Critical",
             "description": "Pricing narrative references a $780,000 annual fee, but no pricing template was submitted; the figure cannot be reconciled to any structured pricing document.",
             "action_required": "Do not use the narrative figure for any comparison; escalate via the GATING Bid Leveling clarification question and wait for the completed template."},
        ],
        "clarification_questions": [
            {"question_id": "Q-001", "supplier_id": "SUP_B", "source_type": "Bid_Leveling_Gap", "req_id": None,
             "description": "Please submit a completed pricing template consistent with the RFP's per-named-user-per-year comparison basis.",
             "priority": "GATING", "recommended_format": "Completed pricing template (xlsx)"},
            {"question_id": "Q-002", "supplier_id": "SUP_A", "source_type": "Low_Confidence", "req_id": "REQ-045",
             "description": "Please confirm current SOC 2 Type II certification status and provide the report or an audit-completion date.",
             "priority": "MEDIUM", "recommended_format": "Written clarification email with attached certification"},
            {"question_id": "Q-003", "supplier_id": "SUP_B", "source_type": "Missing_Response", "req_id": "REQ-102",
             "description": "No response was provided for the data residency requirement's specific EU processing addendum.",
             "priority": "HIGH", "recommended_format": "Written clarification email"},
        ],
        "final_recommendation": {
            "primary": {
                "supplier_id": "SUP_A",
                "evidence": [
                    "Highest requirements coverage of the two suppliers (95.1% Fully Meets) and the highest computed weighted total (4.25 of 5.0).",
                    "The only supplier with a complete, normalized pricing position, clearing the Bid Leveling gate in full.",
                    "Strongest implementation readiness, with a detailed 3-phase, 20-week plan and a named staffing plan.",
                ],
                "conditions": ["Resolve the SOC 2 Type II certification inconsistency (INC-001) before contract execution."],
                "caveats": [
                    "Subject to MSA approval.",
                    "Pending resolution of the SOC 2 Type II clarification question (Q-002).",
                    "Subject to evaluation-engine's official scoring and award decision.",
                ],
            },
            "secondary": None,
            "conditional": [
                {"supplier_id": "SUP_B", "prerequisites": [
                    "Submit a completed pricing template on the RFP's stated comparison basis.",
                    "Clear the Bid Leveling gate (normalized pricing, one-time/recurring split, and TCO).",
                    "Resolve the unreconciled pricing-narrative inconsistency (INC-002).",
                ]},
            ],
            "not_recommended": [],
            "standard_caveats": [
                "Subject to MSA approval.",
                "Pending pricing detail validation for Meridian Contract Systems.",
                "Pending additional reference calls for both suppliers.",
                "Subject to stakeholder consensus through formal evaluation scoring.",
            ],
        },
    }


# ===========================================================================
# Self-test / CLI
# ===========================================================================

def _run_self_test() -> int:
    import tempfile
    import zipfile

    print("=" * 78)
    print("rfp_analysis_report_generator.py self-test")
    print("=" * 78)
    print(f"numeric_kernel.py available: {KERNEL_AVAILABLE}")
    print(f"python-docx available: {DOCX_AVAILABLE}"
          + (f" (version {DOCX_VERSION})" if DOCX_AVAILABLE else ""))
    print()

    results: List[tuple] = []

    def check(label, condition, detail=""):
        results.append((label, bool(condition), detail))
        status = "PASS" if condition else "FAIL"
        line = f"[{status}] {label}"
        if detail:
            line += f"  ({detail})"
        print(line)

    demo = _demo_rfp_analysis_register()

    # --- Step 0: demo register validates and computes as independently ----
    # hand-derived (recomputed here from the demo's own weighted_scoring
    # dimensions, NOT by calling the kernel, so this is a genuine cross-
    # check of compute_ground_truth()'s kernel-mediated result, mirroring
    # evaluation_report_generator.py's "hand-derived expected values ...
    # independently computed from the demo register's own numbers").
    try:
        reg = validate_rfp_analysis_input(copy.deepcopy(demo))
        check("validate_rfp_analysis_input accepts the demo register", True)
    except Exception as e:
        check("validate_rfp_analysis_input accepts the demo register", False, str(e))
        raise

    try:
        gt = compute_ground_truth(reg)
        check("compute_ground_truth runs via numeric_kernel.weighted_score()", True,
              f"weighted_total={ {sid: (round(v, 4) if v is not None else None) for sid, v in gt.weighted_total.items()} }")
    except Exception as e:
        check("compute_ground_truth runs via numeric_kernel.weighted_score()", False, str(e))
        raise

    def _independent_weighted_total(dims_raw, sid):
        covered = [d for d in dims_raw if d["scores"][sid] is not None]
        covered_weight = sum(d["weight_frac"] for d in covered)
        if not covered:
            return None
        return sum(d["scores"][sid] * d["weight_frac"] for d in covered) / covered_weight

    for sid in ("SUP_A", "SUP_B"):
        expected = _independent_weighted_total(demo["weighted_scoring"]["dimensions"], sid)
        check(
            f"Weighted total for {sid} matches an independently recomputed value (not via the kernel)",
            abs(gt.weighted_total[sid] - expected) < 1e-9,
            f"got {gt.weighted_total[sid]}, expected {expected}",
        )
    check("SUP_A adequacy OVERALL is the mean of its 12 canonical scores",
          abs(gt.adequacy_overall["SUP_A"] - (sum(a["score"] for a in demo["suppliers"][0]["adequacy_scores"]) / 12)) < 1e-9)
    check("SUP_B carries a non-zero missing weight fraction (3 PENDING Pricing dimensions, 15%)",
          abs(gt.weighted_missing_frac["SUP_B"] - 0.15) < 1e-9, f"got {gt.weighted_missing_frac['SUP_B']}")
    check("SUP_A carries zero missing weight (fully scored)", abs(gt.weighted_missing_frac["SUP_A"] - 0.0) < 1e-9)
    check("Weighted rank places SUP_A ahead of SUP_B", gt.weighted_rank == ["SUP_A", "SUP_B"])
    check("Bid Leveling gate PASSES on the demo register (SUP_B stays PENDING, the allowed state)",
          gt.bid_leveling_gate_passed is True, f"failures={gt.bid_leveling_gate_failures}")
    check("Category leader for Core Contract Authoring is SUP_A", gt.category_leader["Core Contract Authoring"] == "SUP_A")
    exp_num = _compute_section_numbers(2)
    check("Section numbering: Executive Summary=3, suppliers=4-5, Bid Leveling=6, cross-cut=7-14",
          exp_num == gt.section_numbers, f"got {gt.section_numbers}")

    try:
        run_hardcoded_invariant_checks(reg, gt)
        check("All hard-coded invariant checks PASS on the demo register", True)
    except ReconciliationError as e:
        check("All hard-coded invariant checks PASS on the demo register", False, str(e))
        raise

    # --- Step 1: validation refusal tests ----------------------------------
    def _expect_validation_error(label, broken_register):
        try:
            validate_rfp_analysis_input(broken_register)
            check(label, False, "did not raise")
        except RfpAnalysisValidationError as e:
            check(label, True, str(e)[:160])

    b1 = copy.deepcopy(demo)
    del b1["suppliers"]
    _expect_validation_error("validate_rfp_analysis_input refuses a register missing 'suppliers'", b1)

    b2 = copy.deepcopy(demo)
    b2["suppliers"][0]["coverage"]["total"] = 999
    _expect_validation_error("validate_rfp_analysis_input refuses a coverage tally that does not foot", b2)

    b3 = copy.deepcopy(demo)
    del b3["suppliers"][0]["adequacy_scores"][0]
    _expect_validation_error("validate_rfp_analysis_input refuses an incomplete N.3 adequacy score set", b3)

    b4 = copy.deepcopy(demo)
    b4["weighted_scoring"]["dimensions"][0]["weight_frac"] = 0.99
    _expect_validation_error("validate_rfp_analysis_input refuses weighted_scoring weights that do not sum to 1.0", b4)

    b5 = copy.deepcopy(demo)
    b5["weighted_scoring"]["dimensions"][7]["scores"]["SUP_B"] = 3.0  # Pricing dim, SUP_B has not submitted pricing
    _expect_validation_error("validate_rfp_analysis_input refuses a fabricated Pricing score for a non-pricing-submitting supplier", b5)

    b6 = copy.deepcopy(demo)
    del b6["bid_leveling"]["scope_compliance"][0]["status"]["SUP_B"]
    _expect_validation_error("validate_rfp_analysis_input refuses a scope-compliance row missing a declared supplier", b6)

    b7 = copy.deepcopy(demo)
    b7["suppliers"][0]["profile"]["contract_posture"] = "Hostile"
    _expect_validation_error("validate_rfp_analysis_input refuses an unknown contract_posture enum", b7)

    b8 = copy.deepcopy(demo)
    b8["suppliers"][1]["profile"]["pricing_submitted"] = True  # now inconsistent with leveling_status='Pending Pricing'
    _expect_validation_error("validate_rfp_analysis_input refuses pricing_submitted=True with leveling_status still 'Pending Pricing'", b8)

    b9 = copy.deepcopy(demo)
    b9["commercial_pricing"]["rows"] = b9["commercial_pricing"]["rows"][:5]
    _expect_validation_error("validate_rfp_analysis_input refuses a commercial_pricing table under the 12-row minimum", b9)

    b10 = copy.deepcopy(demo)
    b10["legal_risk"]["rows"] = b10["legal_risk"]["rows"][:5]
    _expect_validation_error("validate_rfp_analysis_input refuses a legal_risk table under the 15-row minimum", b10)

    # --- Step 2: Bid Leveling gate CORRECTLY REJECTS an unleveled pricing --
    # supplier (the named HARD RULE gate; distinct from a validation-shape
    # error, this is a ReconciliationError raised by compute_ground_truth /
    # run_hardcoded_invariant_checks on otherwise well-shaped input).
    b_gate = copy.deepcopy(demo)
    b_gate["bid_leveling"]["normalized_pricing"] = []  # SUP_A submitted pricing but now has no normalized row
    try:
        reg_gate = validate_rfp_analysis_input(b_gate)
        gt_gate = compute_ground_truth(reg_gate)
        run_hardcoded_invariant_checks(reg_gate, gt_gate)
        check("Bid-Leveling-gate invariant CORRECTLY REJECTS a pricing-submitting supplier with no normalized_pricing row", False, "did not raise")
    except ReconciliationError as e:
        check("Bid-Leveling-gate invariant CORRECTLY REJECTS a pricing-submitting supplier with no normalized_pricing row", True, str(e)[:200])

    b_gate2 = copy.deepcopy(demo)
    b_gate2["clarification_questions"] = [q for q in b_gate2["clarification_questions"] if q["source_type"] != "Bid_Leveling_Gap"]
    try:
        reg_gate2 = validate_rfp_analysis_input(b_gate2)
        gt_gate2 = compute_ground_truth(reg_gate2)
        run_hardcoded_invariant_checks(reg_gate2, gt_gate2)
        check("Bid-Leveling-gate invariant CORRECTLY REJECTS an unmatched leveling question (no Bid_Leveling_Gap clarification)", False, "did not raise")
    except ReconciliationError as e:
        check("Bid-Leveling-gate invariant CORRECTLY REJECTS an unmatched leveling question (no Bid_Leveling_Gap clarification)", True, str(e)[:200])

    # --- Step 3: invariant-catches-a-bug tests (direct assertion calls) ----
    def _expect_reconciliation_error(label, fn):
        try:
            fn()
            check(label, False, "did not raise, but should have")
        except ReconciliationError as e:
            check(label, True, str(e)[:160])

    bad_gt1 = copy.deepcopy(gt)
    bad_gt1.adequacy_overall["SUP_A"] = 7.0
    _expect_reconciliation_error(
        "Adequacy-overall-range invariant CORRECTLY REJECTS an out-of-range OVERALL score",
        lambda: _assert_adequacy_overall_range(bad_gt1),
    )

    bad_gt2 = copy.deepcopy(gt)
    bad_gt2.weighted_total["SUP_A"] = 5.5
    _expect_reconciliation_error(
        "Weighted-total-range invariant CORRECTLY REJECTS an out-of-range weighted total",
        lambda: _assert_weighted_total_range(bad_gt2),
    )

    bad_gt3 = copy.deepcopy(gt)
    bad_gt3.weighted_rank = list(reversed(bad_gt3.weighted_rank))
    _expect_reconciliation_error(
        "Weighted-ranking-sorted invariant CORRECTLY REJECTS a reversed ranking",
        lambda: _assert_weighted_ranking_sorted(bad_gt3),
    )

    bad_gt4 = copy.deepcopy(gt)
    bad_gt4.bid_leveling_gate_passed = False
    bad_gt4.bid_leveling_gate_failures = ["synthetic failure for invariant test"]
    _expect_reconciliation_error(
        "Bid-Leveling-gate invariant CORRECTLY REJECTS a False gate_passed flag",
        lambda: _assert_bid_leveling_gate(bad_gt4),
    )

    # --- Step 4: build + save + reopen the real document --------------------
    if not DOCX_AVAILABLE:
        check("python-docx available to write a real .docx file", False,
              "python-docx is NOT installed in this interpreter; DOCX writing could not be exercised.")
        print()
        print("Cannot proceed past this point without python-docx.")
    else:
        check("python-docx available to write a real .docx file", True, f"version {DOCX_VERSION}")

        tmp_dir = tempfile.gettempdir()
        full_path = os.path.join(tmp_dir, "analysis_summary_selftest_full.docx")
        brief_path = os.path.join(tmp_dir, "analysis_summary_selftest_brief.docx")

        try:
            generate_rfp_analysis_report(demo, full_path, mode_override="Full")
            check("generate_rfp_analysis_report() ran end-to-end (Full mode) without raising", True)
        except Exception as e:
            check("generate_rfp_analysis_report() ran end-to-end (Full mode) without raising", False, str(e))
            raise

        try:
            generate_rfp_analysis_report(demo, brief_path, mode_override="Brief")
            check("generate_rfp_analysis_report() ran end-to-end (Brief mode) without raising", True)
        except Exception as e:
            check("generate_rfp_analysis_report() ran end-to-end (Brief mode) without raising", False, str(e))
            raise

        for label_path, path in (("Full", full_path), ("Brief", brief_path)):
            exists = os.path.exists(path)
            size = os.path.getsize(path) if exists else 0
            check(f"{label_path} DOCX file written to {path}", exists and size > 0, f"size={size} bytes")
            check(f"{label_path} DOCX unzips cleanly (valid OOXML zip container)", zipfile.is_zipfile(path))

        try:
            reopened_full = Document(full_path)
            all_text_full = "\n".join(p.text for p in reopened_full.paragraphs)
            for tbl in reopened_full.tables:
                for row in tbl.rows:
                    all_text_full += "\n" + "\n".join(c.text for c in row.cells)

            expected_top_level_headings = [
                "3. Executive Summary", "4. Ashcroft Vantage CLM (illustrative)",
                "5. Meridian Contract Systems (illustrative)", "6. Bid Leveling & Normalization (Gate)",
                "7. Cross-Supplier Comparison Matrix", "8. Requirements Coverage Heatmap",
                "9. Weighted Scoring Matrix", "10. Commercial & Pricing Analysis",
                "11. MSA / Legal Risk Assessment", "12. Inconsistency Register",
                "13. Clarification Questions", "14. Final Recommendation (Proposed)",
            ]
            check("Re-opened Full DOCX contains all 12 fixed-order section headings, correctly numbered for 2 suppliers",
                  all(h in all_text_full for h in expected_top_level_headings))
            check("Re-opened Full DOCX contains both per-vendor TOC-indexed sub-headings", all_text_full.count("Response Summary & Analysis") == 2
                  and all_text_full.count("RFP Section Adequacy Scores") == 2)
            check("Re-opened Full DOCX shows the SUPPLIER RESPONSE ANALYSIS title", "SUPPLIER RESPONSE ANALYSIS" in all_text_full)
            check("Re-opened Full DOCX shows SUP_A's kernel-computed weighted total (4.25)", "4.25" in all_text_full)
            check("Re-opened Full DOCX shows SUP_B's PENDING Pricing-dimension cells", "PENDING" in all_text_full)
            check("Re-opened Full DOCX carries the Bid Leveling GATE CHECK pass language", "GATE CHECK: Bid Leveling Complete, PASSED" in all_text_full)
            check("Re-opened Full DOCX carries the Critical inconsistency description", "cannot be reconciled to any structured pricing document" in all_text_full)
            check("Re-opened Full DOCX carries the GATING clarification question", "GATING" in all_text_full)
            check("Re-opened Full DOCX includes the Conditional Recommendations subsection", "Conditional Recommendations" in all_text_full)
            check("Re-opened Full DOCX includes the Per-Dimension Scoring Rationale subsection (Full only)", "Per-Dimension Scoring Rationale" in all_text_full)
            check("Re-opened Full DOCX includes the Scoring Matrix Improvement Suggestions subsection (Full only)", "Scoring Matrix Improvement Suggestions" in all_text_full)

            n_tables_full = len(reopened_full.tables)
            check("Re-opened Full DOCX contains multiple tables", n_tables_full >= 10, f"n_tables={n_tables_full}")
        except Exception as e:
            check("Re-opened Full DOCX structural spot-checks", False, str(e))

        try:
            reopened_brief = Document(brief_path)
            all_text_brief = "\n".join(p.text for p in reopened_brief.paragraphs)
            for tbl in reopened_brief.tables:
                for row in tbl.rows:
                    all_text_brief += "\n" + "\n".join(c.text for c in row.cells)

            check("Re-opened Brief DOCX still contains all 12 fixed-order section headings", all(h in all_text_brief for h in expected_top_level_headings))
            check("Re-opened Brief DOCX OMITS the Per-Dimension Scoring Rationale subsection", "Per-Dimension Scoring Rationale" not in all_text_brief)
            check("Re-opened Brief DOCX OMITS the Scoring Matrix Improvement Suggestions subsection", "Scoring Matrix Improvement Suggestions" not in all_text_brief)
            check("Re-opened Brief DOCX OMITS the itemized Assumption & Exclusion Register table", "Cost Impact" not in all_text_brief)
            check("Re-opened Brief DOCX still shows the kernel-computed weighted total (4.25)", "4.25" in all_text_brief)
        except Exception as e:
            check("Re-opened Brief DOCX structural spot-checks", False, str(e))

    print()
    print("=" * 78)
    passed = sum(1 for _, ok, _ in results if ok)
    failed = sum(1 for _, ok, _ in results if not ok)
    total = len(results)
    print(f"SUMMARY: {passed}/{total} passed, {failed}/{total} failed")
    if failed:
        print("FAILED CASES:")
        for label, ok, detail in results:
            if not ok:
                print(f"  - {label}: {detail}")
    print("=" * 78)
    return 1 if failed else 0


def main(argv: Optional[List[str]] = None) -> int:
    """CLI entry point. Usage:
        python rfp_analysis_report_generator.py --input register.json --output analysis_summary.docx [--mode brief|full]
        python rfp_analysis_report_generator.py --demo
        python rfp_analysis_report_generator.py --self-test
        python rfp_analysis_report_generator.py                 (no args -> runs the self-test)
    """
    import argparse

    parser = argparse.ArgumentParser(
        description="Generate analysis_summary.docx from a validated RFP analysis register "
                    "(rfp-response-analysis-1c344a). See the module docstring for the input JSON schema."
    )
    parser.add_argument("--input", "-i", help="Path to a JSON file containing the RFP analysis register.")
    parser.add_argument("--output", "-o", default="analysis_summary.docx",
                         help="Output .docx path (default: analysis_summary.docx in the current directory).")
    parser.add_argument("--mode", choices=["brief", "full"], default=None,
                         help="Override report_mode (Brief or Full). Defaults to the register's own report_mode, or Full.")
    parser.add_argument("--demo", action="store_true",
                         help="Run the built-in self-test suite (generates the illustrative demo register's "
                              "DOCX, reopens it, and asserts every expected section, table, and value).")
    parser.add_argument("--self-test", action="store_true", dest="self_test",
                         help="Alias for --demo.")
    args = parser.parse_args(argv)

    if args.demo or args.self_test or not args.input:
        return _run_self_test()

    import json
    try:
        with open(args.input, "r", encoding="utf-8") as f:
            raw = json.load(f)
    except (OSError, ValueError) as e:
        print(f"ERROR: could not read/parse --input {args.input!r}: {e}", file=sys.stderr)
        return 1

    mode_override = args.mode.capitalize() if args.mode else None
    try:
        generate_rfp_analysis_report(raw, args.output, mode_override=mode_override)
    except ImportError as e:
        print(f"ERROR: {e}", file=sys.stderr)
        print("The RFP analysis register was valid and ground truth computed cleanly; "
              "only the DOCX-writing step could not run because python-docx is missing.", file=sys.stderr)
        return 1
    except (RfpAnalysisValidationError, ReconciliationError) as e:
        print(f"ERROR: {e}", file=sys.stderr)
        return 1

    print(f"Wrote {args.output}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
