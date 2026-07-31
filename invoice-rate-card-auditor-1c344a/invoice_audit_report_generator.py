"""
invoice_audit_report_generator.py
invoice-rate-card-auditor - F5: wire a generator for the outputs.

WHY THIS EXISTS
---------------
Before F5, invoice-rate-card-auditor had a kernel (numeric_kernel.py) and, as of
F4, a single-execution engine (invoice_audit_engine.py) that runs PASS_2 through
PASS_4 in one pass. It had no builder: SKILL.md's Deliverables section describes
a JSON findings ledger with a specific header/findings/category_rollup schema
and a DOCX audit report with a fixed title page and nine numbered sections plus
two appendices, but nothing produced either mechanically. Both were
model-assembled from the engine's per-line output, which is exactly the
drift risk F9 (the suite's generator-coverage sweep) calls out for this skill.

This module takes F4's completed AuditResult object (`invoice_audit_engine.audit()`)
plus the same audit_input used to build it, and:

  1. Serializes the JSON findings ledger in the exact schema SKILL.md's
     "Deliverables" section documents (header block + one row per finding,
     enriched with the fields the ledger schema names: contracted_rate,
     invoiced_rate, escalation_cap_rate, hours_invoiced, hours_approved,
     stated_total, expected_total, basis, recommended_action, kernel_calls).
  2. Builds the DOCX audit report (title page, TOC, sections 01-09, Appendix
     A/B) by reading ONLY the ledger dict built in step 1, never re-deriving
     a figure from the input or from prose. That is what makes the DOCX
     traceable line-for-line back to the ledger: every number in the DOCX
     was read out of the same object that was just written to JSON.

Both steps are mechanical. No model judgment enters between "AuditResult
exists" and "ledger.json and audit_report.docx exist on disk."

F5's fields on `Finding` (contracted_rate, invoiced_rate, escalation_cap_rate,
hours_invoiced, hours_approved, stated_total, expected_total,
recommended_action, invoice_number, line_no, resource, role_billed) were added
to invoice_audit_engine.py's Finding dataclass and populated at every finding
call site specifically so this generator would not need to re-parse the
human-readable `detail` string to recover a number. This is the sense in
which "F5 depends on F4's actual output shape": the shape was completed, not
just consumed.

ROW-COUNT RECONCILIATION (F4 and F5's shared verification criterion)
----------------------------------------------------------------------
"Lines in equals lines verified equals lines in the final ledger, no line
silently dropped or duplicated." build_full_ledger() asserts this directly:
every invoice line must appear EXACTLY ONCE across (findings-with-that-line-id)
union (clear_lines), never both, never neither. A violation raises
ReconciliationError (imported from invoice_audit_engine) rather than writing
a ledger that under- or double-counts.

DOCX-TO-LEDGER TRACEABILITY
----------------------------
generate_docx() takes the ledger dict, not the AuditResult and not the raw
audit_input. There is no code path by which the DOCX can show a figure the
ledger does not already carry, because the DOCX builder has no other source
to read from. _assert_docx_traceable_to_ledger() checks this mechanically
after the document is built: every dollar figure found in the rendered text
must equal (to the cent) some numeric value already present in the ledger
object.

House style: Magazine Report (lilly-brand-assets-1c344a/SKILL.md, "## 1.
Magazine Report"), the style this skill's SKILL.md names explicitly (Operating
Rule 9 / Kernel Wiring section). Palette copied verbatim from the inlined
brand-colors.md values quoted there: Lilly Red #E1251B, Lilly Black #212121,
Bold Blue #0F3A85, Bold Brown #521207, Neutral Stone #E4EBF1, Muted Grey
#8A969E for footer text. Structural conventions (manually-styled headings,
raw-OOXML cell shading, a post-build content scan before save) follow
evaluation-engine-1c344a/evaluation_report_generator.py and
executive-summary-package-1c344a/executive_summary_generator.py, the two
established DOCX generators already in this suite.

No em dashes anywhere in generated text (Operating Rule 7, suite-wide hard
rule), enforced both at construction (no f-string in this module uses one)
and by a post-build scan of the assembled Document, same discipline
executive_summary_generator.py's _assert_no_forbidden_content uses.

Stdlib only, plus python-docx (optional import, degrades to "ledger JSON
only" if unavailable, per this module's own fail-closed rule below) and the
vendored numeric_kernel.py indirectly via invoice_audit_engine.
"""

from __future__ import annotations

import copy
import datetime as _dt
import json
import re
import sys
from typing import Any, Dict, List, Optional, Sequence

from invoice_audit_engine import (
    audit,
    AuditResult,
    ReconciliationError,
    KernelUnavailableError,
    BlockingAmbiguityError,
    CATEGORY_RATE,
    CATEGORY_ROLE,
    CATEGORY_ESCALATION,
    CATEGORY_HOURS,
    CATEGORY_DUPLICATE,
    CATEGORY_UNSUPPORTED,
    CATEGORY_MILESTONE,
)

try:
    import docx  # noqa: F401
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import RGBColor
    DOCX_AVAILABLE = True
    _DOCX_IMPORT_ERROR: Optional[Exception] = None
except Exception as _exc:  # pragma: no cover
    DOCX_AVAILABLE = False
    _DOCX_IMPORT_ERROR = _exc


class GeneratorError(Exception):
    """Base class for refusals raised by this generator."""


class DocxUnavailableError(GeneratorError):
    """python-docx is not importable. Per G11's fail-closed pattern this
    module raises rather than silently skipping the DOCX and reporting
    success on the JSON ledger alone."""


EM_DASH = "—"

# House style palette, copied verbatim from lilly-brand-assets-1c344a/SKILL.md
# "## 1. Magazine Report", the style this skill's own SKILL.md names.
LILLY_RED = "E1251B"
LILLY_BLACK = "212121"
BOLD_BLUE = "0F3A85"
BOLD_BROWN = "521207"
NEUTRAL_STONE = "E4EBF1"
MUTED_GREY = "8A969E"
WHITE = "FFFFFF"

# SKILL.md's documented ledger schema uses SCREAMING_SNAKE_CASE category codes;
# the engine's internal CATEGORY_* constants are the human-readable strings
# used for the severity table and category_rollup keys. This maps one to the
# other for the ledger and DOCX only; it never changes the engine's own
# vocabulary (invoice_audit_selftest.py's assertions are unaffected).
_CATEGORY_ENUM = {
    CATEGORY_RATE: "RATE_MISMATCH",
    CATEGORY_ROLE: "ROLE_LEVEL_MISMATCH",
    CATEGORY_ESCALATION: "ESCALATION_CAP_BREACH",
    CATEGORY_HOURS: "HOURS_QUANTITY_DISCREPANCY",
    CATEGORY_DUPLICATE: "DUPLICATE_UNSUPPORTED",
    CATEGORY_UNSUPPORTED: "DUPLICATE_UNSUPPORTED",
    CATEGORY_MILESTONE: "MILESTONE_PAYMENT_MISMATCH",
}

_SECTION_CATEGORY_ORDER = [
    ("03 Rate & Line-Math Findings", "RATE_MISMATCH"),
    ("04 Role & Level Findings", "ROLE_LEVEL_MISMATCH"),
    ("05 Escalation Cap Findings", "ESCALATION_CAP_BREACH"),
    ("06 Hours & Quantity Findings", "HOURS_QUANTITY_DISCREPANCY"),
    ("07 Duplicate & Unsupported Charges", "DUPLICATE_UNSUPPORTED"),
    ("08 Milestone & PO/NTE Findings", "MILESTONE_PAYMENT_MISMATCH"),
]


# ===========================================================================
# 1. Ledger builder (F5, JSON serialization directly from F4's output object)
# ===========================================================================

def build_full_ledger(inp: Dict[str, Any], res: AuditResult,
                       header: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    """Build the documented Findings Ledger schema (SKILL.md "Deliverables",
    item 3) from the engine's own AuditResult, never re-authored by hand.

    `header` supplies the descriptive facts the engine's flattened
    audit_input does not itself carry (supplier name, contract/rate-card
    reference, audit period): these are metadata about the engagement, not
    computed figures, so they are accepted as an optional override rather
    than invented. Per Rule 3 (never fabricate), any field not supplied is
    labeled "NOT PROVIDED" rather than guessed.

    Raises ReconciliationError if the row-count reconciliation fails: every
    invoice line must appear in EXACTLY ONE of (a finding) or (clear_lines),
    never both, never neither. This is F4/F5's shared verification criterion
    ("lines in equals lines verified equals lines in the final ledger").
    """
    header = dict(header or {})
    po = inp.get("po") or {}
    invoice_lines_by_id = {ln["line_id"]: ln for ln in inp["invoice_lines"]}
    invoice_line_ids = set(invoice_lines_by_id)

    finding_line_ids = {f.line_id for f in res.findings if f.line_id in invoice_line_ids}
    clear_ids = set(res.clear_lines)

    overlap = finding_line_ids & clear_ids
    covered = finding_line_ids | clear_ids
    missing = invoice_line_ids - covered
    unexpected = covered - invoice_line_ids
    if overlap or missing or unexpected or len(covered) != len(invoice_line_ids):
        raise ReconciliationError(
            "F5 row-count reconciliation FAILED: every invoice line must appear "
            "exactly once as either a finding line or a clear line. "
            f"lines_in={len(invoice_line_ids)}, covered={len(covered)}, "
            f"overlap(both finding-and-clear)={sorted(overlap)}, "
            f"missing(neither)={sorted(missing)}, "
            f"unexpected(not an input line)={sorted(unexpected)}"
        )

    cumulative_invoiced = None
    if po:
        cumulative_invoiced = round(
            sum(float(l["stated_total"]) for l in inp["invoice_lines"])
            + float(po.get("previously_invoiced", 0) or 0), 2)

    header_block = {
        "supplier": header.get("supplier", "NOT PROVIDED"),
        "contract_reference": header.get("contract_reference", "NOT PROVIDED"),
        "rate_card_reference": header.get("rate_card_reference", "NOT PROVIDED"),
        "po_number": po.get("po_number", "NO_PO_SUPPLIED") if po else "NO_PO_SUPPLIED",
        "audit_period": header.get("audit_period", "NOT PROVIDED"),
        "invoices_audited": sorted({ln["invoice_number"] for ln in inp["invoice_lines"]}),
        "lines_audited": res.reconciliation["lines_in"],
        "exceptions_found": len(res.findings),
        "total_questioned_amount": res.rollup["total_questioned_amount"],
        "confirmed_credit_total": res.rollup["confirmed_potential_credit"],
        "pending_total": res.rollup["pending_supplier_response"],
        "po_nte": po.get("nte") if po else None,
        "cumulative_invoiced": cumulative_invoiced,
        "as_of_date": header.get("as_of_date") or _dt.date.today().isoformat(),
    }

    findings_out: List[Dict[str, Any]] = []
    for i, f in enumerate(res.findings, start=1):
        findings_out.append({
            "finding_id": f"F-{i:04d}",
            "line_id": f.line_id,
            "is_invoice_line": f.line_id in invoice_line_ids,
            "invoice_number": f.invoice_number,
            "line_no": f.line_no,
            "resource": f.resource,
            "role_billed": f.role_billed,
            "category": _CATEGORY_ENUM.get(f.category, f.category),
            "check_type": f.check_type,
            "severity": f.severity,
            "resolution_status": f.resolution_status,
            "contracted_rate": f.contracted_rate,
            "invoiced_rate": f.invoiced_rate,
            "escalation_cap_rate": f.escalation_cap_rate,
            "hours_invoiced": f.hours_invoiced,
            "hours_approved": f.hours_approved,
            "stated_total": f.stated_total,
            "expected_total": f.expected_total,
            "questioned_amount": f.questioned_amount,
            "basis": f.detail,
            "recommended_action": f.recommended_action,
            "kernel_calls": [f.kernel_call] if f.kernel_call else [],
        })

    # Clear lines are enriched with identifying fields (not just the bare
    # line_id) so Appendix A can render "every audited line, exceptions and
    # clean" (SKILL.md Deliverables item 1) without the DOCX builder reaching
    # back into audit_input; it reads only this ledger object.
    clear_out: List[Dict[str, Any]] = []
    for lid in sorted(clear_ids):
        ln = invoice_lines_by_id[lid]
        clear_out.append({
            "line_id": lid,
            "invoice_number": ln["invoice_number"],
            "line_no": ln.get("line_no"),
            "resource": ln.get("resource"),
            "role_billed": ln.get("role_billed"),
            "stated_total": float(ln["stated_total"]),
        })

    ledger = {
        "header": header_block,
        "findings": findings_out,
        "clear_lines": clear_out,
        "needs_model_review": res.needs_model_review,
        "needs_input": res.needs_input,
        "category_rollup": res.rollup["by_category"],
        "rollup": res.rollup,
        "reconciliation": res.reconciliation,
    }
    return ledger


def write_ledger_json(ledger: Dict[str, Any], output_path: str) -> None:
    with open(output_path, "w", encoding="utf-8", newline="\n") as fh:
        json.dump(ledger, fh, indent=2)
        fh.write("\n")


# ===========================================================================
# 2. DOCX builder (python-docx). Reads ONLY the ledger dict, so every figure
#    in the DOCX is, by construction, a figure already present in the ledger.
# ===========================================================================

def _require_docx() -> None:
    if not DOCX_AVAILABLE:
        raise DocxUnavailableError(
            "python-docx is not installed in this Python environment, so the "
            "audit report .docx cannot be written. Per this skill's fail-closed "
            "discipline (G11's rule extended to output generation), this raises "
            "rather than silently skipping the DOCX and reporting success on the "
            "JSON ledger alone. Install it (`pip install python-docx`) or point "
            f"this script at an interpreter that already has it. Original import "
            f"error: {_DOCX_IMPORT_ERROR}"
        )


def _add_horizontal_rule(doc, color_hex: str = LILLY_RED) -> None:
    """A genuine red rule under the title, via a paragraph bottom border.
    python-docx has no high-level API for this; the raw-OOXML recipe is the
    same style of workaround _set_cell_background uses for table shading."""
    para = doc.add_paragraph()
    p = para._p
    p_pr = p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _set_cell_background(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_text(cell, text: str, bold: bool = False, color_hex: str = LILLY_BLACK,
                    size_pt: int = 9) -> None:
    cell.text = str(text)
    para = cell.paragraphs[0]
    run = para.runs[0] if para.runs else para.add_run("")
    run.font.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor.from_string(color_hex)


def _add_heading(doc, text: str, level: int = 1) -> None:
    """Manually-styled heading per docx-design-system.md's Magazine Report
    spec: H1 Calibri 14pt Bold Bold Blue, H2 Calibri 12pt Bold Lilly Black."""
    spec = {1: (14, BOLD_BLUE, 16, 8), 2: (12, LILLY_BLACK, 12, 6)}
    size_pt, color_hex, before_pt, after_pt = spec.get(level, spec[2])
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(before_pt)
    para.paragraph_format.space_after = Pt(after_pt)
    run = para.add_run(text)
    run.font.bold = True
    run.font.size = Pt(size_pt)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor.from_string(color_hex)


def _add_body(doc, text: str, bold: bool = False, italic: bool = False,
              size_pt: int = 10.5, color_hex: str = LILLY_BLACK) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.font.bold = bold
    run.font.italic = italic
    run.font.size = Pt(size_pt)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor.from_string(color_hex)


def _add_table(doc, headers: Sequence[str], rows: Sequence[Sequence[Any]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell_background(hdr_cells[i], LILLY_RED)
        _set_cell_text(hdr_cells[i], h, bold=True, color_hex=WHITE, size_pt=9)
    for row_values in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_values):
            _set_cell_text(cells[i], "" if val is None else str(val), size_pt=9)
    return table


def _money(v: Optional[float]) -> str:
    if v is None:
        return "N/A"
    return f"${v:,.2f}"


def _num(v: Optional[float]) -> str:
    if v is None:
        return ""
    return f"{v:,.2f}"


def _findings_by_category(ledger: Dict[str, Any], category_code: str) -> List[Dict[str, Any]]:
    return [f for f in ledger["findings"] if f["category"] == category_code]


def _finding_rows_for_table(findings: List[Dict[str, Any]]) -> List[Sequence[Any]]:
    return [
        (f["finding_id"], f["invoice_number"] or "", f["line_no"] if f["line_no"] is not None else "",
         f["resource"] or "", f["role_billed"] or "", f["severity"], f["resolution_status"],
         _money(f["questioned_amount"]))
        for f in findings
    ]


def build_document(ledger: Dict[str, Any]):
    """Build the audit_report.docx per SKILL.md Deliverables item 1, reading
    ONLY the ledger dict. Does not save the file. Call
    generate_audit_report_docx() for the full validated pipeline."""
    _require_docx()

    h = ledger["header"]
    doc = Document()

    # Footer, per docx-design-system.md's Footer text spec (Calibri 8pt,
    # Muted Grey).
    footer_para = doc.sections[0].footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        "CONFIDENTIAL, internal use only. Lilly Procurement Skills Suite, "
        "Invoice & Rate-Card Auditor."
    )
    footer_run.font.size = Pt(8)
    footer_run.font.name = "Calibri"
    footer_run.font.color.rgb = RGBColor.from_string(MUTED_GREY)

    # --- Title page ------------------------------------------------------
    title_para = doc.add_paragraph()
    title_para.paragraph_format.space_after = Pt(2)
    title_run = title_para.add_run("INVOICE & RATE-CARD AUDIT")
    title_run.font.bold = True
    title_run.font.size = Pt(26)
    title_run.font.name = "Calibri"
    title_run.font.color.rgb = RGBColor.from_string(LILLY_BLACK)

    subtitle_para = doc.add_paragraph()
    subtitle_para.paragraph_format.space_after = Pt(8)
    subtitle_run = subtitle_para.add_run(f"{h['supplier']}, {h['audit_period']}")
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.name = "Calibri"
    subtitle_run.font.color.rgb = RGBColor.from_string(BOLD_BLUE)

    _add_horizontal_rule(doc, LILLY_RED)

    scope_line = (
        f"{len(h['invoices_audited'])} invoices | {h['lines_audited']} lines audited | "
        f"{h['exceptions_found']} exceptions | {_money(h['total_questioned_amount'])} questioned"
    )
    _add_body(doc, scope_line, bold=True, size_pt=11)

    _add_body(doc, "Prepared by: Invoice & Rate-Card Auditor (Lilly Procurement Skills Suite)",
               size_pt=9.5, color_hex=BOLD_BROWN)
    _add_body(doc, f"As of: {h['as_of_date']}", size_pt=9.5, color_hex=BOLD_BROWN)
    _add_body(doc, "CONFIDENTIAL, INTERNAL USE ONLY", bold=True, size_pt=10, color_hex=LILLY_RED)

    abstract = (
        f"This audit reconciled {len(h['invoices_audited'])} invoice(s) "
        f"({', '.join(h['invoices_audited']) if h['invoices_audited'] else 'none supplied'}) "
        f"against the governing contract and rate card"
        + (f" and PO {h['po_number']}" if h['po_number'] not in (None, "NO_PO_SUPPLIED") else "")
        + f". Of {h['lines_audited']} audited lines, {h['exceptions_found']} exception(s) were found, "
        f"totaling {_money(h['total_questioned_amount'])} questioned "
        f"({_money(h['confirmed_credit_total'])} confirmed, {_money(h['pending_total'])} pending "
        "supplier response). This is a reflect-only, point-in-time audit of the supplied population."
    )
    _add_body(doc, abstract, size_pt=10.5)
    doc.add_page_break()

    # --- Table of contents (rendered list, not a dynamic field; the same
    # simplification the suite's other DOCX generators use) ---------------
    _add_heading(doc, "Table of Contents", level=1)
    toc_items = [
        "01 Executive Summary", "02 Population & Methodology",
        "03 Rate & Line-Math Findings", "04 Role & Level Findings",
        "05 Escalation Cap Findings", "06 Hours & Quantity Findings",
        "07 Duplicate & Unsupported Charges", "08 Milestone & PO/NTE Findings",
        "09 Category Rollup & Recommended Actions",
        "Appendix A: Full Line-Level Ledger", "Appendix B: Assumptions & Data Gaps",
    ]
    for item in toc_items:
        _add_body(doc, item, size_pt=10.5)
    doc.add_page_break()

    # --- 01 Executive Summary --------------------------------------------
    _add_heading(doc, "01 Executive Summary", level=1)
    top_cat = None
    if ledger["category_rollup"]:
        top_cat = max(ledger["category_rollup"].items(), key=lambda kv: kv[1])
    exec_lines = [scope_line]
    if top_cat:
        exec_lines.append(
            f"The largest category by questioned amount is {top_cat[0]} at {_money(top_cat[1])}."
        )
    exec_lines.append(
        f"Confirmed potential credit: {_money(h['confirmed_credit_total'])}. "
        f"Pending supplier response: {_money(h['pending_total'])}."
    )
    for line in exec_lines:
        _add_body(doc, line)

    # --- 02 Population & Methodology --------------------------------------
    _add_heading(doc, "02 Population & Methodology", level=1)
    _add_body(doc, f"Supplier: {h['supplier']}")
    _add_body(doc, f"Contract reference: {h['contract_reference']}")
    _add_body(doc, f"Rate card reference: {h['rate_card_reference']}")
    _add_body(doc, f"PO number: {h['po_number']}")
    _add_body(doc, f"Audit period: {h['audit_period']}")
    _add_body(doc, f"Invoices audited: {', '.join(h['invoices_audited']) or 'none'}")
    _add_body(doc,
        "This is a single-user, reflect-only, point-in-time audit of the supplied "
        "population. It does not write to, or otherwise touch, AP, an ERP, Ariba, or "
        "any system of record, and it does not monitor future invoices."
    )
    _add_body(doc,
        "All rate, line-total, and escalation arithmetic in this report was computed by "
        "calling the vendored numeric_kernel.py (verify_line_math, escalate), never by "
        "model judgment or hand arithmetic (Kernel Wiring, G11)."
    )

    # --- 03 to 08: one section per finding category -----------------------
    for section_title, cat_code in _SECTION_CATEGORY_ORDER:
        _add_heading(doc, section_title, level=1)
        rows = _findings_by_category(ledger, cat_code)
        if rows:
            _add_table(
                doc,
                ["Finding", "Invoice", "Line", "Resource", "Role Billed", "Severity",
                 "Status", "Questioned"],
                _finding_rows_for_table(rows),
            )
            for f in rows:
                _add_body(doc, f"{f['finding_id']}: {f['basis']}", size_pt=9.5,
                           color_hex=MUTED_GREY)
                if f["recommended_action"]:
                    _add_body(doc, f"Recommended action: {f['recommended_action']}",
                               size_pt=9.5, color_hex=BOLD_BROWN)
        else:
            _add_body(doc, "No exceptions found in this category on this run.", italic=True,
                       color_hex=MUTED_GREY)

    # --- 09 Category Rollup & Recommended Actions --------------------------
    _add_heading(doc, "09 Category Rollup & Recommended Actions", level=1)
    if ledger["category_rollup"]:
        _add_table(
            doc, ["Category", "Questioned Amount"],
            [(cat, _money(amt)) for cat, amt in sorted(ledger["category_rollup"].items())],
        )
    _add_body(doc,
        f"Total questioned amount: {_money(h['total_questioned_amount'])} = "
        f"confirmed {_money(h['confirmed_credit_total'])} + pending {_money(h['pending_total'])}.",
        bold=True,
    )
    if h.get("po_nte") is not None:
        _add_body(doc,
            f"PO NTE tracker: cumulative invoiced {_money(h['cumulative_invoiced'])} against "
            f"NTE {_money(h['po_nte'])}."
        )

    # --- Appendix A: full line-level ledger table --------------------------
    doc.add_page_break()
    _add_heading(doc, "Appendix A: Full Line-Level Ledger", level=1)
    exception_line_count = len({f["line_id"] for f in ledger["findings"] if f["is_invoice_line"]})
    _add_body(doc,
        f"Every audited line appears exactly once below, either as an exception or as "
        f"CLEAN: {h['lines_audited']} lines audited = {exception_line_count} line(s) with "
        f"at least one exception (full finding detail below; a line can carry more than "
        f"one finding) + {len(ledger['clear_lines'])} clean line(s) below.",
        size_pt=9.5, color_hex=MUTED_GREY,
    )
    if ledger["clear_lines"]:
        _add_heading(doc, "Clean lines (no exception)", level=2)
        _add_table(
            doc, ["Invoice", "Line", "Resource", "Role Billed", "Stated Total"],
            [(c["invoice_number"], c["line_no"] if c["line_no"] is not None else "",
              c["resource"] or "", c["role_billed"] or "", _money(c["stated_total"]))
             for c in ledger["clear_lines"]],
        )
    _add_heading(doc, "Exception lines (full finding detail)", level=2)
    if ledger["findings"]:
        _add_table(
            doc,
            ["Finding", "Invoice", "Line", "Resource", "Role Billed", "Category", "Severity",
             "Status", "Contracted Rate", "Invoiced Rate", "Escalation Cap", "Hours Invoiced",
             "Hours Approved", "Stated Total", "Expected Total", "Questioned"],
            [
                (f["finding_id"], f["invoice_number"] or "", f["line_no"] if f["line_no"] is not None else "",
                 f["resource"] or "", f["role_billed"] or "", f["category"], f["severity"],
                 f["resolution_status"], _num(f["contracted_rate"]), _num(f["invoiced_rate"]),
                 _num(f["escalation_cap_rate"]), _num(f["hours_invoiced"]), _num(f["hours_approved"]),
                 _money(f["stated_total"]), _money(f["expected_total"]), _money(f["questioned_amount"]))
                for f in ledger["findings"]
            ],
        )

    # --- Appendix B: assumptions and data gaps ------------------------------
    doc.add_page_break()
    _add_heading(doc, "Appendix B: Assumptions & Data Gaps", level=1)
    if ledger["needs_model_review"]:
        _add_heading(doc, "Needs model review (ambiguous matches the code would not guess)", level=2)
        for item in ledger["needs_model_review"]:
            _add_body(doc, f"{item['line_id']}: {item['reason']}", size_pt=9.5)
    if ledger["needs_input"]:
        _add_heading(doc, "Needs additional input (data gap, not a finding)", level=2)
        for item in ledger["needs_input"]:
            _add_body(doc, f"{item['line_id']}: {item['reason']}", size_pt=9.5)
    if not ledger["needs_model_review"] and not ledger["needs_input"]:
        _add_body(doc, "No open assumptions or data gaps on this run.", italic=True,
                   color_hex=MUTED_GREY)

    return doc


_MONEY_RE = re.compile(r"\$([0-9][0-9,]*\.[0-9]{2})")


def _iter_all_text(doc) -> List[str]:
    chunks = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                chunks.append(c.text)
    return chunks


def _assert_no_forbidden_content(doc) -> None:
    """Post-build, pre-save scan (defense in depth, same discipline
    executive_summary_generator.py's _assert_no_forbidden_content uses):
    no em dash anywhere in the rendered text (Operating Rule 7)."""
    all_text = "\n".join(_iter_all_text(doc))
    if EM_DASH in all_text:
        raise ReconciliationError(
            "No-em-dash invariant FAILED: the rendered audit report contains an em "
            "dash character (U+2014); forbidden suite-wide (Operating Rule 7)."
        )


def _ledger_numeric_pool(ledger: Dict[str, Any]) -> set:
    """Every numeric value the ledger itself carries, rounded to the cent.
    Used by _assert_docx_traceable_to_ledger to confirm the DOCX invents
    nothing: any dollar figure rendered must be a member of this set."""
    pool = set()

    def _add(v):
        if isinstance(v, (int, float)) and not isinstance(v, bool):
            pool.add(round(float(v), 2))

    h = ledger["header"]
    for key in ("total_questioned_amount", "confirmed_credit_total", "pending_total",
                "po_nte", "cumulative_invoiced"):
        _add(h.get(key))
    for cat, amt in ledger["category_rollup"].items():
        _add(amt)
    for f in ledger["findings"]:
        for key in ("contracted_rate", "invoiced_rate", "escalation_cap_rate",
                     "stated_total", "expected_total", "questioned_amount"):
            _add(f.get(key))
    for c in ledger["clear_lines"]:
        _add(c.get("stated_total"))
    return pool


def _assert_docx_traceable_to_ledger(doc, ledger: Dict[str, Any]) -> None:
    """F5's own verification criterion: 'no figure in the DOCX that isn't in
    the ledger.' Extracts every dollar-formatted figure from the rendered
    text and confirms each is a value the ledger object itself carries."""
    pool = _ledger_numeric_pool(ledger)
    all_text = "\n".join(_iter_all_text(doc))
    seen = set()
    for m in _MONEY_RE.finditer(all_text):
        value = round(float(m.group(1).replace(",", "")), 2)
        seen.add(value)
        if value not in pool:
            raise ReconciliationError(
                f"DOCX-to-ledger traceability FAILED: the rendered report shows "
                f"${value:,.2f}, which does not match any figure in the ledger "
                "object (header totals, category_rollup, or any finding's "
                "contracted_rate/invoiced_rate/escalation_cap_rate/stated_total/"
                "expected_total/questioned_amount). The DOCX must be built only "
                "from the ledger, never from a separately computed number."
            )


# ===========================================================================
# 3. Full pipeline: audit -> ledger -> DOCX, all from one AuditResult object
# ===========================================================================

def generate_audit_outputs(audit_input: Dict[str, Any], ledger_path: str,
                            docx_path: Optional[str] = None,
                            header: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    """End-to-end: run the F4 engine once, build the F5 ledger from its
    output object, write the JSON ledger, and (if docx_path is given) build
    and save the DOCX audit report from that same ledger object. Raises
    rather than writing partial or unreconciled output."""
    res = audit(audit_input)
    ledger = build_full_ledger(audit_input, res, header=header)
    write_ledger_json(ledger, ledger_path)
    if docx_path is not None:
        doc = build_document(ledger)
        _assert_no_forbidden_content(doc)
        _assert_docx_traceable_to_ledger(doc, ledger)
        doc.save(docx_path)
    return ledger


def main(argv):
    if "--selftest" in argv:
        from invoice_audit_report_generator_selftest import run_selftest
        return run_selftest()
    if not argv:
        print(__doc__)
        print("usage: invoice_audit_report_generator.py <audit_input.json> "
              "--ledger findings_ledger.json [--docx audit_report.docx] "
              "[--header header.json]", file=sys.stderr)
        return 2
    with open(argv[0], encoding="utf-8") as fh:
        inp = json.load(fh)
    ledger_path = None
    docx_path = None
    header = None
    if "--ledger" in argv:
        ledger_path = argv[argv.index("--ledger") + 1]
    if "--docx" in argv:
        docx_path = argv[argv.index("--docx") + 1]
    if "--header" in argv:
        with open(argv[argv.index("--header") + 1], encoding="utf-8") as fh:
            header = json.load(fh)
    if ledger_path is None:
        print("error: --ledger <path> is required", file=sys.stderr)
        return 2
    generate_audit_outputs(inp, ledger_path, docx_path, header=header)
    print(f"wrote {ledger_path}" + (f" and {docx_path}" if docx_path else ""))
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv[1:]))
