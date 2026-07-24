"""
executive_summary_generator.py
Lilly Procurement Skills, executive-summary-package-1c344a deterministic DOCX
generator.

Purpose (mirrors should-cost-builder-1c344a/should_cost_generator.py and
evaluation-engine-1c344a/evaluation_report_generator.py's "generator
scripts, not freehand authoring" pattern): this module takes a validated
executive-summary register (deal facts, pre-composed narrative sections,
financial figures, and approval-chain facts, shaped per this skill's own
"Document Structure" and "Default Document Structure" sections of SKILL.md)
and MECHANICALLY produces the Lilly-approved ATC/ATS Executive Summary
executive_summary.docx: Title, Source Documents Line, Opening Paragraph,
Request Overview (conditional), Background & Justification, Key
Capabilities and Value (conditional), Scope of Agreement, Financial
Summary, a Vendor-Specific section (conditional), Key Contract Benefits,
Key Contract Risks / Considerations, Business Case (conditional), Cost
Efficiency (conditional, OMITTED ENTIRELY when there are no savings, never
a placeholder sentence), Governance & Approvals, and the Data Basis &
Confirmations footer, in exactly that fixed order every run, per SKILL.md's
"Section Order (strict, follow this sequence exactly)".

Two-stage discipline (same as should-cost/evaluation-engine):
  1. The ATC and ATS approval-chain grade sequences are computed by calling
     the vendored frap_chain_kernel.py's compute_chain() exactly once per
     chain (never hand-derived or re-implemented in this module), per
     SKILL.md's "Kernel wiring" section: "Do not hand-recompute the chain
     in prose and do not hand-roll a second copy of the algorithm, call the
     kernel." A kernel refusal (needs_review=True, for example a missing
     deal value, an unrecognized grade, or the CEO capital-vs-operating
     band needing disambiguation) is never overridden or guessed past; this
     module raises rather than fabricating a chain.
  2. Approval-chain NAMES are never computed or invented. Per SKILL.md:
     "Do NOT hardcode any names. Always ask the user to confirm the names
     at each required level. The grade levels are computed; the names are
     user-confirmed." This module accepts an optional name-per-grade map
     for each chain and renders "[To be confirmed]" for any resolved grade
     that has no supplied name, per metadata-fields.md's "Handling Missing
     Fields": "Required field not provided, prompt once; if still missing,
     mark '[To be confirmed]'."

Hard-coded invariants (code-level checks, not comments; mirrors should-cost
and evaluation-engine's run_hardcoded_invariant_checks() pattern), split
across two layers exactly like evaluation_report_generator.py splits
weight-sum checks (validation, pure input arithmetic) from
weighted-contributions-reconcile checks (reconciliation, kernel-dependent):

  Validation-layer (validate_executive_summary_input, no kernel call
  needed):
  - Financial-periods-foot-to-total invariant: Financial Summary table-mode
    period figures must sum to the stated total within $1 tolerance,
    generalizing the Markdown Version's "the year figures MUST sum to the
    stated total... an executive summary whose periods do not foot to the
    total is a fail" rule to every rendering of the table, not only the .md
    conversion.
  - Data-basis-counts-foot invariant: verbatim + user-confirmed + inferred
    field counts must sum to the declared total field count, so the Data
    Basis & Confirmations footer's own coverage line cannot silently drift
    from what it claims to cover.
  - No-em-dash invariant (Operating Rule 7, "Never use em dashes... in ANY
    written output"): every string value anywhere in the raw register is
    scanned for U+2014; the register is refused rather than silently
    stripping or rewriting a violation the caller may not have intended.
  - No-savings-placeholder invariant (Gap Handling: "Missing savings, ask
    user, then omit the section entirely, do not write 'No savings
    identified'"; the v2.3-era fix this task references as "task #25"):
    a cost_efficiency line that reads as a no-savings placeholder is
    refused at validation time, forcing the caller to omit the field
    (set it to null/empty) instead of writing a hedge sentence.

  Reconciliation-layer (run_hardcoded_invariant_checks, needs GroundTruth,
  i.e. depends on the two kernel calls):
  - ATC-chain-resolved and ATS-chain-resolved invariants: both
    frap_chain_kernel.compute_chain() calls must return a non-refused
    Decision before any Governance & Approvals content is built.
  - Approver-names-match-chain-length invariant: the resolved name list for
    each chain (real names or "[To be confirmed]" placeholders) must have
    exactly one entry per grade in that chain's kernel-computed outcome.
  - Financial-Risk-Rating-integrity invariant (Rule 3: "The Financial Risk
    Rating in particular must NOT be defaulted to 'Acceptable' or any other
    value... If the SER score and rating are not in the source or the user
    input, the value is 'NEEDS_INPUT (SER score and rating not
    provided)'"): the displayed rating must be either the exact canonical
    NEEDS_INPUT string (when none was supplied) or the caller-supplied
    value verbatim (never silently altered either way).

  Post-build, pre-save layer (operates on the assembled python-docx
  Document, mirrors "the document is not saved when this fires"):
  - No em dash anywhere in the rendered text (defense in depth against a
    formatting bug in this module itself, not just bad input).
  - The literal phrase "No savings identified" never appears anywhere in
    the rendered text (the regression guard for the task #25 fix).
  - The "Cost Efficiency" heading appears if and only if cost_efficiency
    data was supplied (a bijective presence check, not just an absence
    check, so a future bug that drops real savings data silently is also
    caught).
  - Both "ATS Approver(s):" and "ATC Approver:" governance lines are
    present (Section 14 always carries both blocks, regardless of whether
    the run was framed to the user as an "ATC summary" or an "ATS
    summary").

If any invariant fails, this script RAISES rather than writing a document
that fails its own reconciliation.

Scope note: this script builds the executive_summary.docx generator ONLY,
matching the task's explicit "Word (.docx) generator" scope and the same
one-generator-one-artifact discipline should-cost-builder (XLSX only) and
evaluation-engine (DOCX only) already follow. SKILL.md's Output section
also mandates a companion executive_summary.md (text-only, no tables,
"Produce BOTH files every time"); that Markdown deliverable is a SEPARATE
artifact and is NOT produced by this module.

House-style note (JUDGMENT CALL, flagged): should-cost-builder and
evaluation-engine both render the Lilly "colored dashboard" DOCX house
style (Bold Blue #0F3A85 / Lilly Red #E1251B headings, shaded table
headers, a CONFIDENTIAL title-page banner). This skill's own SKILL.md
explicitly forbids that style for THIS document: "NO colored section
banners. NO red headers. NO dashboard-style layouts," "this document uses
no color in its body," and "Governance fields: Bold label + colon + plain
value on same line" (not a table). This module therefore intentionally
does NOT reuse should-cost/evaluation-engine's color constants or heading
style; it implements SKILL.md's own "Plain ATC/ATS house style" instead
(plain black body text throughout, bold-inline section headers at body
size, light gray table headers used only for the one financial table,
and the single permitted accent color, Bold Grey #8A969E, reserved for the
Source Documents line and the Data Basis & Confirmations footer, the same
"Footer text" role docx-design-system.md defines elsewhere in the suite).

============================================================================
Executive-summary register input contract (JUDGMENT CALL, flagged):
SKILL.md's Document Structure and Default Document Structure sections
describe the OUTPUT document's prose and formatting, not a JSON input
contract for a report generator (the skill is written for a conversational
Read First, Ask Second workflow, not for a generator script). This module
resolves that ambiguity the same way evaluation_report_generator.py
resolved its own equivalent gap: it accepts narrative content ALREADY
composed by the upstream reading/drafting step (Rule 5, "Read the document
before writing," is an upstream responsibility this module does not
attempt to replace by parsing PDFs or DOCX contracts itself), and computes
only the two things SKILL.md explicitly assigns to a deterministic kernel:
the ATC and ATS approval-chain grade sequences.

Register shape:
{
  "deal_name": "string",
  "source_documents": [
    {"name": "string", "detail": "string, e.g. 'executed May 12, 2026' or
       'governing MSA not provided - referenced only'"}, ...
  ],  # at least one required (Rule 5 / BLOCKING FILE INPUTS)
  "opening_paragraph": "string, 3-5 sentences, pre-composed with **markdown
       bold** around the total dollar figure per the Tone examples",
  "total_contract_value_usd": float,  # the single figure that drives both
       approval chains, per the deal-value disambiguation rule; resolved
       upstream when several candidate figures exist
  "request_overview": "string or null (conditional, complex/phased deals only)",
  "background_justification": "string, 3-5 sentences",
  "key_capabilities": ["string", ...] or null (conditional, 5-7 bullets,
       license/platform renewals only),
  "scope_of_agreement": ["string bullet, may use **Label:** markdown", ...],
  "financial_summary": {
    "mode": "simple" | "table",
    "simple_lines": ["string, may use **Label:** markdown", ...],  # simple mode
    "table_rows": [["Year 1", 6470000.0], ...],  # table mode
    "stated_total_usd": float,  # table mode; MUST foot to the row sum
    "total_label": "string (default 'Total')",
    "narrative_note": "string or null"
  },
  "vendor_specific": {"title": "string", "lines": ["string", ...]} or null,
  "key_contract_benefits": ["string", ...],  # non-empty
  "key_contract_risks": ["string", ...],  # non-empty
  "business_case": "string or null (conditional)",
  "cost_efficiency": ["string", ...] or null,  # OMIT (null/empty) when
       there are no documentable savings; never a placeholder sentence
  "governance": {
    "ats_names_by_grade": {"M3/P5/R7-9/S6-7": "Name", ...} (optional),
    "budget_owner": "string", "budget_approved": bool (default true),
    "budgeted_amount_usd": float or null (default: total_contract_value_usd),
    "stakeholders": ["string", ...],
    "business_owner": "string", "business_owner_grade": "string",
    "comments_to_business_owner": "string" (has a fixed default),
    "atc_names_by_grade": {"P5/M3": "Name", ...} (optional),
    "user_grade": "string", "procurement_contact": "string",
    "effective_date": "string", "payment_terms": "string" (default "Net 60"),
    "price_change_mechanism": "string",
    "financial_risk_rating": "string or null (never defaulted; null renders
       the exact canonical NEEDS_INPUT string)",
    "other_contract_elements": "string" (default "None"),
    "is_capital_spend": bool or null (only consulted in the $200M-$1B ATS band),
    "table_source": "live SharePoint" | "vendored snapshot" (REQUIRED, no
       default, per frap_chain_kernel.py's own hard gate),
    "live_atc_table": [["Grade", threshold], ...] or null,
    "live_ats_table_operating": [["Grade", threshold], ...] or null,
    "live_ats_table_capital": [["Grade", threshold], ...] or null
  },
  "data_basis": {
    "total_fields": int, "verbatim_count": int, "user_confirmed_count": int,
    "inferred_fields": [{"field_name": "string", "basis": "string",
                          "confidence": "High"|"Medium"|"Low"}, ...]
  }
}

The ATC and ATS approval-chain GRADE sequences, the resolved approver name
lists (grade plus name-or-placeholder), and the displayed Financial Risk
Rating string are COMPUTED by this generator, never accepted as input,
per SKILL.md's chain-construction rule and Rule 3's never-default clause.
============================================================================
"""

from __future__ import annotations

import copy
import os
import sys
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Tuple

# ---------------------------------------------------------------------------
# Vendored FRAP chain kernel (same directory). Per SKILL.md's "Kernel
# wiring" section: the ATC/ATS approval-chain grade sequence MUST be
# computed by calling frap_chain_kernel.compute_chain(), never hand-derived.
# ---------------------------------------------------------------------------
_KERNEL_IMPORT_ERROR: Optional[Exception] = None
try:
    _THIS_DIR = os.path.dirname(os.path.abspath(__file__))
    if _THIS_DIR not in sys.path:
        sys.path.insert(0, _THIS_DIR)
    from frap_chain_kernel import compute_chain as kernel_compute_chain
    from frap_chain_kernel import Facts as KernelFacts
    from frap_chain_kernel import Decision as KernelDecision
    KERNEL_AVAILABLE = True
except Exception as _exc:  # pragma: no cover - defensive, disclosed at runtime
    KERNEL_AVAILABLE = False
    _KERNEL_IMPORT_ERROR = _exc

    def kernel_compute_chain(*args, **kwargs):  # type: ignore
        raise RuntimeError(
            "frap_chain_kernel.py unavailable; cannot compute the "
            f"kernel-required approval chain. Import error: {_KERNEL_IMPORT_ERROR}"
        )

    @dataclass
    class KernelFacts:  # type: ignore
        deal_value: Optional[float] = None
        user_grade: Optional[str] = None
        business_owner_grade: Optional[str] = None
        is_capital_spend: Optional[bool] = None
        mode: Optional[str] = None
        table_source: Optional[str] = None
        live_atc_table: Optional[List[Tuple[str, float]]] = None
        live_ats_table_operating: Optional[List[Tuple[str, float]]] = None
        live_ats_table_capital: Optional[List[Tuple[str, float]]] = None

    @dataclass
    class KernelDecision:  # type: ignore
        outcome: Optional[List[str]]
        reasoning: str
        needs_review: bool = True

# ---------------------------------------------------------------------------
# python-docx detection. Mirrors should_cost_generator.py's openpyxl
# handling: try the import; if unavailable, raise a clear ImportError at
# document-build time (not at module-import time), so validation and
# ground-truth logic (which need no DOCX library) remain usable regardless.
# ---------------------------------------------------------------------------
try:
    import docx  # noqa: F401
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
    DOCX_VERSION = getattr(docx, "__version__", "unknown")
except Exception as _exc:  # pragma: no cover
    DOCX_AVAILABLE = False
    DOCX_VERSION = None
    _DOCX_IMPORT_ERROR = _exc


class ExecutiveSummaryValidationError(Exception):
    """Raised when the executive-summary register is missing a required
    field or carries a value the skill's Accuracy and Anti-Drift Rules
    forbid. Per Rule 1 ("Every fact in the summary must appear in the
    source document") and Rule 2 ("Do not add terms that are not in the
    contract") this module refuses rather than guessing a missing or
    malformed value."""


class ReconciliationError(Exception):
    """Raised when a hard-coded invariant fails (a refused approval chain,
    a mismatched approver-name list, a silently altered Financial Risk
    Rating, or a forbidden phrase / em dash found in the assembled
    document). The document is not saved when this fires."""


# ===========================================================================
# House style constants (SKILL.md's "Plain ATC/ATS house style," quoted
# verbatim from this skill's own Formatting Rules / Formatting Specifics
# sections: "Body text: Normal 11pt Calibri," "Bold Grey (#8A969E)" for the
# Source Documents line and the Data Basis & Confirmations footer, "gray
# header row" for the one permitted table. Deliberately NOT the should-cost
# / evaluation-engine colored-dashboard palette; see the module docstring's
# "House-style note.")
# ===========================================================================
LILLY_BLACK = "212121"    # all body text, title, section headers, governance
MUTED_GREY = "8A969E"     # Source Documents line + Data Basis footer ONLY
LIGHT_GRAY_FILL = "D9D9D9"  # financial-table header row fill

VALID_FS_MODES = ("simple", "table")
VALID_CONFIDENCE = ("High", "Medium", "Low")
VALID_TABLE_SOURCES = ("live SharePoint", "vendored snapshot")

FINANCIAL_RISK_RATING_NEEDS_INPUT = "NEEDS_INPUT (SER score and rating not provided)"
DEFAULT_COMMENTS_TO_BUSINESS_OWNER = (
    "No action is required. This notification is for your awareness only."
)
DEFAULT_PAYMENT_TERMS = "Net 60"
DEFAULT_OTHER_CONTRACT_ELEMENTS = "None"

# Gap Handling / v2.3 "task #25" fix: never write this sentence. Cost
# Efficiency is either populated with real savings or omitted entirely.
_FORBIDDEN_SAVINGS_PLACEHOLDER_FRAGMENTS = (
    "no savings identified",
    "no savings were identified",
    "not identified",
    "none identified",
)

EM_DASH = "—"


# ===========================================================================
# 1. Executive-summary register: typed input + validation
# ===========================================================================

@dataclass
class SourceDocument:
    name: str
    detail: str


@dataclass
class FinancialSummary:
    mode: str  # "simple" | "table"
    simple_lines: List[str] = field(default_factory=list)
    table_rows: List[Tuple[str, float]] = field(default_factory=list)
    total_label: str = "Total"
    stated_total_usd: Optional[float] = None
    narrative_note: Optional[str] = None


@dataclass
class VendorSpecificSection:
    title: str
    lines: List[str]


@dataclass
class InferredField:
    field_name: str
    basis: str
    confidence: str


@dataclass
class DataBasis:
    total_fields: int
    verbatim_count: int
    user_confirmed_count: int
    inferred_fields: List[InferredField]


@dataclass
class GovernanceInput:
    ats_names_by_grade: Dict[str, str]
    budget_owner: str
    budget_approved: bool
    budgeted_amount_usd: float
    stakeholders: List[str]
    business_owner: str
    business_owner_grade: str
    comments_to_business_owner: str
    atc_names_by_grade: Dict[str, str]
    user_grade: str
    procurement_contact: str
    effective_date: str
    payment_terms: str
    price_change_mechanism: str
    financial_risk_rating: Optional[str]
    other_contract_elements: str
    is_capital_spend: Optional[bool]
    table_source: str
    live_atc_table: Optional[List[Tuple[str, float]]]
    live_ats_table_operating: Optional[List[Tuple[str, float]]]
    live_ats_table_capital: Optional[List[Tuple[str, float]]]


@dataclass
class ExecutiveSummaryInput:
    deal_name: str
    source_documents: List[SourceDocument]
    opening_paragraph: str
    total_contract_value_usd: float
    request_overview: Optional[str]
    background_justification: str
    key_capabilities: Optional[List[str]]
    scope_of_agreement: List[str]
    financial_summary: FinancialSummary
    vendor_specific: Optional[VendorSpecificSection]
    key_contract_benefits: List[str]
    key_contract_risks: List[str]
    business_case: Optional[str]
    cost_efficiency: Optional[List[str]]
    governance: GovernanceInput
    data_basis: DataBasis


REQUIRED_TOP_LEVEL_FIELDS = [
    "deal_name", "source_documents", "opening_paragraph", "total_contract_value_usd",
    "background_justification", "scope_of_agreement", "financial_summary",
    "key_contract_benefits", "key_contract_risks", "governance", "data_basis",
]


def _scan_for_em_dash(value: Any, path: str, hits: List[str]) -> None:
    """Recursively scans every string leaf in a raw (still-untyped) register
    for the em dash character, per Operating Rule 7 ("Never use em dashes...
    in ANY written output"). Runs once, over the whole raw dict, so no
    individual field needs its own hand-wired check."""
    if isinstance(value, str):
        if EM_DASH in value:
            hits.append(path)
    elif isinstance(value, dict):
        for k, v in value.items():
            _scan_for_em_dash(v, f"{path}.{k}", hits)
    elif isinstance(value, (list, tuple)):
        for i, v in enumerate(value):
            _scan_for_em_dash(v, f"{path}[{i}]", hits)


def validate_executive_summary_input(register: Dict[str, Any]) -> ExecutiveSummaryInput:
    """Validate a raw executive-summary register dict and return a typed
    ExecutiveSummaryInput. Refuses (raises ExecutiveSummaryValidationError)
    rather than guessing when a required field is missing or a value
    violates a documented, fixed convention, per Rule 1 and Rule 2."""
    errors: List[str] = []

    em_dash_hits: List[str] = []
    _scan_for_em_dash(register, "register", em_dash_hits)
    if em_dash_hits:
        shown = em_dash_hits[:10]
        more = " ... and more" if len(em_dash_hits) > 10 else ""
        errors.append(
            "Em dash character (U+2014) found; forbidden suite-wide (Operating Rule 7, "
            "'Never use em dashes'). Found at: " + ", ".join(shown) + more
        )

    for f in REQUIRED_TOP_LEVEL_FIELDS:
        if f not in register:
            errors.append(f"Missing required top-level field: '{f}'")
    if errors and any("Missing required top-level field" in e for e in errors):
        # Cannot safely proceed to per-field validation without the shape
        # these fields provide; raise now with everything collected so far.
        raise ExecutiveSummaryValidationError(
            "Executive-summary register is missing required fields or contains a "
            "forbidden character; refusing to guess. " + "; ".join(errors)
        )

    deal_name = str(register["deal_name"]).strip()
    if not deal_name:
        errors.append("'deal_name' must be a non-empty string.")

    # --- Source documents (BLOCKING FILE INPUTS: at least one required) ---
    raw_docs = register.get("source_documents", [])
    if not isinstance(raw_docs, list) or len(raw_docs) == 0:
        errors.append(
            "'source_documents' must be a non-empty list (Rule 5: 'Read the document "
            "before writing'; the summary always names what it was built from)."
        )
        raw_docs = []
    source_documents: List[SourceDocument] = []
    for i, d in enumerate(raw_docs):
        if not isinstance(d, dict) or "name" not in d or "detail" not in d:
            errors.append(f"'source_documents[{i}]' must be an object with 'name' and 'detail'.")
            continue
        name = str(d["name"]).strip()
        detail = str(d["detail"]).strip()
        if not name or not detail:
            errors.append(f"'source_documents[{i}]' name/detail must be non-empty.")
            continue
        source_documents.append(SourceDocument(name=name, detail=detail))

    opening_paragraph = str(register["opening_paragraph"]).strip()
    if not opening_paragraph:
        errors.append("'opening_paragraph' must be a non-empty string.")

    total_contract_value_usd = register["total_contract_value_usd"]
    if not isinstance(total_contract_value_usd, (int, float)) or total_contract_value_usd <= 0:
        errors.append(
            f"'total_contract_value_usd' must be a positive number; got {total_contract_value_usd!r}. "
            "This is the single figure that drives both approval chains (deal-value disambiguation rule)."
        )
        total_contract_value_usd = 0.0

    request_overview = register.get("request_overview")
    request_overview = str(request_overview).strip() if request_overview else None

    background_justification = str(register["background_justification"]).strip()
    if not background_justification:
        errors.append("'background_justification' must be a non-empty string.")

    raw_caps = register.get("key_capabilities")
    key_capabilities: Optional[List[str]] = None
    if raw_caps:
        if not isinstance(raw_caps, list):
            errors.append("'key_capabilities', when present, must be a list of strings.")
        else:
            key_capabilities = [str(x) for x in raw_caps]

    raw_scope = register.get("scope_of_agreement", [])
    if not isinstance(raw_scope, list) or len(raw_scope) == 0:
        errors.append("'scope_of_agreement' must be a non-empty list of bullet strings.")
        raw_scope = []
    scope_of_agreement = [str(x) for x in raw_scope]

    # --- Financial summary ---------------------------------------------
    raw_fs = register.get("financial_summary")
    if not isinstance(raw_fs, dict):
        errors.append("'financial_summary' must be an object with a 'mode' of 'simple' or 'table'.")
        raw_fs = {}
    fs_mode = str(raw_fs.get("mode", "")).strip().lower()
    simple_lines: List[str] = []
    table_rows: List[Tuple[str, float]] = []
    stated_total: Optional[float] = None
    if fs_mode not in VALID_FS_MODES:
        errors.append(f"'financial_summary.mode' must be one of {VALID_FS_MODES}; got {raw_fs.get('mode')!r}.")
    elif fs_mode == "simple":
        raw_lines = raw_fs.get("simple_lines", [])
        if not isinstance(raw_lines, list) or len(raw_lines) == 0:
            errors.append("'financial_summary.simple_lines' must be a non-empty list when mode is 'simple'.")
        else:
            simple_lines = [str(x) for x in raw_lines]
    else:  # table
        raw_rows = raw_fs.get("table_rows", [])
        if not isinstance(raw_rows, list) or len(raw_rows) == 0:
            errors.append("'financial_summary.table_rows' must be a non-empty list when mode is 'table'.")
        else:
            for i, r in enumerate(raw_rows):
                if not (isinstance(r, (list, tuple)) and len(r) == 2):
                    errors.append(f"'financial_summary.table_rows[{i}]' must be a [period, fee_usd] pair.")
                    continue
                period, amt = r
                if not isinstance(amt, (int, float)):
                    errors.append(f"'financial_summary.table_rows[{i}]' fee_usd must be numeric.")
                    continue
                table_rows.append((str(period), float(amt)))
        raw_total = raw_fs.get("stated_total_usd")
        if not isinstance(raw_total, (int, float)):
            errors.append("'financial_summary.stated_total_usd' must be numeric when mode is 'table'.")
        else:
            stated_total = float(raw_total)
            if table_rows:
                summed = sum(a for _, a in table_rows)
                if abs(summed - stated_total) > 1.0:
                    errors.append(
                        f"Financial Summary period figures sum to ${summed:,.0f}, which does not foot to "
                        f"the stated total ${stated_total:,.0f} (tolerance $1). Per the Markdown Version "
                        "conversion rule generalized to every rendering of this table: 'an executive "
                        "summary whose periods do not foot to the total is a fail.'"
                    )
    financial_summary = FinancialSummary(
        mode=fs_mode if fs_mode in VALID_FS_MODES else "simple",
        simple_lines=simple_lines, table_rows=table_rows,
        total_label=str(raw_fs.get("total_label", "Total")),
        stated_total_usd=stated_total,
        narrative_note=(str(raw_fs["narrative_note"]).strip() if raw_fs.get("narrative_note") else None),
    )

    raw_vendor = register.get("vendor_specific")
    vendor_specific: Optional[VendorSpecificSection] = None
    if raw_vendor:
        if not isinstance(raw_vendor, dict) or "title" not in raw_vendor or "lines" not in raw_vendor:
            errors.append("'vendor_specific', when present, must be an object with 'title' and 'lines'.")
        else:
            vendor_specific = VendorSpecificSection(
                title=str(raw_vendor["title"]).strip(), lines=[str(x) for x in raw_vendor["lines"]],
            )

    raw_benefits = register.get("key_contract_benefits", [])
    if not isinstance(raw_benefits, list) or len(raw_benefits) == 0:
        errors.append("'key_contract_benefits' must be a non-empty list of bullet strings.")
        raw_benefits = []
    key_contract_benefits = [str(x) for x in raw_benefits]

    raw_risks = register.get("key_contract_risks", [])
    if not isinstance(raw_risks, list) or len(raw_risks) == 0:
        errors.append("'key_contract_risks' must be a non-empty list of bullet strings.")
        raw_risks = []
    key_contract_risks = [str(x) for x in raw_risks]

    raw_business_case = register.get("business_case")
    business_case = str(raw_business_case).strip() if raw_business_case else None

    # --- Cost efficiency: omit (null/empty), never a placeholder ---------
    raw_cost_eff = register.get("cost_efficiency")
    cost_efficiency: Optional[List[str]] = None
    if raw_cost_eff:
        if not isinstance(raw_cost_eff, list):
            errors.append("'cost_efficiency', when present, must be a list of strings.")
        else:
            lines = [str(x) for x in raw_cost_eff]
            for i, line in enumerate(lines):
                low = line.lower()
                if any(frag in low for frag in _FORBIDDEN_SAVINGS_PLACEHOLDER_FRAGMENTS):
                    errors.append(
                        f"'cost_efficiency[{i}]' reads as a no-savings placeholder ({line!r}). Per Gap "
                        "Handling: 'Missing savings, ask user, then omit the section entirely (do not "
                        "write \"No savings identified\").' Set 'cost_efficiency' to null/empty instead "
                        "of writing a placeholder sentence."
                    )
            if lines:
                cost_efficiency = lines
    # an explicit empty list normalizes to None (section omitted either way)

    # --- Governance --------------------------------------------------------
    raw_gov = register.get("governance")
    if not isinstance(raw_gov, dict):
        errors.append("'governance' must be an object.")
        raw_gov = {}

    for f in ("business_owner", "business_owner_grade", "user_grade", "procurement_contact",
              "effective_date", "price_change_mechanism", "table_source"):
        if not str(raw_gov.get(f, "")).strip():
            errors.append(f"'governance.{f}' must be a non-empty string.")

    table_source = str(raw_gov.get("table_source", "")).strip()
    if table_source not in VALID_TABLE_SOURCES:
        errors.append(
            f"'governance.table_source' must be one of {VALID_TABLE_SOURCES}; got {table_source!r}. "
            "No silent default (frap_chain_kernel.py's own hard gate): the caller must state whether "
            "this run used a live SharePoint fetch or the vendored snapshot."
        )

    def _tuple_table(raw: Any, label: str) -> Optional[List[Tuple[str, float]]]:
        if raw is None:
            return None
        if not isinstance(raw, list):
            errors.append(f"'governance.{label}', when present, must be a list of [grade, threshold] pairs.")
            return None
        out: List[Tuple[str, float]] = []
        for i, row in enumerate(raw):
            if not (isinstance(row, (list, tuple)) and len(row) == 2):
                errors.append(f"'governance.{label}[{i}]' must be a [grade, threshold] pair.")
                continue
            out.append((str(row[0]), float(row[1])))
        return out

    live_atc_table = _tuple_table(raw_gov.get("live_atc_table"), "live_atc_table")
    live_ats_table_operating = _tuple_table(raw_gov.get("live_ats_table_operating"), "live_ats_table_operating")
    live_ats_table_capital = _tuple_table(raw_gov.get("live_ats_table_capital"), "live_ats_table_capital")

    ats_names_by_grade = {str(k): str(v) for k, v in (raw_gov.get("ats_names_by_grade") or {}).items()}
    atc_names_by_grade = {str(k): str(v) for k, v in (raw_gov.get("atc_names_by_grade") or {}).items()}

    budgeted_amount_raw = raw_gov.get("budgeted_amount_usd")
    budgeted_amount_usd = (
        float(budgeted_amount_raw) if isinstance(budgeted_amount_raw, (int, float))
        else total_contract_value_usd  # documented safe default, metadata-fields.md: "From contract"
    )

    financial_risk_rating_raw = raw_gov.get("financial_risk_rating")
    financial_risk_rating = str(financial_risk_rating_raw).strip() if financial_risk_rating_raw else None

    is_capital_spend = raw_gov.get("is_capital_spend")
    if is_capital_spend is not None and not isinstance(is_capital_spend, bool):
        errors.append("'governance.is_capital_spend', when present, must be a boolean.")
        is_capital_spend = None

    stakeholders = [str(x) for x in raw_gov.get("stakeholders", [])]

    governance = GovernanceInput(
        ats_names_by_grade=ats_names_by_grade,
        budget_owner=str(raw_gov.get("budget_owner", "")).strip(),
        budget_approved=bool(raw_gov.get("budget_approved", True)),
        budgeted_amount_usd=budgeted_amount_usd,
        stakeholders=stakeholders,
        business_owner=str(raw_gov.get("business_owner", "")).strip(),
        business_owner_grade=str(raw_gov.get("business_owner_grade", "")).strip(),
        comments_to_business_owner=str(raw_gov.get("comments_to_business_owner") or DEFAULT_COMMENTS_TO_BUSINESS_OWNER),
        atc_names_by_grade=atc_names_by_grade,
        user_grade=str(raw_gov.get("user_grade", "")).strip(),
        procurement_contact=str(raw_gov.get("procurement_contact", "")).strip(),
        effective_date=str(raw_gov.get("effective_date", "")).strip(),
        payment_terms=str(raw_gov.get("payment_terms") or DEFAULT_PAYMENT_TERMS),
        price_change_mechanism=str(raw_gov.get("price_change_mechanism", "")).strip(),
        financial_risk_rating=financial_risk_rating,
        other_contract_elements=str(raw_gov.get("other_contract_elements") or DEFAULT_OTHER_CONTRACT_ELEMENTS),
        is_capital_spend=is_capital_spend,
        table_source=table_source if table_source in VALID_TABLE_SOURCES else "vendored snapshot",
        live_atc_table=live_atc_table,
        live_ats_table_operating=live_ats_table_operating,
        live_ats_table_capital=live_ats_table_capital,
    )

    # --- Data basis ----------------------------------------------------
    raw_db = register.get("data_basis")
    if not isinstance(raw_db, dict):
        errors.append("'data_basis' must be an object.")
        raw_db = {}
    total_fields = raw_db.get("total_fields")
    verbatim_count = raw_db.get("verbatim_count")
    user_confirmed_count = raw_db.get("user_confirmed_count")
    if not all(isinstance(x, int) and x >= 0 for x in (total_fields, verbatim_count, user_confirmed_count)):
        errors.append(
            "'data_basis.total_fields', 'verbatim_count', and 'user_confirmed_count' must all be "
            "non-negative integers."
        )
        total_fields, verbatim_count, user_confirmed_count = 0, 0, 0

    raw_inferred = raw_db.get("inferred_fields", [])
    inferred_fields: List[InferredField] = []
    if not isinstance(raw_inferred, list):
        errors.append("'data_basis.inferred_fields' must be a list.")
    else:
        for i, inf in enumerate(raw_inferred):
            if not isinstance(inf, dict) or not all(k in inf for k in ("field_name", "basis", "confidence")):
                errors.append(f"'data_basis.inferred_fields[{i}]' must have field_name, basis, confidence.")
                continue
            conf = str(inf["confidence"]).strip()
            if conf not in VALID_CONFIDENCE:
                errors.append(f"'data_basis.inferred_fields[{i}].confidence' must be one of {VALID_CONFIDENCE}; got {conf!r}.")
                continue
            inferred_fields.append(InferredField(
                field_name=str(inf["field_name"]).strip(), basis=str(inf["basis"]).strip(), confidence=conf,
            ))

    if isinstance(total_fields, int):
        footed = verbatim_count + user_confirmed_count + len(inferred_fields)
        if footed != total_fields:
            errors.append(
                f"'data_basis' counts do not foot: verbatim ({verbatim_count}) + user-confirmed "
                f"({user_confirmed_count}) + inferred ({len(inferred_fields)}) = {footed}, not the "
                f"declared total_fields ({total_fields}). The Data Basis & Confirmations footer's own "
                "coverage line must not silently disagree with what it claims to cover."
            )

    data_basis = DataBasis(
        total_fields=total_fields, verbatim_count=verbatim_count, user_confirmed_count=user_confirmed_count,
        inferred_fields=inferred_fields,
    )

    if errors:
        raise ExecutiveSummaryValidationError(
            "Executive-summary register failed validation; refusing to guess missing or invalid "
            "fields. Issues found:\n  - " + "\n  - ".join(errors)
        )

    return ExecutiveSummaryInput(
        deal_name=deal_name, source_documents=source_documents, opening_paragraph=opening_paragraph,
        total_contract_value_usd=float(total_contract_value_usd), request_overview=request_overview,
        background_justification=background_justification, key_capabilities=key_capabilities,
        scope_of_agreement=scope_of_agreement, financial_summary=financial_summary,
        vendor_specific=vendor_specific, key_contract_benefits=key_contract_benefits,
        key_contract_risks=key_contract_risks, business_case=business_case, cost_efficiency=cost_efficiency,
        governance=governance, data_basis=data_basis,
    )


# ===========================================================================
# Ground-truth computation (via the vendored kernel only, per the HARD RULE)
# ===========================================================================

@dataclass
class GroundTruth:
    atc_decision: "KernelDecision"
    ats_decision: "KernelDecision"
    atc_chain: List[str]
    ats_chain: List[str]
    atc_approver_names: List[str]
    ats_approver_names: List[str]
    financial_risk_rating_display: str
    financial_periods_sum: Optional[float]
    estimated_word_count: int


def _resolve_names(chain: List[str], names_by_grade: Dict[str, str]) -> List[str]:
    return [names_by_grade.get(grade, "[To be confirmed]") for grade in chain]


def _estimate_word_count(reg: ExecutiveSummaryInput) -> int:
    """Rough, non-authoritative word count used only to surface an
    informational length note (SKILL.md's "Length Management": target
    1.5-2 pages, hard limit 2 pages). This is NOT a page-count computation;
    Word pagination depends on rendering, fonts, and the reader's
    environment, none of which this generator can simulate. It exists only
    to flag an obviously oversized register before the document is opened,
    the same spirit as should-cost/evaluation-engine's own explicitly
    scoped-out, upstream-only computations."""
    chunks: List[str] = [reg.opening_paragraph, reg.background_justification]
    if reg.request_overview:
        chunks.append(reg.request_overview)
    if reg.business_case:
        chunks.append(reg.business_case)
    chunks.extend(reg.scope_of_agreement)
    chunks.extend(reg.key_contract_benefits)
    chunks.extend(reg.key_contract_risks)
    if reg.key_capabilities:
        chunks.extend(reg.key_capabilities)
    if reg.cost_efficiency:
        chunks.extend(reg.cost_efficiency)
    if reg.vendor_specific:
        chunks.extend(reg.vendor_specific.lines)
    if reg.financial_summary.mode == "simple":
        chunks.extend(reg.financial_summary.simple_lines)
    if reg.financial_summary.narrative_note:
        chunks.append(reg.financial_summary.narrative_note)
    text = " ".join(chunks).replace("**", "")
    return len(text.split())


def compute_ground_truth(reg: ExecutiveSummaryInput) -> GroundTruth:
    """Compute the ATC and ATS approval chains by calling ONLY
    frap_chain_kernel.compute_chain(), per SKILL.md's Kernel wiring
    section. This is the reference against which the hard-coded invariants
    are checked before the document is saved."""
    if not KERNEL_AVAILABLE:
        raise RuntimeError(
            "frap_chain_kernel.py could not be imported; this generator refuses to hand-compute the "
            f"approval chain in its place (SKILL.md 'Kernel wiring'). Import error: {_KERNEL_IMPORT_ERROR}"
        )

    gov = reg.governance
    common = dict(
        deal_value=reg.total_contract_value_usd,
        table_source=gov.table_source,
        live_atc_table=gov.live_atc_table,
        live_ats_table_operating=gov.live_ats_table_operating,
        live_ats_table_capital=gov.live_ats_table_capital,
    )

    atc_decision = kernel_compute_chain(KernelFacts(mode="ATC", user_grade=gov.user_grade, **common))
    if atc_decision.needs_review or not atc_decision.outcome:
        raise ReconciliationError(
            "ATC-chain-resolved invariant FAILED: frap_chain_kernel.compute_chain() refused rather than "
            f"returning a chain; refusing to guess past that refusal. Kernel reasoning: {atc_decision.reasoning}"
        )

    ats_decision = kernel_compute_chain(KernelFacts(
        mode="ATS", business_owner_grade=gov.business_owner_grade, is_capital_spend=gov.is_capital_spend,
        **common,
    ))
    if ats_decision.needs_review or not ats_decision.outcome:
        raise ReconciliationError(
            "ATS-chain-resolved invariant FAILED: frap_chain_kernel.compute_chain() refused rather than "
            f"returning a chain; refusing to guess past that refusal. Kernel reasoning: {ats_decision.reasoning}"
        )

    atc_chain = list(atc_decision.outcome)
    ats_chain = list(ats_decision.outcome)
    atc_approver_names = _resolve_names(atc_chain, gov.atc_names_by_grade)
    ats_approver_names = _resolve_names(ats_chain, gov.ats_names_by_grade)

    financial_risk_rating_display = gov.financial_risk_rating or FINANCIAL_RISK_RATING_NEEDS_INPUT

    financial_periods_sum: Optional[float] = None
    if reg.financial_summary.mode == "table":
        financial_periods_sum = sum(amt for _, amt in reg.financial_summary.table_rows)

    return GroundTruth(
        atc_decision=atc_decision, ats_decision=ats_decision, atc_chain=atc_chain, ats_chain=ats_chain,
        atc_approver_names=atc_approver_names, ats_approver_names=ats_approver_names,
        financial_risk_rating_display=financial_risk_rating_display,
        financial_periods_sum=financial_periods_sum, estimated_word_count=_estimate_word_count(reg),
    )


# ===========================================================================
# Hard-coded invariant checks (run BEFORE saving)
# ===========================================================================

def _assert_atc_chain_resolved(gt: GroundTruth) -> None:
    if gt.atc_decision.needs_review or not gt.atc_chain:
        raise ReconciliationError(
            f"ATC-chain-resolved invariant FAILED: outcome is {gt.atc_chain!r}, needs_review="
            f"{gt.atc_decision.needs_review}."
        )


def _assert_ats_chain_resolved(gt: GroundTruth) -> None:
    if gt.ats_decision.needs_review or not gt.ats_chain:
        raise ReconciliationError(
            f"ATS-chain-resolved invariant FAILED: outcome is {gt.ats_chain!r}, needs_review="
            f"{gt.ats_decision.needs_review}."
        )


def _assert_approver_names_match_chain_length(gt: GroundTruth) -> None:
    if len(gt.atc_approver_names) != len(gt.atc_chain):
        raise ReconciliationError(
            f"Approver-names-match-chain-length invariant FAILED (ATC): {len(gt.atc_approver_names)} names "
            f"for a {len(gt.atc_chain)}-level chain."
        )
    if len(gt.ats_approver_names) != len(gt.ats_chain):
        raise ReconciliationError(
            f"Approver-names-match-chain-length invariant FAILED (ATS): {len(gt.ats_approver_names)} names "
            f"for a {len(gt.ats_chain)}-level chain."
        )


def _assert_financial_risk_rating_integrity(reg: ExecutiveSummaryInput, gt: GroundTruth) -> None:
    if reg.governance.financial_risk_rating:
        if gt.financial_risk_rating_display != reg.governance.financial_risk_rating:
            raise ReconciliationError(
                "Financial-Risk-Rating-integrity invariant FAILED: a rating was supplied "
                f"({reg.governance.financial_risk_rating!r}) but the display value was silently altered "
                f"to {gt.financial_risk_rating_display!r}."
            )
    else:
        if gt.financial_risk_rating_display != FINANCIAL_RISK_RATING_NEEDS_INPUT:
            raise ReconciliationError(
                "Financial-Risk-Rating-integrity invariant FAILED (Rule 3): no rating was provided but the "
                f"display value is {gt.financial_risk_rating_display!r}, not the canonical NEEDS_INPUT "
                "string. This field must NEVER be defaulted to 'Acceptable' or any other value."
            )


def run_hardcoded_invariant_checks(reg: ExecutiveSummaryInput, gt: GroundTruth) -> None:
    """Run every kernel-dependent hard-coded invariant. Raises
    ReconciliationError on any failure; callers must not build or save the
    document if this raises."""
    _assert_atc_chain_resolved(gt)
    _assert_ats_chain_resolved(gt)
    _assert_approver_names_match_chain_length(gt)
    _assert_financial_risk_rating_integrity(reg, gt)


# ===========================================================================
# 2. Document builder (python-docx)
# ===========================================================================

def _require_docx() -> None:
    if not DOCX_AVAILABLE:
        raise ImportError(
            "python-docx is not installed in this Python environment, so no .docx file can be written. "
            "Install it (`pip install python-docx`) or point this script at an interpreter that already "
            f"has it. Original import error: {_DOCX_IMPORT_ERROR}"
        )


def _set_cell_background(cell, hex_color: str) -> None:
    """Set a table cell's background fill via raw OOXML; python-docx has no
    high-level cell-shading API."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_text(cell, text: str, bold: bool = False, size_pt: int = 10) -> None:
    cell.text = str(text)
    para = cell.paragraphs[0]
    run = para.runs[0] if para.runs else para.add_run("")
    run.font.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor.from_string(LILLY_BLACK)


def _add_markdown_paragraph(doc, text: str, *, size_pt: int = 11, color_hex: str = LILLY_BLACK,
                             italic: bool = False, style: Optional[str] = None,
                             space_before_pt: int = 0, space_after_pt: int = 6):
    """Renders a paragraph from text that may contain **bold** markdown
    segments, implementing SKILL.md's "Bold inline text only, not colored
    banners, not shaded boxes" rule for section headers, and "Bold label +
    colon + plain value on same line" for governance fields, with a single
    shared helper (this document's formatting spec uses that same
    bold-inline convention throughout its own worked examples, e.g.
    "**ATS Approver(s):** Name One", "**Platform:** [name]"). Assumes
    well-formed pairs of ** markers, matching how SKILL.md's own examples
    are written."""
    para = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before_pt)
    para.paragraph_format.space_after = Pt(space_after_pt)
    parts = text.split("**")
    for i, part in enumerate(parts):
        if part == "":
            continue
        run = para.add_run(part)
        run.font.name = "Calibri"
        run.font.size = Pt(size_pt)
        run.font.italic = italic
        run.font.bold = (i % 2 == 1)
        run.font.color.rgb = RGBColor.from_string(color_hex)
    return para


def _add_section_heading(doc, text: str):
    """Bold inline section header, same size as body, no color, no
    shading, per SKILL.md: 'Section headers: Bold inline text only, NOT
    colored banners, NOT shaded boxes.'"""
    return _add_markdown_paragraph(doc, f"**{text}**", size_pt=11, space_before_pt=10, space_after_pt=4)


def _add_bullet(doc, text: str):
    return _add_markdown_paragraph(doc, text, size_pt=11, style="List Bullet", space_after_pt=2)


def _add_table(doc, headers: List[str], rows: List[Tuple[Any, ...]], bold_last_row: bool = False):
    """Simple clean-bordered table with a light gray header row, per
    SKILL.md: 'Tables: Simple clean borders, gray header row, ONLY for
    financial data, NOT for governance fields.'"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell_background(hdr_cells[i], LIGHT_GRAY_FILL)
        _set_cell_text(hdr_cells[i], h, bold=True, size_pt=10)
    n = len(rows)
    for ridx, row_values in enumerate(rows):
        cells = table.add_row().cells
        is_last = bold_last_row and ridx == n - 1
        for i, val in enumerate(row_values):
            _set_cell_text(cells[i], str(val), bold=is_last, size_pt=10)
    return table


def _usd(value: float) -> str:
    return f"${value:,.0f}"


def build_document(reg: ExecutiveSummaryInput, gt: GroundTruth):
    """Build executive_summary.docx per SKILL.md's Section Order (strict,
    follow this sequence exactly). This function does NOT save the file and
    does NOT run the reconciliation checks; call
    generate_executive_summary_docx() for the full validated pipeline."""
    _require_docx()

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = section.right_margin = Inches(1)
    section.top_margin = section.bottom_margin = Inches(1)

    # --- 1. Title ------------------------------------------------------
    _add_markdown_paragraph(doc, f"**Executive Summary: {reg.deal_name}**", size_pt=14, space_after_pt=2)

    # --- 2. Source Documents Line ----------------------------------------
    if len(reg.source_documents) == 1:
        d = reg.source_documents[0]
        source_line = f"Source Document: {d.name} ({d.detail})."
    else:
        parts = [f"{d.name} ({d.detail})" for d in reg.source_documents]
        source_line = "Source Documents: " + "; ".join(parts) + "."
    _add_markdown_paragraph(doc, source_line, size_pt=9, color_hex=MUTED_GREY, italic=True, space_after_pt=10)

    # --- 3. Opening Paragraph ----------------------------------------------
    _add_markdown_paragraph(doc, reg.opening_paragraph, space_after_pt=8)

    # --- 4. Request Overview (conditional) ----------------------------------
    if reg.request_overview:
        _add_section_heading(doc, "Request Overview")
        _add_markdown_paragraph(doc, reg.request_overview)

    # --- 5. Background & Justification --------------------------------------
    _add_section_heading(doc, "Background & Justification")
    _add_markdown_paragraph(doc, reg.background_justification)

    # --- 6. Key Capabilities and Value (conditional) ------------------------
    if reg.key_capabilities:
        _add_section_heading(doc, "Key Capabilities and Value")
        _add_markdown_paragraph(doc, "Key capabilities and value from this renewal include:")
        for item in reg.key_capabilities:
            _add_bullet(doc, item)

    # --- 7. Scope of Agreement -----------------------------------------------
    _add_section_heading(doc, "Scope of Agreement")
    for item in reg.scope_of_agreement:
        _add_bullet(doc, item)

    # --- 8. Financial Summary -------------------------------------------------
    _add_section_heading(doc, "Financial Summary")
    fs = reg.financial_summary
    if fs.mode == "simple":
        for line in fs.simple_lines:
            _add_bullet(doc, line)
    else:
        rows = [(period, _usd(amt)) for period, amt in fs.table_rows]
        rows.append((fs.total_label, _usd(fs.stated_total_usd or 0.0)))
        _add_table(doc, ["Period", "Fee"], rows, bold_last_row=True)
    if fs.narrative_note:
        _add_markdown_paragraph(doc, fs.narrative_note)

    # --- 9. Vendor-Specific Section (conditional) -----------------------------
    if reg.vendor_specific:
        _add_section_heading(doc, reg.vendor_specific.title)
        for line in reg.vendor_specific.lines:
            _add_bullet(doc, line)

    # --- 10. Key Contract Benefits ---------------------------------------------
    _add_section_heading(doc, "Key Contract Benefits")
    for item in reg.key_contract_benefits:
        _add_bullet(doc, item)

    # --- 11. Key Contract Risks / Considerations --------------------------------
    _add_section_heading(doc, "Key Contract Risks / Considerations")
    for item in reg.key_contract_risks:
        _add_bullet(doc, item)

    # --- 12. Business Case (conditional) ----------------------------------------
    if reg.business_case:
        _add_section_heading(doc, "Business Case")
        _add_markdown_paragraph(doc, reg.business_case)

    # --- 13. Cost Efficiency (conditional; HARD RULE, task #25 fix: OMIT the
    #     entire section when there are no savings, never a placeholder
    #     sentence like "No savings identified.") -----------------------------
    if reg.cost_efficiency:
        _add_section_heading(doc, "Cost Efficiency")
        for item in reg.cost_efficiency:
            _add_bullet(doc, item)
    # else: no heading, no content, no placeholder text. Section does not exist.

    # --- 14. Governance & Approvals ----------------------------------------------
    _add_section_heading(doc, "Governance & Approvals")
    gov = reg.governance

    _add_markdown_paragraph(doc, f"**ATS Approver(s):** {'; '.join(gt.ats_approver_names)}", space_after_pt=2)
    _add_markdown_paragraph(doc, f"**Budget Owner(s):** {gov.budget_owner}", space_after_pt=2)
    _add_markdown_paragraph(
        doc, f"**Budget Approved by Budget Owner:** {'Yes' if gov.budget_approved else 'No'}", space_after_pt=2,
    )
    _add_markdown_paragraph(doc, f"**Budgeted Amount:** {_usd(gov.budgeted_amount_usd)}", space_after_pt=2)
    _add_markdown_paragraph(
        doc, f"**Stakeholders:** {'; '.join(gov.stakeholders) if gov.stakeholders else '[To be confirmed]'}",
        space_after_pt=2,
    )
    _add_markdown_paragraph(doc, f"**Business Owner(s):** {gov.business_owner}", space_after_pt=2)
    _add_markdown_paragraph(
        doc, f"**Comments / Instructions to Business Owner(s):** {gov.comments_to_business_owner}",
        space_after_pt=2,
    )

    _add_markdown_paragraph(
        doc, f"**ATC Approver:** {'; '.join(gt.atc_approver_names)}", space_before_pt=8, space_after_pt=2,
    )
    _add_markdown_paragraph(doc, f"**Procurement Contact:** {gov.procurement_contact}", space_after_pt=2)
    _add_markdown_paragraph(doc, f"**Effective Date:** {gov.effective_date}", space_after_pt=2)
    _add_markdown_paragraph(doc, f"**Payment / Discount Terms:** {gov.payment_terms}", space_after_pt=2)
    _add_markdown_paragraph(doc, f"**Price Change Mechanism:** {gov.price_change_mechanism}", space_after_pt=2)
    _add_markdown_paragraph(doc, f"**Financial Risk Rating:** {gt.financial_risk_rating_display}", space_after_pt=2)
    _add_markdown_paragraph(doc, f"**Other Contract Elements:** {gov.other_contract_elements}", space_after_pt=2)

    # --- 15. Data Basis & Confirmations (footer, always last content) -----------
    db = reg.data_basis
    if db.inferred_fields:
        coverage = (
            f"Data Basis: {db.verbatim_count} of {db.total_fields} fields extracted verbatim from the "
            f"source document; {db.user_confirmed_count} user-confirmed; {len(db.inferred_fields)} "
            "inferred - flagged [CONFIRM] below."
        )
        _add_markdown_paragraph(doc, coverage, size_pt=8, color_hex=MUTED_GREY, space_before_pt=10, space_after_pt=2)
        for inf in db.inferred_fields:
            line = f"**[CONFIRM]** {inf.field_name} - {inf.basis} (Confidence: {inf.confidence})"
            _add_markdown_paragraph(doc, line, size_pt=8, color_hex=MUTED_GREY, space_after_pt=2)
    else:
        coverage = (
            "Data Basis: All fields extracted verbatim from the source document or confirmed by the "
            "user. No inferred fields."
        )
        _add_markdown_paragraph(doc, coverage, size_pt=8, color_hex=MUTED_GREY, space_before_pt=10, space_after_pt=2)

    return doc


def _iter_all_text(doc) -> List[str]:
    chunks = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                chunks.append(c.text)
    return chunks


def _assert_no_forbidden_content(doc, reg: ExecutiveSummaryInput) -> None:
    """Post-build, pre-save scan of the assembled Document. Mirrors the
    "the document is not saved when this fires" discipline, applied here to
    checks that can only be evaluated against the rendered text rather than
    the input register (defense in depth against a bug in THIS module's own
    formatting code, not just bad input)."""
    all_text = "\n".join(_iter_all_text(doc))

    if EM_DASH in all_text:
        raise ReconciliationError(
            "No-em-dash invariant FAILED: the rendered document contains an em dash character (U+2014); "
            "forbidden suite-wide (Operating Rule 7)."
        )

    if "no savings identified" in all_text.lower():
        raise ReconciliationError(
            "No-savings-placeholder invariant FAILED: the rendered document contains a 'No savings "
            "identified' placeholder. Gap Handling requires omitting the Cost Efficiency section "
            "entirely instead (the fix behind task #25)."
        )

    has_cost_efficiency_heading = "Cost Efficiency" in all_text
    if bool(reg.cost_efficiency) != has_cost_efficiency_heading:
        raise ReconciliationError(
            f"Cost-Efficiency-presence invariant FAILED: cost_efficiency data provided="
            f"{bool(reg.cost_efficiency)}, but the Cost Efficiency heading present in the rendered "
            f"document={has_cost_efficiency_heading}. These must always agree."
        )

    if "ATS Approver(s):" not in all_text or "ATC Approver:" not in all_text:
        raise ReconciliationError(
            "Governance-both-blocks-present invariant FAILED: Governance & Approvals must always carry "
            "both the ATS Approver(s) and ATC Approver lines (Section 14), regardless of whether the run "
            "was framed as an ATC summary or an ATS summary."
        )


# ===========================================================================
# Full pipeline: validate -> compute ground truth -> hard-invariant checks
# -> build document -> post-build content scan -> save
# ===========================================================================

def generate_executive_summary_docx(raw_register: Dict[str, Any], output_path: str) -> ExecutiveSummaryInput:
    """End-to-end: validate the raw register, compute ground truth via
    frap_chain_kernel, run the hard-coded invariant checks, build the
    document, run the post-build content scan, and only then save it.
    Raises rather than saving a document that fails validation or
    reconciliation."""
    register = copy.deepcopy(raw_register)
    reg = validate_executive_summary_input(register)
    gt = compute_ground_truth(reg)
    run_hardcoded_invariant_checks(reg, gt)
    doc = build_document(reg, gt)
    _assert_no_forbidden_content(doc, reg)
    doc.save(output_path)
    return reg


# ===========================================================================
# Demo data (clearly illustrative; the two Worked Examples in SKILL.md's
# "Lilly FRAP Approval Threshold Schedules" section are quoted verbatim into
# the two demo registers below, the same "quote the skill's own numeric
# example" discipline should-cost-builder applies to its "60+30+10=100"
# illustration)
# ===========================================================================

def _demo_register_with_savings() -> Dict[str, Any]:
    """ILLUSTRATIVE DEMO DATA, not a real deal. Exercises the MAXIMAL
    skeleton: every conditional section populated (Request Overview, Key
    Capabilities, Vendor-Specific, Business Case, Cost Efficiency), a
    populated Financial Risk Rating, a populated Data Basis footer with a
    [CONFIRM] entry, a table-mode Financial Summary, and 2 source
    documents. The ATC facts (deal value $17,000,000, user grade P4/M2)
    reproduce SKILL.md's own "Worked Example 1" verbatim: 'Chain = start
    (P5/M3) through ceiling (P6/M4), all levels in between, lowest to
    highest = P5/M3; P6/M4.'"""
    return {
        "deal_name": "Northlight Data Platform Renewal",
        "source_documents": [
            {"name": "Northlight Data Platform Master Services Agreement (illustrative)",
             "detail": "executed April 8, 2024"},
            {"name": "Amendment 2, Three-Year Renewal and Expanded Capacity (illustrative)",
             "detail": "executed July 15, 2026"},
        ],
        "opening_paragraph": (
            "Your approval is requested for **$17,000,000** over a 3-year term with Northlight Data "
            "Systems (illustrative) for the renewal and capacity expansion of the enterprise data "
            "platform supporting global commercial analytics. This renewal replaces an expiring "
            "agreement and adds committed capacity to support forecasted growth in reporting volume. "
            "The agreement aligns with the enterprise data modernization roadmap."
        ),
        "total_contract_value_usd": 17_000_000,
        "request_overview": (
            "This request covers the three-year renewal of the existing Northlight Data Platform "
            "agreement plus a committed capacity expansion negotiated alongside the renewal. It does "
            "not cover the separate professional-services statement of work, which is tracked under a "
            "different approval."
        ),
        "background_justification": (
            "Lilly has used the Northlight Data Platform since 2024 to support commercial analytics and "
            "regulatory reporting. The current term expires in September 2026, and reporting volume has "
            "grown beyond the originally licensed capacity. This renewal secures continued access on "
            "improved commercial terms and adds capacity headroom through 2029, avoiding a service "
            "interruption to downstream reporting pipelines."
        ),
        "key_capabilities": [
            "Enterprise-scale data pipeline orchestration across commercial and regulatory reporting workloads",
            "Expanded storage and compute capacity sized to the current three-year forecast",
            "SOC 2 Type II certified hosting with dedicated data residency in the United States",
            "Predictable multi-year pricing with a capped annual escalator",
            "Direct alignment with the enterprise data modernization roadmap",
        ],
        "scope_of_agreement": [
            "**Platform:** Northlight Data Platform (illustrative)",
            "**Users Covered:** 850 named analyst and engineering seats",
            "**Modules Included:** Core Pipeline Orchestration, Regulatory Reporting Extension, Expanded Capacity Tier",
            "**Geography:** United States, with limited read access for the EU commercial analytics team",
            "**Term:** October 1, 2026 through September 30, 2029 (3 years)",
            "**Support Level:** 24x7 Priority Support with a named Technical Account Manager",
        ],
        "financial_summary": {
            "mode": "table",
            "table_rows": [["Year 1", 5_200_000], ["Year 2", 5_700_000], ["Year 3", 6_100_000]],
            "stated_total_usd": 17_000_000,
            "total_label": "Total",
            "narrative_note": (
                "Year 1 reflects a partial migration credit; Years 2 and 3 reflect the full "
                "expanded-capacity rate."
            ),
        },
        "vendor_specific": {
            "title": "Usage Commitments, Discounts & Support",
            "lines": [
                "**Committed Annual Capacity:** 40 TB active storage, with overage billed at the negotiated Year 3 rate",
                "**Renewal Discount:** 12% off list price for the expanded capacity tier, locked for the full 3-year term",
                "**Support Tier:** Priority Support upgraded at no incremental cost from Standard Support under the prior agreement",
            ],
        },
        "key_contract_benefits": [
            "$1,150,000 negotiated reduction (6.3%) versus Northlight's original renewal quote across the 3-year term.",
            "Annual price escalation capped at 3%, versus Northlight's standard 6% list increase.",
            "Priority Support with a named Technical Account Manager added at no incremental cost.",
            "40 TB of committed capacity secured ahead of the forecasted FY27 volume increase, avoiding a mid-term true-up at list price.",
        ],
        "key_contract_risks": [
            "No termination for convenience; early exit requires payment of the remaining committed term value.",
            "Expanded capacity pricing is committed regardless of actual utilization; underuse does not reduce the annual fee.",
            "The separate professional-services statement of work for the Year 1 migration is still under negotiation and is not covered by this approval.",
        ],
        "business_case": (
            "This renewal secures uninterrupted access to the commercial analytics and regulatory "
            "reporting pipelines that depend on the Northlight platform, avoiding a service gap at "
            "contract expiration. The added capacity headroom supports the enterprise data "
            "modernization roadmap through FY29 without a disruptive mid-term renegotiation."
        ),
        "cost_efficiency": [
            "$1,150,000 in negotiated savings (6.3%) versus Northlight's original renewal quote, itemized in Key Contract Benefits above.",
            "Priority Support upgrade valued at approximately $85,000 annually, included at no incremental cost.",
        ],
        "governance": {
            "ats_names_by_grade": {"M3/P5/R7-9/S6-7": "Dana Whitfield"},
            "budget_owner": "Dana Whitfield",
            "budget_approved": True,
            "budgeted_amount_usd": 17_000_000,
            "stakeholders": ["Priya Anand", "Tomasz Lewandowski"],
            "business_owner": "Dana Whitfield",
            "business_owner_grade": "M3/P5/R7-9/S6-7",
            "atc_names_by_grade": {"P5/M3": "Marcus Feld"},
            "user_grade": "P4/M2",
            "procurement_contact": "Alex Rivera",
            "effective_date": "October 1, 2026",
            "payment_terms": "Net 60",
            "price_change_mechanism": (
                "3% annual cap on the expanded-capacity tier; base modules held flat for the 3-year term."
            ),
            "financial_risk_rating": "SER 88, Low Risk (Supplier Evaluation Report, dated June 2026)",
            "other_contract_elements": "None",
            "is_capital_spend": None,
            "table_source": "vendored snapshot",
        },
        "data_basis": {
            "total_fields": 14,
            "verbatim_count": 11,
            "user_confirmed_count": 2,
            "inferred_fields": [
                {"field_name": "Price Change Mechanism",
                 "basis": "inferred from the renewal-pricing schedule in Exhibit C, not a labeled clause",
                 "confidence": "Medium"},
            ],
        },
    }


def _demo_register_no_savings() -> Dict[str, Any]:
    """ILLUSTRATIVE DEMO DATA, not a real deal. Exercises the MINIMAL
    skeleton: every conditional section omitted (Request Overview, Key
    Capabilities, Vendor-Specific, Business Case, Cost Efficiency), a
    missing Financial Risk Rating (exercising the NEEDS_INPUT hard rule), a
    Data Basis footer with zero inferred fields (exercising the 'All
    fields...' default line), a simple-mode Financial Summary, and a
    single source document. The ATC facts (deal value $107,000,000, user
    grade P5/M3) reproduce SKILL.md's own "Worked Example 2" verbatim:
    'Chain = start (P6/M4) through ceiling (CFO), all levels in between,
    lowest to highest = P6/M4; CPO; CFO.'"""
    return {
        "deal_name": "Cascade Cloud Infrastructure Change Order",
        "source_documents": [
            {"name": "Cascade Cloud Infrastructure Master Agreement (illustrative)",
             "detail": "executed January 22, 2023"},
        ],
        "opening_paragraph": (
            "Requesting approval for a change order to amend the existing Cascade Cloud Infrastructure "
            "agreement. This amendment increases committed spend to **$107,000,000** to cover a "
            "multi-region capacity expansion required to support the FY27 platform migration."
        ),
        "total_contract_value_usd": 107_000_000,
        "request_overview": None,
        "background_justification": (
            "The existing Cascade Cloud Infrastructure agreement is being amended to add committed "
            "spend for a multi-region capacity expansion. The expansion is required to support the FY27 "
            "platform migration and was not included in the original committed-use tier."
        ),
        "key_capabilities": None,
        "scope_of_agreement": [
            "**Platform:** Cascade Cloud Infrastructure (illustrative)",
            "**Change Covered:** Multi-region committed-use capacity expansion",
            "**Geography:** United States, EU, and APAC regions",
            "**Term:** Coterminous with the base agreement, expiring January 21, 2028",
        ],
        "financial_summary": {
            "mode": "simple",
            "simple_lines": [
                "**Total Amount Requested:** $107,000,000",
                "**Budgeted Amount:** $107,000,000",
                "**Term:** Coterminous with the base agreement, expiring January 21, 2028",
                "**Payment Terms:** Net 45",
            ],
        },
        "vendor_specific": None,
        "key_contract_benefits": [
            "Multi-region committed-use rate locked at the base agreement's existing discount tier, avoiding a higher on-demand rate for the expansion capacity.",
            "No change to the base agreement's existing termination or audit rights.",
        ],
        "key_contract_risks": [
            "Committed-use spend is fixed regardless of actual regional utilization during the FY27 migration ramp.",
            "This change order increases total committed spend; a downstream true-up review is scheduled at the FY27 migration midpoint.",
        ],
        "business_case": None,
        "cost_efficiency": None,
        "governance": {
            "ats_names_by_grade": {},
            "budget_owner": "Priya Anand",
            "budget_approved": True,
            "budgeted_amount_usd": None,
            "stakeholders": [],
            "business_owner": "Priya Anand",
            "business_owner_grade": "M6",
            "atc_names_by_grade": {},
            "user_grade": "P5/M3",
            "procurement_contact": "Alex Rivera",
            "effective_date": "Date of last signature",
            "payment_terms": "Net 45",
            "price_change_mechanism": (
                "Committed-use rate per the base agreement's existing discount schedule; no new "
                "escalation mechanism introduced by this amendment."
            ),
            "financial_risk_rating": None,
            "other_contract_elements": None,
            "is_capital_spend": None,
            "table_source": "vendored snapshot",
        },
        "data_basis": {
            "total_fields": 11,
            "verbatim_count": 9,
            "user_confirmed_count": 2,
            "inferred_fields": [],
        },
    }


# ===========================================================================
# Self-test / CLI
# ===========================================================================

def _run_self_test() -> int:
    import tempfile
    import zipfile

    print("=" * 78)
    print("executive_summary_generator.py self-test")
    print("=" * 78)
    print(f"frap_chain_kernel.py available: {KERNEL_AVAILABLE}")
    print(f"python-docx available: {DOCX_AVAILABLE}" + (f" (version {DOCX_VERSION})" if DOCX_AVAILABLE else ""))
    print()

    results: List[tuple] = []

    def check(label, condition, detail=""):
        results.append((label, bool(condition), detail))
        status = "PASS" if condition else "FAIL"
        line = f"[{status}] {label}"
        if detail:
            line += f"  ({detail})"
        print(line)

    demo_savings = _demo_register_with_savings()
    demo_no_savings = _demo_register_no_savings()

    # --- Step 0: both demo registers validate and compute -----------------
    try:
        reg1 = validate_executive_summary_input(copy.deepcopy(demo_savings))
        check("validate_executive_summary_input accepts the with-savings demo register", True)
    except Exception as e:
        check("validate_executive_summary_input accepts the with-savings demo register", False, str(e))
        raise

    try:
        reg2 = validate_executive_summary_input(copy.deepcopy(demo_no_savings))
        check("validate_executive_summary_input accepts the no-savings demo register", True)
    except Exception as e:
        check("validate_executive_summary_input accepts the no-savings demo register", False, str(e))
        raise

    try:
        gt1 = compute_ground_truth(reg1)
        check("compute_ground_truth runs via frap_chain_kernel.compute_chain() (with-savings)", True,
              f"ATC={gt1.atc_chain}, ATS={gt1.ats_chain}")
    except Exception as e:
        check("compute_ground_truth runs via frap_chain_kernel.compute_chain() (with-savings)", False, str(e))
        raise

    try:
        gt2 = compute_ground_truth(reg2)
        check("compute_ground_truth runs via frap_chain_kernel.compute_chain() (no-savings)", True,
              f"ATC={gt2.atc_chain}, ATS={gt2.ats_chain}")
    except Exception as e:
        check("compute_ground_truth runs via frap_chain_kernel.compute_chain() (no-savings)", False, str(e))
        raise

    # Hand-derived expected values. The ATC chains reproduce SKILL.md's own
    # quoted Worked Examples 1 and 2 verbatim (not independently invented,
    # unlike the ATS chains and the single-level-collapse check below,
    # which SKILL.md does not carry a worked example for).
    check(
        "ATC chain (with-savings, $17M/P4/M2) matches SKILL.md Worked Example 1 verbatim: P5/M3; P6/M4",
        gt1.atc_chain == ["P5/M3", "P6/M4"], f"got {gt1.atc_chain}",
    )
    check(
        "ATC chain (no-savings, $107M/P5/M3) matches SKILL.md Worked Example 2 verbatim: P6/M4; CPO; CFO",
        gt2.atc_chain == ["P6/M4", "CPO", "CFO"], f"got {gt2.atc_chain}",
    )
    check(
        "ATS chain (with-savings, $17M/M3-P5-R7-9-S6-7) independently derived",
        gt1.ats_chain == ["M3/P5/R7-9/S6-7", "M4/P6/R10-11", "M5/R12", "M6", "M7"], f"got {gt1.ats_chain}",
    )
    check(
        "ATS chain (no-savings, $107M/M6) independently derived",
        gt2.ats_chain == ["M6", "M7", "CEO"], f"got {gt2.ats_chain}",
    )
    check(
        "ATC approver names resolve real name + '[To be confirmed]' placeholder (with-savings)",
        gt1.atc_approver_names == ["Marcus Feld", "[To be confirmed]"], f"got {gt1.atc_approver_names}",
    )
    check(
        "ATC approver names are all '[To be confirmed]' when no names were supplied (no-savings)",
        gt2.atc_approver_names == ["[To be confirmed]", "[To be confirmed]", "[To be confirmed]"],
        f"got {gt2.atc_approver_names}",
    )
    check(
        "Financial Risk Rating displays the user-supplied value verbatim (with-savings)",
        gt1.financial_risk_rating_display == "SER 88, Low Risk (Supplier Evaluation Report, dated June 2026)",
    )
    check(
        "Financial Risk Rating HARD RULE: falls back to the exact canonical NEEDS_INPUT string, never "
        "'Acceptable' or any other default (no-savings)",
        gt2.financial_risk_rating_display == FINANCIAL_RISK_RATING_NEEDS_INPUT,
        f"got {gt2.financial_risk_rating_display!r}",
    )
    check(
        "Budgeted Amount safe-defaults to total_contract_value_usd when omitted (no-savings)",
        reg2.governance.budgeted_amount_usd == 107_000_000, f"got {reg2.governance.budgeted_amount_usd}",
    )

    # Independently-derived kernel check (SKILL.md carries no worked
    # example for the "start grade already at/above ceiling" single-level
    # collapse branch; this call exercises it directly against the kernel,
    # not through a full demo document).
    collapse_decision = kernel_compute_chain(KernelFacts(
        deal_value=9_400_000, user_grade="P4/M2", mode="ATC", table_source="vendored snapshot",
    ))
    check(
        "Single-level-chain-collapse branch: $9.4M/P4/M2 (ATC) collapses to ['P5/M3'] "
        "(start grade already at/above ceiling)",
        collapse_decision.outcome == ["P5/M3"], f"got {collapse_decision.outcome}",
    )

    try:
        run_hardcoded_invariant_checks(reg1, gt1)
        check("All hard-coded invariant checks PASS on the with-savings demo register", True)
    except ReconciliationError as e:
        check("All hard-coded invariant checks PASS on the with-savings demo register", False, str(e))
        raise

    try:
        run_hardcoded_invariant_checks(reg2, gt2)
        check("All hard-coded invariant checks PASS on the no-savings demo register", True)
    except ReconciliationError as e:
        check("All hard-coded invariant checks PASS on the no-savings demo register", False, str(e))
        raise

    # --- Step 1: validation refusal tests ----------------------------------
    def _expect_validation_error(label, broken_register):
        try:
            validate_executive_summary_input(broken_register)
            check(label, False, "did not raise")
        except ExecutiveSummaryValidationError as e:
            check(label, True, str(e)[:160])

    b1 = copy.deepcopy(demo_savings)
    del b1["deal_name"]
    _expect_validation_error("validate_executive_summary_input refuses a register missing 'deal_name'", b1)

    b2 = copy.deepcopy(demo_savings)
    b2["background_justification"] = "This deal is fine — no further review needed."
    _expect_validation_error("validate_executive_summary_input refuses an em dash anywhere in the register", b2)

    b3 = copy.deepcopy(demo_savings)
    b3["financial_summary"]["table_rows"][0][1] = 999_999_999  # blows the footing
    _expect_validation_error(
        "validate_executive_summary_input refuses Financial Summary periods that do not foot to the stated total", b3,
    )

    b4 = copy.deepcopy(demo_savings)
    b4["cost_efficiency"] = ["No savings identified for this renewal."]
    _expect_validation_error(
        "validate_executive_summary_input refuses a 'No savings identified' placeholder in cost_efficiency", b4,
    )

    b5 = copy.deepcopy(demo_savings)
    b5["data_basis"]["total_fields"] = 999
    _expect_validation_error("validate_executive_summary_input refuses data_basis counts that do not foot", b5)

    b6 = copy.deepcopy(demo_savings)
    del b6["governance"]["table_source"]
    _expect_validation_error(
        "validate_executive_summary_input refuses a missing governance.table_source (no silent default)", b6,
    )

    b7 = copy.deepcopy(demo_savings)
    b7["data_basis"]["inferred_fields"][0]["confidence"] = "Extremely Sure"
    _expect_validation_error("validate_executive_summary_input refuses an unknown confidence enum", b7)

    b8 = copy.deepcopy(demo_savings)
    b8["key_contract_benefits"] = []
    _expect_validation_error("validate_executive_summary_input refuses an empty key_contract_benefits list", b8)

    b9 = copy.deepcopy(demo_savings)
    b9["source_documents"] = []
    _expect_validation_error("validate_executive_summary_input refuses an empty source_documents list", b9)

    b10 = copy.deepcopy(demo_savings)
    b10["financial_summary"]["mode"] = "narrative"
    _expect_validation_error("validate_executive_summary_input refuses an invalid financial_summary.mode", b10)

    # --- Step 2: kernel-refusal and reconciliation-catches-a-bug tests -----
    try:
        compute_ground_truth(validate_executive_summary_input({
            **copy.deepcopy(demo_savings),
            "governance": {**demo_savings["governance"], "user_grade": "Z9-NOT-A-GRADE"},
        }))
        check("compute_ground_truth CORRECTLY REJECTS an unrecognized ATC grade (kernel refusal surfaced)", False, "did not raise")
    except ReconciliationError as e:
        check("compute_ground_truth CORRECTLY REJECTS an unrecognized ATC grade (kernel refusal surfaced)", True, str(e)[:160])

    try:
        broken_ceo_band = copy.deepcopy(demo_no_savings)
        broken_ceo_band["total_contract_value_usd"] = 500_000_000
        broken_ceo_band["governance"]["is_capital_spend"] = None
        compute_ground_truth(validate_executive_summary_input(broken_ceo_band))
        check(
            "compute_ground_truth CORRECTLY REJECTS a $500M deal with is_capital_spend unset "
            "(CEO band disambiguation required)", False, "did not raise",
        )
    except ReconciliationError as e:
        check(
            "compute_ground_truth CORRECTLY REJECTS a $500M deal with is_capital_spend unset "
            "(CEO band disambiguation required)", True, str(e)[:160],
        )

    def _expect_reconciliation_error(label, fn):
        try:
            fn()
            check(label, False, "did not raise, but should have")
        except ReconciliationError as e:
            check(label, True, str(e)[:160])

    bad_gt1 = copy.deepcopy(gt1)
    bad_gt1.atc_approver_names = ["Only One Name"]  # length no longer matches the 2-level chain
    _expect_reconciliation_error(
        "Approver-names-match-chain-length invariant CORRECTLY REJECTS a truncated name list",
        lambda: _assert_approver_names_match_chain_length(bad_gt1),
    )

    bad_gt2 = copy.deepcopy(gt2)
    bad_gt2.financial_risk_rating_display = "Acceptable"  # the exact forbidden silent default
    _expect_reconciliation_error(
        "Financial-Risk-Rating-integrity invariant CORRECTLY REJECTS a silently defaulted 'Acceptable' rating",
        lambda: _assert_financial_risk_rating_integrity(reg2, bad_gt2),
    )

    bad_gt3 = copy.deepcopy(gt1)
    bad_gt3.financial_risk_rating_display = "SER 12, Low Risk"  # silently altered from what reg1 supplied
    _expect_reconciliation_error(
        "Financial-Risk-Rating-integrity invariant CORRECTLY REJECTS a silently altered user-supplied rating",
        lambda: _assert_financial_risk_rating_integrity(reg1, bad_gt3),
    )

    bad_gt4 = copy.deepcopy(gt1)
    bad_gt4.atc_chain = []
    bad_gt4.atc_decision = KernelDecision(outcome=None, reasoning="corrupted for test", needs_review=True)
    _expect_reconciliation_error(
        "ATC-chain-resolved invariant CORRECTLY REJECTS a refused/empty chain",
        lambda: _assert_atc_chain_resolved(bad_gt4),
    )

    # --- Step 3: build + save + reopen the real documents -------------------
    if not DOCX_AVAILABLE:
        check("python-docx available to write a real .docx file", False,
              "python-docx is NOT installed in this interpreter; DOCX writing could not be exercised.")
        print()
        print("Cannot proceed past this point without python-docx.")
    else:
        check("python-docx available to write a real .docx file", True, f"version {DOCX_VERSION}")

        tmp_dir = tempfile.gettempdir()
        savings_path = os.path.join(tmp_dir, "executive_summary_selftest_with_savings.docx")
        no_savings_path = os.path.join(tmp_dir, "executive_summary_selftest_no_savings.docx")

        try:
            generate_executive_summary_docx(demo_savings, savings_path)
            check("generate_executive_summary_docx() ran end-to-end (with-savings) without raising", True)
        except Exception as e:
            check("generate_executive_summary_docx() ran end-to-end (with-savings) without raising", False, str(e))
            raise

        try:
            generate_executive_summary_docx(demo_no_savings, no_savings_path)
            check("generate_executive_summary_docx() ran end-to-end (no-savings) without raising", True)
        except Exception as e:
            check("generate_executive_summary_docx() ran end-to-end (no-savings) without raising", False, str(e))
            raise

        for label, path in (("With-savings", savings_path), ("No-savings", no_savings_path)):
            exists = os.path.exists(path)
            size = os.path.getsize(path) if exists else 0
            check(f"{label} DOCX file written to {path}", exists and size > 0, f"size={size} bytes")
            check(f"{label} DOCX unzips cleanly (valid OOXML zip container)", zipfile.is_zipfile(path))

        try:
            reopened_savings = Document(savings_path)
            text_savings = "\n".join(_iter_all_text(reopened_savings))

            check(
                "Re-opened with-savings DOCX carries the title",
                "Executive Summary: Northlight Data Platform Renewal" in text_savings,
            )
            check(
                "Re-opened with-savings DOCX carries the plural Source Documents line (2 docs)",
                "Source Documents: Northlight Data Platform Master Services Agreement (illustrative) "
                "(executed April 8, 2024); Amendment 2, Three-Year Renewal and Expanded Capacity "
                "(illustrative) (executed July 15, 2026)." in text_savings,
            )
            for heading in ("Request Overview", "Background & Justification", "Key Capabilities and Value",
                             "Scope of Agreement", "Financial Summary", "Usage Commitments, Discounts & Support",
                             "Key Contract Benefits", "Key Contract Risks / Considerations", "Business Case",
                             "Cost Efficiency", "Governance & Approvals"):
                check(f"Re-opened with-savings DOCX contains the '{heading}' heading", heading in text_savings)
            check("Re-opened with-savings DOCX shows the Financial Summary table total", "$17,000,000" in text_savings)
            check("Re-opened with-savings DOCX shows the resolved ATC approver names", "Marcus Feld" in text_savings and "[To be confirmed]" in text_savings)
            check("Re-opened with-savings DOCX shows the populated Financial Risk Rating", "SER 88, Low Risk" in text_savings)
            check("Re-opened with-savings DOCX shows the Data Basis coverage line with a [CONFIRM] entry",
                  "Data Basis: 11 of 14 fields" in text_savings and "[CONFIRM]" in text_savings)
            check("Re-opened with-savings DOCX does NOT contain 'No savings identified'",
                  "no savings identified" not in text_savings.lower())
            check("Re-opened with-savings DOCX contains no em dash character", EM_DASH not in text_savings)

            n_tables_savings = len(reopened_savings.tables)
            check("Re-opened with-savings DOCX contains exactly 1 table (the Financial Summary table)",
                  n_tables_savings == 1, f"n_tables={n_tables_savings}")
        except Exception as e:
            check("Re-opened with-savings DOCX structural spot-checks", False, str(e))

        try:
            reopened_no_savings = Document(no_savings_path)
            text_no_savings = "\n".join(_iter_all_text(reopened_no_savings))

            check(
                "Re-opened no-savings DOCX carries the singular Source Document line (1 doc)",
                "Source Document: Cascade Cloud Infrastructure Master Agreement (illustrative) "
                "(executed January 22, 2023)." in text_no_savings,
            )
            for heading in ("Background & Justification", "Scope of Agreement", "Financial Summary",
                             "Key Contract Benefits", "Key Contract Risks / Considerations",
                             "Governance & Approvals"):
                check(f"Re-opened no-savings DOCX contains the '{heading}' heading", heading in text_no_savings)
            for absent in ("Request Overview", "Key Capabilities and Value", "Business Case", "Cost Efficiency"):
                check(f"Re-opened no-savings DOCX OMITS the '{absent}' section entirely (no data provided)",
                      absent not in text_no_savings)
            check("Re-opened no-savings DOCX does NOT contain 'No savings identified'",
                  "no savings identified" not in text_no_savings.lower())
            check("Re-opened no-savings DOCX shows the canonical NEEDS_INPUT Financial Risk Rating string",
                  FINANCIAL_RISK_RATING_NEEDS_INPUT in text_no_savings)
            check("Re-opened no-savings DOCX shows the zero-inferred-fields Data Basis default line",
                  "Data Basis: All fields extracted verbatim from the source document or confirmed by "
                  "the user. No inferred fields." in text_no_savings)
            check("Re-opened no-savings DOCX shows the ATC chain names all as '[To be confirmed]'",
                  text_no_savings.count("[To be confirmed]") >= 6)  # 3 ATC + 3 ATS placeholders
            check("Re-opened no-savings DOCX contains no em dash character", EM_DASH not in text_no_savings)

            check("No-savings DOCX has fewer tables than the with-savings DOCX (simple-mode Financial Summary)",
                  len(reopened_no_savings.tables) < n_tables_savings,
                  f"no_savings={len(reopened_no_savings.tables)}, with_savings={n_tables_savings}")
        except Exception as e:
            check("Re-opened no-savings DOCX structural spot-checks", False, str(e))

    print()
    if gt1 is not None and DOCX_AVAILABLE:
        note_threshold = 900
        for label, g in (("with-savings", gt1), ("no-savings", gt2)):
            flag = "note: register is long, re-check the 2-page limit in Word" if g.estimated_word_count > note_threshold else "within the informal length heuristic"
            print(f"Length heuristic ({label}): ~{g.estimated_word_count} words ({flag})")

    print()
    print("=" * 78)
    passed = sum(1 for _, ok, _ in results if ok)
    failed = sum(1 for _, ok, _ in results if not ok)
    total = len(results)
    print(f"SUMMARY: {passed}/{total} passed, {failed}/{total} failed")
    if failed:
        print("FAILED CASES:")
        for label, ok, detail in results:
            if not ok:
                print(f"  - {label}: {detail}")
    print("=" * 78)
    return 1 if failed else 0


def main(argv: Optional[List[str]] = None) -> int:
    """CLI entry point. Usage:
        python executive_summary_generator.py --input register.json --output executive_summary.docx
        python executive_summary_generator.py --demo
        python executive_summary_generator.py --self-test
        python executive_summary_generator.py                 (no args -> runs the self-test)
    """
    import argparse

    parser = argparse.ArgumentParser(
        description="Generate executive_summary.docx from a validated executive-summary register "
                    "(executive-summary-package-1c344a). See the module docstring for the input JSON schema."
    )
    parser.add_argument("--input", "-i", help="Path to a JSON file containing the executive-summary register.")
    parser.add_argument("--output", "-o", default="executive_summary.docx",
                         help="Output .docx path (default: executive_summary.docx in the current directory).")
    parser.add_argument("--demo", action="store_true",
                         help="Run the built-in self-test suite (generates both illustrative demo registers' "
                              "DOCX files, reopens them, and asserts every expected section, table, and value).")
    parser.add_argument("--self-test", action="store_true", dest="self_test", help="Alias for --demo.")
    args = parser.parse_args(argv)

    if args.demo or args.self_test or not args.input:
        return _run_self_test()

    import json
    try:
        with open(args.input, "r", encoding="utf-8") as f:
            raw = json.load(f)
    except (OSError, ValueError) as e:
        print(f"ERROR: could not read/parse --input {args.input!r}: {e}", file=sys.stderr)
        return 1

    try:
        generate_executive_summary_docx(raw, args.output)
    except ImportError as e:
        print(f"ERROR: {e}", file=sys.stderr)
        print("The executive-summary register was valid and the approval chains computed cleanly; "
              "only the DOCX-writing step could not run because python-docx is missing.", file=sys.stderr)
        return 1
    except (ExecutiveSummaryValidationError, ReconciliationError) as e:
        print(f"ERROR: {e}", file=sys.stderr)
        return 1

    print(f"Wrote {args.output}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
