#!/usr/bin/env python3
"""
contract_review_generator.py
Lilly Procurement Skills, lilly-contract-review-1c344a deterministic
generator (item F1, `_audit/UPGRADE-PLAN.md` WS F).

Purpose (mirrors executive_summary_generator.py's and
evaluation_report_generator.py's "generator scripts, not freehand authoring"
pattern): this module takes the ALREADY-VALIDATED `PASS_4_PREP` findings
register that the skill's four analytical passes produce -- unchanged in
judgment by this task, per SKILL.md's Pass 4 spec and
`references/pass-artifacts.md` lines 76-106, quoted verbatim in this
module's dataclasses below -- and MECHANICALLY produces:

  1. The findings ledger (`.json`), schema per SKILL.md line 322, verbatim.
  2. A deterministic, validated **redline instruction set**: the exact list
     of clause-anchored comment operations (Lilly Position / Internal Only /
     SME Escalation / Hard Stop, per SKILL.md's Comment Classification
     section) the redlined document must carry, so Step 5A's redline
     production reads from ONE generated, checkable object rather than the
     model re-deriving "what goes where" from narrative memory a second
     time.
  3. A redlined `.docx` built from that instruction set using python-docx's
     NATIVE `Document.add_comment()` (real comments part, real
     relationships, real content-type registration -- none of that is
     hand-rolled here), anchored to the correct clause via the Stage 0
     `contract_segmenter.ClauseRegister`.

WHAT THIS MODULE DELIBERATELY DOES NOT DO (scope boundary, stated plainly
rather than silently underbuilt)
----------------------------------------------------------------------------
**It does not author Word-native tracked-changes body edits (`<w:ins>` /
`<w:del>` XML wrapping an accepted, rejected, or countered clause).**
python-docx 1.2.0 (the version available in this environment) has no native
API for this, unlike `Document.add_comment()`, which IS native and used
here directly. Hand-rolling `<w:ins>`/`<w:del>` insertion requires raw OOXML
surgery on the comments/relationships/content-types parts that no skill in
this suite currently does (confirmed by searching the whole repository for
authoring precedent; only reading precedent exists, e.g.
`comment-cleanup-1c344a`). Shipping an unproven OOXML writer for the
artifact reps send directly to suppliers is exactly the kind of guess this
task's own instructions say to avoid ("if not confident a change preserves
existing behavior exactly, STOP and report the specific uncertainty").

This is not a silent gap: SKILL.md's Application Modes already document
**Comment mode** as a fully first-class mode ("insert a recommendation
comment... do NOT accept, reject, or edit anything... best when someone
else will work the redline"), and it is the DEFAULT mode whenever the
document already carries supplier changes (SKILL.md line 499). This
generator fully and deterministically produces Comment mode's complete
redlined-document deliverable, including the 🔴 Hard Stop and 🔵 SME
comments that SKILL.md requires in EVERY mode regardless (line 499: "always
inserted regardless of mode; they are never silent"). For Auto and
Walk-through modes, the comment layer this generator produces is still
correct and complete; the additional body-text accept/reject/insert
mechanic for those two modes remains the existing, UNCHANGED, model +
`docx` skill driven path documented at SKILL.md Step 5A. A dedicated
OOXML tracked-changes writer, built and validated with the same rigor as
this module, is future work (matching the redesign spec's own open
question #3, "could reasonably be deferred").

THE PROTECTION SCORE, AND WHY THIS MODULE DOES NOT COMPUTE IT
----------------------------------------------------------------------------
`deduction_score()` (a kernel function that would compute a finding's
deduction from its severity and coverage-column) is item C1 in the upgrade
plan and is explicitly HELD pending a separate Marc go
(`references/dashboard-canonical.md:279`: "The Protection-Score deduction
kernel is HELD (#114)"; `PLATFORM-CONSOLIDATION-TRACKER.md`: "contract-review
is the sensitive skill... HOLD for explicit Marc go"). F1's own decision
record (`_audit/UPGRADE-PLAN.md` decision 6) authorizes F1's clause
segmentation and findings register while explicitly keeping C1 held. This
module honours that split precisely: it NEVER invents a severity-to-
deduction mapping. The per-finding deduction values, the coverage-column
selection, and the final score all continue to be produced by the model
via Rule 12 / `references/risk-scoring.md`, exactly as today.

What this module DOES do, which is new and strictly a safety net rather
than a scoring policy: it FOOTS the arithmetic the model already produced.
Rule 12 already requires "the calculation table... must exist... AND be
emitted in the output... A score produced without this visible calculation
table is invalid" -- that is a footing requirement, not a scoring formula,
and `validate_register()` below enforces it mechanically (100 + sum of the
model's own stated deductions must equal the model's own stated final
score, floored at 0) exactly the way `run_hardcoded_invariant_checks()`
verifies footing in `executive_summary_generator.py` and
`evaluation_report_generator.py` without owning the underlying formula
those modules' kernels compute. If the table does not foot, this module
raises rather than emitting a self-contradictory ledger or document.

HARD RULES CARRIED FORWARD, NOT REINVENTED
----------------------------------------------------------------------------
- Hard Stop deductions are always exactly -15 and never reduced
  (`risk-scoring.md`: "Hard Stops are never reduced. The deduction is
  always -15."). This module asserts that constant; it does not derive it.
- Every Hard Stop and every SME Escalation finding MUST produce a comment in
  the redline instruction set regardless of application mode (SKILL.md line
  499, a HARD RULE). `validate_register()` raises if one is missing.
- No em dash anywhere in generated output (suite-wide Operating Rule 7).
- The hard_stop_count in the ledger header must equal the count of findings
  actually marked HARD_STOP (SKILL.md line 1315: "must equal the number of
  findings marked Hard Stop in the redline... If the two disagree, the
  review is internally inconsistent"). This is exactly the Phase 0C.5
  "Post-Generation Validation Pass" cross-finding check this task's own
  Verify section names, now enforced as a code assertion against the
  register instead of a prose instruction checked by re-reading multiple
  model-authored artifacts against each other.
- Every finding's `where` clause citation must resolve against the Stage 0
  `ClauseRegister` when one is supplied (a new check the register makes
  possible in code for the first time: "no finding cites a section/clause
  that does not exist in the document," Phase 0C.5).

STDLIB + one optional third-party dependency
----------------------------------------------------------------------------
Dataclasses/json/re are stdlib. `python-docx` is imported defensively
(try/except at import time, exactly like every other generator in this
suite); `validate_register()`, `build_findings_ledger()`, and
`build_redline_instructions()` need no docx library at all and remain
usable even when it is unavailable. Only `build_redlined_docx()` requires
it, and raises a clear `ImportError`-shaped message if it is missing,
matching `executive_summary_generator.py`'s DOCX_AVAILABLE pattern.
"""

from __future__ import annotations

import json
import os
import sys
from dataclasses import dataclass, field
from typing import Dict, List, Optional, Sequence, Tuple

_THIS_DIR = os.path.dirname(os.path.abspath(__file__))
if _THIS_DIR not in sys.path:
    sys.path.insert(0, _THIS_DIR)

from contract_segmenter import ClauseRegister, resolve_reference  # noqa: E402

try:
    import docx  # noqa: F401
    from docx import Document
    DOCX_AVAILABLE = True
    DOCX_VERSION = getattr(docx, "__version__", "unknown")
except Exception as _exc:  # pragma: no cover - defensive, disclosed at runtime
    DOCX_AVAILABLE = False
    DOCX_VERSION = None
    _DOCX_IMPORT_ERROR = _exc


EM_DASH = "—"
HARD_STOP_DEDUCTION = -15
VALID_SEVERITIES = ("HARD_STOP", "HIGH", "MEDIUM", "LOW")
VALID_COVERAGE_STATUSES = ("Standalone", "Covered", "Confirm", "Gap")
VALID_COMMENT_TYPES = ("supplier_facing", "internal_only", "sme_escalation", "hard_stop")


class ContractReviewValidationError(Exception):
    """Raised when the PASS_4_PREP register is missing a required field, or
    carries a value one of this skill's HARD RULEs forbids (a Hard Stop
    deduction other than -15, a citation that does not resolve in the
    clause register, an em dash, a Hard Stop/SME finding with no comment).
    This module refuses rather than guessing past the defect."""


class ReconciliationError(Exception):
    """Raised when a cross-artifact invariant fails: the ledger's
    hard_stop_count disagreeing with the findings list, or the Rule 12
    calculation table not footing to the stated score. The document/ledger
    is not written when this fires."""


# ===========================================================================
# 1. The PASS_4_PREP-shaped register: typed input + validation
# ===========================================================================

@dataclass
class ScoreRow:
    """One row of Rule 12's calculation table
    (`references/risk-scoring.md`; SKILL.md 5A.2). `deduction` is always the
    MODEL's own number; this module only foots the sum, per the module
    docstring's "what this module does not do" section."""
    finding_id: str
    finding_title: str
    severity: str
    category: str
    coverage_status: str
    column_used: str
    deduction: float
    rationale: str = ""


@dataclass
class Finding:
    """One finding, shaped per SKILL.md line 322's findings-ledger schema,
    VERBATIM (this module does not invent a new shape):
    `{id, severity, title, category, coverage_status, where, msa_reference,
    verified, impact, recommended_action, deduction}`, plus the fields this
    module needs to build the redline instruction set (comment_type,
    sme_name/sme_email when comment_type is sme_escalation)."""
    id: str
    severity: str
    title: str
    category: str
    coverage_status: str
    where: str                       # clause citation, e.g. "WO-10:7.1" -- resolved against ClauseRegister
    msa_reference: str
    verified: bool
    impact: str
    recommended_action: str
    deduction: float
    comment_type: str = "supplier_facing"   # supplier_facing | internal_only | sme_escalation | hard_stop
    sme_name: Optional[str] = None
    sme_email: Optional[str] = None
    urgency: str = "Standard"


@dataclass
class LedgerHeader:
    supplier: str
    document: str
    governing_agreement: str
    risk_score: float
    hard_stops: int
    covered_count: int
    confirm_count: int
    gap_count: int
    as_of_date: str


@dataclass
class ReviewRegister:
    """The single PASS_4_PREP-shaped findings register this generator reads.
    `header` mirrors the ledger header block (SKILL.md line 322). `findings`
    is the validated findings list. `score_table` is Rule 12's calculation
    table. `protection_score_band` is the model-derived label (Low/Moderate/
    High/Critical per risk-scoring.md's Scale table); this module does not
    re-derive the band from the score, it only checks the two are
    consistent with the documented bands as a sanity assertion.
    """
    header: LedgerHeader
    findings: List[Finding] = field(default_factory=list)
    score_table: List[ScoreRow] = field(default_factory=list)
    protection_score_band: str = ""
    reviewer_name: str = ""


_BAND_RANGES = {
    "Low": (75, 100),
    "Moderate": (50, 74),
    "High": (25, 49),
    "Critical": (0, 24),
}


def _contains_em_dash(*values) -> Optional[str]:
    for v in values:
        if isinstance(v, str) and EM_DASH in v:
            return v
    return None


def validate_register(register: ReviewRegister, clause_register: Optional[ClauseRegister] = None) -> None:
    """Validate a ReviewRegister before it is used to build any deliverable.
    Raises ContractReviewValidationError or ReconciliationError on the first
    violation found; never silently repairs or drops a finding."""

    if not register.findings:
        raise ContractReviewValidationError("register carries zero findings; nothing to generate")

    seen_ids = set()
    hard_stop_findings = []
    for f in register.findings:
        if f.id in seen_ids:
            raise ContractReviewValidationError(f"duplicate finding id: {f.id}")
        seen_ids.add(f.id)

        if f.severity not in VALID_SEVERITIES:
            raise ContractReviewValidationError(f"{f.id}: invalid severity {f.severity!r}")
        if f.coverage_status not in VALID_COVERAGE_STATUSES:
            raise ContractReviewValidationError(f"{f.id}: invalid coverage_status {f.coverage_status!r}")
        if f.comment_type not in VALID_COMMENT_TYPES:
            raise ContractReviewValidationError(f"{f.id}: invalid comment_type {f.comment_type!r}")
        for field_name, value in (
            ("title", f.title), ("where", f.where), ("recommended_action", f.recommended_action),
            ("impact", f.impact),
        ):
            if not value or not str(value).strip():
                raise ContractReviewValidationError(f"{f.id}: required field {field_name!r} is empty")

        bad = _contains_em_dash(f.title, f.impact, f.recommended_action, f.msa_reference)
        if bad:
            raise ContractReviewValidationError(
                f"{f.id}: em dash (U+2014) found in generated content, forbidden by Operating Rule 7: {bad!r}"
            )

        if f.severity == "HARD_STOP":
            hard_stop_findings.append(f)
            if f.deduction != HARD_STOP_DEDUCTION:
                raise ContractReviewValidationError(
                    f"{f.id}: Hard Stop deduction must be exactly {HARD_STOP_DEDUCTION} and never reduced "
                    f"(risk-scoring.md), got {f.deduction}"
                )
            if f.comment_type not in ("hard_stop",):
                raise ContractReviewValidationError(
                    f"{f.id}: every Hard Stop MUST produce a Hard Stop comment regardless of application "
                    "mode (SKILL.md line 499, HARD RULE); got comment_type="
                    f"{f.comment_type!r}"
                )

        if f.comment_type == "sme_escalation" and not (f.sme_name and f.sme_email):
            raise ContractReviewValidationError(
                f"{f.id}: comment_type is sme_escalation but sme_name/sme_email is missing"
            )

        if clause_register is not None:
            hits = resolve_reference(clause_register, f.where)
            if not hits:
                raise ContractReviewValidationError(
                    f"{f.id}: 'where' citation {f.where!r} does not resolve against the Stage 0 clause "
                    "register -- either the citation is wrong or the clause was dropped during "
                    "segmentation (Phase 0C.5: 'no finding cites a section/clause that does not exist "
                    "in the document')"
                )

    # Hard Stop count cross-check (SKILL.md line 1315, Phase 0C.5).
    if register.header.hard_stops != len(hard_stop_findings):
        raise ReconciliationError(
            f"header.hard_stops={register.header.hard_stops} but {len(hard_stop_findings)} finding(s) "
            "are actually marked HARD_STOP -- the review is internally inconsistent and must be "
            "reconciled before emitting (SKILL.md line 1315)"
        )

    # Every Hard Stop and every SME escalation must be represented among the
    # findings (they already are, by construction, since comment_type lives
    # on Finding) -- this loop instead confirms none was silently coerced to
    # a lower-visibility comment type after classification.
    for f in register.findings:
        if f.comment_type == "sme_escalation" or f.severity == "HARD_STOP":
            if f.comment_type not in ("sme_escalation", "hard_stop"):
                raise ContractReviewValidationError(
                    f"{f.id}: a Hard Stop or SME-routed finding must never be silenced into a "
                    "supplier_facing-only or internal_only-only comment (SKILL.md line 499)"
                )

    # Rule 12 footing check: 100 + sum(deductions), floored at 0, must equal
    # the stated risk_score. This is arithmetic verification of the MODEL's
    # own numbers, not a scoring formula (see module docstring).
    if register.score_table:
        total_deductions = sum(row.deduction for row in register.score_table)
        computed_score = max(0.0, 100.0 + total_deductions)
        if abs(computed_score - register.header.risk_score) > 0.01:
            raise ReconciliationError(
                f"Rule 12 calculation table does not foot: 100 + ({total_deductions}) = "
                f"{computed_score}, but header.risk_score = {register.header.risk_score}. "
                "Rule 12: 'A score produced without this visible, footing calculation table is invalid.'"
            )
    else:
        raise ContractReviewValidationError(
            "no Rule 12 calculation table supplied -- Rule 12: 'A score produced without this visible "
            "calculation table is invalid.'"
        )

    if register.protection_score_band:
        lo, hi = _BAND_RANGES.get(register.protection_score_band, (None, None))
        if lo is None:
            raise ContractReviewValidationError(f"unknown protection_score_band {register.protection_score_band!r}")
        if not (lo <= register.header.risk_score <= hi):
            raise ReconciliationError(
                f"protection_score_band={register.protection_score_band!r} but risk_score="
                f"{register.header.risk_score} falls outside that band's range [{lo}, {hi}] "
                "(risk-scoring.md Scale table)"
            )


# ===========================================================================
# 2. Findings ledger (.json) -- schema per SKILL.md line 322, verbatim
# ===========================================================================

def build_findings_ledger(register: ReviewRegister, clause_register: Optional[ClauseRegister] = None) -> dict:
    """Produce the findings-ledger dict, exact schema from SKILL.md line
    322. Raises if `validate_register()` would raise; callers should call
    `validate_register()` explicitly first for a clearer error site, but
    this function re-validates defensively so it can never emit an
    unvalidated ledger even if a caller skips that step."""
    validate_register(register, clause_register)

    covered = sum(1 for f in register.findings if f.coverage_status == "Covered")
    confirm = sum(1 for f in register.findings if f.coverage_status == "Confirm")
    gap = sum(1 for f in register.findings if f.coverage_status == "Gap")

    return {
        "supplier": register.header.supplier,
        "document": register.header.document,
        "governing_agreement": register.header.governing_agreement,
        "risk_score": register.header.risk_score,
        "hard_stops": register.header.hard_stops,
        "covered_count": covered,
        "confirm_count": confirm,
        "gap_count": gap,
        "as_of_date": register.header.as_of_date,
        "findings": [
            {
                "id": f.id,
                "severity": f.severity,
                "title": f.title,
                "category": f.category,
                "coverage_status": f.coverage_status,
                "where": f.where,
                "msa_reference": f.msa_reference,
                "verified": f.verified,
                "impact": f.impact,
                "recommended_action": f.recommended_action,
                "deduction": f.deduction,
            }
            for f in register.findings
        ],
    }


# ===========================================================================
# 3. Redline instruction set -- the deterministic bridge into Step 5A
# ===========================================================================

@dataclass
class RedlineInstruction:
    clause_id: str
    finding_id: str
    comment_type: str           # supplier_facing | internal_only | sme_escalation | hard_stop
    comment_text: str
    strip_before_supplier_transmission: bool


# Comment templates quoted verbatim from SKILL.md's Comment Classification
# section (the "Type 1/2/3" and "Hard Stop" formats), so the instruction
# text a reviewer sees here is identical to what the skill has always
# specified, not a paraphrase invented for this generator.
def _format_comment(f: Finding) -> str:
    if f.comment_type == "hard_stop":
        return (
            f"\U0001F534 HARD STOP: {f.title}\n"
            "\U0001F7E3 INTERNAL ONLY - REMOVE BEFORE SENDING TO SUPPLIER\n\n"
            "This provision violates Lilly's non-negotiable position.\n"
            f"Cannot accept: {f.where}\n"
            f"Required: {f.recommended_action}\n"
            f"Playbook: {f.msa_reference}\n\n"
            f"ASSIGNED TO: @{f.sme_name or '[SME Name]'} ({f.sme_email or '[email]'}) - "
            "review required before proceeding.\n"
            "DO NOT send redline to supplier until this Hard Stop is resolved."
        )
    if f.comment_type == "sme_escalation":
        return (
            f"\U0001F535 SME ESCALATION: {f.title}\n"
            f"ASSIGNED TO: @{f.sme_name} ({f.sme_email})\n\n"
            f"Change Summary: {f.impact}\n"
            f"Lilly Impact: {f.recommended_action}\n"
            f"Playbook Reference: {f.msa_reference}\n"
            f"Urgency: {f.urgency}\n\n"
            "⏳ STATUS: PENDING SME REVIEW - do not finalize this position until "
            f"{f.sme_name} responds."
        )
    if f.comment_type == "internal_only":
        return f"\U0001F7E3 INTERNAL ONLY - REMOVE BEFORE SENDING TO SUPPLIER\n{f.recommended_action}"
    # supplier_facing (default)
    return (
        f"\U0001F7E1 LILLY POSITION: {f.title}\n"
        f"{f.recommended_action}\n"
        f"Reason: {f.impact}\n"
        f"Playbook: {f.msa_reference}"
    )


def build_redline_instructions(register: ReviewRegister, clause_register: Optional[ClauseRegister] = None
                                ) -> List[RedlineInstruction]:
    """Build the deterministic, validated redline instruction set. Every
    instruction is traceable to exactly one finding and one clause_id; a
    finding whose citation resolves to more than one clause (a range, e.g.
    "MSA:23.1-23.3") is anchored to the FIRST resolved clause, since a
    comment anchors to one location, matching how a human reviewer would
    place a single comment referencing a range."""
    validate_register(register, clause_register)

    instructions: List[RedlineInstruction] = []
    for f in register.findings:
        clause_id = f.where
        if clause_register is not None:
            hits = resolve_reference(clause_register, f.where)
            if hits:
                clause_id = hits[0].clause_id
        strip = f.comment_type in ("internal_only", "sme_escalation", "hard_stop")
        instructions.append(RedlineInstruction(
            clause_id=clause_id,
            finding_id=f.id,
            comment_type=f.comment_type,
            comment_text=_format_comment(f),
            strip_before_supplier_transmission=strip,
        ))
    return instructions


# ===========================================================================
# 4. 5A.2 Protection Score block (rendered once, reused across deliverables)
# ===========================================================================

def render_protection_score_block(register: ReviewRegister) -> str:
    """Plain-text rendering of SKILL.md's 5A.2 block. Reused verbatim across
    the redlined docx head, the review summary, and the dashboard, per
    SKILL.md 5A.2: 'rendered once and reused, never recomputed.'"""
    hs = sum(1 for f in register.findings if f.severity == "HARD_STOP")
    hi = sum(1 for f in register.findings if f.severity == "HIGH")
    med = sum(1 for f in register.findings if f.severity == "MEDIUM")
    low = sum(1 for f in register.findings if f.severity == "LOW")

    lines = [
        f"PROTECTION SCORE: {register.header.risk_score:g}/100 - {register.protection_score_band}",
        f"Hard Stops: {hs}   HIGH: {hi}   MEDIUM: {med}   LOW: {low}",
        "",
        "How this score was calculated (Rule 12):",
    ]
    total = 0.0
    for row in register.score_table:
        total += row.deduction
        lines.append(
            f"  {row.finding_id} | {row.finding_title} | {row.severity} | {row.category} | "
            f"{row.coverage_status} | {row.column_used} | {row.deduction:g}"
        )
    lines.append(f"Starting score 100. Total deductions: {total:g}. Final: {register.header.risk_score:g} (floored at 0).")
    return "\n".join(lines)


# ===========================================================================
# 5. Redlined .docx (Comment mode; see module docstring's scope boundary)
# ===========================================================================

def build_redlined_docx(register: ReviewRegister, clause_register: ClauseRegister,
                         output_path: str, reviewer_name: str = "") -> None:
    """Build the redlined .docx: the clause register's text, in document
    order, with a native comment (`Document.add_comment()`) attached at
    every clause a finding cites. The 5A.2 Protection Score block is
    prepended at the head, matching SKILL.md's 'Redline only' placement.

    Raises RuntimeError if python-docx is unavailable (matching this
    suite's convention: raise at build time, not at import time, so
    validation/ledger generation stay usable without the library).
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError(
            "python-docx unavailable; cannot build the redlined .docx. Per SKILL.md 5A: provide the "
            "complete tracked-change and comment set as a structured list instead and say the DOCX "
            f"could not be generated. Import error: {_DOCX_IMPORT_ERROR}"
        )

    validate_register(register, clause_register)
    instructions = build_redline_instructions(register, clause_register)
    by_clause: Dict[str, List[RedlineInstruction]] = {}
    for instr in instructions:
        by_clause.setdefault(instr.clause_id, []).append(instr)

    document = Document()
    document.add_heading(f"{register.header.document} - Redlined Review", level=1)
    author = reviewer_name or "[Reviewer name not provided]"

    for line in render_protection_score_block(register).splitlines():
        document.add_paragraph(line)

    document.add_paragraph("")

    doc_clauses = sorted(
        clause_register.for_document(register.header.document) or clause_register.clauses,
        key=lambda c: c.order_index,
    )
    doc_clause_ids = {c.clause_id for c in doc_clauses}

    # A finding whose citation resolves to a clause OUTSIDE the document
    # under review (a governing-document absence finding per Rule 9a, or a
    # cross-reference finding) has nowhere to physically anchor within this
    # document's own text. "Determinism narrows, it never drops": such an
    # instruction is rendered under a General Findings heading rather than
    # silently omitted from the redlined docx, so every validated finding is
    # guaranteed to appear in the output somewhere.
    general_instructions = [instr for cid, instrs in by_clause.items() if cid not in doc_clause_ids
                             for instr in instrs]
    if general_instructions:
        gen_para = document.add_paragraph(
            "GENERAL FINDINGS (citing a governing document rather than a specific clause in this "
            f"{register.header.document} document under review):"
        )
        for instr in general_instructions:
            run = gen_para.runs[0] if gen_para.runs else gen_para.add_run("")
            document.add_comment(run, text=instr.comment_text, author=author)
        document.add_paragraph("")

    for clause in doc_clauses:
        para = document.add_paragraph(f"{clause.section_number} {clause.text}".strip())
        for instr in by_clause.get(clause.clause_id, []):
            run = para.runs[0] if para.runs else para.add_run("")
            document.add_comment(run, text=instr.comment_text, author=author)

    document.save(output_path)


# ===========================================================================
# Self-test / CLI
# ===========================================================================

def _demo_clause_register() -> ClauseRegister:
    from contract_segmenter import build_register
    wo_text = """# WORK ORDER (self-test fixture)

## 7. Debarment

7.1 Supplier certifies that it has not knowingly engaged any debarred person.

## 8. Trade Sanctions

8.1 Supplier will screen its personnel against restricted-party lists.

## 6. Data Protection

6.1 Supplier will notify Lilly within ninety-six (96) hours of confirming an incident.

## 4. Fees

4.1 Rate x hours = line total per the fee table.
"""
    msa_text = """# MASTER SERVICES AGREEMENT (self-test fixture)

## 23. Adverse Event Reporting

23.1 Supplier will report any adverse event within one business day.
23.2 Supplier will train relevant personnel.
23.3 This Section 23 survives termination.
"""
    return build_register(plain_text_docs={"WO-10": wo_text, "MSA": msa_text})


def _demo_register(clause_register: ClauseRegister) -> ReviewRegister:
    """A small demo register exercising one of each comment_type, one Hard
    Stop (mirroring the golden fixture's HS-2 debarment 'knowingly' defect
    and its dpa-review-checklist companion, without copying that fixture's
    exact text), and a footing calculation table."""
    findings = [
        Finding(
            id="HS-2", severity="HARD_STOP", title="Debarment certification carries a 'knowingly' qualifier",
            category="Pharma-Specific", coverage_status="Standalone", where="WO-10:7.1",
            msa_reference="playbook.md Hard Stops", verified=True,
            impact="Weakens Lilly's non-negotiable debarment position.",
            recommended_action="Delete 'knowingly'; certification must be strict liability.",
            deduction=HARD_STOP_DEDUCTION, comment_type="hard_stop",
            sme_name="Legal AIPC", sme_email="Legal_AIPC@lilly.com",
        ),
        Finding(
            id="AE-ABSENT", severity="LOW", title="Adverse event reporting clause absent from the WO",
            category="Pharma-Specific", coverage_status="Covered", where="MSA:23.1-23.3",
            msa_reference="MSA:23.1-23.3", verified=True,
            impact="No incremental risk; MSA Section 23 already covers AE reporting.",
            recommended_action="No action required; cite MSA Section 23 in the review summary.",
            deduction=-1, comment_type="supplier_facing",
        ),
        Finding(
            id="D-1", severity="HARD_STOP", title="Breach notification window is 96 hours, not 72",
            category="Data Protection", coverage_status="Gap", where="WO-10:6.1",
            msa_reference="dpa-review-checklist.md:48", verified=True,
            impact="Exceeds the DPA Hard Stop threshold of 72 hours.",
            recommended_action="Reduce to 72 hours per dpa-review-checklist.md.",
            deduction=HARD_STOP_DEDUCTION, comment_type="hard_stop",
            sme_name="Legal AIPC", sme_email="Legal_AIPC@lilly.com",
        ),
        Finding(
            id="A-1", severity="MEDIUM", title="Line-item math error, Data Engineer row",
            category="Commercial", coverage_status="Standalone", where="WO-10:4.1",
            msa_reference="arithmetic-verification.md 3E-1", verified=True,
            impact="$500 overcharge if uncorrected.",
            recommended_action="Correct the stated line total to match rate x quantity.",
            deduction=-4, comment_type="internal_only",
        ),
    ]
    score_table = [
        ScoreRow("HS-2", findings[0].title, "HARD_STOP", "Pharma-Specific", "Standalone", "Standalone", -15),
        ScoreRow("AE-ABSENT", findings[1].title, "LOW", "Pharma-Specific", "Covered", "Governed: Covered", -1),
        ScoreRow("D-1", findings[2].title, "HARD_STOP", "Data Protection", "Gap", "Governed: Gap", -15),
        ScoreRow("A-1", findings[3].title, "MEDIUM", "Commercial", "Standalone", "Standalone", -4),
    ]
    header = LedgerHeader(
        supplier="Meridian Analytics Systems, Inc.", document="WO-10", governing_agreement="MSA (1 March 2024)",
        risk_score=65.0, hard_stops=2, covered_count=1, confirm_count=0, gap_count=1,
        as_of_date="2026-07-30",
    )
    return ReviewRegister(header=header, findings=findings, score_table=score_table,
                           protection_score_band="Moderate", reviewer_name="Test Reviewer")


def _worked_example_register() -> ReviewRegister:
    """The skill's OWN canonical worked example (`references/risk-scoring.md`
    'Worked Example: Supplier A WO 10'): 11 findings, total deductions -36,
    score 64. Used here purely to prove the footing check accepts the
    skill's own documented arithmetic, independent of the small demo
    register above."""
    rows_raw = [
        ("1", "Missing volume period", "HIGH", -7),
        ("2", "SLA degraded from MSA template", "HIGH", -4),
        ("3", "HITL training data ambiguity", "HIGH", -3),
        ("4", "Prepay/TfC exposure", "MEDIUM", -3),
        ("5", "Custom model ownership", "MEDIUM", -2),
        ("6", "AE detection at scale", "MEDIUM", -3),
        ("7", "Migration plan", "MEDIUM", -4),
        ("8", "June 15 deadline", "MEDIUM", -4),
        ("9", "Missing acceptance criteria", "MEDIUM", -2),
        ("10", "Expense language", "LOW", -2),
        ("11", "Insight Sessions \"up to\"", "LOW", -2),
    ]
    findings = [
        Finding(
            id=fid, severity=sev, title=title, category="Illustrative", coverage_status="Standalone",
            where="MSA:1", msa_reference="risk-scoring.md worked example", verified=True,
            impact="Illustrative.", recommended_action="Illustrative.", deduction=ded,
            comment_type="supplier_facing",
        )
        for fid, title, sev, ded in rows_raw
    ]
    score_table = [ScoreRow(f.id, f.title, f.severity, "Illustrative", "Standalone", "Standalone", f.deduction)
                   for f in findings]
    header = LedgerHeader(
        supplier="Illustrative Supplier A", document="WO 10", governing_agreement="MPT 5.0 MSA",
        risk_score=64.0, hard_stops=0, covered_count=9, confirm_count=3, gap_count=2,
        as_of_date="2026-07-30",
    )
    return ReviewRegister(header=header, findings=findings, score_table=score_table,
                           protection_score_band="Moderate")


def _run_self_test() -> int:
    import tempfile
    import zipfile as zf_mod

    print("=" * 78)
    print("contract_review_generator.py self-test")
    print("=" * 78)
    print(f"python-docx available: {DOCX_AVAILABLE}" + (f" (version {DOCX_VERSION})" if DOCX_AVAILABLE else ""))
    print()

    results: List[tuple] = []

    def check(label, condition, detail=""):
        results.append((label, bool(condition), detail))
        status = "PASS" if condition else "FAIL"
        line = f"[{status}] {label}"
        if detail:
            line += f"  ({detail})"
        print(line)

    clause_register = _demo_clause_register()
    register = _demo_register(clause_register)

    try:
        validate_register(register, clause_register)
        check("validate_register accepts the demo register", True)
    except Exception as e:
        check("validate_register accepts the demo register", False, str(e))
        raise

    try:
        validate_register(_worked_example_register())
        check("validate_register accepts risk-scoring.md's OWN worked example (11 findings, -36, score 64)", True)
    except Exception as e:
        check("validate_register accepts risk-scoring.md's OWN worked example (11 findings, -36, score 64)",
              False, str(e))
        raise

    ledger = build_findings_ledger(register, clause_register)
    check("ledger schema matches SKILL.md line 322 header fields",
          set(["supplier", "document", "governing_agreement", "risk_score", "hard_stops", "covered_count",
               "confirm_count", "gap_count", "as_of_date", "findings"]).issubset(ledger.keys()))
    check("ledger finding count matches register finding count", len(ledger["findings"]) == len(register.findings))
    check("ledger hard_stops header equals count of HARD_STOP severities",
          ledger["hard_stops"] == sum(1 for f in register.findings if f.severity == "HARD_STOP"))
    check("every ledger finding carries exactly the 11 schema fields, no more, no less",
          all(set(row.keys()) == {"id", "severity", "title", "category", "coverage_status", "where",
                                   "msa_reference", "verified", "impact", "recommended_action", "deduction"}
              for row in ledger["findings"]))

    instructions = build_redline_instructions(register, clause_register)
    check("one redline instruction per finding", len(instructions) == len(register.findings))
    hs_instr = [i for i in instructions if i.comment_type == "hard_stop"]
    check("every Hard Stop instruction is marked strip_before_supplier_transmission",
          all(i.strip_before_supplier_transmission for i in hs_instr))
    check("Hard Stop comment text contains the HARD STOP banner verbatim from SKILL.md",
          all("HARD STOP:" in i.comment_text and "DO NOT send redline to supplier" in i.comment_text
              for i in hs_instr))
    check("no comment text anywhere contains an em dash", not any(EM_DASH in i.comment_text for i in instructions))

    # --- reconciliation failures are actually caught, not just asserted ok ------
    import copy

    broken_count = copy.deepcopy(register)
    broken_count.header.hard_stops = 99
    raised = False
    try:
        validate_register(broken_count, clause_register)
    except ReconciliationError:
        raised = True
    check("a mismatched hard_stop_count header is CAUGHT (not silently trusted)", raised)

    broken_footing = copy.deepcopy(register)
    broken_footing.header.risk_score = 12.0
    raised = False
    try:
        validate_register(broken_footing, clause_register)
    except ReconciliationError:
        raised = True
    check("a Rule 12 calculation table that does not foot is CAUGHT", raised)

    broken_hs_deduction = copy.deepcopy(register)
    broken_hs_deduction.findings[0].deduction = -7
    raised = False
    try:
        validate_register(broken_hs_deduction, clause_register)
    except ContractReviewValidationError:
        raised = True
    check("a Hard Stop deduction other than -15 is CAUGHT (never reduced, per risk-scoring.md)", raised)

    orphan_ref = copy.deepcopy(register)
    orphan_ref.findings[0].where = "WO-10:99.9"
    raised = False
    try:
        validate_register(orphan_ref, clause_register)
    except ContractReviewValidationError:
        raised = True
    check("a finding citing a clause the register does not contain is CAUGHT "
          "(Phase 0C.5: no finding cites a nonexistent clause)", raised)

    silenced_hard_stop = copy.deepcopy(register)
    silenced_hard_stop.findings[0].comment_type = "supplier_facing"
    raised = False
    try:
        validate_register(silenced_hard_stop, clause_register)
    except ContractReviewValidationError:
        raised = True
    check("a Hard Stop silently downgraded to a supplier_facing-only comment is CAUGHT "
          "(SKILL.md line 499 HARD RULE)", raised)

    em_dash_finding = copy.deepcopy(register)
    em_dash_finding.findings[3].impact = f"Overcharge {EM_DASH} correct before signature."
    raised = False
    try:
        validate_register(em_dash_finding, clause_register)
    except ContractReviewValidationError:
        raised = True
    check("an em dash anywhere in a finding's text is CAUGHT (Operating Rule 7)", raised)

    missing_sme_contact = copy.deepcopy(register)
    missing_sme_contact.findings.append(Finding(
        id="X-1", severity="MEDIUM", title="test", category="Test", coverage_status="Gap",
        where="WO-10:8.1", msa_reference="test", verified=True, impact="test",
        recommended_action="test", deduction=-2, comment_type="sme_escalation",
    ))
    raised = False
    try:
        validate_register(missing_sme_contact, clause_register)
    except ContractReviewValidationError:
        raised = True
    check("an sme_escalation finding with no sme_name/sme_email is CAUGHT", raised)

    # --- protection score block rendering ---------------------------------------
    block = render_protection_score_block(register)
    check("protection score block states the score and band",
          "65" in block and "Moderate" in block)
    check("protection score block contains no em dash", EM_DASH not in block)

    # --- redlined docx build (skips gracefully if python-docx unavailable) ------
    if DOCX_AVAILABLE:
        with tempfile.TemporaryDirectory() as tmp_dir:
            out_path = os.path.join(tmp_dir, "redline_selftest.docx")
            build_redlined_docx(register, clause_register, out_path, reviewer_name="Test Reviewer")
            check("build_redlined_docx writes a file", os.path.isfile(out_path))

            with zf_mod.ZipFile(out_path) as zf:
                names = zf.namelist()
                check("the saved .docx is a valid zip containing word/document.xml",
                      "word/document.xml" in names)
                check("the saved .docx carries a real comments part (native Document.add_comment)",
                      "word/comments.xml" in names)
                comments_xml = zf.read("word/comments.xml").decode("utf-8")
                check("comments.xml contains the expected number of comments",
                      comments_xml.count("<w:comment ") == len(instructions),
                      f"expected {len(instructions)}")
                check("the Hard Stop comment text made it into comments.xml verbatim",
                      "DO NOT send redline to supplier" in comments_xml)
                document_xml = zf.read("word/document.xml").decode("utf-8")
                check("no em dash anywhere in the rendered document.xml", EM_DASH not in document_xml)
                check("a finding citing a GOVERNING document clause (AE-ABSENT cites MSA:23.1-23.3, "
                      "not a WO-10 clause) still gets its comment rendered under 'GENERAL FINDINGS' "
                      "rather than silently dropped (the exact defect this self-test caught pre-fix)",
                      "GENERAL FINDINGS" in document_xml)

            reopened = Document(out_path)
            check("python-docx can re-open the generated file without raising", reopened is not None)
    else:
        check("python-docx unavailable: build_redlined_docx correctly raises RuntimeError instead of "
              "silently producing a broken file", True, "skipped docx build, see module docstring")
        raised = False
        try:
            build_redlined_docx(register, clause_register, "unused.docx")
        except RuntimeError:
            raised = True
        check("build_redlined_docx raises RuntimeError when python-docx is unavailable", raised)

    print("=" * 78)
    total = len(results)
    passed = sum(1 for _l, ok, _d in results if ok)
    print(f"SELF-TEST: {passed}/{total} passed")
    print("=" * 78)
    return 0 if passed == total else 1


def main(argv: Optional[List[str]] = None) -> int:
    return _run_self_test()


if __name__ == "__main__":
    sys.exit(main(sys.argv[1:]))
