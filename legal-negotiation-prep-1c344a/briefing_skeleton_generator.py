#!/usr/bin/env python3
"""
briefing_skeleton_generator.py — the briefing STRUCTURE (F9 build 4, the split).

F9 split this skill deliberately, and the split is the whole design:

    BUILD   the briefing skeleton, the tier tables, the position counts, the
            Position Map & MSA Coverage panel. Structure and arithmetic.
    PROSE   the negotiation argument itself. That is the skill's whole value, and a
            generator authoring negotiation rationale would produce worse output than
            the model.

So this emits a document with every heading, every table and every count already correct,
and clearly marked placeholders where the argument goes. It never writes the argument.
The same split `executive-summary-package` already implements.

`tier_kernel.py` computes the tiering (G11 HARD RULE). This module does not re-derive a
single tier; it calls the kernel and lays the results out.

WHAT IT REFUSES TO DO
---------------------
  * override a BINDING kernel result                        -> BindingOverrideError
  * silently tier a position with missing inputs            -> surfaced as REVIEW, never guessed
  * emit counts that do not reconcile with the position list-> CountMismatchError
  * a position with no name                                 -> BriefingError
  * python-docx unavailable                                 -> DocxUnavailableError

The binding rule is the important one. `tier_kernel` marks a RED LINE (Tier 1) result
`binding=True` and its own trace says the result "can never be softened by the LLM,
regardless of any other input". A briefing that renders a Lilly non-negotiable as anything
softer hands the negotiator a position Lilly never agreed to hold, which is worse than
having no briefing.
"""
from __future__ import annotations

import json
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from tier_kernel import (                                      # noqa: E402  (G11)
    REVIEW,
    TIER_EASY_CONCEDE,
    TIER_HOLD_FIRM,
    TIER_RED_LINE,
    TIER_STRATEGIC_TRADE,
    TermAttrs,
    assign_tier,
)

try:
    from docx import Document
    from docx.shared import Pt
    _DOCX = True
except ImportError:                                            # pragma: no cover
    _DOCX = False

TIER_ORDER = [TIER_RED_LINE, TIER_HOLD_FIRM, TIER_STRATEGIC_TRADE, TIER_EASY_CONCEDE, REVIEW]

TIER_ACTION = {
    TIER_RED_LINE:         "Hold absolutely. No concession.",
    TIER_HOLD_FIRM:        "Hold. Concede only against a named trade.",
    TIER_STRATEGIC_TRADE:  "Trade deliberately, for something named.",
    TIER_EASY_CONCEDE:     "Concede early, to buy movement elsewhere.",
    REVIEW:                "NEEDS INPUT. Not tiered: inputs are missing.",
}

# Where the model writes. Each placeholder names what belongs there, so an unfilled
# briefing is obviously unfinished rather than quietly thin.
PROSE_SECTIONS = [
    ("Executive Summary",
     "[PROSE: the negotiation's shape in 4-6 sentences. What Lilly wants, what the "
     "supplier will resist, and where this is won or lost.]"),
    ("Leverage Map",
     "[PROSE: what Lilly holds, what the supplier holds, and what changes hands first.]"),
    ("Fallback Sequencing",
     "[PROSE: the order concessions are offered in, and what each one must buy.]"),
    ("Predicted Supplier Pushback",
     "[PROSE: per position, the argument the supplier will make and the answer to it. "
     "Cite the source for any claim about this supplier's past behaviour.]"),
]


class BriefingError(Exception):
    """Base for every refusal."""


class BindingOverrideError(BriefingError):
    pass


class CountMismatchError(BriefingError):
    pass


class DocxUnavailableError(BriefingError):
    pass


def tier_positions(positions):
    """Run every position through the kernel. Returns rows carrying the kernel's own trace.

    A position whose inputs are incomplete comes back as REVIEW with `missing_inputs`
    populated. It is NOT defaulted into a tier: an invented tier is a negotiating
    instruction nobody authorised.
    """
    rows = []
    for i, p in enumerate(positions or []):
        name = (p.get("position_name") or "").strip()
        if not name:
            raise BriefingError("positions[%d] has no position_name" % i)

        attrs = TermAttrs(**{f: p.get(f) for f in TermAttrs.__dataclass_fields__})
        result = assign_tier(attrs)

        claimed = (p.get("claimed_tier") or "").strip()
        if claimed and result.binding and claimed != result.decision:
            raise BindingOverrideError(
                "position %r: the kernel assigned %r and marked it BINDING, but the input "
                "claims %r. A binding tier cannot be softened. Rendering a Lilly "
                "non-negotiable as anything less hands the negotiator a position Lilly "
                "never agreed to hold." % (name, result.decision, claimed)
            )

        rows.append({
            "position_name": name,
            "tier": result.decision,
            "tier_number": result.tier_number,
            "action": TIER_ACTION.get(result.decision, ""),
            "source": result.source,
            "binding": bool(result.binding),
            "missing_inputs": list(result.missing_inputs or []),
            "trace": list(result.trace or []),
            "msa_coverage": p.get("msa_coverage") or "Not stated",
        })
    return rows


def counts_by_tier(rows):
    """Position counts per tier, plus the reconciliation the briefing states."""
    counts = {t: 0 for t in TIER_ORDER}
    for r in rows:
        counts[r["tier"]] = counts.get(r["tier"], 0) + 1
    total = sum(counts.values())
    if total != len(rows):
        raise CountMismatchError(
            "tier counts sum to %d but there are %d positions. A briefing whose own totals "
            "do not foot cannot be trusted on anything else it counts." % (total, len(rows))
        )
    return counts, total


def build_skeleton(rows, meta):
    """The document model: headings, tables, counts, prose placeholders. No argument."""
    counts, total = counts_by_tier(rows)
    review = [r for r in rows if r["tier"] == REVIEW]
    return {
        "title": "%s Negotiation Briefing" % (meta.get("supplier") or "[Supplier]"),
        "meta": {k: meta.get(k) for k in ("supplier", "contract_type", "contract_value",
                                          "version", "prepared_for")},
        "position_map": [
            {"position_name": r["position_name"], "tier": r["tier"],
             "msa_coverage": r["msa_coverage"], "binding": r["binding"]}
            for r in rows
        ],
        "tier_table": [
            {"tier": t, "tier_number": next((r["tier_number"] for r in rows if r["tier"] == t), None),
             "count": counts.get(t, 0), "action": TIER_ACTION.get(t, ""),
             "positions": [r["position_name"] for r in rows if r["tier"] == t]}
            for t in TIER_ORDER if counts.get(t, 0)
        ],
        "counts": counts,
        "total_positions": total,
        "needs_input": [{"position_name": r["position_name"],
                         "missing_inputs": r["missing_inputs"]} for r in review],
        "prose_sections": [{"heading": h, "placeholder": p} for h, p in PROSE_SECTIONS],
    }


def write_docx(skeleton, path):
    """Render the skeleton. Every number is already decided; only prose is left."""
    if not _DOCX:
        raise DocxUnavailableError(
            "python-docx is unavailable, so the briefing skeleton cannot be built. "
            "Refusing to emit a markdown file under a .docx name."
        )
    doc = Document()
    doc.add_heading(skeleton["title"], level=0)

    m = skeleton["meta"]
    line = " | ".join(str(m[k]) for k in ("supplier", "contract_type", "contract_value")
                      if m.get(k))
    if line:
        doc.add_paragraph(line)

    doc.add_heading("Position Map and MSA Coverage", level=1)
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    for c, h in enumerate(["Position", "Tier", "MSA coverage", "Binding"]):
        cell = t.rows[0].cells[c]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for row in skeleton["position_map"]:
        cells = t.add_row().cells
        cells[0].text = row["position_name"]
        cells[1].text = row["tier"]
        cells[2].text = str(row["msa_coverage"])
        cells[3].text = "BINDING" if row["binding"] else ""

    doc.add_heading("Positions by Tier", level=1)
    t2 = doc.add_table(rows=1, cols=4)
    t2.style = "Table Grid"
    for c, h in enumerate(["Tier", "Count", "Action", "Positions"]):
        cell = t2.rows[0].cells[c]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for row in skeleton["tier_table"]:
        cells = t2.add_row().cells
        cells[0].text = row["tier"]
        cells[1].text = str(row["count"])
        cells[2].text = row["action"]
        cells[3].text = "; ".join(row["positions"])
    doc.add_paragraph("Total positions: %d" % skeleton["total_positions"])

    if skeleton["needs_input"]:
        doc.add_heading("Not tiered: inputs missing", level=1)
        doc.add_paragraph(
            "These positions were NOT assigned a tier. The kernel refuses to guess, and a "
            "tier nobody computed is a negotiating instruction nobody authorised.")
        for n in skeleton["needs_input"]:
            doc.add_paragraph("%s - missing: %s"
                              % (n["position_name"], ", ".join(n["missing_inputs"]) or "unstated"),
                              style="List Bullet")

    for sec in skeleton["prose_sections"]:
        doc.add_heading(sec["heading"], level=1)
        p = doc.add_paragraph(sec["placeholder"])
        for r in p.runs:
            r.italic = True
            r.font.size = Pt(10)

    doc.save(path)
    return path


def main(argv):
    if not argv:
        print(__doc__.strip())
        print("\nusage: briefing_skeleton_generator.py <positions.json> [out.docx]")
        return 0
    with open(argv[0], encoding="utf-8") as fh:
        payload = json.load(fh)
    try:
        rows = tier_positions(payload.get("positions"))
        skeleton = build_skeleton(rows, payload.get("meta") or {})
        if len(argv) > 1:
            write_docx(skeleton, argv[1])
    except BriefingError as e:
        print("REFUSED: %s: %s" % (type(e).__name__, e), file=sys.stderr)
        return 2
    print(json.dumps({"counts": skeleton["counts"],
                      "total": skeleton["total_positions"],
                      "needs_input": [n["position_name"] for n in skeleton["needs_input"]]},
                     indent=2))
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv[1:]))
