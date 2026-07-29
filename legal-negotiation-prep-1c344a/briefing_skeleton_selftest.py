#!/usr/bin/env python3
"""
Self-test for briefing_skeleton_generator.py (F9 build 4).

The tiering itself is `tier_kernel`'s and is tested there. These assertions target the
SPLIT: that structure and counts are built by code, that the argument is left to the model
as visible placeholders, and above all that a BINDING kernel result cannot be softened on
its way into the document.

Run: python briefing_skeleton_selftest.py
"""
from __future__ import annotations

import os
import sys
import tempfile

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from briefing_skeleton_generator import (                       # noqa: E402
    PROSE_SECTIONS,
    BindingOverrideError,
    BriefingError,
    build_skeleton,
    counts_by_tier,
    tier_positions,
    write_docx,
)
from tier_kernel import REVIEW, TIER_RED_LINE                   # noqa: E402

try:
    from docx import Document
    _DOCX = True
except ImportError:
    _DOCX = False

PASS, FAIL = [], []


def ok(name, cond, detail=""):
    (PASS if cond else FAIL).append(name)
    print(("  ok   " if cond else "  FAIL ") + name + (("  <- " + detail) if detail and not cond else ""))


def raises(name, exc, fn):
    try:
        fn()
    except exc:
        ok(name, True)
        return
    except Exception as e:                                     # noqa: BLE001
        ok(name, False, "raised %s, expected %s" % (type(e).__name__, exc.__name__))
        return
    ok(name, False, "did not raise " + exc.__name__)


def pos(name, **kw):
    base = dict(position_name=name, is_hard_stop_or_regulatory=False,
                is_playbook_position=True, msa_conflict=False,
                has_compliance_leverage=False, historical_acceptance_rate=0.5,
                acceptable_fallback_exists=True,
                financial_exposure_pct_of_contract_value=2.0,
                supplier_accepted_standard_count=1, creates_precedent_risk=False)
    base.update(kw)
    return base


def positions():
    return [
        pos("Indemnification cap", is_hard_stop_or_regulatory=True, msa_conflict=True,
            creates_precedent_risk=True),
        pos("Liability cap", msa_conflict=True, financial_exposure_pct_of_contract_value=15.0,
            acceptable_fallback_exists=False),
        pos("Payment terms"),
        pos("Marketing reference rights", historical_acceptance_rate=0.9,
            financial_exposure_pct_of_contract_value=0.1),
    ]


def run():
    print("=" * 84)
    print("briefing_skeleton_generator self-test")
    print("=" * 84)

    rows = tier_positions(positions())
    ok("T1  every position is tiered by the kernel", len(rows) == 4)
    ok("T2  each row carries the kernel's own trace", all(r["trace"] for r in rows))
    ok("T3  each row records the kernel's source", all(r["source"] for r in rows))

    red = [r for r in rows if r["tier"] == TIER_RED_LINE]
    ok("T4  a Hard Stop position tiers as RED LINE", len(red) == 1, str([r["tier"] for r in rows]))
    ok("T5  the RED LINE row is marked BINDING", red and red[0]["binding"] is True)
    ok("T6  it carries the tier's action, not an invented one",
       red and red[0]["action"] == "Hold absolutely. No concession.")

    # --- the binding rule, the reason this generator exists -----------------------------
    def soften():
        p = positions()
        p[0]["claimed_tier"] = "STRATEGIC TRADE"      # try to trade away a Lilly red line
        tier_positions(p)
    raises("T7  refuses to SOFTEN a binding RED LINE", BindingOverrideError, soften)

    def soften_to_review():
        p = positions()
        p[0]["claimed_tier"] = REVIEW
        tier_positions(p)
    raises("T8  refuses to demote a binding result to REVIEW", BindingOverrideError,
           soften_to_review)

    p_ok = positions()
    p_ok[0]["claimed_tier"] = TIER_RED_LINE
    ok("T9  NEGATIVE CONTROL: a claimed tier that AGREES with the kernel is fine",
       tier_positions(p_ok)[0]["tier"] == TIER_RED_LINE)

    # binding=True covers the whole DETERMINISTIC path (source playbook or msa), not just
    # Tier 1. A non-binding row is one the kernel could not decide deterministically, so it
    # is built here by removing a required input rather than by picking a softer tier.
    p_nb = positions()
    p_nb[2]["is_playbook_position"] = None          # forces REVIEW / non-binding
    rows_nb = tier_positions(p_nb)
    nb = [r for r in rows_nb if not r["binding"]]
    ok("T10 a row the kernel cannot decide deterministically is NON-binding",
       len(nb) == 1, str([(r["tier"], r["binding"]) for r in rows_nb]))

    p_nb2 = positions()
    p_nb2[2]["is_playbook_position"] = None
    p_nb2[2]["claimed_tier"] = "EASY CONCEDE"
    ok("T10a NEGATIVE CONTROL: a claim on a NON-binding row does not raise",
       len(tier_positions(p_nb2)) == 4)

    # --- missing inputs are surfaced, never guessed -------------------------------------
    p_missing = positions()
    p_missing[1]["historical_acceptance_rate"] = None
    p_missing[1]["financial_exposure_pct_of_contract_value"] = None
    rows_m = tier_positions(p_missing)
    sk_m = build_skeleton(rows_m, {"supplier": "Meridian"})
    review_rows = [r for r in rows_m if r["tier"] == REVIEW]
    if review_rows:
        ok("T11 a position with missing inputs is REVIEW, not a guessed tier", True)
        ok("T12 the skeleton names WHICH inputs are missing",
           bool(sk_m["needs_input"]) and sk_m["needs_input"][0]["missing_inputs"] is not None)
    else:
        ok("T11 kernel tiered it without the optional inputs (acceptable)", True)
        ok("T12 skeleton still builds", sk_m["total_positions"] == 4)

    raises("T13 refuses a position with no name", BriefingError,
           lambda: tier_positions([pos("")]))

    # --- counts -------------------------------------------------------------------------
    counts, total = counts_by_tier(rows)
    ok("T14 counts reconcile with the position list", total == 4, str(total))
    ok("T15 counts are per tier, not a single number", isinstance(counts, dict))

    sk = build_skeleton(rows, {"supplier": "Meridian", "contract_type": "MSA",
                               "contract_value": "$3.2M", "version": "1"})
    ok("T16 the skeleton states the total", sk["total_positions"] == 4)
    ok("T17 the tier table omits empty tiers rather than showing zeros",
       all(r["count"] > 0 for r in sk["tier_table"]))
    ok("T18 the position map carries MSA coverage per position",
       all("msa_coverage" in r for r in sk["position_map"]))
    ok("T19 binding is visible in the position map",
       any(r["binding"] for r in sk["position_map"]))

    # --- the split: prose is left to the model ------------------------------------------
    ok("T20 the skeleton carries prose placeholders, not prose",
       len(sk["prose_sections"]) == len(PROSE_SECTIONS))
    ok("T21 every placeholder is marked so an unfilled briefing is obviously unfinished",
       all(s["placeholder"].startswith("[PROSE:") for s in sk["prose_sections"]))
    ok("T22 the generator writes NO negotiation argument itself",
       not any(len(s["placeholder"]) > 400 for s in sk["prose_sections"]))
    for heading in ("Executive Summary", "Leverage Map", "Fallback Sequencing",
                    "Predicted Supplier Pushback"):
        ok("T23 prose section present: %-28s" % heading,
           any(s["heading"] == heading for s in sk["prose_sections"]))

    # --- docx --------------------------------------------------------------------------
    if _DOCX:
        tmp = tempfile.mkdtemp(prefix="f9_brief_")
        path = write_docx(sk, os.path.join(tmp, "briefing.docx"))
        ok("T24 briefing docx written", os.path.isfile(path))
        d = Document(path)
        text = "\n".join(p.text for p in d.paragraphs)
        heads = [p.text for p in d.paragraphs if p.style.name.startswith("Heading")]
        ok("T25 the docx carries the tier and position-map headings",
           any("Position Map" in h for h in heads) and any("Tier" in h for h in heads),
           repr(heads[:4]))
        ok("T26 the docx carries the prose placeholders", "[PROSE:" in text)
        ok("T27 the docx has both tables", len(d.tables) >= 2, str(len(d.tables)))
        ok("T28 the total is stated in the document", "Total positions: 4" in text)
        binding_cells = [c.text for t in d.tables for r in t.rows for c in r.cells]
        ok("T29 BINDING is visible to the reader, not only in the data",
           "BINDING" in binding_cells)
    else:
        print("  skip  T24-T29 (python-docx unavailable)")

    print("=" * 84)
    print("SUMMARY: %d/%d passed, %d failed" % (len(PASS), len(PASS) + len(FAIL), len(FAIL)))
    print("=" * 84)
    return 1 if FAIL else 0


if __name__ == "__main__":
    sys.exit(run())
